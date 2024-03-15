import logging
import os

from PIL import Image
from docx import Document
from docx.shared import Inches

from whispervid.support_file import clean_and_format_string

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def create_docx_for_failed_video(titles, error_message, save_directory):
    """
    Create a DOCX file with the given title in the specified directory.
    The function checks if a file with the same name already exists,
    and if not, it creates a new DOCX file containing the title.

    Parameters:
    titles (str): The title of the video.
    error_message (str): Error message to be included in the document (currently unused).
    save_directory (str): Directory path where the DOCX file will be saved.
    """
    sanitized_title = clean_and_format_string(titles)
    filename = os.path.join(save_directory, f'{sanitized_title}.docx')

    if not os.path.exists(filename):
        doc = Document()
        doc.add_paragraph(titles)
        # doc.add_paragraph(error_message)  # Uncomment if needed
        doc.save(filename)
        logging.info(f"Document '{filename}' created.")
    else:
        logging.info(f"Document '{filename}' already exists. Skipping.")


def calculate_image_dimensions(image_path, max_width, max_height):
    """
    Calculate new dimensions for an image, maintaining the aspect ratio,
    with the constraint that the image width and height do not exceed
    the specified max_width and max_height.

    Parameters:
    image_path (str): Path of the image file.
    max_width (int): Maximum allowed width.
    max_height (int): Maximum allowed height.

    Returns:
    Tuple (float, float): New width and height of the image.
    """
    with Image.open(image_path) as img:
        width, height = img.size

    width_ratio = max_width / width
    height_ratio = max_height / height
    scale_factor = min(width_ratio, height_ratio)

    return width * scale_factor, height * scale_factor

def add_images_to_word_document(df, word_file_name, max_width, max_height, video_urls, num_images=10):
    """
    Add images and their corresponding transcriptions to a Word document.
    If an image fails to be added, it logs an error and continues with the next image.

    Parameters:
    df (DataFrame): A DataFrame containing image paths and transcriptions.
    word_file_name (str): The name of the Word document to be saved.
    max_width (int): Maximum width for the images in the document.
    max_height (int): Maximum height for the images in the document.
    video_urls (str): URL of the video from which the images are extracted.
    num_images (int): Number of images to add (default is 10).
    """
    logging.info('Saving all transcriptions and images in a Word document.')

    new_width, new_height = calculate_image_dimensions(df.iloc[0]['image_path'], max_width, max_height)
    doc = Document()
    doc.add_paragraph('The video was downloaded from the following URL:')
    doc.add_paragraph(video_urls)
    doc.add_paragraph("\n")

    for index, row in df.iterrows():
        try:
            image_path = row['image_path']
            doc.add_picture(image_path, width=Inches(new_width), height=Inches(new_height))
        except Exception as e:
            logging.error(f"An error occurred while adding image: {e}")
            continue

        doc.add_paragraph(row['Transcription']['text'])
        doc.add_paragraph("\n")

    doc.save(word_file_name)
    logging.info(f'Completed saving all transcriptions and images in the Word document.')
    logging.info(f'The file is saved as {word_file_name}')