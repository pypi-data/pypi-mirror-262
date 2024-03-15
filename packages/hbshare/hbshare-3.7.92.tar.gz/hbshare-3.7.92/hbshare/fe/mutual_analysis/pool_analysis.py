import datetime
import pandas as pd
import numpy as np
from hbshare.fe.XZ import db_engine
from hbshare.fe.XZ import functionality
from hbshare.fe.mutual_analysis import  jj_picturing as jjpic
from ortools.linear_solver import  pywraplp
from PIL import Image
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import statsmodels.api as sm




util=functionality.Untils()
hbdb=db_engine.HBDB()
localdb=db_engine.PrvFunDB().engine
plot=functionality.Plot(1000,500)



def get_core_pool():

    sql="select jjdm,jjjc from st_fund.t_st_gm_jjxx where hxbz='1' and cpfl='2' "
    pool=hbdb.db2df(sql,db='funduser')

    return  pool

def pool_his_fromexcel2db():

    df_in=pd.read_excel(r"C:\Users\xuhuai.zhe\Downloads\公募股多核心池20220429.xlsx",
                     sheet_name='核心池-主动股多').sort_values('调入时间').drop(['备注','研究员'],axis=1)
    df_out=pd.read_excel(r"C:\Users\xuhuai.zhe\Downloads\公募股多核心池20220429.xlsx",
                     sheet_name='出池').sort_values('调出时间').drop('备注',axis=1)
    df_in['调入时间']=df_in['调入时间'].astype(str).str[0:7].str.replace('-', '')
    df_in['调出时间']=''
    df_out['调出时间'] = df_out['调出时间'].astype(str).str[0:7].str.replace('-', '')
    df_out['调入时间'] = df_out['调入时间'].astype(str).str[0:7].str.replace('-', '')
    df=pd.concat([df_in,df_out],axis=0).sort_values('调入时间')
    df.reset_index(drop=True,inplace=True)
    df['基金代码']=df['基金代码'].astype(str).str[0:6]

    poo_his=pd.DataFrame()
    for year in ['2020','2021','2022']:
        for month in ['03','06','09','12']:
            tempdf=df[(df['调入时间']<=year+month)&((df['调出时间']=='')|(df['调出时间']>year+month))][['基金代码','基金名称','基金经理']]
            tempdf['asofdate']=year+month
            poo_his=pd.concat([poo_his,tempdf],axis=0)

    poo_his.reset_index(drop=True,inplace=True)
    poo_his=poo_his[poo_his['asofdate']<=str(datetime.datetime.now())[0:7].replace('-','')]
    sql='delete from core_pool_history'
    localdb.execute(sql)
    poo_his.to_sql('core_pool_history',con=localdb,if_exists='append',index=False)

def get_core_poolfromlocaldb(asofdate=None):

    if(asofdate is not None):
        sql="select * from core_pool_history where asofdate='{}'".format(asofdate)
    else:
        sql="select * from core_pool_history"
    pool=pd.read_sql(sql,con=localdb).rename(columns={'基金代码':'jjdm',
                                                      '基金名称':'jjjc',
                                                      '基金经理':'manager'})

    return pool

def get_jj_picture(jjdm_con=None,if_inverse=False):

    if(if_inverse):
        inornot='not in '
    else:
        inornot='in'

    value_df_list = [pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(),
                     pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()]

    industry_df_list = [pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()]


    stock_df_list = [pd.DataFrame(), pd.DataFrame()]


    if(jjdm_con is not None):
        value_p, value_p_hbs, value_sp, value_sp_hbs, size_p, size_p_hbs, size_sp, size_sp_hbs, \
        industry_p, industry_sp, theme_p, theme_sp, industry_detail_df_list, stock_p, stock_tp=\
            jjpic.get_pic_from_localdb("jjdm {1} ({0})".format(jjdm_con,inornot), if_percentage=False)
    else:
        value_p, value_p_hbs, value_sp, value_sp_hbs, size_p, size_p_hbs, size_sp, size_sp_hbs, \
        industry_p, industry_sp, theme_p, theme_sp, industry_detail_df_list, stock_p, stock_tp=\
            jjpic.get_pic_from_localdb("1=1", if_percentage=False)

    value_df_list[0]=value_p
    value_df_list[1] = value_p_hbs
    value_df_list[2] = value_sp
    value_df_list[3] = value_sp_hbs
    value_df_list[4] = size_p
    value_df_list[5] = size_p_hbs
    value_df_list[6] = size_sp
    value_df_list[7] = size_sp_hbs

    industry_df_list[0]=industry_p
    industry_df_list[1] = industry_sp
    industry_df_list[2] = theme_p
    industry_df_list[3] = theme_sp
    for j in range(len(industry_detail_df_list)):
        # industry_detail_df_list[j]['jjdm'] = jjdm
        industry_detail_df_list[j]['industry_level'] = j + 1
        industry_df_list[4] = pd.concat([industry_df_list[4], industry_detail_df_list[j]], axis=0)

    stock_df_list[0]=stock_p
    stock_df_list[1] = stock_tp

    return value_df_list,industry_df_list,stock_df_list

def features_calculation(value_df_list, industry_df_list, stock_df_list):

    industry_style=industry_df_list[0].groupby('一级行业类型').count()['jjdm'].to_frame('num')
    industry_style=industry_style/len(industry_df_list[0])*100
    industry_style_dis = (industry_df_list[4].groupby(['industry_level', '行业名称']).mean() * 100).reset_index()[
        ['industry_level', '行业名称', '占持仓比例(时序均值)', '占持仓比例(时序均值)排名']]
    #take industry with weight more than 5%
    industry_style_dis=industry_style_dis[industry_style_dis['占持仓比例(时序均值)'] >= (5/industry_style_dis['industry_level'])]
    class_list=[]
    for i in range(3):
        class_list.append(industry_style_dis[industry_style_dis['industry_level'] == (i + 1)] \
            .sort_values('占持仓比例(时序均值)', ascending=False).set_index('行业名称').drop('industry_level', axis=1))

    style_type_dis = value_df_list[1].groupby('风格类型').count()['jjdm'].to_frame('风格类型') \
                     / len(value_df_list[1]) * 100
    style_incline_dis=value_df_list[1][value_df_list[1]['风格类型']=='专注'][['风格类型','风格偏好']]
    style_incline_dis=style_incline_dis.groupby('风格偏好').count()/len(style_incline_dis)*100
    style_weight_dis=value_df_list[1][['成长绝对暴露(持仓)',
       '价值绝对暴露(持仓)','成长暴露排名(持仓)', '价值暴露排名(持仓)']].mean(axis=0).to_frame('成长价值权重分布')*100

    style_type_dis2 = value_df_list[5].groupby('规模风格类型').count()['jjdm'].to_frame('规模风格类型') \
                      / len(value_df_list[5]) * 100
    style_incline_dis2=value_df_list[5][value_df_list[5]['规模风格类型']=='专注'][['规模风格类型','规模偏好']]
    style_incline_dis2=style_incline_dis2.groupby('规模偏好').count()/len(style_incline_dis2)*100
    style_weight_dis2=value_df_list[5][['大盘绝对暴露(持仓)',
       '中盘绝对暴露(持仓)', '小盘绝对暴露(持仓)','大盘暴露排名(持仓)',
       '中盘暴露排名(持仓)', '小盘暴露排名(持仓)']].mean(axis=0).to_frame('成长价值权重分布')*100

    stock_style_a=stock_df_list[0].groupby('个股风格A').count()['jjdm'].to_frame('个股风格类型A')\
                  /len(stock_df_list[0])*100
    stock_style_b = stock_df_list[0].groupby('个股风格B').count()['jjdm'].to_frame('个股风格类型B')\
                    /len(stock_df_list[0])*100
    stock_left = stock_df_list[1].groupby('左侧标签').count()['jjdm'].to_frame('左侧类型分布')\
                 / len(stock_df_list[1]) * 100
    stock_df_list[1].loc[stock_df_list[1]['新股次新股偏好'] == '', '新股次新股偏好'] = '无'
    stock_new=stock_df_list[1].groupby('新股次新股偏好').count()['jjdm'].to_frame('新股偏好分布')\
              / len(stock_df_list[1]) * 100
    stock_fin2=(stock_df_list[1][['平均持有时间(出持仓前)_rank', '左侧概率(出重仓前,半年线)_rank','新股概率(出持仓前)_rank'
        ,'出重仓前平均收益率_rank',
       '出全仓前平均收益率_rank',]]).mean(axis=0)
    stock_fin=(stock_df_list[0][['PE_rank', 'PB_rank',
       'ROE_rank', '股息率_rank','平均仓位']]).mean(axis=0)
    stock_fin=pd.concat([stock_fin,stock_fin2],axis=0).to_frame('个股持仓特征')
    stock_fin.index = [x.replace('rank', '排名') for x in stock_fin.index]
    stock_fin=stock_fin*100


    return industry_style,class_list,style_type_dis,style_incline_dis,\
           style_weight_dis,style_type_dis2,style_incline_dis2,style_weight_dis2,\
           stock_style_a,stock_style_b,stock_left,stock_new,stock_fin

def pool_industry(industry_style,class_list,save_local_file=False):

    if(len(industry_style.columns)>1):
        plot.plotly_jjpic_bar(industry_style, '行业风格分布',save_local_file=save_local_file)
    else:
        plot.plotly_pie(industry_style, '行业风格分布',save_local_file=save_local_file)

    for i in range(3):

        class_list[i].columns=[x.replace('占持仓比例(时序均值)','') for x in class_list[i].columns]

        plot.plotly_jjpic_bar(class_list[i][class_list[i].columns[0:int(len(class_list[i].columns)/2)]].fillna(0),"{}级行业持仓分布_占持仓比例".format(i+1),save_local_file=save_local_file)
        plot.plotly_jjpic_bar(class_list[i][class_list[i].columns[int(len(class_list[i].columns)/2):]].fillna(0), "{}级行业持仓分布_占持仓比例排名".format(i + 1),save_local_file=save_local_file)

def pool_style(style_type_dis,style_incline_dis,
               style_weight_dis,style_type_dis2,style_incline_dis2,style_weight_dis2,save_local_file=False):

    if(len(style_type_dis.columns)>1):
        plot.plotly_jjpic_bar(style_type_dis, '风格类型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_incline_dis, '风格专注型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis.iloc[0:2], '成长价值持仓占比',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis.iloc[2:], '成长价值持仓占比排名',save_local_file=save_local_file)

        plot.plotly_jjpic_bar(style_type_dis2, '规模风格类型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_incline_dis2, '专注型风格分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis2.iloc[0:3], '大中小盘持仓占比',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis2.iloc[3:], '大中小盘持仓占比排名',save_local_file=save_local_file)
    else:
        plot.plotly_pie(style_type_dis,'风格类型分布',save_local_file=save_local_file)
        plot.plotly_pie(style_incline_dis,'风格专注型分布',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis.iloc[0:2],'成长价值持仓占比',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis.iloc[2:],'成长价值持仓占比排名',save_local_file=save_local_file)


        plot.plotly_pie(style_type_dis2,'规模风格类型分布',save_local_file=save_local_file)
        plot.plotly_pie(style_incline_dis2,'专注型风格分布',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis2.iloc[0:3],'大中小盘持仓占比',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis2.iloc[3:],'大中小盘持仓占比排名',save_local_file=save_local_file)

def pool_other(stock_style_a,stock_style_b,stock_left,stock_new,stock_fin,save_local_file=False):

    if(len(stock_style_a.columns)>1):
        plot.plotly_jjpic_bar(stock_style_a,'个股风格类型分布A',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(stock_style_b,'个股风格类型分布B',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(stock_left ,'左侧类型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(stock_new,'新股偏好分布',save_local_file=save_local_file)
    else:
        plot.plotly_pie(stock_style_a,'个股风格类型分布A',save_local_file=save_local_file)
        plot.plotly_pie(stock_style_b,'个股风格类型分布B',save_local_file=save_local_file)
        plot.plotly_pie(stock_left ,'左侧类型分布',save_local_file=save_local_file)
        plot.plotly_pie(stock_new,'新股偏好分布',save_local_file=save_local_file)
    plot.plotly_jjpic_bar(stock_fin,'个股持仓特征',save_local_file=save_local_file)

def pool_picturing_engine(industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin,save_local_file=False):

    pool_industry(industry_style,class_list,save_local_file)
    pool_style(style_type_dis,style_incline_dis,
               style_weight_dis,style_type_dis2,style_incline_dis2,style_weight_dis2,save_local_file)
    pool_other(stock_style_a,stock_style_b,stock_left,stock_new,stock_fin,save_local_file)


def get_potient_pool(jjdm_con):

    value_df_list,industry_df_list,stock_df_list=get_jj_picture(jjdm_con,if_inverse=True)

    return value_df_list,industry_df_list,stock_df_list

def optimizer_allocation(inputdf,target_col,con_ub,con_lb,turnover_num):



    solver=pywraplp.Solver.CreateSolver('SCIP')
    x=[]
    #create variables
    for i in range(len(inputdf)):
        x.append(solver.IntVar(0,1,"x_{}".format(i)))


    cons_list=inputdf.columns.tolist()
    cons_list.remove(target_col)
    #get the number of cons
    num_con=len(cons_list)


    con_ub=list(turnover_num*np.array(con_ub))
    con_lb=list(turnover_num*np.array(con_lb))


    # #set constrains bound
    cons_dict=dict()
    #set the turnover constraints
    cons_dict['turnover_num']=solver.Constraint(turnover_num, turnover_num)
    #set the other constraints
    for m in range(num_con):
        cons_dict['cons_{}'.format(cons_list[m])] = \
            solver.Constraint(con_lb[m], con_ub[m])

    for i in range(len(x)):
        cons_dict['turnover_num'].SetCoefficient(x[i], 1)
        for m in range(num_con):
            cons_dict['cons_{}'.format(cons_list[m])].SetCoefficient(x[i], inputdf[cons_list[m]].iloc[i])

    objective = solver.Objective()
    for i in range(len(x)):
        objective.SetCoefficient(x[i], inputdf[target_col].iloc[i])
    objective.SetMaximization()
    #
    result_status = solver.Solve()

    solution=[y.solution_value() for y in x]

    if(result_status == pywraplp.Solver.OPTIMAL):
        print("optimal solution is found and the optimal result is {}"
              .format(solver.Objective().Value()/turnover_num))


    return  solution

def optimizer_allocation_fromlast_holding(inputdf,target_col,allowed_hsl,con_ub,con_lb,turnover_num):


    static__num=turnover_num-np.floor(turnover_num*allowed_hsl)
    inputdf['old_sol']=inputdf['old_sol'].astype(float)

    solver=pywraplp.Solver.CreateSolver('SCIP')
    x=[]
    #create variables
    for i in range(len(inputdf)):
        x.append(solver.IntVar(0,1,"x_{}".format(i)))


    cons_list=inputdf.columns.tolist()
    cons_list.remove(target_col)
    cons_list.remove('old_sol')
    #get the number of cons
    num_con=len(cons_list)


    con_ub=list(turnover_num*np.array(con_ub))
    con_lb=list(turnover_num*np.array(con_lb))


    # #set constrains bound
    cons_dict=dict()
    #set the turnover constraints
    cons_dict['turnover_num']=solver.Constraint(turnover_num, turnover_num)
    # set the turnover ratio constrains
    cons_dict['turnover_ratio'] = solver.Constraint(static__num, turnover_num)
    #set the other constraints
    for m in range(num_con):
        cons_dict['cons_{}'.format(cons_list[m])] = \
            solver.Constraint(con_lb[m], con_ub[m])

    for i in range(len(x)):
        cons_dict['turnover_num'].SetCoefficient(x[i], 1)
        cons_dict['turnover_ratio'].SetCoefficient(x[i], inputdf['old_sol'].iloc[i])
        for m in range(num_con):
            cons_dict['cons_{}'.format(cons_list[m])].SetCoefficient(x[i], inputdf[cons_list[m]].iloc[i])

    objective = solver.Objective()
    for i in range(len(x)):
        objective.SetCoefficient(x[i], inputdf[target_col].iloc[i])
    objective.SetMaximization()
    #
    result_status = solver.Solve()

    solution=[y.solution_value() for y in x]

    if(result_status == pywraplp.Solver.OPTIMAL):
        print("optimal solution is found and the optimal result is {}"
              .format(solver.Objective().Value()/turnover_num))


    return  solution

def optimizer_portfoli_weight(inputdf,target_col,con_ub,con_lb,weight_range):



    solver=pywraplp.Solver.CreateSolver('SCIP')
    x=[]
    #create variables
    for i in range(len(inputdf)):
        x.append(solver.NumVar(weight_range[0],weight_range[1],"x_{}".format(i)))


    cons_list=inputdf.columns.tolist()
    cons_list.remove(target_col)
    #get the number of cons
    num_con=len(cons_list)


    con_ub=list(np.array(con_ub))
    con_lb=list(np.array(con_lb))


    # #set constrains bound
    cons_dict=dict()
    #set the turnover constraints
    cons_dict['total_weight']=solver.Constraint(1, 1)
    #set the other constraints
    for m in range(num_con):
        cons_dict['cons_{}'.format(cons_list[m])] = \
            solver.Constraint(con_lb[m], con_ub[m])

    for i in range(len(x)):
        cons_dict['total_weight'].SetCoefficient(x[i], 1)
        for m in range(num_con):
            cons_dict['cons_{}'.format(cons_list[m])].SetCoefficient(x[i], inputdf[cons_list[m]].iloc[i])

    objective = solver.Objective()
    for i in range(len(x)):
        objective.SetCoefficient(x[i], inputdf[target_col].iloc[i])
    objective.SetMaximization()
    #
    result_status = solver.Solve()

    solution=[y.solution_value() for y in x]

    if(result_status == pywraplp.Solver.OPTIMAL):
        print("optimal solution is found and the optimal result is {}"
              .format(solver.Objective().Value()))


    return  solution

def optimizer_portfoli_weight_fromlast_holding(inputdf,target_col,allowed_hsl,con_ub,con_lb,weight_range):



    solver=pywraplp.Solver.CreateSolver('SCIP')
    x=[]

    fund_num=len(inputdf)

    #create  buy variables
    for i in range(fund_num):
        x.append(solver.NumVar(0,1,"x_buy_{}".format(i)))
    # create sell variables
    for i in range(fund_num):
        x.append(solver.NumVar(0,1,"x_sell_{}".format(i)))


    cons_list=inputdf.columns.tolist()
    cons_list.remove(target_col)
    cons_list.remove('old_sol')
    #get the number of cons
    num_con=len(cons_list)

    #adjust the up and lower bound
    for i in range(num_con):
        con_ub[i]=con_ub[i]-sum(inputdf[cons_list[i]]*inputdf['old_sol'])
        con_lb[i] = con_lb[i] - sum(inputdf[cons_list[i]] * inputdf['old_sol'])


    con_ub=list(np.array(con_ub))
    con_lb=list(np.array(con_lb))


    # #set constrains bound
    cons_dict=dict()
    #set the turnover constraints
    cons_dict['total_weight']=solver.Constraint(0, allowed_hsl*2)
    #set the other constraints
    for m in range(num_con):
        cons_dict['cons_{}'.format(cons_list[m])] = \
            solver.Constraint(con_lb[m], con_ub[m])

    #set the variables constraints
    for j in range(fund_num):
        cons_dict['cons_v_{}'.format(j)] = \
            solver.Constraint(weight_range[0]-inputdf['old_sol'].iloc[j], weight_range[1]-inputdf['old_sol'].iloc[j])


    #set constraints coff
    for i in range(fund_num):
        cons_dict['total_weight'].SetCoefficient(x[i], 1)
        cons_dict['cons_v_{}'.format(i)].SetCoefficient(x[i], 1)
        for m in range(num_con):
            cons_dict['cons_{}'.format(cons_list[m])].SetCoefficient(x[i], inputdf[cons_list[m]].iloc[i])

    for i in range(fund_num):
        cons_dict['total_weight'].SetCoefficient(x[i+fund_num], 1)
        cons_dict['cons_v_{}'.format(i)].SetCoefficient(x[i+fund_num], -1)
        for m in range(num_con):
            cons_dict['cons_{}'.format(cons_list[m])].SetCoefficient(x[i+fund_num], -1*inputdf[cons_list[m]].iloc[i])


    objective = solver.Objective()
    for i in range(len(inputdf)):
        objective.SetCoefficient(x[i], inputdf[target_col].iloc[i])
    for i in range(len(inputdf)):
        objective.SetCoefficient(x[i+fund_num], -1*inputdf[target_col].iloc[i])

    objective.SetMaximization()
    #
    result_status = solver.Solve()

    solution=[y.solution_value() for y in x]
    real_solution=inputdf['old_sol']+solution[0:fund_num]-solution[fund_num:]

    if(result_status == pywraplp.Solver.OPTIMAL):
        print("optimal solution is found and the optimal result is {}"
              .format(solver.Objective().Value()))


    return  real_solution.values.tolist()

def pool_picturing_opt():

    pool = get_core_pool()
    pool=get_core_poolfromlocaldb('202203')
    #picture the original pool
    jjdm_con = util.list_sql_condition(pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)
    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    pool_picturing_engine(industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin)

    # new_jjdm_list=optimizer_allocation(value_df_list, industry_df_list, stock_df_list,pool,industry_style, class_list, style_type_dis, style_incline_dis, \
    # style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    # stock_style_a, stock_style_b, stock_left, stock_new, stock_fin)

    #pic the new pool
    #new_jjdm_con=util.list_sql_condition(new_jjdm_list)
    new_jjdm_con="'001705','001410','001856','450009','166006'," \
                 "'001975','003095','005241','161606','005827','002708'," \
                 "'005968','005267','519133','002340','000577','163415'," \
                 "'000729','000756','000991','001208','001487','001532'," \
                 "'001718','001877','002258','002871','002943','004350'," \
                 "'004784','005161','005457','005669','005825','006257','007130','519702'"
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(new_jjdm_con)
    industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    industry_style_new=pd.merge(industry_style,industry_style_new,how='outer',left_index=True,right_index=True)
    industry_style_new.columns=['旧池','新池']

    for i in range(3):
        class_list_new[i]=pd.merge(class_list[i],class_list_new[i],how='outer',
                                   left_index=True,right_index=True)
        class_list_new[i].columns=[x.replace('x','旧') for x in class_list_new[i].columns]
        class_list_new[i].columns = [x.replace('y', '新') for x in class_list_new[i].columns]
        class_list_new[i]=class_list_new[i][class_list_new[i].columns.sort_values()]


    style_type_dis_new=pd.merge(style_type_dis,style_type_dis_new,how='outer',
                                left_index=True,right_index=True).fillna(0)
    style_type_dis_new.columns=['旧池','新池']

    style_incline_dis_new=pd.merge(style_incline_dis,style_incline_dis_new,how='outer',
                                   left_index=True,right_index=True).fillna(0)
    style_incline_dis_new.columns=['旧池','新池']

    style_weight_dis_new=pd.merge(style_weight_dis,style_weight_dis_new,how='outer',
                                  left_index=True,right_index=True).fillna(0)
    style_weight_dis_new.columns=['旧池','新池']

    style_type_dis2_new=pd.merge(style_type_dis2,style_type_dis2_new,how='outer',
                                 left_index=True,right_index=True).fillna(0)
    style_type_dis2_new.columns=['旧池','新池']

    style_incline_dis2_new=pd.merge(style_incline_dis2,style_incline_dis2_new,
                                    how='outer',left_index=True,right_index=True).fillna(0)
    style_incline_dis2_new.columns=['旧池','新池']

    style_weight_dis2_new=pd.merge(style_weight_dis2,style_weight_dis2_new,
                                   how='outer',left_index=True,right_index=True).fillna(0)
    style_weight_dis2_new.columns=['旧池','新池']

    stock_style_a_new=pd.merge(stock_style_a,stock_style_a_new,
                               how='outer',left_index=True,right_index=True).fillna(0)
    stock_style_a_new.columns=['旧池','新池']

    stock_style_b_new=pd.merge(stock_style_b,stock_style_b_new,how='outer'
                               ,left_index=True,right_index=True).fillna(0)
    stock_style_b_new.columns=['旧池','新池']

    stock_left_new=pd.merge(stock_left,stock_left_new,how='outer'
                            ,left_index=True,right_index=True).fillna(0)
    stock_left_new.columns=['旧池','新池']

    stock_new_new=pd.merge(stock_new,stock_new_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_new_new.columns=['旧池','新池']

    stock_fin_new=pd.merge(stock_fin,stock_fin_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_fin_new.columns=['旧池','新池']


    pool_picturing_engine(industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new)

def pool_picturing(pool=None,returndata=False,save_local_file=False):

    if(pool is None):
        pool=get_core_poolfromlocaldb('202203')

    #picture the original pool
    jjdm_con = util.list_sql_condition(pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)
    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    plot.plotly_table(pool, 1000, '池列表', save_local_file=True)
    pool_picturing_engine(industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin,save_local_file)

    if(returndata):
        return industry_style, class_list, style_type_dis, style_incline_dis, \
        style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
        stock_style_a, stock_style_b, stock_left, stock_new, stock_fin

def pool_comparision(date1,date2):

    pool=get_core_poolfromlocaldb(date1)

    #picture the original pool
    jjdm_con = util.list_sql_condition(pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)
    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    # picture the new pool
    new_pool = get_core_poolfromlocaldb(date2)
    new_jjdm_con=util.list_sql_condition(new_pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(new_jjdm_con)
    industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    #get the return for in and out jj for the next 1 quarter and year
    in_jjdm=list(set(new_pool['jjdm'].unique().tolist()).difference(set(pool['jjdm'].unique().tolist())))
    out_jjdm=list(set(pool['jjdm'].unique().tolist()).difference(set(new_pool['jjdm'].unique().tolist())))
    start_date=util._shift_date(date2+'31')
    in_ret=pd.DataFrame(data=in_jjdm,columns=['jjdm'])
    out_ret = pd.DataFrame(data=out_jjdm, columns=['jjdm'])

    end_date_y=util._shift_date(str(int(date2[0:4])+1)+date2[4:]+'30')
    if(date2[4:]!='12'):
        end_date_q=util._shift_date(date2[0:4]+("0"+str(int(date2[4:])+3))[-2:]+'30')
    else:
        end_date_q=util._shift_date(str(int(date2[0:4])+1)+'0331')

    if(end_date_y>=datetime.datetime.today().strftime('%Y%m%d')[0:6]):
        end_date_y=util._shift_date(str(hbdb.db2df("select max(jzrq) as jzrq from st_fund.t_st_gm_rhb", db='funduser')['jzrq'][0]))
    sql = "select jjdm,jzrq,fqdwjz from st_fund.t_st_gm_rhb where (jzrq='{0}' or  jzrq='{1}') and jjdm in({2})" \
        .format(end_date_y, start_date, util.list_sql_condition(in_jjdm+out_jjdm))
    gap_return_y = hbdb.db2df(sql, db='funduser')
    gap_return_y=pd.merge(gap_return_y[gap_return_y['jzrq']==int(start_date)]
                          ,gap_return_y[gap_return_y['jzrq']==int(end_date_y)],on='jjdm')
    gap_return_y['ret']=gap_return_y['fqdwjz_y']/gap_return_y['fqdwjz_x']-1


    if(end_date_q>=datetime.datetime.today().strftime('%Y%m%d')[0:6]):
        end_date_q = util._shift_date(str(hbdb.db2df("select max(jzrq) as jzrq from st_fund.t_st_gm_rhb", db='funduser')['jzrq'][0]))
    sql = "select jjdm,jzrq,fqdwjz from st_fund.t_st_gm_rhb where (jzrq='{0}' or  jzrq='{1}') and jjdm in({2})" \
        .format(end_date_q, start_date, util.list_sql_condition(in_jjdm+out_jjdm))
    gap_return_q = hbdb.db2df(sql, db='funduser')
    gap_return_q=pd.merge(gap_return_q[gap_return_q['jzrq']==int(start_date)]
                          ,gap_return_q[gap_return_q['jzrq']==int(end_date_q)],on='jjdm')
    gap_return_q['ret']=gap_return_q['fqdwjz_y']/gap_return_q['fqdwjz_x']-1

    in_ret=pd.merge(in_ret,gap_return_q[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_q'})
    in_ret=pd.merge(in_ret,gap_return_y[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_y'})
    in_ret=in_ret.mean()
    out_ret=pd.merge(out_ret,gap_return_q[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_q'})
    out_ret=pd.merge(out_ret,gap_return_y[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_y'})
    out_ret = out_ret.mean()

    #join the two pool picture
    industry_style_new=pd.merge(industry_style,industry_style_new,how='outer',left_index=True,right_index=True)
    industry_style_new.columns=['旧池','新池']

    for i in range(3):
        class_list_new[i]=pd.merge(class_list[i],class_list_new[i],how='outer',
                                   left_index=True,right_index=True)
        class_list_new[i].columns=[x.replace('x','旧') for x in class_list_new[i].columns]
        class_list_new[i].columns = [x.replace('y', '新') for x in class_list_new[i].columns]
        class_list_new[i]=class_list_new[i][class_list_new[i].columns.sort_values()]


    style_type_dis_new=pd.merge(style_type_dis,style_type_dis_new,how='outer',
                                left_index=True,right_index=True).fillna(0)
    style_type_dis_new.columns=['旧池','新池']

    style_incline_dis_new=pd.merge(style_incline_dis,style_incline_dis_new,how='outer',
                                   left_index=True,right_index=True).fillna(0)
    style_incline_dis_new.columns=['旧池','新池']

    style_weight_dis_new=pd.merge(style_weight_dis,style_weight_dis_new,how='outer',
                                  left_index=True,right_index=True).fillna(0)
    style_weight_dis_new.columns=['旧池','新池']

    style_type_dis2_new=pd.merge(style_type_dis2,style_type_dis2_new,how='outer',
                                 left_index=True,right_index=True).fillna(0)
    style_type_dis2_new.columns=['旧池','新池']

    style_incline_dis2_new=pd.merge(style_incline_dis2,style_incline_dis2_new,
                                    how='outer',left_index=True,right_index=True).fillna(0)
    style_incline_dis2_new.columns=['旧池','新池']

    style_weight_dis2_new=pd.merge(style_weight_dis2,style_weight_dis2_new,
                                   how='outer',left_index=True,right_index=True).fillna(0)
    style_weight_dis2_new.columns=['旧池','新池']

    stock_style_a_new=pd.merge(stock_style_a,stock_style_a_new,
                               how='outer',left_index=True,right_index=True).fillna(0)
    stock_style_a_new.columns=['旧池','新池']

    stock_style_b_new=pd.merge(stock_style_b,stock_style_b_new,how='outer'
                               ,left_index=True,right_index=True).fillna(0)
    stock_style_b_new.columns=['旧池','新池']

    stock_left_new=pd.merge(stock_left,stock_left_new,how='outer'
                            ,left_index=True,right_index=True).fillna(0)
    stock_left_new.columns=['旧池','新池']

    stock_new_new=pd.merge(stock_new,stock_new_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_new_new.columns=['旧池','新池']

    stock_fin_new=pd.merge(stock_fin,stock_fin_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_fin_new.columns=['旧池','新池']



    pool_picturing_engine(industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new)


    return  in_ret,out_ret

def pic_pool_history(returndata=False,save_local_file=False):

    pool_his=get_core_poolfromlocaldb()

    jjdm_con = util.list_sql_condition(pool_his['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)

    industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new=pd.DataFrame(),[pd.DataFrame(),pd.DataFrame(),pd.DataFrame()],\
                                                                                       pd.DataFrame(),pd.DataFrame(),\
                                                                                       pd.DataFrame(),pd.DataFrame(),pd.DataFrame(),\
                                                                                       pd.DataFrame(),pd.DataFrame(),pd.DataFrame(),\
                                                                                       pd.DataFrame(),pd.DataFrame(),pd.DataFrame()

    new_col_list = []
    for asofdate in pool_his['asofdate'].unique():
        jjdm_df=pool_his[pool_his['asofdate']==asofdate]['jjdm'].to_frame('jjdm_x')
        new_col_list.append(str(asofdate))
        value_df_list_temp=[]
        industry_df_list_temp=[]
        stock_df_list_temp=[]
        for i in range(len(value_df_list)):
            value_df_list_temp.append(pd.merge(jjdm_df,value_df_list[i],
                                           how='left',left_on='jjdm_x',right_on='jjdm').drop('jjdm_x',axis=1))
        for i in range(len(industry_df_list)):
            industry_df_list_temp.append(pd.merge(jjdm_df,industry_df_list[i],
                                           how='left',left_on='jjdm_x',right_on='jjdm').drop('jjdm_x',axis=1))
        for i in range(len(stock_df_list)):
            stock_df_list_temp.append(pd.merge(jjdm_df,stock_df_list[i],
                                           how='left',left_on='jjdm_x',right_on='jjdm').drop('jjdm_x',axis=1))

        industry_style, class_list, style_type_dis, style_incline_dis, \
        style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
        stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
            =features_calculation(value_df_list_temp, industry_df_list_temp, stock_df_list_temp)

        industry_style_new = pd.merge(industry_style_new,industry_style,  how='outer', left_index=True,
                                      right_index=True)


        for i in range(3):

            class_list[i].columns = [x+"_"+str(asofdate) for x in class_list[i].columns]
            class_list_new[i] = pd.merge( class_list_new[i],class_list[i], how='outer',
                                         left_index=True, right_index=True)
            class_list_new[i] = class_list_new[i][class_list_new[i].columns.sort_values()]

        style_type_dis_new = pd.merge( style_type_dis_new,style_type_dis, how='outer',
                                      left_index=True, right_index=True).fillna(0)


        style_incline_dis_new = pd.merge(style_incline_dis_new,style_incline_dis, how='outer',
                                         left_index=True, right_index=True).fillna(0)


        style_weight_dis_new = pd.merge( style_weight_dis_new,style_weight_dis, how='outer',
                                        left_index=True, right_index=True).fillna(0)


        style_type_dis2_new = pd.merge(style_type_dis2_new,style_type_dis2,  how='outer',
                                       left_index=True, right_index=True).fillna(0)


        style_incline_dis2_new = pd.merge( style_incline_dis2_new,style_incline_dis2,
                                          how='outer', left_index=True, right_index=True).fillna(0)


        style_weight_dis2_new = pd.merge(style_weight_dis2_new,style_weight_dis2,
                                         how='outer', left_index=True, right_index=True).fillna(0)


        stock_style_a_new = pd.merge(stock_style_a_new,stock_style_a,
                                     how='outer', left_index=True, right_index=True).fillna(0)


        stock_style_b_new = pd.merge( stock_style_b_new,stock_style_b, how='outer'
                                     , left_index=True, right_index=True).fillna(0)


        stock_left_new = pd.merge( stock_left_new,stock_left, how='outer'
                                  , left_index=True, right_index=True).fillna(0)


        stock_new_new = pd.merge(stock_new_new, stock_new, how='outer',
                                 left_index=True, right_index=True).fillna(0)


        stock_fin_new = pd.merge(stock_fin_new,stock_fin,  how='outer',
                                 left_index=True, right_index=True).fillna(0)


    industry_style_new.columns = new_col_list
    style_type_dis_new.columns = new_col_list
    style_incline_dis_new.columns = new_col_list
    style_weight_dis_new.columns = new_col_list
    style_type_dis2_new.columns = new_col_list
    style_incline_dis2_new.columns =new_col_list
    style_weight_dis2_new.columns = new_col_list
    stock_style_a_new.columns = new_col_list
    stock_style_b_new.columns = new_col_list
    stock_left_new.columns = new_col_list
    stock_new_new.columns = new_col_list
    stock_fin_new.columns = new_col_list


    pool_picturing_engine(industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new,save_local_file=save_local_file)


    if(returndata):
        return industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new

def get_pool_history_report():
    
    def style_analysis(df, additional_str1='',additional_str2=''):
        style_style = "从基金{0}来看，其中:".format(additional_str1)
        for style in df.index:
            style_style += "{0}，".format(str(style)+additional_str2)
            style_style+="历史最高为{0}，历史最低为{1}。最新一期为{2}，最早一期为{3}"\
                .format((str(round(df.loc[style].max(), 2)) + "%,"),(str(round(df.loc[style].min(), 2)) + "%,"),
                        (str(round(df.loc[style].iloc[-1], 2)) + "%,"),(str(round(df.loc[style].iloc[0], 2)) + "%,"))

        return style_style

    pic_name_list = ['行业风格分布.png', '1级行业持仓分布_占持仓比例.png',
                     '1级行业持仓分布_占持仓比例排名.png',
                     '2级行业持仓分布_占持仓比例.png',
                     '2级行业持仓分布_占持仓比例排名.png',
                     '3级行业持仓分布_占持仓比例.png',
                     '3级行业持仓分布_占持仓比例排名.png',
                     '风格类型分布.png',
                     '风格专注型分布.png',
                     '成长价值持仓占比.png',
                     '成长价值持仓占比排名.png',
                     '规模风格类型分布.png',
                     '专注型风格分布.png',
                     '大中小盘持仓占比.png',
                     '大中小盘持仓占比排名.png',
                     '个股风格类型分布A.png',
                     '个股风格类型分布B.png',
                     '左侧类型分布.png',
                     '新股偏好分布.png',
                     '个股持仓特征.png',
                     ]

    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin =\
        pic_pool_history(returndata=True,save_local_file=True)


    time_length=int(len(class_list[0].columns)/2)
    report_date=class_list[0].columns[0:time_length]
    report_date_rank=class_list[0].columns[time_length:]

    data_list=[industry_style , class_list[0][report_date],class_list[0][report_date_rank],
               class_list[1][report_date],class_list[1][report_date_rank],
               class_list[2][report_date],class_list[2][report_date_rank], style_type_dis , style_incline_dis , \
    style_weight_dis[0:2] ,style_weight_dis.iloc[2:], style_type_dis2 , style_incline_dis2 , style_weight_dis2[0:3] ,style_weight_dis2[3:], \
    stock_style_a , stock_style_b , stock_left , stock_new , stock_fin]
    key_word_list=['行业配置类型','一级行业分布','一级行业分布排名','二级行业分布','二级行业分布排名','三级行业分布','三级行业分布排名',
                   '风格类型','中属于风格专注的基金','持仓中的风格占比','持仓中的风格占比排名','规模风格类型','中属于风格专注的基金','持仓中的规模风格占比','持仓中的规模风格占比排名',
                   '个股配置风格','个股策略','个股左侧特征','对次新股的偏好分布','其他个股持仓特征']
    key_word_list2 = ['型占比', '行业占比', '行业占比相对排名', '行业占比', '行业占比相对排名',
                      '行业占比', '行业占比相对排名','型占比', '型占比','占比','排名','型占比','型占比',
                      '占比','排名','型占比','型占比','型占比','型占比','']

    doc = Document()
    fold = r"E:\GitFolder\hbshare\fe\mutual_analysis"

    title_num=2
    name=''
    pic_paragraphs_list = []
    string_paragraphs_list = []
    paragraphs_count = 0

    doc.add_paragraph("")
    paragraphs_count += 1

    p = doc.add_paragraph("按照季度，对公墓核心池历史进行了统计分析，结果如下")
    p.style.font.size = Pt(10)
    p.paragraph_format.first_line_indent = p.style.font.size * 2
    string_paragraphs_list.append(paragraphs_count)
    paragraphs_count += 1

    for i in range(len(pic_name_list)):
        pic=pic_name_list[i]
        file_path = fold + "\\" + pic
        desc = style_analysis(data_list[i].fillna(0), key_word_list[i],key_word_list2[i])

        doc.add_paragraph("")
        paragraphs_count += 1

        try:
            doc.add_picture(file_path, width=Inches(6), height=Inches(4))
            pic_paragraphs_list.append(paragraphs_count)
            paragraphs_count += 1

        except Exception as e:
            pic_temp = Image.open(file_path)
            pic_temp.save(pic_temp)
            doc.add_picture(file_path, width=Inches(6), height=Inches(4))

        p = doc.add_paragraph(desc)
        p.style.font.size = Pt(10)
        p.paragraph_format.first_line_indent = p.style.font.size * 2
        string_paragraphs_list.append(paragraphs_count)
        paragraphs_count += 1

    for j in pic_paragraphs_list:
        doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for j in string_paragraphs_list:
        doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.save(r"E:\GitFolder\hbshare\fe\mutual_analysis\公募核心池历史报告.docx")
    print('')

def calculate_ir_based_on_index(jjdm_list, start_date, end_date, index_code):

    sql="select jjdm,jzrq,hbdr from st_fund.t_st_gm_rhb where jjdm in ({0}) and jzrq>='{1}' and jzrq<='{2}'"\
        .format(util.list_sql_condition(jjdm_list),start_date,end_date)
    jj_ret=hbdb.db2df(sql,db='funduser')

    sql="select zqdm,jyrq,hbdr from st_market.t_st_zs_rhb where zqdm='{0}' and jyrq>='{1}' and jyrq<='{2}' "\
        .format(index_code,start_date,end_date)
    index_ret=hbdb.db2df(sql,db='alluser')

    ir_df=pd.merge(index_ret,jj_ret,how='left',
                   left_on='jyrq',right_on='jzrq')
    ir_df['extra_ret']=ir_df['hbdr_y']-ir_df['hbdr_x']
    ir_df= \
        (ir_df.groupby('jjdm').mean()['extra_ret']/
         ir_df.groupby('jjdm').std()['extra_ret']).to_frame('info_ratio').reset_index()


    return ir_df


class Poolvs885001:


    @staticmethod
    def get_fund_current_holding_features(jjdm_List,asofdate,if_885001):


        theme_col = ['大金融', '消费', 'TMT', '周期', '制造']
        theme_map = dict(zip(theme_col,
                             [['银行', '非银金融', '房地产'],
                              ['食品饮料', '家用电器', '医药生物', '社会服务', '农林牧渔', '商贸零售', '美容护理'],
                              ['通信', '计算机', '电子', '传媒', '国防军工'],
                              ['钢铁', '有色金属', '建筑装饰', '建筑材料', '基础化工', '石油石化', '煤炭', '环保', '公用事业'],
                              ['交通运输', '机械设备', '汽车', '纺织服饰', '轻工制造', '电力设备']
                              ]
                             ))
        lista = []
        listb = []
        for theme in theme_col:
            for col in theme_map[theme]:
                lista.append(col)
                listb.append(theme)
        ind2thememap = pd.DataFrame()
        ind2thememap['industry_name'] = lista
        ind2thememap['theme'] = listb

        if(if_885001):
            table='885001_'
        else:
            table=''

        sql="SELECT * from {2}hbs_industry_class1_exp where jjdm in ({0}) and  jsrq='{1}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate,table)
        industry_date=pd.read_sql(sql,con=localdb)

        jj_size_data=industry_date.drop_duplicates(['jjdm'])[['jjdm','jjjzc']]
        jj_size_data['jjjzc']=jj_size_data['jjjzc']/jj_size_data['jjjzc'].sum()

        industry_date=pd.merge(industry_date,ind2thememap
                               ,how='left',left_on='yjxymc',right_on='industry_name')

        theme_date=industry_date.groupby(['jjdm','theme']).sum()['zjbl'].reset_index().pivot_table('zjbl','jjdm','theme').fillna(0)
        industry_date=industry_date.pivot_table('zjbl','jjdm','yjxymc').fillna(0)


        sql="SELECT * from {2}hbs_industry_class2_exp where jjdm in ({0}) and  jsrq='{1}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate,table)
        industry2_date=pd.read_sql(sql,con=localdb)
        industry2_date = industry2_date.pivot_table('zjbl', 'jjdm', 'ejxymc').fillna(0)


        theme_list=theme_date.columns.tolist()
        industry_list=industry_date.columns.tolist()
        industry_list2 = industry2_date.columns.tolist()
        # industry_list3 = industry3_date.columns.tolist()
        industry_date['total'] = industry_date.sum(axis=1)
        industry2_date['total'] = industry2_date.sum(axis=1)
        # industry3_date['total'] = industry3_date.sum(axis=1)
        theme_date['total'] = theme_date.sum(axis=1)

        #read fund style,size,industry,theme exp history from local db
        sql="SELECT * from {2}hbs_style_exp where jjdm in ({0}) and  jsrq='{1}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate,table)
        style_data=pd.read_sql(sql,con=localdb).pivot_table('zjbl','jjdm','style_type').fillna(0)
        style_data['total']=style_data['价值']+style_data['成长']
        style_data=pd.merge(style_data,industry_date['total'].to_frame('know_weight')
                            ,left_index=True,right_index=True)


        sql="SELECT * from {2}hbs_size_exp where jjdm in ({0}) and  jsrq='{1}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate,table)
        size_data=pd.read_sql(sql,con=localdb).pivot_table('zjbl','jjdm','size_type').fillna(0)
        size_col=size_data.columns.tolist()
        size_data['total']=size_data.sum(axis=1)
        size_data=pd.merge(size_data,industry_date['total'].to_frame('know_weight')
                            ,left_index=True,right_index=True)


        #if want zsbl then run the below code; if want zjbl skip the below code
        # for col in industry_list:
        #     industry_date[col]=industry_date[col]/industry_date['total']
        # for col in industry_list2:
        #     industry2_date[col]=industry2_date[col]/industry2_date['total']
        # for col in theme_list:
        #     theme_date[col]=theme_date[col]/theme_date['total']


        for col in ['价值','成长']:
            style_data[col]=style_data[col]*(style_data['know_weight']/style_data['total'])
        for col in size_col:
            size_data[col]=size_data[col]*(size_data['know_weight']/size_data['total'])


        cons_df=style_data.copy().drop(['total','know_weight'],axis=1)
        cons_df=pd.merge(cons_df,size_data.drop(['total','know_weight'],axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,industry_date.drop('total',axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,industry2_date.drop('total',axis=1),how='left',on='jjdm')
        # cons_df=pd.merge(cons_df,industry3_date.drop('total',axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,theme_date.drop('total',axis=1),how='left',on='jjdm')
        cons_df.columns=[x+'_当期' for x in cons_df.columns]

        cons_df=pd.merge(cons_df,jj_size_data,how='left',on='jjdm')

        return cons_df

    def pool_features_summary(self,jjdm_List,asofdate,if_885001=False,weight_average=False):

        current_holding_features=\
            self.get_fund_current_holding_features(jjdm_List,asofdate,if_885001)

        col_list = current_holding_features.columns.tolist()
        col_list.remove('jjjzc')
        col_list.remove('jjdm')

        if(weight_average):

            for col in col_list:

                current_holding_features[col]=\
                    current_holding_features[col]*current_holding_features['jjjzc']

            summary_result=current_holding_features[col_list].sum(axis=0)
        else:
            summary_result = current_holding_features[col_list].mean(axis=0)


        return summary_result

class Customize_pool:

    def __init__(self):


        theme_col = ['大金融', '消费', 'TMT', '周期', '制造']
        theme_map = dict(zip(theme_col,
                             [['银行', '非银金融', '房地产'],
                              ['食品饮料', '家用电器', '医药生物', '社会服务', '农林牧渔', '商贸零售', '美容护理', '纺织服饰'],
                              ['通信', '计算机', '电子', '传媒'],
                              ['钢铁', '有色金属', '建筑装饰', '建筑材料', '基础化工', '石油石化', '煤炭'],
                              ['交通运输', '机械设备', '汽车', '轻工制造', '电力设备','环保', '公用事业', '国防军工']
                              ]
                             ))
        lista = []
        listb = []
        for theme in theme_col:
            for col in theme_map[theme]:
                lista.append(col)
                listb.append(theme)
        ind2thememap = pd.DataFrame()
        ind2thememap['industry_name'] = lista
        ind2thememap['theme'] = listb

        self.ind2thememap=ind2thememap


    @staticmethod
    def get_pool_pic(jjdm_list,asofdate):

        #get the jj pic from local db

        # value_p, value_p_hbs, value_sp, value_sp_hbs, size_p, size_p_hbs, size_sp, size_sp_hbs, \
        # industry_p, industry_sp, theme_p, theme_sp, industry_detail_df_list, stock_p, stock_tp, \
        # jj_performance, industry_contribution_list, ticker_con, theme_exp, ind2_exp \
        #     , industry_contribution_perweight_list, style_exp, size_exp = \
        #     jjpic.get_pic_from_localdb(
        #         "jjdm in ({0})".format(util.list_sql_condition(jjdm_list))
        #         ,asofdate=asofdate, if_percentage=False)

        date1=pd.read_sql("select max(asofdate) as asofdate from jjpic_value_p_hbs where asofdate<='{}'"
                          .format(asofdate),con=localdb)['asofdate'].iloc[0]
        date2=pd.read_sql("select max(asofdate) as asofdate from jjpic_industry_p where asofdate<='{}'"
                          .format(asofdate),con=localdb)['asofdate'].iloc[0]

        value_p_hbs=pd.read_sql("select * from jjpic_value_p_hbs where jjdm in ({0}) and asofdate='{1}'"
                                .format(util.list_sql_condition(jjdm_list),date1),con=localdb)
        size_p_hbs=pd.read_sql("select * from jjpic_size_p_hbs where jjdm in ({0}) and asofdate='{1}'"
                                .format(util.list_sql_condition(jjdm_list),date1),con=localdb)
        stock_p=pd.read_sql("select * from jjpic_stock_p where jjdm in ({0}) and asofdate='{1}'"
                                .format(util.list_sql_condition(jjdm_list),date1),con=localdb)

        industry_p=pd.read_sql("select * from jjpic_industry_p where jjdm in ({0}) and asofdate='{1}'"
                                .format(util.list_sql_condition(jjdm_list),date2),con=localdb)
        theme_p=pd.read_sql("select * from jjpic_theme_p where jjdm in ({0}) and asofdate='{1}'"
                                .format(util.list_sql_condition(jjdm_list),date2),con=localdb)

        tempdf = pd.merge(value_p_hbs, size_p_hbs, how='inner', on='jjdm')
        tempdf = pd.merge(tempdf, theme_p, how='inner', on='jjdm')
        tempdf = pd.merge(tempdf, industry_p, how='inner', on='jjdm')
        tempdf = pd.merge(tempdf, stock_p, how='inner', on='jjdm')

        return  tempdf.drop('基金简称_x',axis=1)

    @staticmethod
    def customize_pool(tempdf,target_col,turnover_num,last_solution,allowed_hsl,
                       asofdate,int_solution,cons_col=None,ub=None,lb=None):

        if(cons_col is not None):
            if(int_solution):

                if(len(last_solution)>0):
                    # last_hold_features=tempdf[tempdf['jjdm'].isin(last_solution)][cons_col].mean()
                    tempdf['old_sol']=0
                    tempdf.loc[tempdf['jjdm'].isin(last_solution['jjdm'].tolist()),'old_sol']=1

                    opt_sol =optimizer_allocation_fromlast_holding(inputdf=(tempdf[cons_col+[target_col]+['old_sol']]).dropna()
                                                   , target_col=target_col,allowed_hsl=allowed_hsl
                                                   , con_ub=ub
                                                   , con_lb=lb, turnover_num=turnover_num)
                else:

                    opt_sol = optimizer_allocation(inputdf=(tempdf[cons_col+[target_col]]).dropna()
                                                   , target_col=target_col
                                                   , con_ub=ub
                                                   , con_lb=lb, turnover_num=turnover_num)


            else:
                if(len(last_solution)>0):
                    # last_hold_features=tempdf[tempdf['jjdm'].isin(last_solution)][cons_col].mean()
                    #get the new weight for old solution
                    asofdate2=(datetime.datetime.strptime(asofdate, '%Y%m%d')
                                     +datetime.timedelta(days=75)).strftime('%Y%m%d')
                    sql = "select jjdm,hb1y  ,tjyf,rqzh from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' and hb1y!=99999" \
                        .format(util.list_sql_condition(last_solution['jjdm'].tolist())
                                , last_solution['asofdate'][0][0:6]
                                , asofdate2[0:6])
                    poolret = hbdb.db2df(sql, db='funduser')
                    # implyment T+2
                    sql = "select jjdm,hbdr as pool_ret,jzrq  from st_fund.t_st_gm_rhb where jzrq>= '{0}' and jzrq<='{1}' and jjdm in ({2}) " \
                        .format(str(poolret['tjyf'].iloc[0]) + '01'
                                , str(poolret['tjyf'].iloc[0]) + '15'
                                , util.list_sql_condition(last_solution['jjdm'].tolist()))
                    removed_ret = hbdb.db2df(sql, db='funduser')
                    removed_ret = pd.merge(removed_ret, last_solution[['jjdm', 'flag']]
                                           , how='left', on='jjdm')
                    removed_ret = \
                        (removed_ret.pivot_table('pool_ret', 'jzrq', 'jjdm').iloc[0:2] / 100 + 1).cumprod().iloc[-1]
                    poolret = pd.merge(poolret, removed_ret.to_frame('T+2_fee'), how='left', on='jjdm')
                    poolret['hb1y']=poolret['hb1y']/100+1
                    poolret['cum_ret']=poolret.groupby('jjdm').cumprod()['hb1y']
                    poolret=poolret.drop_duplicates('jjdm',keep='last')
                    poolret['cum_ret']=poolret['cum_ret']/poolret['T+2_fee']
                    last_solution=pd.merge(last_solution,poolret[['jjdm','cum_ret']])
                    last_solution['flag']=last_solution['flag']*last_solution['cum_ret']
                    last_solution['flag']=last_solution['flag']/last_solution['flag'].sum()

                    tempdf=\
                        pd.merge(tempdf,last_solution,how='left',on='jjdm').fillna(0).rename(columns=
                                                                                             {'flag':'old_sol'})

                    opt_sol =optimizer_portfoli_weight_fromlast_holding(inputdf=(tempdf[cons_col+[target_col]+['old_sol']]).dropna()
                                                   , target_col=target_col,allowed_hsl=0.5
                                                   , con_ub=ub
                                                   , con_lb=lb, weight_range=[0,0.04])
                else:

                    opt_sol = optimizer_portfoli_weight(inputdf=(tempdf[cons_col+[target_col]]).dropna()
                                                   , target_col=target_col
                                                   , con_ub=ub
                                                   , con_lb=lb, weight_range=[0,0.04])

            output_pool = tempdf[['jjdm']+cons_col+[target_col]].dropna()
            output_pool['flag'] = opt_sol
            output_pool = output_pool[output_pool['flag'] >0 ]
            output_pool = pd.merge(output_pool[['jjdm','flag']], tempdf, how='left', on='jjdm')

        else:
            output_pool=tempdf.sort_values(target_col,ascending=False)[0:turnover_num]



        output_pool['asofdate']=asofdate

        if('集中度(持仓)_x' in output_pool.columns.tolist()):
            # output_pool['基金简称']=output_pool['基金简称_x']
            return output_pool.rename(columns={'集中度(持仓)_x': '集中度(持仓)_风格',
                                        '换手率(持仓)_x': '换手率(持仓)_风格',
                                        '集中度(持仓)_y': '集中度(持仓)_规模',
                                        '换手率(持仓)_y': '换手率(持仓)_规模',
                                        })[
                ['jjdm','flag', '基金简称', '风格类型', '风格偏好', '成长绝对暴露(持仓)', '价值绝对暴露(持仓)', '集中度(持仓)_风格', '换手率(持仓)_风格'
                    , '规模风格类型', '规模偏好', '大盘绝对暴露(持仓)', '中小盘绝对暴露(持仓)', '集中度(持仓)_规模', '换手率(持仓)_规模'
                    , '大金融', '消费', 'TMT', '周期', '制造'
                    ,'一级行业类型', '一级行业集中度', '一级行业换手率', '前五大行业',target_col,'asofdate']]
        else:
            output_pool[['基金简称', '风格类型', '风格偏好', '成长绝对暴露(持仓)', '价值绝对暴露(持仓)', '集中度(持仓)_风格', '换手率(持仓)_风格'
                    , '规模风格类型', '规模偏好', '大盘绝对暴露(持仓)', '中小盘绝对暴露(持仓)', '集中度(持仓)_规模', '换手率(持仓)_规模'
                    , '大金融', '消费', 'TMT', '周期', '制造'
                    ,'一级行业类型', '一级行业集中度', '一级行业换手率', '前五大行业']]=''
            return output_pool[['jjdm','flag', '基金简称', '风格类型', '风格偏好', '成长绝对暴露(持仓)', '价值绝对暴露(持仓)', '集中度(持仓)_风格', '换手率(持仓)_风格'
                    , '规模风格类型', '规模偏好', '大盘绝对暴露(持仓)', '中小盘绝对暴露(持仓)', '集中度(持仓)_规模', '换手率(持仓)_规模'
                    , '大金融', '消费', 'TMT', '周期', '制造'
                    ,'一级行业类型', '一级行业集中度', '一级行业换手率', '前五大行业',target_col,'asofdate']]

    @staticmethod
    def customize_fund(inputdf,cons_col=None,ub=None,lb=None
                       ,cons_cols2=None,operator=None,condition=None):

        output_df=inputdf.copy()
        if(cons_col is not None):
            for i in range(len(cons_col)):
                output_df=output_df[((output_df[cons_col[i]]>=lb[i])
                                    &(output_df[cons_col[i]]<=ub[i]))|(output_df[cons_col[i]].isnull())]

        if(cons_cols2 is not None):
            for i in range(len(cons_cols2)):
                if(operator[i]=='='):
                    output_df=output_df[(output_df[cons_cols2[i]]==condition[i])
                                        |(output_df[cons_cols2[i]].isnull())]
                elif(operator[i]=='!='):
                    output_df=output_df[(output_df[cons_cols2[i]]!=condition[i])
                                        |(output_df[cons_cols2[i]].isnull())]
                else:
                    print('given operator is not supported ')
                    raise Exception

        return  output_df


    def get_fund_current_holding_features(self,jjdm_List,asofdate,if_zgbl=True,other_table=None):

        if(other_table is not None):
            table=other_table
        else:
            table=''

        sql="SELECT * from {3}hbs_industry_class1_exp where jjdm in ({0}) and  jsrq>='{1}' and jsrq<='{2}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate[0:6]+'01',asofdate[0:6]+'31',table)
        industry_date=pd.read_sql(sql,con=localdb)
        industry_date=pd.merge(industry_date,self.ind2thememap
                               ,how='left',left_on='yjxymc',right_on='industry_name')


        theme_date=industry_date.groupby(['jjdm','theme']).sum()['zjbl'].reset_index().pivot_table('zjbl','jjdm','theme').fillna(0)
        industry_date=industry_date.pivot_table('zjbl','jjdm','yjxymc').fillna(0)


        sql="SELECT * from {3}hbs_industry_class2_exp where jjdm in ({0}) and jsrq>='{1}' and jsrq<='{2}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate[0:6]+'01',asofdate[0:6]+'31',table)
        industry2_date=pd.read_sql(sql,con=localdb)
        industry2_date = industry2_date.pivot_table('zjbl', 'jjdm', 'ejxymc').fillna(0)


        theme_list=theme_date.columns.tolist()
        industry_list=industry_date.columns.tolist()
        industry_list2 = industry2_date.columns.tolist()
        # industry_list3 = industry3_date.columns.tolist()
        industry_date['total'] = industry_date.sum(axis=1)
        industry2_date['total'] = industry2_date.sum(axis=1)
        # industry3_date['total'] = industry3_date.sum(axis=1)
        theme_date['total'] = theme_date.sum(axis=1)


        sql="select jjdm,gptzzjb from st_fund.t_st_gm_zcpz where jsrq>='{0}' and jsrq<='{1}' and jjdm in ({2})"\
            .format(asofdate[0:6]+'01',asofdate[0:6]+'31',util.list_sql_condition(jjdm_List))
        hkzb=hbdb.db2df(sql,db='funduser').drop_duplicates('jjdm')
        industry_date=pd.merge(industry_date,hkzb,how='left',on='jjdm')
        industry_date['港股占比']=(industry_date['gptzzjb']-industry_date['total'])/100
        industry_date.set_index('jjdm',inplace=True)

        #read fund style,size,industry,theme exp history from local db
        sql="SELECT * from {3}hbs_style_exp where jjdm in ({0}) and   jsrq>='{1}' and jsrq<='{2}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate[0:6]+'01',asofdate[0:6]+'31',table)
        style_data=pd.read_sql(sql,con=localdb).pivot_table('zjbl','jjdm','style_type').fillna(0)
        style_data['total']=style_data.sum(axis=1)
        style_col=style_data.columns.tolist()
        # style_data=pd.merge(style_data,industry_date['total'].to_frame('know_weight')
        #                     ,left_index=True,right_index=True)


        sql="SELECT * from {3}hbs_size_exp where jjdm in ({0}) and   jsrq>='{1}' and jsrq<='{2}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate[0:6]+'01',asofdate[0:6]+'31',table)
        size_data=pd.read_sql(sql,con=localdb).pivot_table('zjbl','jjdm','size_type').fillna(0)
        size_col=size_data.columns.tolist()
        size_data['total']=size_data.sum(axis=1)
        # size_data=pd.merge(size_data,industry_date['total'].to_frame('know_weight')
        #                     ,left_index=True,right_index=True)

        if(if_zgbl):
            #if want zsbl then run the below code; if want zjbl skip the below code
            for col in industry_list:
                industry_date[col]=industry_date[col]/industry_date['total']
            for col in industry_list2:
                industry2_date[col]=industry2_date[col]/industry2_date['total']
            for col in theme_list:
                theme_date[col]=theme_date[col]/theme_date['total']

            for col in style_col:
                style_data[col] = style_data[col] / style_data['total']
            for col in size_col:
                size_data[col] = size_data[col] / size_data['total']


        cons_df=style_data.copy().drop(['total'],axis=1)
        cons_df=pd.merge(cons_df,size_data.drop(['total'],axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,industry_date.drop(['total','gptzzjb'],axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,industry2_date.drop('total',axis=1),how='left',on='jjdm')
        # cons_df=pd.merge(cons_df,industry3_date.drop('total',axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,theme_date.drop('total',axis=1),how='left',on='jjdm')
        cons_df.columns=[x+'_当期' for x in cons_df.columns]

        return cons_df


    @staticmethod
    def get_target_df(jjdm_List,asofdate,realdate,target_col):

        if ('new_factor' in target_col):
            from hbshare.fe.factor_analysis import multi_factors_model
            com_factors_list = target_col.split('@')[1:]
            mfm = multi_factors_model.Multi_factor_model(
                com_factors_list, fre='Q')
            target_df = mfm.get_the_combined_factors(method='avg', factor_list=com_factors_list
                                                     , factor_symbol=[1]*len(com_factors_list))

            target_df=target_df[(target_df['date']>=asofdate[0:6]+'01')
                                &(target_df['date']<=asofdate[0:6]+'31')]

        elif('info_ratio' in target_col):
            zblb=target_col.split('info_ratio_')[1]

            if(zblb=='2201'):
                start_date=str(int(asofdate[0:4])-1)+asofdate[4:6]+'28'
            else:
                start_date = str(int(asofdate[0:4]) - 1) + asofdate[4:6]+'28'

            # target_df=calculate_ir_based_on_index(jjdm_List,
            #                                       start_date,
            #                                       asofdate,index_code='885001')

            # sql="select jjdm,info_ratio from factor_infor_ratio where jjdm in ({0}) and date>='{1}' and date <='{2}' and zblb='{3}' and info_ratio!=99999 "\
            #     .format(util.list_sql_condition(jjdm_List)
            #             ,asofdate[0:6]+'01',asofdate[0:6]+'31'
            #             ,zblb )
            # target_df=pd.read_sql(sql,con=localdb)


            max_date=\
                hbdb.db2df("select max(tjrq) as max_date from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq>='{1}' and tjrq<='{2}' and zblb='{3}' and zbnp!=99999 "
                                .format(util.list_sql_condition(jjdm_List)
                        ,asofdate[0:6]+'01',asofdate[0:6]+'31'
                        ,zblb),db='funduser')['max_date'].iloc[0]
            sql=\
                "select jjdm,zbnp as info_ratio,tjrq from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq='{1}' and zblb='{2}' and zbnp!=99999"\
                    .format(util.list_sql_condition(jjdm_List)
                        ,max_date
                        ,zblb )
            target_df = hbdb.db2df(sql,db='funduser')

            size = pd.read_sql("select jjdm,jjzzc  from factor_size where date>='{0}' and date<='{1}'"
                               .format(str(realdate)[0:6]+'01',str(realdate)[0:6]+'31'), con=localdb)
            size['size_rank'] = size['jjzzc'].rank(pct=True)
            size = size[(size['size_rank'] >= 0.2) & (size['size_rank'] < 0.8)]
            target_df=pd.merge(target_df,size,how='inner',on='jjdm')
            from sklearn import preprocessing
            target_df[['info_ratio','jjzzc']]=preprocessing.scale(target_df[['info_ratio','jjzzc']])
            target_df['info_ratio']=target_df['info_ratio']-target_df['jjzzc']
            target_df.drop(['jjzzc','size_rank'],axis=1,inplace=True)



        elif('manager_alpha' in target_col):
            sql="select jjdm,new_factor as manager_alpha from factor_manager_alpha where jjdm in ({0}) and reg_date>='{1}' and reg_date <='{2}' "\
                .format(util.list_sql_condition(jjdm_List)
                        ,asofdate[0:6]+'01',asofdate[0:6]+'31' )
            target_df=pd.read_sql(sql,con=localdb).drop_duplicates('jjdm')

        elif("win_" in target_col):

            fre=target_col.split('@')[1]
            target_col=target_col.split('@')[0]

            sql="SELECT jjdm,{0} from factor_{3}_win_pro where count>=16 and date>='{1}' and date <='{2}'"\
                .format(target_col,asofdate[0:6]+'01',asofdate[0:6]+'31',fre)
            target_df=pd.read_sql(sql,con=localdb).drop_duplicates('jjdm')

        elif('%' in target_col):

            fre=target_col.split('@')[1]
            target_col = target_col.split('@')[0]

            sql="SELECT jjdm,`{0}%` from factor_{3}_rank_quantile where count>=12 and date>='{1}' and date <='{2}'"\
                .format(target_col,asofdate[0:6]+'01',asofdate[0:6]+'31',fre)
            target_df=pd.read_sql(sql,con=localdb)

        else:
            print('input target_col is not available for function get_target_df')
            raise Exception

        return target_df

    @staticmethod
    def return_comparision(bmk_pool,pool,pic_name,laps_days=0,int_solution=True):


        date_list=bmk_pool['asofdate'].unique().tolist()
        date_list2=[(datetime.datetime.strptime(x, '%Y%m%d') +datetime.timedelta(days=laps_days)).strftime('%Y%m%d') for x in date_list]
        date_list.sort()
        date_list2.sort()
        # date_list+=[datetime.datetime.today().strftime('%Y%m%d')]
        # date_list2+=[datetime.datetime.today().strftime('%Y%m%d')]
        # date_list+=['20220930']
        # date_list2+=['20220930']

        def old_T2_fee_reduction(poolret, pool, pool_ret, date_list, int_solution, old_pool, date_list2, hsl_list, hsl,
                                 i):

            # implyment T+2
            sql = "select jjdm,hbdr as pool_ret,jzrq  from st_fund.t_st_gm_rhb where jzrq>= '{0}' and jzrq<='{1}' and jjdm in ({2}) " \
                .format(str(poolret['tjyf'].iloc[0]) + '01'
                        , str(poolret['tjyf'].iloc[0]) + '15'
                        , util.list_sql_condition(pool[pool['asofdate'] == date_list[i]]['jjdm'].tolist()))

            removed_ret = hbdb.db2df(sql, db='funduser')
            removed_ret = pd.merge(removed_ret, pool[pool['asofdate'] == date_list[i]][['jjdm', 'flag']]
                                   , how='left', on='jjdm')

            if (int_solution):
                poolret['flag'] = poolret['flag'] / len(pool[pool['asofdate'] == date_list[i]])

            start_yf = poolret['tjyf'].min()

            removed_ret = \
                (removed_ret.pivot_table('pool_ret', 'jzrq', 'jjdm').iloc[0:2] / 100 + 1).cumprod().iloc[-1]
            poolret = pd.merge(poolret, removed_ret.to_frame('T+2_fee'), how='left', on='jjdm')

            if (len(old_pool) > 0):

                old_pool['cum_ret'] = old_pool['hb1y'] / 100 + 1
                old_pool['cum_ret'] = old_pool.groupby('jjdm').cumprod()['cum_ret']
                snapshot = old_pool.drop_duplicates('jjdm', keep='last')
                snapshot['cum_ret'] = snapshot['cum_ret'] * snapshot['flag']
                snapshot['old_weight'] = snapshot['cum_ret'] / snapshot['cum_ret'].sum()

                snapshot = pd.merge(snapshot, poolret[poolret['tjyf'] == start_yf][['jjdm', 'flag']],
                                    how='left', on='jjdm').fillna(0)

                # remove reblance fee

                snapshot['weight_change'] = snapshot['flag_y'] - snapshot['old_weight']
                sell_fee = (snapshot[snapshot['weight_change'] < 0]['weight_change']).sum() * 0.5 * -1
                buy_fee = (snapshot[snapshot['weight_change'] > 0]['weight_change']).sum() * 0.15

                old_pool['ret_times_w'] = old_pool['hb1y'] * old_pool['flag']

                old_pool = old_pool.groupby('tjyf').sum()['ret_times_w'].to_frame('pool_ret')
                old_pool.loc[old_pool.index[-1], 'pool_ret'] = old_pool.iloc[-1]['pool_ret'] - sell_fee - buy_fee

                pool_ret.append(old_pool)

                # for the new pool
                poolret = pd.merge(poolret, snapshot[['jjdm', 'old_weight']], how='left', on='jjdm')
                # remove t+2 fee
                poolret.loc[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull()), 'hb1y'] = \
                    ((poolret[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull())]['hb1y'] / 100 + 1) / \
                     (poolret[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull())]['T+2_fee']) - 1) * 100

                # remove buying fee
                poolret.loc[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull()), 'hb1y'] = \
                    poolret[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull())]['hb1y'] - 0.15
                old_pool = poolret.copy()

                hsl = snapshot['weight_change'].abs().sum() + \
                      poolret[poolret['old_weight'].isnull()].drop_duplicates('jjdm')['flag'].sum()

                hsl_list.loc[date_list2[i], '换手率'] = hsl

                if (i == len(date_list) - 1 - 1):
                    old_pool['ret_times_w'] = old_pool['hb1y'] * old_pool['flag']
                    old_pool = old_pool.groupby('tjyf').sum()['ret_times_w'].to_frame('pool_ret')
                    pool_ret.append(old_pool)

            else:
                # remove buying fee
                poolret.loc[poolret['tjyf'] == start_yf, 'hb1y'] = \
                    poolret[poolret['tjyf'] == start_yf]['hb1y'] - 0.15
                old_pool = poolret.copy()

                hsl_list.loc[date_list2[i], '换手率'] = hsl

            return pool_ret, old_pool, hsl_list


        def T7_fee_reduction(poolret, pool, pool_ret, date_list, int_solution, old_pool, date_list2, hsl_list, hsl,
                                 i):

            start_yf = poolret['tjyf'].min()
            if (int_solution):
                poolret['flag'] = poolret['flag'] / len(pool[pool['asofdate'] == date_list[i]])

            if (len(old_pool) > 0):

                # implyment T+2
                sql = "select jjdm,hbdr as pool_ret,jzrq  from st_fund.t_st_gm_rhb where jzrq>= '{0}' and jzrq<='{1}' and jjdm in ({2}) " \
                    .format(str(poolret['tjyf'].iloc[0]) + '01'
                            , str(poolret['tjyf'].iloc[0]) + '20'
                            , util.list_sql_condition(pool[pool['asofdate'] == date_list[i]]['jjdm'].tolist()))

                removed_ret = hbdb.db2df(sql, db='funduser')
                removed_ret = pd.merge(removed_ret, pool[pool['asofdate'] == date_list[i]][['jjdm', 'flag']]
                                       , how='left', on='jjdm')



                removed_ret = \
                    (removed_ret.pivot_table('pool_ret', 'jzrq', 'jjdm').iloc[0:7] / 100/5 + 1).cumprod().iloc[-1]
                poolret = pd.merge(poolret, removed_ret.to_frame('T+2_fee'), how='left', on='jjdm')


                old_pool['cum_ret'] = old_pool['hb1y'] / 100 + 1
                old_pool['cum_ret'] = old_pool.groupby('jjdm').cumprod()['cum_ret']
                snapshot = old_pool.drop_duplicates('jjdm', keep='last')
                snapshot['cum_ret'] = snapshot['cum_ret'] * snapshot['flag']
                snapshot['old_weight'] = snapshot['cum_ret'] / snapshot['cum_ret'].sum()

                snapshot = pd.merge(snapshot, poolret[poolret['tjyf'] == start_yf][['jjdm', 'flag']],
                                    how='left', on='jjdm').fillna(0)

                # remove reblance fee

                snapshot['weight_change'] = snapshot['flag_y'] - snapshot['old_weight']
                sell_fee = (snapshot[snapshot['weight_change'] < 0]['weight_change']).sum() * 0.5 * -1
                buy_fee = (snapshot[snapshot['weight_change'] > 0]['weight_change']).sum() * 0.15

                old_pool['ret_times_w'] = old_pool['hb1y'] * old_pool['flag']

                old_pool = old_pool.groupby('tjyf').sum()['ret_times_w'].to_frame('pool_ret')
                old_pool.loc[old_pool.index[-1], 'pool_ret'] = old_pool.iloc[-1]['pool_ret'] - sell_fee - buy_fee

                pool_ret.append(old_pool)

                # for the new pool
                poolret = pd.merge(poolret, snapshot[['jjdm', 'old_weight']], how='left', on='jjdm')
                # remove t+2 fee
                poolret.loc[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull()), 'hb1y'] = \
                    ((poolret[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull())]['hb1y'] / 100 + 1) / \
                     (poolret[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull())]['T+2_fee']) - 1) * 100

                # remove buying fee
                poolret.loc[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull()), 'hb1y'] = \
                    poolret[(poolret['tjyf'] == start_yf) & (poolret['old_weight'].isnull())]['hb1y'] - 0.15
                old_pool = poolret.copy()

                hsl = snapshot['weight_change'].abs().sum() + \
                      poolret[poolret['old_weight'].isnull()].drop_duplicates('jjdm')['flag'].sum()

                hsl_list.loc[date_list2[i], '换手率'] = hsl

                if (i == len(date_list) - 1 - 1):
                    old_pool['ret_times_w'] = old_pool['hb1y'] * old_pool['flag']
                    old_pool = old_pool.groupby('tjyf').sum()['ret_times_w'].to_frame('pool_ret')
                    pool_ret.append(old_pool)

            else:
                # remove buying fee
                poolret.loc[poolret['tjyf'] == start_yf, 'hb1y'] = \
                    poolret[poolret['tjyf'] == start_yf]['hb1y'] - 0.15
                old_pool = poolret.copy()

                hsl_list.loc[date_list2[i], '换手率'] = hsl

            return pool_ret, old_pool, hsl_list

        pool_ret=[]
        universie_ret=[]
        old_pool=[]
        hsl_list=pd.DataFrame(index=date_list2,columns=['换手率'])
        hsl=1
        for i in range(len(date_list)-1):

            sql = "select avg(hb1y) as universie_ret ,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' and hb1y!=99999 group by tjyf" \
                .format(util.list_sql_condition(bmk_pool[bmk_pool['asofdate']==date_list[i]]['jjdm'].tolist())
                        ,date_list2[i][0:6]
                        ,date_list2[i+1][0:6])
            universie_ret.append(hbdb.db2df(sql, db='funduser'))


            sql = "select jjdm,hb1y ,tjyf,rqzh from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' and hb1y!=99999" \
                .format(util.list_sql_condition(pool[pool['asofdate'] == date_list[i]]['jjdm'].tolist())
                        , date_list2[i][0:6]
                        , date_list2[i + 1][0:6])
            if(date_list[i+1]=='20220630'):
                print('')
            poolret = hbdb.db2df(sql, db='funduser')
            poolret = pd.merge(poolret, pool[pool['asofdate'] == date_list[i]][['jjdm', 'flag']]
                               , how='left', on='jjdm')

            # implyment T+2
            # sql = "select jjdm,hbdr as pool_ret,jzrq  from st_fund.t_st_gm_rhb where jzrq>= '{0}' and jzrq<='{1}' and jjdm in ({2}) " \
            #     .format(str(poolret['tjyf'].iloc[0]) + '01'
            #             , str(poolret['tjyf'].iloc[0]) + '15'
            #             , util.list_sql_condition(pool[pool['asofdate'] == date_list[i]]['jjdm'].tolist()))
            #
            # removed_ret = hbdb.db2df(sql, db='funduser')
            # removed_ret = pd.merge(removed_ret, pool[pool['asofdate'] == date_list[i]][['jjdm', 'flag']]
            #                        , how='left', on='jjdm')
            #
            # if (int_solution):
            #     poolret['flag']=poolret['flag']/len(pool[pool['asofdate'] == date_list[i]])
            #
            # start_yf=poolret['tjyf'].min()
            #
            # removed_ret = \
            #     (removed_ret.pivot_table('pool_ret', 'jzrq', 'jjdm').iloc[0:2] / 100 + 1).cumprod().iloc[-1]
            # poolret = pd.merge(poolret, removed_ret.to_frame('T+2_fee'), how='left', on='jjdm')
            #
            #
            # if (len(old_pool) > 0):
            #
            #     old_pool['cum_ret']=old_pool['hb1y']/100+1
            #     old_pool['cum_ret']=old_pool.groupby('jjdm').cumprod()['cum_ret']
            #     snapshot=old_pool.drop_duplicates('jjdm',keep='last')
            #     snapshot['cum_ret']=snapshot['cum_ret']*snapshot['flag']
            #     snapshot['old_weight']=snapshot['cum_ret']/snapshot['cum_ret'].sum()
            #
            #     snapshot=pd.merge(snapshot,poolret[poolret['tjyf']==start_yf][['jjdm','flag']],
            #                       how='left',on='jjdm').fillna(0)
            #
            #     #remove reblance fee
            #
            #     snapshot['weight_change']=snapshot['flag_y']-snapshot['old_weight']
            #     sell_fee=(snapshot[snapshot['weight_change']<0]['weight_change']).sum()*0.5*-1
            #     buy_fee = (snapshot[snapshot['weight_change'] > 0]['weight_change']).sum() * 0.15
            #
            #     old_pool['ret_times_w']=old_pool['hb1y']*old_pool['flag']
            #
            #     old_pool=old_pool.groupby('tjyf').sum()['ret_times_w'].to_frame('pool_ret')
            #     old_pool.loc[old_pool.index[-1],'pool_ret']=old_pool.iloc[-1]['pool_ret']-sell_fee-buy_fee
            #
            #     pool_ret.append(old_pool)
            #
            #     #for the new pool
            #     poolret=pd.merge(poolret,snapshot[['jjdm','old_weight']],how='left',on='jjdm')
            #     # remove t+2 fee
            #     poolret.loc[(poolret['tjyf']==start_yf)&(poolret['old_weight'].isnull()),'hb1y']= \
            #         ((poolret[(poolret['tjyf']==start_yf)&(poolret['old_weight'].isnull())]['hb1y']/100+1)/\
            #         (poolret[(poolret['tjyf']==start_yf)&(poolret['old_weight'].isnull())]['T+2_fee'])-1)*100
            #
            #     #remove buying fee
            #     poolret.loc[(poolret['tjyf']==start_yf)&(poolret['old_weight'].isnull()),'hb1y']=\
            #         poolret[(poolret['tjyf']==start_yf)&(poolret['old_weight'].isnull())]['hb1y']-0.15
            #     old_pool=poolret.copy()
            #
            #     hsl =snapshot['weight_change'].abs().sum()+\
            #          poolret[poolret['old_weight'].isnull()].drop_duplicates('jjdm')['flag'].sum()
            #
            #     hsl_list.loc[date_list2[i], '换手率'] = hsl
            #
            #     if(i==len(date_list)-1-1):
            #         old_pool['ret_times_w'] = old_pool['hb1y'] * old_pool['flag']
            #         old_pool = old_pool.groupby('tjyf').sum()['ret_times_w'].to_frame('pool_ret')
            #         pool_ret.append(old_pool)
            #
            # else:
            #
            #     # remove T+2 fee
            #     # poolret.loc[poolret['tjyf']==start_yf,'hb1y']= \
            #     #     ((poolret[poolret['tjyf']==start_yf]['hb1y']/100+1)/\
            #     #     (poolret[poolret['tjyf']==start_yf]['T+2_fee'])-1)*100
            #
            #     #remove buying fee
            #     poolret.loc[poolret['tjyf']==start_yf,'hb1y']=\
            #         poolret[poolret['tjyf']==start_yf]['hb1y']-0.15
            #     old_pool=poolret.copy()
            #
            #     hsl_list.loc[date_list2[i], '换手率'] = hsl

            pool_ret, old_pool, hsl_list=\
                T7_fee_reduction(poolret,pool,pool_ret,date_list,int_solution,old_pool,date_list2,hsl_list,hsl,i)

        pool_ret=pd.concat(pool_ret,axis=0)
        universie_ret = pd.concat(universie_ret, axis=0)

        sql = "select hb1y as 偏股混指数 ,tjyf from st_market.t_st_zs_yhb where zqdm='885001' and  tjyf>{0} and tjyf<={1} and hb1y!=99999"\
            .format(date_list2[0][0:6],date_list2[-1][0:6])
        index_ret = hbdb.db2df(sql, db='alluser')

        ret = pd.merge(index_ret, pool_ret, how='left', on='tjyf')
        ret = pd.merge(ret, universie_ret, how='left', on='tjyf')
        ret['tjyf'] = ret['tjyf'].astype(str)
        ret['nav'] = (ret['pool_ret'] / 100 + 1).cumprod()
        print('max draw back is {}'.format((ret['nav']/ret['nav'].rolling(80,1).max()-1).min()))

        month_count=len(ret)

        #calculate vol and sharp
        vol_df=ret.std()/100*np.sqrt(12)

        ret.set_index('tjyf', inplace=True)
        ret = ret / 100 + 1
        #calculate the yearly winning pro

        ret['year']=ret.index.astype(str).str[0:4]
        ret.loc[ret.index.astype(str).str[4:6].isin(['04','05','06','07','08','09']),'half_year'] =ret[ret.index.astype(str).str[4:6].isin(['04','05','06','07','08','09'])]['year']+'1'
        ret.loc[ret.index.astype(str).str[4:6].isin(['10','11','12']),'half_year'] =ret[ret.index.astype(str).str[4:6].isin(['10','11','12'])]['year']+'2'
        ret.loc[ret.index.astype(str).str[4:6].isin(['01','02','03']),'half_year'] =(ret[ret.index.astype(str).str[4:6].isin(['01','02','03'])]['year'].astype(int)-1).astype(str)+'2'

        yearly_ret=ret.groupby('year').cumprod()
        yearly_ret['year']=yearly_ret.index.astype(str).str[0:4]
        yearly_ret.drop_duplicates('year', keep='last',inplace=True)

        yearly_ret['pool_vs_index']=yearly_ret['pool_ret']>yearly_ret['偏股混指数']
        yearly_ret['pool_mins_index'] = yearly_ret['pool_ret'] - yearly_ret['偏股混指数']
        yearly_ret['pool_vs_bmk'] = yearly_ret['pool_ret'] > yearly_ret['universie_ret']
        yearly_ret['pool_mins_bmk'] = yearly_ret['pool_ret'] - yearly_ret['universie_ret']
        yearly_ret['bmk_vs_index'] = yearly_ret['universie_ret'] > yearly_ret['偏股混指数']

        half_yearly_ret=ret.groupby('half_year').cumprod()
        half_yearly_ret['half_year']=ret['half_year']
        half_yearly_ret.drop_duplicates('half_year', keep='last',inplace=True)

        print('半年度胜率：池相对于基准胜率为 {0}%，相对于偏股混指数为{1}%，跑输基准年份为{2}，跑输指数年份为{3}'
              .format(100*(half_yearly_ret['pool_ret']>half_yearly_ret['universie_ret']).sum()/len(half_yearly_ret)
                      ,100*(half_yearly_ret['pool_ret']>half_yearly_ret['偏股混指数']).sum()/len(half_yearly_ret)
                      ,util.list_sql_condition(half_yearly_ret[half_yearly_ret['pool_ret']<half_yearly_ret['universie_ret']].index.tolist())
                      ,util.list_sql_condition(half_yearly_ret[half_yearly_ret['pool_ret']<half_yearly_ret['偏股混指数']].index.tolist())))

        print('池相对于基准胜率为 {0}%，相对于偏股混指数为{1}%，跑输基准年份为{2}，跑输指数年份为{3}'
              .format(100*yearly_ret['pool_vs_bmk'].sum()/len(yearly_ret)
                      ,100*yearly_ret['pool_vs_index'].sum()/len(yearly_ret)
                      ,util.list_sql_condition(yearly_ret[~yearly_ret['pool_vs_bmk']]['year'].tolist())
                      ,util.list_sql_condition(yearly_ret[~yearly_ret['pool_vs_index']]['year'].tolist())))

        print('池与指数每年收益差')
        print(yearly_ret[['pool_mins_index']])
        print('池与bmk每年收益')
        print(yearly_ret[['pool_mins_bmk']])

        print('基准相对于偏股混指数胜率为 {0}%，跑输指数年份为{1}'
              .format(100*yearly_ret['bmk_vs_index'].sum()/len(yearly_ret)
                      ,util.list_sql_condition(yearly_ret[~yearly_ret['bmk_vs_index']]['year'].tolist())))

        print('换手率历史：')
        print(hsl_list)
        print('平均年换手为{}'.format((hsl_list['换手率'].sum()-1)/((len(hsl_list)-2)/2)))
        ret.drop(['year','half_year'],axis=1,inplace=True)

        for col in ret:
            ret[col] = ret[col].cumprod()
            print('{0} 年化收益率为{1}%,年化波动率为{3}%,年化夏普为{2}'
                  .format(col
                          ,(np.power(ret[col].iloc[-1],12/month_count)-1)*100
                          ,(np.power(ret[col].iloc[-1],12/month_count)-1)/vol_df[col]
                          ,vol_df[col]*100)
                  )
        plot.plotly_line_style(ret, pic_name, save_local_file=True)

        ret['year'] = ret.index.astype(str).str[0:4]
        ret=ret.reset_index()
        initial_price=ret.drop_duplicates('year', keep='last')
        initial_price['year'] = (initial_price['year'].astype(int) + 1).astype(str)
        ret = pd.merge(ret, initial_price, how='left', on='year').fillna(1)
        ret=ret.set_index('tjyf_x')
        for col in ['偏股混指数','pool_ret','universie_ret']:
            ret[col]=ret[col+'_x']/ret[col+'_y']
        for year in ret['year'].unique().tolist():
            plot.plotly_line_style(ret[ret['year']==year][['偏股混指数','pool_ret','universie_ret']]
                                   , year+'业绩走势对比', save_local_file=False)

    #the function for creating certain pool

    def generate_equilibrium_pool(self,input_df,cons_df,asofdate,last_date,turnover_num,target_col
                                  ,cons_col,last_solution,allowed_hsl=0.5,int_solution=True,nav_adj=True,gap=0.01):

        # sql="select * from pool_history where asofdate='{0}' and pool_name in ('HB周期','HBTMT','HB医药','HB大金融', 'HB消费', 'HB制造')"\
        #     .format(asofdate)
        # pre_given_jjlist=pd.read_sql(sql,con=localdb)
        # pre_given_jjlist = pd.merge(pre_given_jjlist, cons_df, how='left', on='jjdm')
        pre_given_jjlist=[]

        cons_df=cons_df.mean(axis=0)

        max=cons_df.loc[cons_col[6:]]+gap


        min=cons_df.loc[cons_col[6:]]-gap


        if(asofdate>='20220630' and (asofdate[4:6]=='06' or asofdate[4:6]=='12')):
            pic=self.get_pool_pic(input_df['jjdm'].unique().tolist()
                                  ,(datetime.datetime.strptime(asofdate, '%Y%m%d')
                                     +datetime.timedelta(days=120)).strftime('%Y%m%d')).drop_duplicates('jjdm')
            input_df=pd.merge(input_df,pic,how='left',on='jjdm')
            input_df=self.customize_fund( inputdf=input_df,cons_col=['换手率(持仓)_x','换手率(持仓)_y','一级行业换手率']
                                    ,ub=[1,1,1],lb=[0.1,0.1,0.1])


        nav_driven_weight_change = 0
        if(asofdate[4:6]=='03' or asofdate[4:6]=='09'):
            if(nav_adj):
                import sympy
                # # #for nav
                if (asofdate[4:6] == '03'):
                    end_date = asofdate[0:4] + '0701'
                    start_date = str(int(asofdate[0:4]) - 1) + '1001'
                else:
                    end_date = str(int(asofdate[0:4]) + 1) + '0101'
                    start_date = asofdate[0:4] + '0401'

                sql = """select zqdm,jyrq,hbdr from st_market.t_st_zs_rhb
                where zqdm in ('885001','399370','399371','CBA00301')
                and jyrq>='{0}' and jyrq<='{1}'""".format(start_date, end_date)
                index_ret = hbdb.db2df(sql, db='alluser').pivot_table('hbdr', 'jyrq', 'zqdm')

                L=[sympy.Matrix(index_ret.loc[asofdate:]['399370'].tolist()),
                   sympy.Matrix(index_ret.loc[asofdate:]['399371'].tolist())]
                L=sympy.GramSchmidt(L)
                index_ret.loc[asofdate:,'399371']=L[1][0:]
                index_ret['399371']=index_ret['399371'].astype(float)

                L=[sympy.Matrix(index_ret.loc[:last_date]['399370'].tolist()),
                   sympy.Matrix(index_ret.loc[:last_date]['399371'].tolist())]
                L=sympy.GramSchmidt(L)
                index_ret.loc[:last_date,'399371']=L[1][0:]
                index_ret['399371']=index_ret['399371'].astype(float)


                result1 =\
                    sm.OLS(index_ret.loc[asofdate:]['885001'].values,
                                index_ret.loc[asofdate:][['399370', '399371', 'CBA00301']].values).fit().params.tolist()

                result0 = \
                sm.OLS(index_ret.loc[:last_date]['885001'].values,
                                                  index_ret.loc[:last_date][['399370', '399371', 'CBA00301']].values).fit().params.tolist()


                nav_driven_weight_change = result1[0]-result0[0]

                # nav_driven_weight_change=1
                # if(asofdate[4:6]=='03'):
                #     end_date=asofdate[0:4]+'0701'
                #     start_date=str(int(asofdate[0:4])-1)+'1201'
                # else:
                #     end_date=asofdate[0:4]+'1001'
                #     start_date = asofdate[0:4]+ '0601'
                # sql="select * from hbs_style_exp_key where jsrq >='{0}' and jsrq<='{1}'"\
                #     .format(start_date,end_date)
                # key_style=pd.read_sql(sql,con=localdb)
                # key_style=pd.merge(key_style,
                #                    key_style.groupby(['jsrq','jjdm']).sum()['zjbl'].to_frame('total').reset_index(),
                #                    how='left',on=['jsrq','jjdm'])
                # key_style['zjbl']=key_style['zjbl']/key_style['total']
                # key_style=key_style.groupby(['jsrq','style_type']).mean()[['zjbl']].reset_index()
                # nav_driven_weight_change=\
                #     key_style[key_style['style_type']=='成长'].iloc[-1]['zjbl']-\
                #     key_style[key_style['style_type']=='成长'].iloc[0]['zjbl']
        if( nav_adj and( (asofdate[4:6]=='06' or asofdate[4:6]=='12') or( abs(nav_driven_weight_change)>0.1)) or not nav_adj):

            # cons_col.remove(cons_col[37])
            #(cons_df.loc[cons_col[36:37]].values+1).tolist()
            #+(cons_df.loc[cons_col[36:37]].values-1).tolist()

            # ub = \
            #     (cons_df.loc[cons_col[0:1]].values + 0.03 - nav_driven_weight_change * 0.5).tolist() + \
            #     (cons_df.loc[cons_col[1:2]].values + 0.03).tolist() + (cons_df.loc[cons_col[2:3]].values + 0.03 + nav_driven_weight_change * 0.5).tolist()+\
            #     (cons_df.loc[cons_col[3:6]].values + 0.03).tolist() + \
            #     (cons_df.loc[cons_col[6:37]].values + gap ).tolist() + \
            #     (cons_df.loc[cons_col[37:]].values + gap).tolist()
            #
            # lb = \
            #     (cons_df.loc[cons_col[0:1]].values - 0.03 - nav_driven_weight_change * 0.5).tolist() + \
            #     (cons_df.loc[cons_col[1:2]].values - 0.03 ).tolist() + (cons_df.loc[cons_col[2:3]].values - 0.03 + nav_driven_weight_change * 0.5).tolist()+\
            #     (cons_df.loc[cons_col[3:6]].values - 0.03).tolist() + \
            #     (cons_df.loc[cons_col[6:37]].values- gap ).tolist() + \
            #     (cons_df.loc[cons_col[37:]].values - gap).tolist()

            #customize
            ub = \
                (cons_df.loc[cons_col[0:1]].values + 0.03 - nav_driven_weight_change * 0.5).tolist() + \
                (cons_df.loc[cons_col[1:2]].values + 0.03).tolist() + (cons_df.loc[cons_col[2:3]].values + 0.03 + nav_driven_weight_change * 0.5).tolist()+\
                (cons_df.loc[cons_col[3:6]].values + 0.03).tolist() + \
                (cons_df.loc[cons_col[6:37]].values + gap ).tolist() + \
                (cons_df.loc[cons_col[37:]].values + gap).tolist()

            lb = \
                (cons_df.loc[cons_col[0:1]].values - 0 - nav_driven_weight_change * 0.5).tolist() + \
                (cons_df.loc[cons_col[1:2]].values - 0.03 ).tolist() + (cons_df.loc[cons_col[2:3]].values - 0.03 + nav_driven_weight_change * 0.5).tolist()+\
                (cons_df.loc[cons_col[3:6]].values - 0.03).tolist() + \
                (cons_df.loc[cons_col[6:37]].values- gap ).tolist() + \
                (cons_df.loc[cons_col[37:]].values - gap).tolist()



            if(len(pre_given_jjlist)>0):

                w2 = len(pre_given_jjlist) / turnover_num/2
                w1 = 1 - w2
                turnover_num = turnover_num - len(pre_given_jjlist)


                #get the theme current weight
                theme_weight=cons_df[[ 'TMT_当期', '制造_当期', '周期_当期','大金融_当期', '消费_当期','医药生物_当期']]
                theme_weight['消费_当期']=theme_weight['消费_当期']-theme_weight['医药生物_当期']
                theme_weight.index=["HB"+str(x).replace('_当期','') for x in theme_weight.index ]
                theme_weight=theme_weight/(len(pre_given_jjlist)/len(theme_weight))

                pre_given_jjlist.loc[pre_given_jjlist['pool_name']=='HB医药','pool_name']='HB医药生物'
                pre_given_jjlist=pd.merge(pre_given_jjlist,theme_weight.to_frame('weight')
                                          ,how='left',left_on='pool_name',right_index=True)
                for col in cons_col:
                    pre_given_jjlist[col]=\
                        pre_given_jjlist[col]*pre_given_jjlist['weight']
                adjust_cons=w2*pre_given_jjlist[cons_col].sum(axis=0).values

                ub=list((np.array(ub)-adjust_cons)/w1)
                lb=list((np.array(lb)-adjust_cons)/w1)
                input_df=input_df[~input_df['jjdm'].isin(pre_given_jjlist['jjdm'].tolist())]

            opti_pool=self.customize_pool(tempdf=input_df,target_col=target_col
                              ,turnover_num=turnover_num,asofdate=asofdate,int_solution=int_solution,cons_col=cons_col
                              ,ub=ub,lb=lb,last_solution=last_solution,allowed_hsl=allowed_hsl)
        else:
            opti_pool=[]


        if(len(opti_pool)>0):
            last_solution = opti_pool[['jjdm', 'flag']]

        if(len(pre_given_jjlist)>0 and len(opti_pool)>0):
            opti_pool['flag']=w1/len(opti_pool)
            pre_given_jjlist['flag']=pre_given_jjlist['weight']*w2
            opti_pool=pd.concat([opti_pool,pre_given_jjlist[['jjdm', 'asofdate','flag']]],axis=0)


        return opti_pool,last_solution

    def generate_abs_equilibrium_pool(self,input_df,cons_df,asofdate,target_col,turnover_num=2):

        child_pool_cap=50
        #filter style fund
        #filter growth fund
        tempdf=input_df[input_df['换手率排名(持仓)_x']<=0.3].sort_values('成长绝对暴露(持仓)'
                                                                  ,ascending=False)[0:child_pool_cap][['jjdm']]
        tempdf['growth_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)

        #filter value fund
        tempdf=input_df[(input_df['换手率排名(持仓)_x']<=0.3)
                        &(input_df['价值绝对暴露(持仓)']>input_df['成长绝对暴露(持仓)'])].sort_values('价值绝对暴露(持仓)'
                                                                  ,ascending=False)[0:child_pool_cap][['jjdm']]
        tempdf['value_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)

        #filter style equilibrium fund
        tempdf=input_df[input_df['换手率排名(持仓)_x']<=0.3].sort_values('集中度排名(持仓)_x'
                                                                  ,ascending=True)[0:child_pool_cap][['jjdm']]
        tempdf['style_eq_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)


        #filter size fund
        #filter big size fund
        tempdf=input_df[input_df['换手率排名(持仓)_y']<=0.3].sort_values('大盘绝对暴露(持仓)'
                                                                  ,ascending=False)[0:child_pool_cap][['jjdm']]
        tempdf['big_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)

        #filter med size  fund
        tempdf=input_df[input_df['换手率排名(持仓)_y']<=0.3].sort_values('中盘绝对暴露(持仓)'
                                                                  ,ascending=False)[0:child_pool_cap][['jjdm']]
        tempdf['med_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)

        #filter small size fund
        tempdf=input_df[input_df['换手率排名(持仓)_y']<=0.3].sort_values('小盘绝对暴露(持仓)'
                                                                  ,ascending=False)[0:child_pool_cap][['jjdm']]
        tempdf['small_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)


        #filter size equilibrium fund
        tempdf=input_df[input_df['换手率排名(持仓)_x']<=0.3].sort_values('集中度排名(持仓)_y'
                                                                  ,ascending=True)[0:child_pool_cap][['jjdm']]
        tempdf['size_eq_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)


        #filter strategy fund
        #filter down 2 top  fund
        tempdf=input_df[input_df['一级行业集中度']<=0.5].sort_values('个股集中度'
                                                                  ,ascending=False)[0:child_pool_cap][['jjdm']]
        tempdf['down2top_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)

        #filter top 2 down  fund
        tempdf=input_df[input_df['个股集中度']<=0.5].sort_values('一级行业集中度'
                                                                  ,ascending=False)[0:child_pool_cap][['jjdm']]
        tempdf['top2down_flag']=1
        input_df=pd.merge(input_df,tempdf,how='left',on='jjdm').fillna(0)


        type_flag_list=input_df.columns[input_df.columns.str.contains('_flag')].tolist()
        cons_col=input_df.columns.tolist()[99:130]+input_df.columns.tolist()[-14:-9]


        #get the optimal pool based on condition
        #we are trying to get a pool with 2 of each type of fund defined above
        # and requires that the total industry and themem exp match market exp

        # opti_pool = self.customize_pool(tempdf=input_df, target_col=target_col
        #                                 , turnover_num=len(type_flag_list) * turnover_num, asofdate=asofdate,
        #                                 cons_col=type_flag_list
        #                                 , ub=[turnover_num] * len(type_flag_list)
        #                                 , lb=[turnover_num] * len(type_flag_list)
        #                                 )

        opti_pool=self.customize_pool(tempdf=input_df,target_col=target_col
                          ,turnover_num=len(type_flag_list)*turnover_num,asofdate=asofdate,cons_col=cons_col+type_flag_list
                          ,ub=(cons_df.mean(axis=0).loc[cons_col].values+0.02).tolist()+[1/len(type_flag_list)]*len(type_flag_list)
                          ,lb=(cons_df.mean(axis=0).loc[cons_col].values-0.02).tolist()+[1/len(type_flag_list)]*len(type_flag_list)
                                      )



        return opti_pool

class Mutual_fund_report():

    def __init__(self,start_date,end_date,asofdate='20211231'):

        #get the theme map relation
        theme_col = ['大金融', '消费', 'TMT', '周期', '制造']
        theme_map = dict(zip(theme_col,
                             [['银行', '非银金融', '房地产'],
                              ['食品饮料', '家用电器', '医药生物', '社会服务', '农林牧渔', '商贸零售', '美容护理'],
                              ['通信', '计算机', '电子', '传媒', '国防军工'],
                              ['钢铁', '有色金属', '建筑装饰', '建筑材料', '基础化工', '石油石化', '煤炭', '环保', '公用事业'],
                              ['交通运输', '机械设备', '汽车', '纺织服饰', '轻工制造', '电力设备']
                              ]
                             ))
        lista = []
        listb = []
        for theme in theme_col:
            for col in theme_map[theme]:
                lista.append(col)
                listb.append(theme)
        ind2thememap = pd.DataFrame()
        ind2thememap['industry_name'] = lista
        ind2thememap['theme'] = listb
        self.ind2thememap=ind2thememap


        saved_folder=r"E:\GitFolder\docs\公募基金周报\\"
        pic_list=[]
        data_list=[]

        self.end_date=end_date
        self.folder=saved_folder

        self.main_index=self.get_main_index_weekly_return_table(start_date,end_date)
        self.stock_fund=self.get_fund_index_weekly_return_table(end_date)
        self.ind_ret=self.get_industry_weekly_return_table(end_date)
        self.theme_ret=self.get_wind_theme_weekly_return_table()
        self.style_ret=self.get_style_weekly_return_table(end_date)

        plot.plotly_table(self.main_index[0],1000,'主要指数月收益',save_local_file=saved_folder)
        pic_list.append(saved_folder+'主要指数月收益'+'.png')
        data_list.append(self.main_index[0])
        plot.plotly_line_style(self.main_index[1],'主要指数今年以来收益',save_local_file=saved_folder)
        pic_list.append(saved_folder+'主要指数今年以来收益'+'.png')
        data_list.append(self.main_index[1])

        plot.plotly_jjpic_bar(self.stock_fund[0],'偏股型合基金月收益率（%）',save_local_file=saved_folder)
        pic_list.append(saved_folder + '偏股型合基金月收益率（%）' + '.png')
        data_list.append(self.stock_fund[0])
        plot.plotly_line_style(self.stock_fund[1],'偏股型合基金今年以来收益',save_local_file=saved_folder)
        pic_list.append(saved_folder+'偏股型合基金今年以来收益'+'.png')
        data_list.append(self.stock_fund[1])


        plot.plotly_jjpic_bar(pd.concat([self.ind_ret[0].iloc[0:5],
                                         self.ind_ret[0].iloc[-5:]],axis=0),'申万一级行业月收益率（%）前5与后5',save_local_file=saved_folder)
        pic_list.append(saved_folder + '申万一级行业月收益率（%）前5与后5' + '.png')
        data_list.append(pd.concat([self.ind_ret[0].iloc[0:5],
                                    self.ind_ret[0].iloc[-5:]],axis=0))
        plot.plotly_line_style(self.ind_ret[1],'申万一级行业今年以来收益',save_local_file=saved_folder)
        pic_list.append(saved_folder+'申万一级行业今年以来收益'+'.png')
        data_list.append(self.ind_ret[1])

        plot.plotly_jjpic_bar(pd.concat([self.theme_ret[0],self.theme_ret[1]],axis=0),'WIND概念月收益率（%）前5与后5',save_local_file=saved_folder)
        pic_list.append(saved_folder + 'WIND概念月收益率（%）前5与后5' + '.png')
        data_list.append(pd.concat([self.theme_ret[0],self.theme_ret[1]],axis=0))


        plot.plotly_jjpic_bar(self.style_ret[0], '巨潮风格指数月收益率（%）',save_local_file=saved_folder)
        pic_list.append(saved_folder + '巨潮风格指数月收益率（%）' + '.png')
        data_list.append(self.style_ret[0])
        plot.plotly_line_style(self.style_ret[1],'巨潮风格指数今年以来收益',save_local_file=saved_folder)
        pic_list.append(saved_folder+'巨潮风格指数今年以来收益'+'.png')
        data_list.append(self.style_ret[1])

        bmk_features=self.get_fund_current_holding_features(util.get_bmk_funds_list(asofdate),asofdate)
        pool_features=bmk_features.mean(axis=0)
        pool_features=pd.concat([pool_features[0:5], pool_features[-5:], pool_features[36:37]], axis=0)
        pool_features=pool_features.to_frame('市场风格暴露')

        self.benmark_info=self.get_benchmark_ret(end_date,asofdate)
        self.core_info=self.get_core_pool_ret(end_date,'君瑞',bmk_features,pool_features)
        pool_features=self.core_info[-1]

        ind_weight=self.core_info[-2].mean(axis=0).reset_index(drop=False)
        ind_weight['index'] = [x.replace('_当期', '') for x in ind_weight['index']]
        ind_weight=pd.merge(ind_weight,self.ind_ret[0].reset_index(drop=False)
                            ,how='left',left_on='index',right_on='证券代码').drop('index',axis=1)
        ind_weight['ind_ret']=ind_weight['月涨幅%']*ind_weight[0]
        manager_ret=self.core_info[2].iloc[-1]['AVG(hb1y)']-ind_weight['ind_ret'].sum()
        ind_weight=ind_weight.sort_values('ind_ret', ascending=False)[['证券代码', 'ind_ret']]
        # ind_weight=pd.concat([ind_weight.iloc[0:5],ind_weight.iloc[-5:]],axis=0)
        ind_weight['ind_ret']=(ind_weight['ind_ret']/100).map("{:.2%}".format)
        ret_contribution=pd.DataFrame(columns=['月收益','经理收益'])
        ret_contribution['月收益']=\
            [(self.core_info[2]['AVG(hb1y)']/100).map("{:.2%}".format).iloc[-1]]
        ret_contribution['经理收益']=["{:.2%}".format(manager_ret/100)]
        ind_weight=ind_weight.set_index('证券代码').T
        ind_weight.index=[0]
        ret_contribution=pd.concat([ret_contribution
                                            ,ind_weight],axis=1)


        self.value_info=self.get_core_pool_ret(end_date,'价值',bmk_features,pool_features)
        pool_features=self.value_info[-1]

        self.growth_info=self.get_core_pool_ret(end_date,'成长',bmk_features,pool_features)
        pool_features=self.growth_info[-1]

        self.quant=self.get_core_pool_ret(end_date,'动量',bmk_features,pool_features)
        pool_features=self.quant[-1]

        self.finance = self.get_core_pool_ret(end_date, 'HB大金融',bmk_features,pool_features)
        pool_features=self.finance[-1]

        self.consume = self.get_core_pool_ret(end_date, 'HB消费',bmk_features,pool_features)
        pool_features=self.consume[-1]

        self.tmt = self.get_core_pool_ret(end_date, 'HBTMT',bmk_features,pool_features)
        pool_features=self.tmt[-1]

        self.cycle = self.get_core_pool_ret(end_date, 'HB周期',bmk_features,pool_features)
        pool_features=self.cycle[-1]

        self.manu = self.get_core_pool_ret(end_date, 'HB制造',bmk_features,pool_features)
        pool_features=self.manu[-1]

        self.industry_rotation = self.get_core_pool_ret(end_date, '行业轮动',bmk_features,pool_features)
        pool_features=self.industry_rotation[-1]
        theme_weight=pool_features['市场风格暴露'].iloc[-6:-1]
        pool_features['HB主题池风格暴露']=\
            pool_features['HBTMT风格暴露']*theme_weight.iloc[0]+pool_features['HB制造风格暴露']*theme_weight.iloc[1]+\
        pool_features['HB周期风格暴露']*theme_weight.iloc[2]+pool_features['HB大金融风格暴露']*theme_weight.iloc[3]+\
        pool_features['HB消费风格暴露']*theme_weight.iloc[4]

        for col in pool_features.columns:
            pool_features[col]=pool_features[col].map("{:.2%}".format)

        table_width_pic=1800
        table_width_performance = 1000
        pool_avg_ret=[]
        pool_name_list=['偏股型基金','君瑞','价值池','成长池','动量','行业轮动'
            ,'HB制造','HB周期','HB消费','HBTMT','HB大金融']


        data_map=dict(zip(pool_name_list
                          ,[self.benmark_info,self.core_info,
                            self.value_info,self.growth_info,
                            self.quant,self.industry_rotation,self.manu,
                            self.cycle,self.consume,
                            self.tmt,self.finance]))
        pool_avg_ret_year=[]
        for name in data_map.keys():

            pool_avg_ret.append(data_map[name][0].loc['2101'])
            if(len(pool_avg_ret_year)==0):
                pool_avg_ret_year=data_map[name][2].set_index('rqzh')
            else:
                pool_avg_ret_year = pd.merge(pool_avg_ret_year, data_map[name][2], how='inner', on='rqzh')
            pic_data=self.get_jj_simple_pic(data_map[name][1])
            plot.plotly_table(pic_data[['基金名称', '最近一月', '最近一季', '最近一年', '成立以来']]
                              ,table_width_performance,name+'绩优基金业绩',save_local_file=saved_folder)
            pic_list.append(saved_folder + name+'绩优基金业绩' + '.png')
            data_list.append(pic_data)
            plot.plotly_table(pic_data[['基金名称', '一级行业类型', '前五大行业', '风格类型',
           '风格偏好', '规模风格类型', '规模偏好']]
                              ,table_width_pic,name+'绩优基金画像',save_local_file=saved_folder)
            pic_list.append(saved_folder + name+'绩优基金画像' + '.png')
            data_list.append(pic_data)


        pool_avg_ret_year['rqzh']= pool_avg_ret_year['rqzh'].astype(str)
        pool_avg_ret_year.set_index('rqzh',inplace=True)
        pool_avg_ret_year.columns=pool_name_list
        pool_avg_ret_year['HB主题池']=\
            pool_avg_ret_year['HBTMT']*theme_weight.iloc[0]+pool_avg_ret_year['HB制造']*theme_weight.iloc[1]+\
        pool_avg_ret_year['HB周期']*theme_weight.iloc[2]+pool_avg_ret_year['HB大金融']*theme_weight.iloc[3]+\
        pool_avg_ret_year['HB消费']*theme_weight.iloc[4]
        pool_name_list.append('HB主题池')

        pool_ret_data=pd.DataFrame(data=pool_avg_ret_year.iloc[-1].values,columns=['月涨幅（%）'],index=pool_name_list)

        length=len(pool_avg_ret_year)
        std=(pool_avg_ret_year[pool_name_list]/100).std()
        for col  in pool_name_list:
            pool_avg_ret_year[col + 'winning']=pool_avg_ret_year[col]>pool_avg_ret_year['偏股型基金']

        for col in pool_name_list:
            pool_avg_ret_year[col]=(pool_avg_ret_year[col]/100+1).cumprod()



        plot.plotly_jjpic_bar(pool_ret_data[['月涨幅（%）']], '基金池月涨幅(%)', save_local_file=saved_folder)
        pic_list.append(saved_folder + '基金池月涨幅(%)' + '.png')
        data_list.append(pool_ret_data)

        plot.plotly_line_style(pool_avg_ret_year[pool_name_list], '基金池今年以来走势', save_local_file=saved_folder)
        pic_list.append(saved_folder + '基金池今年以来走势' + '.png')
        data_list.append(pool_avg_ret_year)

        pool_features.index=[x.replace('_当期','') for x in pool_features.index.tolist()]
        plot.plotly_table(pool_features.reset_index(),120*len(pool_features.columns)
                          ,'池暴露对比', save_local_file=saved_folder)
        pic_list.append(saved_folder + '池暴露对比' + '.png')
        data_list.append(pool_features)

        plot.plotly_table(ret_contribution,120*len(pool_features.columns)
                          ,'君瑞收益拆解', save_local_file=saved_folder)
        pic_list.append(saved_folder + '君瑞收益拆解' + '.png')
        data_list.append(ret_contribution)


        for col  in pool_name_list:
            pool_avg_ret_year[col + 'max']=pool_avg_ret_year[col].rolling(length,1).max()
            pool_avg_ret_year[col + 'max']=pool_avg_ret_year[col + 'max']/pool_avg_ret_year[col]-1


        date_length=(datetime.datetime.strptime(end_date, '%Y%m%d')\
                    -datetime.datetime.strptime((end_date[0:4]+'0101'), '%Y%m%d')).days
        std=std*np.sqrt(52/length)

        pool_year_stats=pd.DataFrame()
        pool_year_stats['池类型']=pool_name_list
        pool_year_stats['月胜率(%)']=(pool_avg_ret_year[[x + 'winning' for x in pool_name_list]].sum() / length).values
        pool_year_stats['年化收益(%)']=((pool_avg_ret_year[pool_name_list].iloc[-1]).pow(365/date_length)-1).values
        pool_year_stats['截止本月收益(%)'] = (pool_avg_ret_year[pool_name_list].iloc[-1] - 1).values
        pool_year_stats['年化夏普']=pool_year_stats['年化收益(%)']/std.values
        pool_year_stats['最大回撤(%)'] = (pool_avg_ret_year[[x + 'max' for x in pool_name_list]].max()).values
        for col in ['月胜率(%)','最大回撤(%)','年化收益(%)','截止本月收益(%)']:
            pool_year_stats[col]=pool_year_stats[col].map("{:.2%}".format)

        plot.plotly_table(pool_year_stats,table_width_pic, '基金池业绩统计表', save_local_file=saved_folder)
        pic_list.append(saved_folder + '基金池业绩统计表' + '.png')
        data_list.append(pool_ret_data)


        self.pic_list=pic_list
        self.data_list=data_list

    def get_fund_current_holding_features(self,jjdm_List,asofdate,if_zgbl=True,other_table=None):

        if(other_table is not None):
            table=other_table
        else:
            table=''

        sql="SELECT * from {2}hbs_industry_class1_exp where jjdm in ({0}) and  jsrq='{1}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate,table)
        industry_date=pd.read_sql(sql,con=localdb)
        industry_date=pd.merge(industry_date,self.ind2thememap
                               ,how='left',left_on='yjxymc',right_on='industry_name')


        theme_date=industry_date.groupby(['jjdm','theme']).sum()['zjbl'].reset_index().pivot_table('zjbl','jjdm','theme').fillna(0)
        industry_date=industry_date.pivot_table('zjbl','jjdm','yjxymc').fillna(0)


        # sql="SELECT * from {2}hbs_industry_class2_exp where jjdm in ({0}) and  jsrq='{1}'"\
        #     .format(util.list_sql_condition(jjdm_List),asofdate,table)
        # industry2_date=pd.read_sql(sql,con=localdb)
        # industry2_date = industry2_date.pivot_table('zjbl', 'jjdm', 'ejxymc').fillna(0)


        theme_list=theme_date.columns.tolist()
        industry_list=industry_date.columns.tolist()
        # industry_list2 = industry2_date.columns.tolist()
        # industry_list3 = industry3_date.columns.tolist()
        industry_date['total'] = industry_date.sum(axis=1)
        # industry2_date['total'] = industry2_date.sum(axis=1)
        # industry3_date['total'] = industry3_date.sum(axis=1)
        theme_date['total'] = theme_date.sum(axis=1)


        sql="select jjdm,gptzzjb from st_fund.t_st_gm_zcpz where jsrq>='{0}' and jsrq<='{1}' and jjdm in ({2})"\
            .format(asofdate[0:6]+'01',asofdate[0:6]+'31',util.list_sql_condition(jjdm_List))
        hkzb=hbdb.db2df(sql,db='funduser')
        industry_date=pd.merge(industry_date,hkzb,how='left',on='jjdm')
        industry_date['港股占比']=(industry_date['gptzzjb']-industry_date['total'])/100
        industry_date.set_index('jjdm',inplace=True)

        #read fund style,size,industry,theme exp history from local db
        sql="SELECT * from {2}hbs_style_exp where jjdm in ({0}) and  jsrq='{1}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate,table)
        style_data=pd.read_sql(sql,con=localdb).pivot_table('zjbl','jjdm','style_type').fillna(0)
        style_data['total']=style_data['价值']+style_data['成长']
        style_data=pd.merge(style_data,industry_date['total'].to_frame('know_weight')
                            ,left_index=True,right_index=True)


        sql="SELECT * from {2}hbs_size_exp where jjdm in ({0}) and  jsrq='{1}'"\
            .format(util.list_sql_condition(jjdm_List),asofdate,table)
        size_data=pd.read_sql(sql,con=localdb).pivot_table('zjbl','jjdm','size_type').fillna(0)
        size_col=size_data.columns.tolist()
        size_data['total']=size_data.sum(axis=1)
        size_data=pd.merge(size_data,industry_date['total'].to_frame('know_weight')
                            ,left_index=True,right_index=True)

        if(if_zgbl):
            #if want zsbl then run the below code; if want zjbl skip the below code
            for col in industry_list:
                industry_date[col]=industry_date[col]/industry_date['total']
            # for col in industry_list2:
            #     industry2_date[col]=industry2_date[col]/industry2_date['total']
            for col in theme_list:
                theme_date[col]=theme_date[col]/theme_date['total']

            for col in ['价值', '成长']:
                style_data[col] = style_data[col] / style_data['total']
            for col in size_col:
                size_data[col] = size_data[col] / size_data['total']

        else:
            for col in ['价值','成长']:
                style_data[col]=style_data[col]*(style_data['know_weight']/style_data['total'])
            for col in size_col:
                size_data[col]=size_data[col]*(size_data['know_weight']/size_data['total'])


        cons_df=style_data.copy().drop(['total','know_weight'],axis=1)
        cons_df=pd.merge(cons_df,size_data.drop(['total','know_weight'],axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,industry_date.drop(['total','gptzzjb'],axis=1),how='left',on='jjdm')
        # cons_df=pd.merge(cons_df,industry2_date.drop('total',axis=1),how='left',on='jjdm')
        # cons_df=pd.merge(cons_df,industry3_date.drop('total',axis=1),how='left',on='jjdm')
        cons_df=pd.merge(cons_df,theme_date.drop('total',axis=1),how='left',on='jjdm')
        cons_df.columns=[x+'_当期' for x in cons_df.columns]

        return cons_df

    @staticmethod
    def get_main_index_weekly_return_table(start_date,end_date):

        main_index_list=['000001','399001','000016'
            ,'000300','000905','399006']
        main_index_name=['上证指数','深证指数','上证50'
            ,'沪深300','中证500','创业板指']
        name_map=dict(zip(main_index_list,main_index_name))

        sql="select zqdm as 证券代码,zqmc as 证券简称 ,jyrq,kpjg as 开盘价,spjg as 收盘价 ,zgjg as 最高价 ,zdjg as 最低价 ,cjsl as 月成交量_亿 ,cjjs as 月成交额_亿,pe from st_market.t_st_zs_hqql where zqdm in ({0}) and jyrq>='{1}' and jyrq<='{2}' "\
            .format(util.list_sql_condition(main_index_list),start_date,end_date)
        index_return=hbdb.db2df(sql,db='alluser')

        sql="select zqdm,hb1y,rqzh from st_market.t_st_zs_yhb where rqzh>={0} and rqzh<={1} and zqdm in ({2})"\
            .format(end_date[0:4]+'0101' ,end_date,util.list_sql_condition(main_index_list))
        weekly_ret=hbdb.db2df(sql,db='alluser')
        weekly_ret['hb1y']=weekly_ret['hb1y']/100


        cjl=index_return.groupby(['证券代码','证券简称']).sum()[['月成交量_亿','月成交额_亿']]/100000000
        price_list=[index_return.groupby(['证券代码','证券简称']).max()[['最高价']]
                          ,index_return.groupby(['证券代码','证券简称']).min()[['最低价']]
                          ,index_return.drop_duplicates(['证券代码','证券简称'],keep='last')[['证券代码','证券简称','收盘价']]
                          ,index_return.drop_duplicates(['证券代码','证券简称'],keep='first')[['证券代码','证券简称','开盘价']]]

        for df in price_list:
            cjl=pd.merge(cjl,df,how='left',on=['证券代码','证券简称'])

        cjl=pd.merge(cjl,weekly_ret.drop_duplicates('zqdm', keep='last'),
                     how='left',left_on='证券代码',right_on='zqdm').drop(['rqzh','zqdm'],axis=1).rename(columns={'hb1y':'月涨幅(%)'})

        cjl['月涨幅(%)']=(cjl['月涨幅(%)']).map("{:.2%}".format)

        weekly_ret['hb1y'] = weekly_ret['hb1y']+1
        weekly_ret['nav']=weekly_ret.groupby('zqdm').cumprod()['hb1y']
        weekly_ret['rqzh'] = weekly_ret['rqzh'].astype(str)
        weekly_ret=weekly_ret.pivot_table('nav','rqzh','zqdm')
        weekly_ret.rename(columns=name_map,inplace=True)


        return cjl[['证券代码', '证券简称','月涨幅(%)','收盘价', '开盘价','最高价', '最低价','月成交量_亿', '月成交额_亿']],weekly_ret

    @staticmethod
    def get_fund_index_weekly_return_table(end_date):

        # get the quarter end date
        sql = "SELECT jyrq JYRQ FROM st_main.t_st_gg_jyrl WHERE jyrq >= {0} and jyrq<={1} and sfjj=0 and sfym=1".format(
            end_date[0:4]+'0101', end_date)
        calander = hbdb.db2df(sql, db='alluser').sort_values('JYRQ')['JYRQ'].tolist()

        fund_index=['HM0013','HM0034','HM0035','HM0037']
        fund_name=['普通股票型','平衡混合型','灵活配置型','偏股型合型']
        name_map=dict(zip(fund_index,fund_name))

        sql="select zsdm as 指数代码,hb1y,jzrq  from st_fund.t_st_gm_clzs where zsdm in ({0}) and jzrq in ({1}) and hb1y!=99999 "\
            .format(util.list_sql_condition(fund_index),util.list_sql_condition(calander))
        weekly_ret=hbdb.db2df(sql,db='funduser')
        index_return=weekly_ret[weekly_ret['jzrq']==end_date]
        index_return.index=fund_name
        index_return.rename(columns={'hb1y':'月涨幅(%)'},inplace=True)

        weekly_ret['hb1y']=weekly_ret['hb1y']/100
        weekly_ret['hb1y'] = weekly_ret['hb1y']+1
        weekly_ret['nav']=weekly_ret.groupby('指数代码').cumprod()['hb1y']
        weekly_ret=weekly_ret.pivot_table('nav','jzrq','指数代码')
        weekly_ret.rename(columns=name_map,inplace=True)

        return index_return[['月涨幅(%)']],weekly_ret

    @staticmethod
    def get_industry_weekly_return_table(end_date):

        industry_index_list=['801010','801030','801040','801050','801080','801110',
                              '801120','801130','801140','801150','801160','801170',
                              '801180','801200','801210','801230','801710','801720',
                              '801730','801740','801750','801760','801770','801780',
                              '801790','801880','801890','801950','801960','801970','801980']
        industry_name_list=['农林牧渔','基础化工','钢铁','有色金属','电子','家用电器','食品饮料',
                           '纺织服饰','轻工制造','医药生物','公用事业','交通运输','房地产','商贸零售',
                           '社会服务','综合','建筑材料','建筑装饰','电力设备','国防军工','计算机','传媒',
                           '通信','银行','非银金融','汽车','机械设备','煤炭','石油石化','环保','美容护理']
        name_map=dict(zip(industry_index_list,industry_name_list))


        sql="select zqdm as 证券代码,hb1y,rqzh from st_market.t_st_zs_yhb where zqdm in ({0}) and rqzh<='{1}' and rqzh>='{2}'"\
            .format(util.list_sql_condition(industry_index_list),end_date,end_date[0:4]+'0101')
        weekly_ret=hbdb.db2df(sql,db='alluser')
        index_return=weekly_ret[weekly_ret['rqzh']==end_date]
        index_return['证券代码']=[name_map[x] for x  in index_return['证券代码']]
        index_return.set_index('证券代码',inplace=True)
        index_return.sort_values('hb1y', ascending=False,inplace=True)
        index_return.rename(columns={'hb1y':'月涨幅%'},inplace=True)

        weekly_ret['hb1y']=weekly_ret['hb1y']/100
        weekly_ret['hb1y'] = weekly_ret['hb1y']+1
        weekly_ret['nav']=weekly_ret.groupby('证券代码').cumprod()['hb1y']
        weekly_ret['rqzh'] = weekly_ret['rqzh'].astype(str)
        weekly_ret=weekly_ret.pivot_table('nav','rqzh','证券代码')
        weekly_ret.rename(columns=name_map,inplace=True)

        return index_return[['月涨幅%']],\
               weekly_ret[index_return.index[[0,1,2,3,4,-5,-4,-3,-2,-1]].tolist()]

    @staticmethod
    def get_style_weekly_return_table(end_date):

        style_index_list=['399370', '399371','399314','399315','399316']
        style_name_list=['成长','价值','大盘','中盘','小盘']
        name_map=dict(zip(style_index_list,style_name_list))

        sql="select zqdm as 证券代码,hb1y,rqzh from st_market.t_st_zs_yhb where zqdm in ({0}) and rqzh<='{1}' and rqzh>='{2}'"\
            .format(util.list_sql_condition(style_index_list),end_date,end_date[0:4]+'0101')
        weekly_ret=hbdb.db2df(sql,db='alluser')
        index_return=weekly_ret[weekly_ret['rqzh']==end_date]
        index_return['证券代码']=[name_map[x] for x  in index_return['证券代码']]
        index_return.set_index('证券代码',inplace=True)
        index_return.sort_values('hb1y', ascending=False,inplace=True)
        index_return.rename(columns={'hb1y':'月涨幅%'},inplace=True)


        weekly_ret['hb1y']=weekly_ret['hb1y']/100
        weekly_ret['hb1y'] = weekly_ret['hb1y']+1
        weekly_ret['nav']=weekly_ret.groupby('证券代码').cumprod()['hb1y']
        weekly_ret['rqzh'] = weekly_ret['rqzh'].astype(str)
        weekly_ret=weekly_ret.pivot_table('nav','rqzh','证券代码')
        weekly_ret.rename(columns=name_map,inplace=True)


        return index_return[['月涨幅%']],weekly_ret

    @staticmethod
    def get_wind_theme_weekly_return_table():

        data=pd.read_excel(r"E:\GitFolder\docs\公募基金周报\概念指数收益.xlsx")
        data.columns=['证券代码','证券简称','月涨跌幅(%)']
        data=data.sort_values('月涨跌幅(%)',ascending=False)
        data=data.set_index('证券简称')

        return  data.iloc[0:5][['月涨跌幅(%)']],data.iloc[-5:][['月涨跌幅(%)']]

    @staticmethod
    def get_benchmark_ret(end_date,asofdate):

        benchmark_list=util.get_mutual_stock_funds(asofdate)

        sql="select jjdm,zbnp,jzrq,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb in ('2101','2103','2201','2999')   "\
            .format(util.list_sql_condition(benchmark_list)
                    ,end_date)
        benchmark_ret=(hbdb.db2df(sql,db='funduser').pivot_table('zbnp','jjdm','zblb')/100).sort_values('2101',ascending=False)


        sql="select rqzh,hb1y from st_market.t_st_zs_yhb where zqdm='885001' and rqzh>='{0}' and rqzh<='{1}' and hb1y !=99999   "\
            .format(str(int(end_date[0:4]))+'0101',end_date)
        weekly_ret = hbdb.db2df(sql, db='alluser')
        weekly_ret['rqzh']=weekly_ret['rqzh'].astype(int)

        median_ret=benchmark_ret[['2101']].median(axis=0)

        # sql="select rqzh,AVG(hb1y) from st_fund.t_st_gm_yhb where jjdm in ({0}) and rqzh>='{1}' and rqzh<='{2}' and hb1y !=99999 group by rqzh"\
        #     .format(util.list_sql_condition(benchmark_list),str(int(end_date[0:4]))+'0101',end_date)
        # weekly_ret=hbdb.db2df(sql,db='funduser')

        best_5=benchmark_ret.iloc[0:5]


        name_map=dict(zip(['2101','2103','2201','2999'],['最近一月','最近一季','最近一年','成立以来']))
        best_5.rename(columns=name_map,inplace=True)

        jjjc=hbdb.db2df("select jjdm,jjjc from st_fund.t_st_gm_jjxx where jjdm in ({})"
                                .format(util.list_sql_condition(best_5.index.tolist()))
                        ,db='funduser')
        best_5=pd.merge(best_5,jjjc,how='left',on='jjdm').rename(columns={'jjjc':'基金名称'})


        return  median_ret,best_5[['基金名称','最近一月','最近一季','最近一年','成立以来']],weekly_ret

    @staticmethod
    def get_core_pool_ret(end_date,dbname,input_bmk_features,pool_features):

        quant_pool_history=pd.read_sql("select jjdm,asofdate from pool_history where asofdate>='{1}' and asofdate<='{2}' and pool_name='{0}'"
                                       .format(dbname,str(int(end_date[0:4])-1)+'0101',end_date),con=localdb)
        max_asofdate=quant_pool_history['asofdate'].max()
        jjdm_list=quant_pool_history[quant_pool_history['asofdate']==max_asofdate]['jjdm'].tolist()

        #get the pool features
        bmk_features=input_bmk_features.reset_index()
        this_pool_features=bmk_features[bmk_features['jjdm'].isin(jjdm_list)].set_index('jjdm').mean(axis=0)
        this_pool_features=\
            pd.concat([this_pool_features[0:5], this_pool_features[-5:], this_pool_features[36:37]], axis=0)
        pool_features=pd.merge(pool_features,this_pool_features.to_frame(dbname+'风格暴露')
                               ,how='inner',left_index=True,right_index=True)

        indus_weight=\
            bmk_features[bmk_features['jjdm'].isin(jjdm_list)].set_index('jjdm')[bmk_features.columns[6:37]]


        sql="select jjdm,zbnp,jzrq,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb in ('2101','2103','2201','2999')   "\
            .format(util.list_sql_condition(jjdm_list)
                    ,end_date)
        quant_pool_ret=hbdb.db2df(sql,db='funduser').pivot_table('zbnp','jjdm','zblb')

        sql="select jjdm,hbdr from st_fund.t_st_gm_rhb where jjdm in ({}) and jzrq='20221010'"\
            .format(util.list_sql_condition(jjdm_list))
        remove_ret=hbdb.db2df(sql,db='funduser').set_index('jjdm')
        quant_pool_ret=pd.merge(quant_pool_ret,remove_ret,how='left',on='jjdm')
        quant_pool_ret=quant_pool_ret/100

        for col in ['2101','2103','2201','2999']:
            quant_pool_ret[col]=(quant_pool_ret[col]+1)/(quant_pool_ret['hbdr']+1)-1

        quant_pool_ret.drop('hbdr',axis=1,inplace=True)
        quant_return=quant_pool_ret[['2101']].mean(axis=0)

        name_map=dict(zip(['2101','2103','2201','2999'],['最近一月','最近一季','最近一年','成立以来']))
        quant_pool_ret.rename(columns=name_map,inplace=True)

        jjjc=hbdb.db2df("select jjdm,jjjc from st_fund.t_st_gm_jjxx where jjdm in ({})"
                                .format(util.list_sql_condition(jjdm_list))
                        ,db='funduser')
        quant_pool_ret=pd.merge(quant_pool_ret,jjjc,how='left',on='jjdm').rename(columns={'jjjc':'基金名称'})

        quant_pool_list = []
        for i in range(4):
            quant_pool_list.append(jjjc.iloc[i*5:(i+1)*5].reset_index(drop=True)[['jjjc']].rename(columns={'jjjc':'基金名称'}))
        #shall be removed later
        quant_pool_history['asofdate']='20220401'

        asofdate_list=quant_pool_history['asofdate'].unique().tolist()+[end_date]
        weekly_ret=[]
        for i in range(len(asofdate_list)-1):

            sql="select rqzh,AVG(hb1y) from st_fund.t_st_gm_yhb where jjdm in ({0}) and rqzh>'{1}' and rqzh<='{2}' and hb1y !=99999 group by rqzh"\
                .format(util.list_sql_condition(quant_pool_history[quant_pool_history['asofdate']==asofdate_list[i]]['jjdm'].tolist())
                        ,asofdate_list[i],asofdate_list[i+1])
            weekly_ret.append(hbdb.db2df(sql,db='funduser'))

        weekly_ret=pd.concat(weekly_ret,axis=0).reset_index(drop=True)

        weekly_ret.loc[weekly_ret['rqzh']==20221031,'AVG(hb1y)']=quant_return.values[0]*100-0.05176

        return  quant_return\
            ,quant_pool_ret.sort_values('最近一月',ascending=False)[['基金名称','最近一月','最近一季','最近一年','成立以来']][0:5]\
            ,weekly_ret,indus_weight,pool_features

    @staticmethod
    def get_value_pool_ret(end_date):


        latest_date=pd.read_sql("select max(asofdate) as mas_date from jjpic_value_p_hbs"
                                ,con=localdb)['mas_date'][0]

        value_pool=pd.read_sql("select * from jjpic_value_p_hbs where asofdate='{}'  "
                              .format(latest_date),con=localdb).drop('asofdate',axis=1)
        growth_pool = value_pool.sort_values('成长暴露排名(持仓)', ascending=False)[0:20]
        value_pool=value_pool.sort_values('价值暴露排名(持仓)',ascending=False)[0:20]

        sql="select jjdm,zbnp,jzrq,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb in ('2007','2103','2201','2999')   "\
            .format(util.list_sql_condition(value_pool['jjdm'].tolist())
                    ,end_date)
        value_pool_ret=hbdb.db2df(sql,db='funduser').pivot_table('zbnp','jjdm','zblb')
        sql="select jjdm,zbnp,jzrq,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb in ('2007','2103','2201','2999')   "\
            .format(util.list_sql_condition(growth_pool['jjdm'].tolist())
                    ,end_date)
        growth_pool_ret=hbdb.db2df(sql,db='funduser').pivot_table('zbnp','jjdm','zblb')


        sql="select rqzh,AVG(hb1z) from st_fund.t_st_gm_zhb where jjdm in ({0}) and rqzh>='{1}' and rqzh<='{2}' and hb1z !=99999 group by rqzh"\
            .format(util.list_sql_condition(value_pool['jjdm'].tolist()),str(int(end_date[0:4]))+'0101',end_date)
        weekly_ret_value=hbdb.db2df(sql,db='funduser')
        sql="select rqzh,AVG(hb1z) from st_fund.t_st_gm_zhb where jjdm in ({0}) and rqzh>='{1}' and rqzh<='{2}' and hb1z !=99999 group by rqzh"\
            .format(util.list_sql_condition(growth_pool['jjdm'].tolist()),str(int(end_date[0:4]))+'0101',end_date)
        weekly_ret_growth=hbdb.db2df(sql,db='funduser')

        value_pool_ret=value_pool_ret/100
        growth_pool_ret=growth_pool_ret/100

        value_pool_return=value_pool_ret[['2007']].mean(axis=0)
        growth_pool_return=growth_pool_ret[['2007']].mean(axis=0)

        name_map=dict(zip(['2007','2103','2201','2999'],['最近一周','最近一季','最近一年','成立以来']))
        value_pool_ret.rename(columns=name_map,inplace=True)
        growth_pool_ret.rename(columns=name_map,inplace=True)

        jjjc=hbdb.db2df("select jjdm,jjjc from st_fund.t_st_gm_jjxx where jjdm in ({})"
                                .format(util.list_sql_condition(value_pool['jjdm'].tolist()+growth_pool['jjdm'].tolist()))
                        ,db='funduser')
        value_pool_ret=pd.merge(value_pool_ret,jjjc,how='left',on='jjdm').rename(columns={'jjjc':'基金名称'})
        growth_pool_ret = pd.merge(growth_pool_ret, jjjc, how='left', on='jjdm').rename(columns={'jjjc':'基金名称'})

        value_pool=value_pool[['基金简称', '价值绝对暴露(持仓)']]
        value_pool['价值绝对暴露(持仓)']=value_pool['价值绝对暴露(持仓)'].map("{:.2%}".format)

        growth_pool = growth_pool[['基金简称', '成长绝对暴露(持仓)']]
        growth_pool['成长绝对暴露(持仓)'] = growth_pool['成长绝对暴露(持仓)'].map("{:.2%}".format)

        value_pool_list=[]
        growth_pool_list = []
        for i in range(4):
            value_pool_list.append(value_pool.iloc[i*5:(i+1)*5].reset_index(drop=True))
            growth_pool_list.append(growth_pool.iloc[i*5:(i+1)*5].reset_index(drop=True))



        return value_pool_return,pd.concat(value_pool_list,axis=1)\
            ,value_pool_ret.sort_values('最近一周',ascending=False)[0:5][['基金名称','最近一周','最近一季','最近一年','成立以来']]\
            ,growth_pool_return,pd.concat(growth_pool_list,axis=1)\
            ,growth_pool_ret.sort_values('最近一周',ascending=False)[0:5][['基金名称','最近一周','最近一季','最近一年','成立以来']]\
            ,weekly_ret_value,weekly_ret_growth

    @staticmethod
    def get_advance_pool_ret(end_date):

        latest_date=pd.read_sql("select max(asofdate) as latest_date from pool_advance_man where asofdate<='{}'".format(end_date)
                                ,con=localdb)['latest_date'][0]
        pool_list=pd.read_sql("select jjdm from pool_advance_man where asofdate='{}'".format(latest_date)
                              ,con=localdb)['jjdm'].tolist()
        #
        # pool_list=['001856','000924','001410','001476','001975','005968','270028','377240']

        sql="select jjdm,zbnp,jzrq,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb in ('2007','2103','2201','2999')   "\
            .format(util.list_sql_condition(pool_list)
                    ,end_date)
        advance_pool_ret=hbdb.db2df(sql,db='funduser').pivot_table('zbnp','jjdm','zblb')
        advance_pool_ret=advance_pool_ret/100
        advance_return=advance_pool_ret[['2007']].mean(axis=0)

        name_map=dict(zip(['2007','2103','2201','2999'],['最近一周','最近一季','最近一年','成立以来']))
        advance_pool_ret.rename(columns=name_map,inplace=True)

        jjjc=hbdb.db2df("select jjdm,jjjc,glrm from st_fund.t_st_gm_jjxx where jjdm in ({})"
                                .format(util.list_sql_condition(pool_list))
                        ,db='funduser')
        advance_pool_ret=pd.merge(advance_pool_ret,jjjc,how='left',on='jjdm').rename(columns={'jjjc':'基金名称'})

        sql="select rqzh,AVG(hb1z) from st_fund.t_st_gm_zhb where jjdm in ({0}) and rqzh>='{1}' and rqzh<='{2}' and hb1z !=99999 group by rqzh"\
            .format(util.list_sql_condition(pool_list),str(int(end_date[0:4]))+'0101',end_date)
        weekly_ret=hbdb.db2df(sql,db='funduser')

        return  advance_return\
            ,advance_pool_ret.sort_values('最近一周',ascending=False)[['基金名称','最近一周','最近一季','最近一年','成立以来']][0:5]\
            ,weekly_ret

    @staticmethod
    def get_quant_pool_ret(end_date,dbname):

        quant_pool_history=pd.read_sql("select jjdm,asofdate from {0} where asofdate>='{1}' and asofdate<='{2} '"
                                       .format(dbname,str(int(end_date[0:4])-1)+'1201',end_date),con=localdb)
        max_asofdate=quant_pool_history['asofdate'].max()
        jjdm_list=quant_pool_history[quant_pool_history['asofdate']==max_asofdate]['jjdm'].tolist()

        sql="select jjdm,zbnp,jzrq,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb in ('2007','2103','2201','2999')   "\
            .format(util.list_sql_condition(jjdm_list)
                    ,end_date)
        quant_pool_ret=hbdb.db2df(sql,db='funduser').pivot_table('zbnp','jjdm','zblb')
        quant_pool_ret=quant_pool_ret/100
        quant_return=quant_pool_ret[['2007']].mean(axis=0)

        name_map=dict(zip(['2007','2103','2201','2999'],['最近一周','最近一季','最近一年','成立以来']))
        quant_pool_ret.rename(columns=name_map,inplace=True)

        jjjc=hbdb.db2df("select jjdm,jjjc from st_fund.t_st_gm_jjxx where jjdm in ({})"
                                .format(util.list_sql_condition(jjdm_list))
                        ,db='funduser')
        quant_pool_ret=pd.merge(quant_pool_ret,jjjc,how='left',on='jjdm').rename(columns={'jjjc':'基金名称'})

        quant_pool_list = []
        for i in range(4):
            quant_pool_list.append(jjjc.iloc[i*5:(i+1)*5].reset_index(drop=True)[['jjjc']].rename(columns={'jjjc':'基金名称'}))

        asofdate_list=quant_pool_history['asofdate'].unique().tolist()+[end_date]
        weekly_ret=[]
        for i in range(len(asofdate_list)-1):

            sql="select rqzh,AVG(hb1z) from st_fund.t_st_gm_zhb where jjdm in ({0}) and rqzh>'{1}' and rqzh<='{2}' and hb1z !=99999 group by rqzh"\
                .format(util.list_sql_condition(quant_pool_history[quant_pool_history['asofdate']==asofdate_list[i]]['jjdm'].tolist())
                        ,asofdate_list[i],asofdate_list[i+1])
            weekly_ret.append(hbdb.db2df(sql,db='funduser'))

        weekly_ret=pd.concat(weekly_ret,axis=0).reset_index(drop=True)

        return  quant_return,pd.concat(quant_pool_list,axis=1)\
            ,quant_pool_ret.sort_values('最近一周',ascending=False)[['基金名称','最近一周','最近一季','最近一年','成立以来']][0:5]\
            ,weekly_ret

    @staticmethod
    def get_jj_simple_pic(outdf):

        for col in ['最近一月', '最近一季', '最近一年', '成立以来']:
            outdf[col]=outdf[col].map("{:.2%}".format)


        jjjc_list=outdf['基金名称'].tolist()

        asofdate=pd.read_sql("select max(asofdate) as mas_date from jjpic_value_p_hbs"
                                ,con=localdb)['mas_date'][0]
        value_pic="select `基金简称`,`风格类型`,`风格偏好` from jjpic_value_p_hbs where `基金简称` in ({0}) and asofdate='{1}'"\
            .format(util.list_sql_condition(jjjc_list),asofdate)
        value_pic=pd.read_sql(value_pic,con=localdb)

        asofdate=pd.read_sql("select max(asofdate) as mas_date from jjpic_size_p_hbs"
                                ,con=localdb)['mas_date'][0]
        size_pic="select `基金简称`,`规模风格类型`,`规模偏好` from jjpic_size_p_hbs where `基金简称` in ({0}) and asofdate='{1}'"\
            .format(util.list_sql_condition(jjjc_list),asofdate)
        size_pic = pd.read_sql(size_pic, con=localdb)

        asofdate=pd.read_sql("select max(asofdate) as mas_date from jjpic_industry_p"
                                ,con=localdb)['mas_date'][0]
        ind_pic="select `基金简称`,`一级行业类型`,`前五大行业` from jjpic_industry_p where `基金简称` in ({0}) and asofdate='{1}'"\
            .format(util.list_sql_condition(jjjc_list),asofdate)
        ind_pic = pd.read_sql(ind_pic, con=localdb)


        outdf = pd.merge(outdf, ind_pic, how='left',left_on='基金名称',right_on='基金简称').drop('基金简称',axis=1)
        outdf=pd.merge(outdf,value_pic,how='left',left_on='基金名称',right_on='基金简称').drop('基金简称',axis=1)
        outdf=pd.merge(outdf,size_pic,how='left',left_on='基金名称',right_on='基金简称').drop('基金简称',axis=1)

        return  outdf

    @staticmethod
    def simple_desc(data,key_word):

        if(key_word=='主要指数月收益'):
            desc = ""
            data=data.set_index('证券简称')
            for index_name in data.index:
                split_data=data.loc[index_name]['月涨幅(%)'].split('-')
                if(len(split_data)>1):
                    desc += index_name + "月内下跌" + split_data[1]+'，'
                else:
                    desc += index_name + "月内上涨" + split_data[0]+'，'
            desc=desc[0:-1]+'。'

        elif(key_word=='偏股型合基金月收益率（%）'):
            desc='上月偏股型合型基金中，'
            for index_name in data.index:
                split_data=data.loc[index_name]['月涨幅(%)']
                if(split_data<0):
                    desc += index_name + "下跌" + str(abs(split_data))+'%，'
                else:
                    desc += index_name + "上涨" + str(abs(split_data))+'%，'
            desc = desc[0:-1] + '。'

        elif('申万一级行业月收益' in key_word):
            desc='其中'+data.index[0]+"涨幅最大，而"+data.index[-1]+"跌幅最大。具体的，"
            for index_name in data.index:
                split_data=data.loc[index_name]['月涨幅%']
                if (split_data < 0):
                    desc += index_name + "下跌" + str(abs(split_data)) + '%，'
                else:
                    desc += index_name + "上涨" + str(abs(split_data)) + '%，'
            desc = desc[0:-1] + '。'

        elif('WIND概念月收益' in key_word):
            desc='上月在万得 771 个概念板块指数中'+data.index[0]+"涨幅最大，而"+data.index[-1]+"跌幅最大。具体的，"
            for index_name in data.index:
                split_data=data.loc[index_name]['月涨跌幅(%)']
                if (split_data < 0):
                    desc += index_name + "下跌" + str(abs(split_data)) + '%，'
                else:
                    desc += index_name + "上涨" + str(abs(split_data)) + '%，'
            desc = desc[0:-1] + '。'

        elif ('巨潮风格指数月收益' in key_word):
            desc =''
            for index_name in [ '价值','成长','大盘', '中盘','小盘']:
                split_data = data.loc[index_name]['月涨幅%']
                if (split_data < 0):
                    desc += index_name + "下跌" + str(abs(split_data)) + '%，'
                else:
                    desc += index_name + "上涨" + str(abs(split_data)) + '%，'
            desc = desc[0:-1] + '。'

        elif ('绩优偏股型基金业绩' in key_word):
            desc = '偏股型合个基方面，' + data['基金名称'][0]+","+data['基金名称'][1]+'等基金表现优异。在表现最优异的五只个基中，'

        elif ('绩优核心池基金业绩' in key_word):
            desc = '在核心池中，' +data['基金名称'][0] + "," + data['基金名称'][1] + '等基金表现优异。在表现最优异的五只基金中，'

        elif ('绩优价值池基金业绩' in key_word):
            desc = '在价值基金池中，' + data['基金名称'][0] + "," + data['基金名称'][1] + '等基金表现优异。在表现最优异的五只基金中，'

        elif ('绩优成长池基金业绩' in key_word):
            desc = '在成长基金池中，' + data['基金名称'][0] + "," + data['基金名称'][1] + '等基金表现优异。在表现最优异的五只基金中，'

        elif ('绩优先进制造基金业绩' in key_word):
            desc = '在先进制造基金池中，' + data['基金名称'][0] + "," + data['基金名称'][1] + '等基金表现优异。在表现最优异的五只基金中，'

        else:
            desc=''


        return desc

    def generate_weekly_report_draft(self):

        pic_name_list=self.pic_list
        end_date=self.end_date
        data_list=self.data_list

        doc = Document()

        pic_paragraphs_list = []
        string_paragraphs_list = []
        paragraphs_count = 0

        doc.add_paragraph("")
        paragraphs_count += 1

        p = doc.add_paragraph("按照季度，对公墓核心池历史进行了统计分析，结果如下")
        p.style.font.size = Pt(11)
        p.paragraph_format.first_line_indent = p.style.font.size * 2
        string_paragraphs_list.append(paragraphs_count)
        paragraphs_count += 1

        for i in range(len(pic_name_list)):
            pic = pic_name_list[i]
            data=data_list[i]
            key_word=pic.split('\\\\')[1].split('.png')[0]
            file_path = pic
            desc = self.simple_desc(data,key_word)

            doc.add_paragraph("")
            paragraphs_count += 1

            p = doc.add_paragraph(desc)
            p.style.font.size = Pt(11)
            p.paragraph_format.first_line_indent = p.style.font.size * 2
            string_paragraphs_list.append(paragraphs_count)
            paragraphs_count += 1

            if('画像' in  key_word):
                pic_width=Inches(6.25)
                pic_height=Inches(1.5)
            elif('业绩' in key_word):
                pic_width=Inches(8/1.3)
                pic_height=Inches(2/1.3)
            else:
                pic_width=Inches(8/1.3)
                pic_height=Inches(4/1.3)

            try:

                doc.add_picture(file_path,width=pic_width, height=pic_height)
                pic_paragraphs_list.append(paragraphs_count)
                paragraphs_count += 1

            except Exception as e:
                pic_temp = Image.open(file_path)
                pic_temp.save(pic_temp)
                doc.add_picture(file_path,width=pic_width, height=pic_height)

        for j in pic_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for j in string_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.save(self.folder+"基金池月报草稿.docx")

        print('Done')

def back_test(asofdate_list):

    eq_pool=[]
    momnet_pool=[]
    bmk_df=[]
    #laps_day for quarterly trade
    laps_day=15
    # #laps_day for half yearly  trade
    # laps_day=75


    cp = Customize_pool()
    nav_adj=False
    int_solution=True
    last_solution=\
        pd.read_excel("E:\GitFolder\docs\金工基金池\新算法测试\调仓记录.xlsx"
                      ,sheet_name='Sheet3')
    last_solution['jjdm']=\
        [("000000"+str(x))[-6:] for x in last_solution['jjdm']]
    # last_solution=\
    #     pd.read_excel(r"E:\GitFolder\docs\金工基金池\调仓记录.xlsx",sheet_name='Sheet2')
    # last_solution['jjdm']=\
    #     last_solution['期末权重'].iloc[0:29]
    # last_solution=last_solution.iloc[0:29]
    # last_solution['jjdm']=\
    #     [("000000"+str(x))[-6:] for x in last_solution['jjdm'].astype(int)]



    # last_solution=[]
    last_date=''

    for asofdate in asofdate_list:

        print(asofdate)
        #get the bmk list for each asofdate
        bmk=pd.DataFrame()
        # sql = "select jjdm from pool_history where asofdate='{0}' and pool_name in ('HB周期','HBTMT','HB医药','HB大金融', 'HB消费', 'HB制造')" \
        #     .format(asofdate)
        # extra_jjdm_list=pd.read_sql(sql,con=localdb)['jjdm'].tolist()
        extra_jjdm_list=[]


        if(len(last_solution)>0):
            bmk['jjdm']=list(set(util.get_bmk_funds_list(asofdate)+last_solution['jjdm'].tolist()))+extra_jjdm_list
        else:
            bmk['jjdm'] =util.get_bmk_funds_list(asofdate)+extra_jjdm_list

        bmk['asofdate']=asofdate
        bmk_df.append(bmk)


        # if(asofdate[4:6]=='03' or asofdate[4:6]=='09'):
        #     cons_df=cp.get_fund_current_holding_features(bmk['jjdm'].tolist(), last_date,if_zgbl=True)
        # else:
        #     cons_df = cp.get_fund_current_holding_features(bmk['jjdm'].tolist(), asofdate,if_zgbl=True)

        cons_df = cp.get_fund_current_holding_features(bmk['jjdm'].tolist(), asofdate,if_zgbl=True,other_table='mimic_')
        # cons_df = cp.get_fund_current_holding_features(bmk['jjdm'].tolist(), asofdate,if_zgbl=True)

        #get the jj from theme pool
        # jjdm_list=\
        #     pd.read_sql("select jjdm from pool_history where pool_name in ('advance_manu','eng_chem','med','consume','fin')"
        #                 ,con=localdb)['jjdm'].tolist()

        #get the information ratio from local DB
        target_col='info_ratio'
        target_df=cp.get_target_df(bmk['jjdm'].unique().tolist()
                                   ,(datetime.datetime.strptime(asofdate, '%Y%m%d')
                                     +datetime.timedelta(days=laps_day)).strftime('%Y%m%d')
                                   ,asofdate,target_col+'_2201')
        # target_col='manager_alpha'
        # target_df=cp.get_target_df(bmk['jjdm'].unique().tolist()
        #                            ,(datetime.datetime.strptime(asofdate, '%Y%m%d')
        #                              +datetime.timedelta(days=75)).strftime('%Y%m%d')
        #                            ,target_col)


        # target_df=cp.get_target_df(jjdm_list
        #                            ,(datetime.datetime.strptime(asofdate, '%Y%m%d')
        #                              +datetime.timedelta(days=75)).strftime('%Y%m%d')
        #                            ,target_col+'_2201')
        #get the winning pro from local DB
        # target_col='win_median'
        # #
        # target_df=cp.get_target_df(bmk['jjdm'].unique().tolist()
        #                            ,(datetime.datetime.strptime(asofdate, '%Y%m%d')+datetime.timedelta(days=75)).strftime('%Y%m%d')
        #                            ,target_col+'@quarter')

        # target_df=target_df[target_df[target_col]>=0.55]

        # #get the quantile from local DB
        # target_col='25%'
        # #
        # target_df=cp.get_target_df(bmk['jjdm'].unique().tolist()
        #                            ,(datetime.datetime.strptime(asofdate, '%Y%m%d')
        #                              +datetime.timedelta(days=75)).strftime('%Y%m%d')
        #                            ,target_col+'@half_year')
        #

        #get the combined factors from local DB
        # target_col='new_factor'
        # #
        # target_df=cp.get_target_df(bmk['jjdm'].unique().tolist()
        #                            ,(datetime.datetime.strptime(asofdate, '%Y%m%d')+datetime.timedelta(days=75)).strftime('%Y%m%d')
        #                            ,target_col+"@info_ratio_2201@info_ratio_2106@info_ratio_2101")\
        #remove last 50% target col
        # target_df['rank'] = target_df.rank(method='min')['info_ratio'] / len(target_df)
        # target_df=target_df[target_df['rank']>0.1].drop('rank',axis=1)

        input_df = target_df.copy()
        input_df=input_df[input_df[target_col].notnull()]
        input_df = pd.merge(input_df, cons_df, how='outer', on='jjdm')
        input_df[target_col]=input_df[target_col].fillna(-99999)

        #remove small company
        # sql="select jjdm,glrm from st_fund.t_st_gm_jjxx where jjdm in ({0})"\
        #     .format(util.list_sql_condition(input_df['jjdm'].tolist()))
        # glrm=hbdb.db2df(sql,db='funduser')
        # sql="select glrm,fhbgm from st_fund.t_st_gm_gsfhbgm where jsrq>='{0}' and jsrq<='{1}' and glrm in ({2})"\
        #     .format(asofdate[0:6]+'01',asofdate[0:6]+'31',util.list_sql_condition(glrm['glrm'].tolist()))
        # glgm=hbdb.db2df(sql,db='funduser')
        # company_list=pd.merge(glrm,glgm,
        #                       how='left',on='glrm').drop_duplicates('glrm').sort_values('fhbgm',ascending=False)
        # company_list=company_list.iloc[0:int(np.floor(len(company_list)/4))]
        # glrm=pd.merge(glrm,company_list,how='inner',on='glrm')['jjdm_x'].tolist()
        # input_df=input_df[input_df['jjdm'].isin(glrm)]


        #remove closed fund
        sql="select fbq,jjdm from st_fund.t_st_gm_dkxx "
        remove_list=hbdb.db2df(sql,db='funduser')['jjdm'].unique().tolist()
        input_df=input_df[~input_df['jjdm'].isin(remove_list)]

        #remove the fund that can not sell at anytime
        sql = "select jjdm,jjmc from st_fund.t_st_gm_jjxx where jjmc is not null and cpfl='2' and (ejfl='13' or ejfl='35' or ejfl='37')"
        remove_list=hbdb.db2df(sql,db='funduser')
        remove_list=remove_list[remove_list['jjmc'].str.contains('持有期')]['jjdm'].tolist()
        input_df = input_df[~input_df['jjdm'].isin(remove_list)]

        # remove fund that can not be bought from howbuy
        sql = "select jjdm,xsbz from st_fund.t_st_gm_jjxx where jjdm in ({0}) and xsbz='0'"\
            .format(util.list_sql_condition(bmk['jjdm'].tolist()))
        remove_list = hbdb.db2df(sql, db='funduser')['jjdm'].tolist()
        input_df = input_df[~input_df['jjdm'].isin(remove_list)]


        #remove limit buy upper bound
        sql = "select jjdm,qsrq,zzrq,flmczx,flfysx,syjesx from st_fund.t_st_gm_jjflxx where zzrq>='{0}' and qsrq<='{0}'  and flmczx='1001' "\
            .format((datetime.datetime.strptime(asofdate, '%Y%m%d')
                                     +datetime.timedelta(days=laps_day+25)).strftime('%Y%m%d'))
        remove_list = hbdb.db2df(sql, db='funduser')
        remove_list=remove_list.sort_values('qsrq',ascending=False).drop_duplicates('jjdm',keep='last')
        remove_list=remove_list[remove_list['flfysx']<=250000]['jjdm'].unique().tolist()
        input_df = input_df[~input_df['jjdm'].isin(remove_list)]


        # abs equilibrium pool
        # opti_pool=cp.generate_abs_equilibrium_pool(input_df=input_df,cons_df,asofdate=asofdate
        #                                       ,target_col=target_col)
        # opti_pool2=cp.customize_pool(tempdf=input_df,target_col=target_col
        #                   ,turnover_num=30,asofdate=asofdate,int_solution=True)



        #纯主题池
        # sql="select * from pool_history where asofdate='{0}' and pool_name in ('HB周期','HBTMT','HB医药','HB大金融', 'HB消费', 'HB制造')"\
        #     .format(asofdate)
        # pre_given_jjlist=pd.read_sql(sql,con=localdb)
        # pre_given_jjlist = pd.merge(pre_given_jjlist, cons_df, how='left', on='jjdm')
        # cons_df=cons_df.mean(axis=0)
        #
        # theme_weight = cons_df[['TMT_当期', '制造_当期', '周期_当期', '大金融_当期', '消费_当期', '医药生物_当期']]
        # theme_weight['消费_当期'] = theme_weight['消费_当期'] - theme_weight['医药生物_当期']
        # theme_weight.index = ["HB" + str(x).replace('_当期', '') for x in theme_weight.index]
        # theme_weight = theme_weight / (len(pre_given_jjlist) / len(theme_weight))
        #
        # pre_given_jjlist.loc[pre_given_jjlist['pool_name'] == 'HB医药', 'pool_name'] = 'HB医药生物'
        # pre_given_jjlist = pd.merge(pre_given_jjlist, theme_weight.to_frame('weight')
        #                             , how='left', left_on='pool_name', right_index=True)
        # eq_pool.append(pre_given_jjlist)
        # pd.concat(eq_pool, axis=0).to_excel('纯主题池2.xlsx', index=False)


        opti_pool,last_solution=cp.generate_equilibrium_pool(input_df=input_df,cons_df=cons_df,asofdate=asofdate,last_date=last_date,turnover_num=20
                                              ,target_col=target_col,last_solution=last_solution,allowed_hsl=0.65
                                               ,cons_col=cons_df.columns.to_list()[0:37]+cons_df.columns.to_list()[-5:],int_solution=int_solution,nav_adj=nav_adj,gap=0.02)
        # opti_pool,last_solution=cp.generate_equilibrium_pool(input_df=input_df,cons_df=cons_df,asofdate=asofdate,last_date=last_date,turnover_num=30
        #                                       ,target_col=target_col,last_solution=last_solution,allowed_hsl=0.8
        #                                        ,cons_col=cons_df.columns.to_list()[0:37]+cons_df.columns.to_list()[38:],int_solution=int_solution,nav_adj=nav_adj,gap=0.01)
    #
        market_style=cons_df[cons_df.columns.to_list()[0:37]+cons_df.columns.to_list()[-5:]]
        print(market_style.loc[opti_pool['jjdm']].mean()-market_style.mean())

        last_date=asofdate
        if(len(opti_pool)>0):
            eq_pool.append(opti_pool)
            last_solution['asofdate']=(datetime.datetime.strptime(asofdate, '%Y%m%d')
                                     +datetime.timedelta(days=laps_day)).strftime('%Y%m%d')
        else:
            bmk_df=bmk_df[0:-1]
    # pd.concat(test_df,axis=1).to_excel('重仓bmk.xlsx')
    pd.concat(eq_pool, axis=0).to_excel('均衡池v1.xlsx', index=False)
    cp.return_comparision(pd.concat(bmk_df, axis=0),
                          pd.concat(eq_pool, axis=0),
                          '均衡池v1业绩比较', int_solution=int_solution, laps_days=laps_day)

def return_test(pool_name,int_solution,laps_days=75):

    cp = Customize_pool()
    bmk_df=[]

    pdf = pd.read_excel(r"E:\GitFolder\docs\金工基金池\{}.xlsx".format(pool_name))
    pdf['jjdm'] = ['000000' + str(x) for x in pdf['jjdm']]
    pdf['jjdm'] = pdf['jjdm'].str[-6:]
    pdf['asofdate']=pdf['asofdate'].astype(str)
    for asofdate in pdf['asofdate'].unique().tolist():
        bmk=pd.DataFrame()
        bmk['jjdm'] = util.get_bmk_funds_list(asofdate)
        bmk['asofdate']=asofdate
        bmk_df.append(bmk)

    cp.return_comparision(pd.concat(bmk_df, axis=0), pdf
                          , '均衡池v1业绩比较', int_solution=int_solution,laps_days=laps_days)

def simple_filter(asofdate):

    jjdm_list=util.get_bmk_funds_list_20220824(asofdate,5)

    sql="select jjdm,ryxm,rydm,rzrq,lrrq from st_fund.t_st_gm_jjjl where jjdm in ({}) and ryzw='基金经理' "\
        .format(util.list_sql_condition(jjdm_list))
    manager_info=hbdb.db2df(sql,db='funduser')

    manager_info.loc[manager_info['lrrq']>=int(asofdate),'lrrq']=int(asofdate)

    manager_info['rzrq']=\
        [datetime.datetime.strptime(str(x), '%Y%m%d') for x in manager_info['rzrq']]
    manager_info['lrrq']=\
        [datetime.datetime.strptime(str(x), '%Y%m%d') for x in manager_info['lrrq']]

    manager_info['duration']=manager_info['lrrq']-manager_info['rzrq']
    manager_info['duration']=[x.days for x in manager_info['duration']]
    manager_dm=\
        list(set(manager_info[manager_info['duration']>=360*5]['rydm']).\
            intersection(manager_info[manager_info['lrrq']==manager_info['lrrq'].max()]['rydm']))
    selected_manager=manager_info[(manager_info['rydm'].isin(manager_dm))
                 &(manager_info['duration']>=360*5)]

    sql="select jjdm,zbnp,jzrq from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and zblb='2205' and jzrq='{1}' and zbnp!=99999"\
        .format(util.list_sql_condition(jjdm_list),asofdate)
    jj_rank=hbdb.db2df(sql,db='funduser')
    jj_rank['rank']=jj_rank['zbnp'].rank(ascending=False)/len(jj_rank)
    winning_fund=jj_rank[(jj_rank['jjdm'].isin(selected_manager['jjdm'].tolist()))
                         &(jj_rank['rank']<=1/3)].sort_values('rank')['jjdm'].tolist()[0:30]

    return  winning_fund

if __name__ == '__main__':


    #
    # mr=Mutual_fund_report(start_date='20231031',end_date='20231130',
    #                       asofdate='20230630')
    # mr.generate_weekly_report_draft()

    # get_pool_history_report()

    # #
    # simple_pool=pd.concat(simple_pool,axis=0)
    # simple_pool.to_excel(r"E:\GitFolder\docs\金工基金池\简单池.xlsx",index=False)
    # return_test('纯主题池', int_solution=False)
    # return_test('动量池v有滞后', int_solution=True)
    # return_test('均衡池v1_剔除规模后50%的公司', int_solution=True)
    # asofdate_list=['20151231', '20160331', '20160630', '20160930', '20161230'
    #     , '20170331', '20170630', '20170930', '20171229'
    #     , '20180331', '20180629', '20180930', '20181228'
    #     , '20190331', '20190628', '20190930', '20191231'
    #     , '20200331', '20200630', '20200930', '20201231'
    #     , '20210331', '20210630', '20210930', '20211231','20220331','20220630','20220930','20221230','20230331']

    # asofdate_list=['20171229' , '20180629','20181228',
    #                '20190628',  '20191231'
    #     ,  '20200630', '20201231'
    #     ,  '20210630',  '20211231','20220630']
    asofdate_list=[
        '20181228', '20190331', '20190628', '20190930', '20191231'
            , '20200331', '20200630', '20200930', '20201231', '20210331',
        '20210630', '20210930', '20211231','20220331','20220630','20220930','20221230','20230331']
    # asofdate_list=[
    #     '20220630','20220930','20221230','20230331']
    asofdate_list=['20231231']

    # return_test("新算法测试\\半年度调仓_年度信息比率",int_solution=True)
    # return_test("新算法测试\\半年度调仓_半年度信息比率",int_solution=True)
    # return_test("新算法测试\\半年度调仓_季度信息比率",int_solution=True)
    #return_test("新算法测试\\纯因子模型",int_solution=True,laps_days=0)
    #return_test("新算法测试\\均衡池v1", int_solution=True, laps_days=75)
    back_test(asofdate_list=asofdate_list)
    # ['20181228'
    # , '20190331', '20190628', '20190930', '20191231'
    # , '20200331', '20200630', '20200930', '20201231'
    # , '20210331', '20210630', '20210930', '20211231']

    print('')


    #pool_his_fromexcel2db()
    #pic_pool_history()
    #pool_comparision('202103', '202203')

    #pool_picturing