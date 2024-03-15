# -*- coding: utf-8 -*-

from datetime import datetime, timedelta
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from hbshare.fe.xwq.analysis.orm.fedb import FEDB
from hbshare.fe.xwq.analysis.orm.hbdb import HBDB
import docx
import os
import numpy as np
import pandas as pd
from matplotlib.ticker import FuncFormatter
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import seaborn as sns
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False
sns.set_style('white', {'font.sans-serif': ['simhei', 'Arial']})
line_color_list = ['#F04950', '#6268A2', '#959595', '#333335', '#EE703F', '#7E4A9B', '#8A662C',
                  '#44488E', '#BA67E9', '#3FAEEE']
bar_color_list = ['#C94649', '#EEB2B4', '#E1777A', '#D57C56', '#E39A79', '#DB8A66', '#E5B88C',
                  '#8588B7', '#B4B6D1', '#55598D', '#628497', '#A9C6CB', '#866EA9', '#B79BC7',
                  '#7D7D7E', '#CACACA', '#A7A7A8', '#606063', '#C4C4C4', '#99999B', '#B7B7B7']
area_color_list = ['#D55659', '#E1777A', '#DB8A66', '#E5B88C', '#EEB2B4', '#D57C56', '#E39A79',
                   '#8588B7', '#626697', '#866EA9', '#B79BC7', '#B4B6D1', '#628497', '#A9C6CB',
                   '#7D7D7E', '#A7A7A8', '#99999B', '#B7B7B7', '#CACACA', '#969696', '#C4C4C4']
new_color_list = ['#F04950', '#959595', '#6268A2', '#333335', '#D57C56', '#628497']
industry_theme_dic = {'银行': '大金融', '非银金融': '大金融', '房地产': '大金融',
                      '食品饮料': '消费', '家用电器': '消费', '医药生物': '消费', '社会服务': '消费', '农林牧渔': '消费', '商贸零售': '消费', '美容护理': '消费',
                      '通信': 'TMT', '计算机': 'TMT', '电子': 'TMT', '传媒': 'TMT', '国防军工': 'TMT',
                      '交通运输': '制造', '机械设备': '制造', '汽车': '制造', '纺织服饰': '制造', '轻工制造': '制造', '电力设备': '制造',
                      '钢铁': '周期', '有色金属': '周期', '建筑装饰': '周期', '建筑材料': '周期', '基础化工': '周期', '石油石化': '周期', '煤炭': '周期', '公用事业': '周期', '环保': '周期',
                      '综合': '其他'}


def to_percent(temp, position):
    return '%1.0f'%(temp) + '%'

def to_100percent(temp, position):
    return '%1.0f'%(temp * 100) + '%'

def to_percent_r2(temp, position):
    return '%0.1f'%(temp) + '%'

def to_100percent_r2(temp, position):
    return '%0.1f'%(temp * 100) + '%'

def get_change_dates_loc(change_dates, all_dates):
    change_dates_loc = []
    for change_date in change_dates:
        for idx, date in enumerate(all_dates[:-1]):
            if change_date >= all_dates[idx] and change_date <= all_dates[idx + 1]:
                loc = idx + (datetime.strptime(change_date, '%Y%m%d') - datetime.strptime(all_dates[idx], '%Y%m%d')) / (datetime.strptime(all_dates[idx + 1], '%Y%m%d') - datetime.strptime(all_dates[idx], '%Y%m%d'))
                change_dates_loc.append(loc)
                break
    return change_dates_loc

def get_date(start_date, end_date):
    calendar_df = HBDB().read_cal(start_date, end_date)
    calendar_df = calendar_df.rename(columns={'JYRQ': 'CALENDAR_DATE', 'SFJJ': 'IS_OPEN', 'SFZM': 'IS_WEEK_END', 'SFYM': 'IS_MONTH_END'})
    calendar_df['CALENDAR_DATE'] = calendar_df['CALENDAR_DATE'].astype(str)
    calendar_df = calendar_df.sort_values('CALENDAR_DATE')
    calendar_df['IS_OPEN'] = calendar_df['IS_OPEN'].astype(int).replace({0: 1, 1: 0})
    calendar_df['YEAR_MONTH'] = calendar_df['CALENDAR_DATE'].apply(lambda x: x[:6])
    calendar_df['MONTH'] = calendar_df['CALENDAR_DATE'].apply(lambda x: x[4:6])
    calendar_df['MONTH_DAY'] = calendar_df['CALENDAR_DATE'].apply(lambda x: x[4:])
    calendar_df = calendar_df[(calendar_df['CALENDAR_DATE'] >= start_date) & (calendar_df['CALENDAR_DATE'] <= end_date)]
    trade_df = calendar_df[calendar_df['IS_OPEN'] == 1].rename(columns={'CALENDAR_DATE': 'TRADE_DATE'})
    trade_df = trade_df[(trade_df['TRADE_DATE'] >= start_date) & (trade_df['TRADE_DATE'] <= end_date)]
    report_df = calendar_df.drop_duplicates('YEAR_MONTH', keep='last').rename(columns={'CALENDAR_DATE': 'REPORT_DATE'})
    report_df = report_df[report_df['MONTH_DAY'].isin(['0331', '0630', '0930', '1231'])]
    report_df = report_df[(report_df['REPORT_DATE'] >= start_date) & (report_df['REPORT_DATE'] <= end_date)]
    report_trade_df = calendar_df[calendar_df['IS_OPEN'] == 1].rename(columns={'CALENDAR_DATE': 'TRADE_DATE'})
    report_trade_df = report_trade_df.sort_values('TRADE_DATE').drop_duplicates('YEAR_MONTH', keep='last')
    report_trade_df = report_trade_df[report_trade_df['MONTH'].isin(['03', '06', '09', '12'])]
    report_trade_df = report_trade_df[(report_trade_df['TRADE_DATE'] >= start_date) & (report_trade_df['TRADE_DATE'] <= end_date)]
    calendar_trade_df = calendar_df[['CALENDAR_DATE']].merge(trade_df[['TRADE_DATE']], left_on=['CALENDAR_DATE'], right_on=['TRADE_DATE'], how='left')
    calendar_trade_df['TRADE_DATE'] = calendar_trade_df['TRADE_DATE'].fillna(method='ffill')
    calendar_trade_df = calendar_trade_df[(calendar_trade_df['TRADE_DATE'] >= start_date) & (calendar_trade_df['TRADE_DATE'] <= end_date)]
    return calendar_df, report_df, trade_df, report_trade_df, calendar_trade_df


class QuantitativeReport:
    def __init__(self, file_path, manager_code, fund_code=None):
        self.manager_code = manager_code
        self.fund_code = fund_code
        self.manager_fund = HBDB().read_fund_given_manager_code(self.manager_code)
        self.manager_fund = self.manager_fund.rename(columns={'jjdm': 'FUND_CODE', 'rydm': 'MANAGER_CODE', 'ryxm': 'MANAGER_NAME', 'ryzw': 'POSITION', 'ryzt': 'STATUS', 'bdrq': 'CHANGE_DATE', 'rzrq': 'BEGIN_DATE', 'lrrq': 'END_DATE', 'lryy': 'REASON'})
        self.manager_fund['FUND_CODE'] = self.manager_fund['FUND_CODE'].astype(str)
        self.manager_fund['BEGIN_DATE'] = self.manager_fund['BEGIN_DATE'].astype(int).astype(str)
        self.manager_fund['END_DATE'] = self.manager_fund['END_DATE'].astype(int).astype(str)
        manager_fund_scale_list = []
        for code in self.manager_fund['FUND_CODE'].unique().tolist():
            manager_fund_scale_code = HBDB().read_fund_scale_given_code(code)
            manager_fund_scale_code = manager_fund_scale_code if len(manager_fund_scale_code) > 0 else pd.DataFrame(columns=['zcjz', 'bblb', 'jsrq', 'jjdm', 'ggrq'])
            manager_fund_scale_code = manager_fund_scale_code[manager_fund_scale_code['bblb'] == 13]
            manager_fund_scale_code = manager_fund_scale_code.sort_values(['jjdm', 'jsrq', 'ggrq']).drop_duplicates(['jjdm', 'jsrq'], keep='last')
            manager_fund_scale_list.append(manager_fund_scale_code)
        self.manager_fund_scale = pd.concat(manager_fund_scale_list)
        self.manager_fund_scale = self.manager_fund_scale.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zcjz': 'AUM'})
        self.manager_fund_scale['REPORT_DATE'] = self.manager_fund_scale['REPORT_DATE'].astype(str)
        manager_fund_position_list = []
        for code in self.manager_fund['FUND_CODE'].unique().tolist():
            manager_fund_position_code = HBDB().read_fund_position_given_code(code)
            manager_fund_position_list.append(manager_fund_position_code)
        self.manager_fund_position = pd.concat(manager_fund_position_list)
        self.manager_fund_position = self.manager_fund_position.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'jjzzc': 'TOTAL', 'gptzsz': 'STOCK', 'zqzsz': 'BOND', 'jjtzszhj': 'FUND', 'hbzjsz': 'CASH'})
        self.manager_fund_position['REPORT_DATE'] = self.manager_fund_position['REPORT_DATE'].astype(str)
        self.manager_fund_position['STOCK_RATIO'] = self.manager_fund_position['STOCK'].fillna(0.0) / self.manager_fund_position['TOTAL'] * 100.0
        self.manager_fund_position['BOND_RATIO'] = self.manager_fund_position['BOND'].fillna(0.0) / self.manager_fund_position['TOTAL'] * 100.0
        self.manager_fund_position['FUND_RATIO'] = self.manager_fund_position['FUND'].fillna(0.0) / self.manager_fund_position['TOTAL'] * 100.0
        self.manager_fund_position['CASH_RATIO'] = self.manager_fund_position['CASH'].fillna(0.0) / self.manager_fund_position['TOTAL'] * 100.0
        self.manager_fund_position['OTHER_RATIO'] = 100.0 - self.manager_fund_position[['STOCK_RATIO', 'BOND_RATIO', 'FUND_RATIO', 'CASH_RATIO']].sum(axis=1)
        self.manager_name = self.manager_fund['MANAGER_NAME'].values[0]
        self.file_path = file_path + self.manager_name + '/'
        if not os.path.exists(self.file_path):
            os.makedirs(self.file_path)
        self.fund = HBDB().read_fund_info()
        self.fund = self.fund.rename(columns={'jjdm': 'FUND_CODE', 'jjmc': 'FUND_FULL_NAME', 'jjjc': 'FUND_SHORT_NAME', 'clrq': 'BEGIN_DATE', 'zzrq': 'END_DATE', 'jjfl': 'FUND_TYPE_1ST', 'ejfl': 'FUND_TYPE_2ND', 'kffb': 'OPEN_CLOSE'})
        self.fund['FUND_CODE'] = self.fund['FUND_CODE'].astype(str)
        self.fund['FUND_TYPE_1ST'] = self.fund['FUND_TYPE_1ST'].replace({'f': 'FOF型', '1': '股票型', '2': '债券型', '3': '混合型', '4': '另类投资型', '7': '货币型', '9': 'QDII型'})
        self.fund['FUND_TYPE_2ND'] = self.fund['FUND_TYPE_2ND'].replace(
            {'f3': '债券型FOF', 'f4': '货币型FOF', 'f1': '股票型FOF', 'f2': '混合型FOF', 'f5': '另类投资FOF',
             '13': '普通股票型', '14': '股票型', '15': '增强指数型', '16': '被动指数型',
             '21': '被动指数型债券', '22': '短期纯债型', '23': '混合债券型一级', '24': '混合债券型二级', '25': '增强指数型债券', '26': '债券型', '27': '中长期纯债型', '28': '可转换债券型',
             '34': '平衡混合型', '35': '灵活配置型', '36': '混合型', '37': '偏股混合型', '38': '偏债混合型',
             '41': '股票多空', '42': '商品型', '43': 'REITs',
             '91': '股票型QDII', '93': '混合型QDII', '94': '债券型QDII', '95': '另类型QDII'})
        self.report_dates = []
        for year in range(1990, datetime.today().year + 1):
            self.report_dates += [str(year) + '0331', str(year) + '0630', str(year) + '0930', str(year) + '1231']
        self.report_dates = sorted(self.report_dates)
        self.representative_fund = self.get_representative_fund()
        self.fund_code_list = self.representative_fund['FUND_CODE'].unique().tolist()
        self.change_dates = self.representative_fund['END_DATE'].unique().tolist()
        if len(self.change_dates) == 1:
            self.change_dates = []
        else:
             self.change_dates = [(datetime.strptime(date, '%Y%m%d') + timedelta(45)).strftime('%Y%m%d') for date in self.change_dates][:-1]
        self.calendar_df, self.report_df, self.trade_df, self.report_trade_df, self.calendar_trade_df = get_date('19000101', datetime.today().strftime('%Y%m%d'))

    def get_representative_fund(self):
        manager_fund = self.manager_fund.copy(deep=True)
        manager_fund['END_DATE'] = manager_fund['END_DATE'].replace('29991231', datetime.today().date().strftime('%Y%m%d'))
        # manager_fund = manager_fund[manager_fund['POSITION'] == '基金经理']
        manager_fund = manager_fund[manager_fund['FUND_CODE'].isin(self.fund[self.fund['FUND_TYPE_2ND'].isin(['普通股票型', '灵活配置型', '偏股混合型'])]['FUND_CODE'].unique().tolist())]
        if self.fund_code is not None:
            manager_fund = manager_fund[manager_fund['FUND_CODE'] == self.fund_code][['FUND_CODE', 'BEGIN_DATE', 'END_DATE']]
            if len(manager_fund) > 1:
                representative_fund = pd.DataFrame(index=range(1), columns=['FUND_CODE', 'BEGIN_DATE', 'END_DATE'])
                representative_fund.loc[0, 'FUND_CODE'] = manager_fund['FUND_CODE'].iloc[0]
                representative_fund.loc[0, 'BEGIN_DATE'] = manager_fund['BEGIN_DATE'].iloc[0]
                representative_fund.loc[0, 'END_DATE'] = manager_fund['END_DATE'].iloc[-1]
            else:
                representative_fund = manager_fund.copy(deep=True)
        else:
            report_dates = [date for date in self.report_dates if date > min(manager_fund['BEGIN_DATE']) and date < max(manager_fund['END_DATE'])]
            report_dates = [min(manager_fund['BEGIN_DATE'])] + report_dates + [max(manager_fund['END_DATE'])]
            manager_fund_scale = self.manager_fund_scale.copy(deep=True)
            manager_fund_scale = manager_fund_scale.pivot(index='REPORT_DATE', columns='FUND_CODE', values='AUM')
            manager_fund_scale = manager_fund_scale.reindex(report_dates).sort_index().fillna(method='ffill')
            representative_fund_list = []
            for i in range(len(report_dates) - 1):
                last_date = report_dates[i]
                date = report_dates[i + 1]
                manager_fund_scale_date = manager_fund_scale[manager_fund_scale.index == date]
                manager_fund_scale_date = manager_fund_scale_date.unstack().reset_index()
                manager_fund_scale_date.columns = ['FUND_CODE', 'REPORT_DATE', 'AUM']
                manager_fund_date = manager_fund[(manager_fund['BEGIN_DATE'] < date) & (manager_fund['END_DATE'] >= date)]
                manager_fund_date = manager_fund_date.merge(manager_fund_scale_date[['FUND_CODE', 'AUM']], on=['FUND_CODE'], how='left')
                manager_fund_date['PERIOD'] = manager_fund_date['BEGIN_DATE'].apply(lambda x: (datetime.strptime(date, '%Y%m%d') - datetime.strptime(x, '%Y%m%d')).days)
                manager_fund_date['SCORE'] = manager_fund_date['PERIOD'] / max(manager_fund_date['PERIOD']) + manager_fund_date['AUM'] / max(manager_fund_date['AUM'])
                manager_fund_position_date = self.manager_fund_position[self.manager_fund_position['REPORT_DATE'].isin([last_date, date])]
                manager_fund_position_date = manager_fund_position_date[['FUND_CODE', 'STOCK_RATIO']].groupby('FUND_CODE').mean().reset_index()
                manager_fund_date = manager_fund_date.merge(manager_fund_position_date, on=['FUND_CODE'], how='left')
                manager_fund_date = manager_fund_date[manager_fund_date['STOCK_RATIO'] >= 60]
                manager_fund_date = manager_fund_date.sort_values(['SCORE', 'FUND_CODE'], ascending=[False, True]).head(1)
                representative_fund_date = manager_fund_date[['FUND_CODE']]
                representative_fund_date['BEGIN_DATE'] = last_date
                representative_fund_date['END_DATE'] = date
                representative_fund_list.append(representative_fund_date)
            representative_fund = pd.concat(representative_fund_list)
            representative_fund = representative_fund.reset_index().drop('index', axis=1)
            i = 0
            while not i == len(representative_fund):
                fund_code = representative_fund['FUND_CODE'].iloc[i]
                fund_idx = representative_fund[representative_fund['FUND_CODE'] == fund_code].index[-1]
                representative_fund.loc[i: fund_idx, 'FUND_CODE'] = fund_code
                i = fund_idx + 1
            representative_fund_cp = representative_fund.copy(deep=True)
            representative_fund = pd.DataFrame(index=range(len(representative_fund_cp['FUND_CODE'].unique().tolist())), columns=['FUND_CODE', 'BEGIN_DATE', 'END_DATE'])
            for i, fund_code in enumerate(representative_fund_cp['FUND_CODE'].unique().tolist()):
                representative_fund.loc[i, 'FUND_CODE'] = fund_code
                representative_fund.loc[i, 'BEGIN_DATE'] = representative_fund_cp[representative_fund_cp['FUND_CODE'] == fund_code]['BEGIN_DATE'].min()
                representative_fund.loc[i, 'END_DATE'] = representative_fund_cp[representative_fund_cp['FUND_CODE'] == fund_code]['END_DATE'].max()
            representative_fund = representative_fund.sort_values('BEGIN_DATE')
        return representative_fund

    def report_info(self):
        """
        风险提示及重要声明
        """
        risk = '投资有风险。基金的过往业绩并不预示其未来表现。基金管理人管理的其他基金的业绩并不构成基金业绩表现的保证。相关数据仅供参考，不构成投资建议。投资人请详阅基金合同等法律文件，了解产品风险收益特征，根据自身资产状况、风险承受能力审慎决策，独立承担投资风险。本资料仅为宣传用品，本机构及工作人员不存在直接或间接主动推介相关产品的行为。'
        statement1 = '本文件中的信息基于已公开的信息、数据及尽调访谈等，好买基金或好买基金研究中心（以下简称“本公司”）对这些信息的及时性、准确性及完整性不做任何保证，也不保证所包含的信息不会发生变更。文件中的内容仅供参考，不代表任何确定性的判断。本文件及其内容均不构成投资建议，也没有考虑个别客户特殊的投资目标、财务状况或需要。获得本文件的机构或个人据此做出投资决策，应自行承担投资风险。'
        statement2 = '本文件版权为本公司所有。未经本公司书面许可，任何机构或个人不得以翻版、复制、 发表、引用或再次分发他人等任何形式侵犯本公司版权。本文件中的信息均为保密信息，未经本公司事先同意，不得以任何目的，复制或传播本文本中所含信息，亦不可向任何第三方披露。'
        return risk, statement1, statement2

    def get_manager_info(self):
        """
        基金经理信息
        """
        manager_info = HBDB().read_manager_info_given_manager_code(self.manager_code)
        manager_info = manager_info.rename(columns={'rydm': 'MANAGER_CODE', 'ryxm': 'MANAGER_NAME', 'ryxb': 'SEX', 'ryxl': 'EDU', 'ryjj': 'INTRODUCTION'})
        manager_info = manager_info['INTRODUCTION'].values[0].replace('先生：', '先生，').replace('女士：', '女士，').replace('/r', '').replace('/n', '')
        manager_fund = self.manager_fund.merge(self.fund[['FUND_CODE', 'FUND_FULL_NAME', 'FUND_SHORT_NAME', 'FUND_TYPE_2ND']], on=['FUND_CODE'], how='left')
        latest_scale = self.manager_fund_scale[self.manager_fund_scale['REPORT_DATE'] == self.manager_fund_scale['REPORT_DATE'].max()]
        manager_fund = manager_fund.merge(latest_scale[['FUND_CODE', 'AUM']], on=['FUND_CODE'], how='left')
        manager_fund = manager_fund.sort_values(['BEGIN_DATE', 'END_DATE', 'FUND_SHORT_NAME'])
        manager_fund['FUND_TYPE_2ND'] = manager_fund['FUND_TYPE_2ND'].replace(np.nan, '-')
        manager_fund['END_DATE'] = manager_fund['END_DATE'].replace('29991231', '-')
        manager_fund['REASON'] = manager_fund['REASON'].replace(np.nan, '-')
        manager_fund['AUM'] = manager_fund['AUM'].apply(lambda x: round(x / 100000000.0, 2))
        manager_fund['AUM'] = manager_fund['AUM'].replace(0.0, '-')
        manager_fund = manager_fund[['FUND_SHORT_NAME', 'FUND_TYPE_2ND', 'POSITION', 'BEGIN_DATE', 'END_DATE', 'REASON', 'AUM']]
        manager_fund.columns = ['基金名称', '基金类型', '职位', '任职日期', '离任日期', '离任原因', '最新规模（A、C份额合计，亿元）']
        manager_fund = manager_fund.T.reset_index().T
        manager_info_text = '{0}具体管理产品情况如下：'.format(manager_info)
        return manager_info_text, manager_fund

    def get_manager_scale(self):
        """
        基金经理管理规模
        """
        manager_fund_scale = HBDB().read_fund_manager_product_given_code(self.manager_code)
        manager_fund_scale = manager_fund_scale.rename(columns={'rydm': 'MANAGER_CODE', 'jsrq': 'REPORT_DATE', 'zgcpsl': 'FUND_NUM', 'zgcpgm': 'FUND_SCALE', 'dbjj': 'FUND_REPRESENT'})
        manager_fund_scale['REPORT_DATE'] = manager_fund_scale['REPORT_DATE'].astype(str)
        manager_fund_scale = manager_fund_scale.sort_values('REPORT_DATE')
        manager_fund_scale['FUND_SCALE'] = round(manager_fund_scale['FUND_SCALE'] / 100000000.0, 2)
        manager_fund_scale = manager_fund_scale[['REPORT_DATE', 'FUND_SCALE', 'FUND_NUM']]
        manager_fund_scale = manager_fund_scale.sort_values('REPORT_DATE')

        fig, ax1 = plt.subplots(figsize=(12, 6))
        ax2 = ax1.twinx()
        bar_width = 0.3
        ax1.bar(np.arange(len(manager_fund_scale)) - 0.5 * bar_width, manager_fund_scale['FUND_SCALE'].values, bar_width, label='在管产品规模（亿元）', color=bar_color_list[0])
        ax2.bar(np.arange(len(manager_fund_scale)) + 0.5 * bar_width, manager_fund_scale['FUND_NUM'].values, bar_width, label='在管产品数量（只，右轴）', color=bar_color_list[7])
        h1, l1 = ax1.get_legend_handles_labels()
        h2, l2 = ax2.get_legend_handles_labels()
        plt.legend(handles=h1 + h2, labels=l1 + l2, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax1.set_xticks(np.arange(len(manager_fund_scale)))
        ax1.set_xticklabels(labels=manager_fund_scale['REPORT_DATE'], rotation=45)
        ax1.set_xlabel('')
        ax2.set_xlabel('')
        ax1.set_ylabel('')
        ax2.set_ylabel('')
        plt.title('在管产品情况', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=False)
        plt.savefig('{0}{1}_scale.png'.format(self.file_path, self.manager_name))

        latest_date = manager_fund_scale['REPORT_DATE'].iloc[-1]
        latest_scale = manager_fund_scale['FUND_SCALE'].iloc[-1]
        fund_date = self.manager_fund[(self.manager_fund['BEGIN_DATE'] < latest_date) & (self.manager_fund['END_DATE'] >= latest_date)]
        scale_date = self.manager_fund_scale[self.manager_fund_scale['REPORT_DATE'] == latest_date]
        scale_date = scale_date[scale_date['FUND_CODE'].isin(fund_date['FUND_CODE'].unique().tolist())]
        scale_date = scale_date.merge(self.fund[['FUND_CODE', 'FUND_TYPE_2ND']], on=['FUND_CODE'], how='left')
        scale_date = scale_date[['FUND_TYPE_2ND', 'AUM']].groupby('FUND_TYPE_2ND').sum().reset_index()
        scale_date['AUM'] = round(scale_date['AUM'] / 100000000.0, 2)
        scale_date['AUM'] = scale_date['AUM'].astype(str)
        scale_date_text = '{0}基金管理规模分别为{1}亿元'.format('、'.join(scale_date['FUND_TYPE_2ND'].unique().tolist()), '、'.join(scale_date['AUM'].unique().tolist())) if len(scale_date) > 1 else '{0}基金管理规模为{1}亿元'.format(scale_date['FUND_TYPE_2ND'].values[0], scale_date['AUM'].values[0])

        manager_scale_text = '{0}作为基金经理，其在管产品规模和在管产品数量（考虑A、C份额）及变化情况如下，截至最新报告期，{1}管理的产品规模达到了{2}亿元，其中{3}。'.format(self.manager_name, self.manager_name, latest_scale, scale_date_text)
        return manager_scale_text

    def get_repr_fund(self):
        repr_fund = self.representative_fund.copy(deep=True)
        repr_fund = repr_fund.replace(datetime.today().date().strftime('%Y%m%d'), '-')
        repr_fund = repr_fund.merge(self.fund[['FUND_CODE', 'FUND_SHORT_NAME']], on=['FUND_CODE'], how='left')
        repr_fund = repr_fund[['FUND_SHORT_NAME', 'FUND_CODE', 'BEGIN_DATE', 'END_DATE']]
        repr_fund = repr_fund.sort_values('BEGIN_DATE')
        repr_fund.columns = ['代表基金名称', '代表基金代码', '起始时间', '结束时间']
        repr_fund = repr_fund.T.reset_index().T

        repr_fund_text = '我们选取{0}作为基金经理管理期间的代表产品特征来刻画其投资行为，根据任职时长和基金净资产，其代表产品如下：'.format(self.manager_name)
        return repr_fund_text, repr_fund

    def get_manager_position(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_position_history_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_position_history = HBDB().read_fund_position_given_code(fund_code)
            hbs_position_history = hbs_position_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'jjzzc': 'TOTAL', 'gptzsz': 'STOCK', 'zqzsz': 'BOND', 'jjtzszhj': 'FUND', 'hbzjsz': 'CASH'})
            hbs_position_history['REPORT_DATE'] = hbs_position_history['REPORT_DATE'].astype(str)
            hbs_position_history['REPORT_DATE'] = hbs_position_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_position_history = hbs_position_history[(hbs_position_history['REPORT_DATE'] > begin_date) & (hbs_position_history['REPORT_DATE'] <= end_date)]
            hbs_position_history['STOCK_RATIO'] = hbs_position_history['STOCK'].fillna(0.0) / hbs_position_history['TOTAL'] * 100.0
            hbs_position_history['BOND_RATIO'] = hbs_position_history['BOND'].fillna(0.0) / hbs_position_history['TOTAL'] * 100.0
            hbs_position_history['FUND_RATIO'] = hbs_position_history['FUND'].fillna(0.0) / hbs_position_history['TOTAL'] * 100.0
            hbs_position_history['CASH_RATIO'] = hbs_position_history['CASH'].fillna(0.0) / hbs_position_history['TOTAL'] * 100.0
            hbs_position_history['OTHER_RATIO'] = 100.0 - hbs_position_history[['STOCK_RATIO', 'BOND_RATIO', 'FUND_RATIO', 'CASH_RATIO']].sum(axis=1)
            hbs_position_history['MANAGER_NAME'] = self.manager_name
            hbs_position_history = hbs_position_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'STOCK_RATIO', 'BOND_RATIO', 'FUND_RATIO', 'CASH_RATIO', 'OTHER_RATIO']]
            hbs_position_history_list.append(hbs_position_history)

        hbs_position_history = pd.concat(hbs_position_history_list)
        change_dates = [date for date in self.change_dates if date >= min(hbs_position_history['REPORT_DATE']) and date <= max(hbs_position_history['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_position_history['REPORT_DATE'].unique().tolist()))
        hbs_position_history_disp = hbs_position_history[['REPORT_DATE', 'STOCK_RATIO', 'BOND_RATIO', 'FUND_RATIO', 'CASH_RATIO', 'OTHER_RATIO']].set_index('REPORT_DATE')
        hbs_position_history_disp = hbs_position_history_disp.sort_index()
        hbs_position_history_disp.columns = ['股票', '债券', '基金', '现金', '其他']
        fig, ax = plt.subplots(figsize=(12, 6))
        hbs_position_history_disp.plot.area(ax=ax, stacked=True, color=[area_color_list[0], area_color_list[1], area_color_list[7], area_color_list[8], area_color_list[15]])
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=5)
        plt.xticks(rotation=45)
        ax.set_xlabel('')
        ax.set_ylabel('')
        ax.set_xticks(np.arange(len(hbs_position_history)))
        ax.set_xticklabels(labels=hbs_position_history['REPORT_DATE'], rotation=45)
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=100.0, linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.title('仓位占比', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_position.png'.format(self.file_path, self.manager_name))

        stock_position_mean = round(hbs_position_history['STOCK_RATIO'].mean(), 2)
        stock_position_std = round((hbs_position_history['STOCK_RATIO'] / 100.0).std() * 100.0, 2)
        std_mark = '稳定' if stock_position_std < 5 else '灵活'
        position_text = '在仓位配比上，{0}对各类资产的配置情况如下，其中对股票资产的配置比例均值为{1}%，标准差为{2}%，仓位配置较为{3}。'.format(self.manager_name, stock_position_mean, stock_position_std, std_mark)
        return position_text

    def get_manager_size(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_size_sc_history_list, hbs_size_sc_rank_history_list, hbs_size_history_list = [], [], []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_size_sc_history = FEDB().read_hbs_size_style_property_history_given_code(fund_code)
            hbs_size_sc_history = hbs_size_sc_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_size': 'SIZE_SHIFT', 'c_level_size': 'SIZE_CEN'})
            hbs_size_sc_history['REPORT_DATE'] = hbs_size_sc_history['REPORT_DATE'].astype(str)
            hbs_size_sc_history['REPORT_DATE'] = hbs_size_sc_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_size_sc_history = hbs_size_sc_history[(hbs_size_sc_history['REPORT_DATE'] > begin_date) & (hbs_size_sc_history['REPORT_DATE'] <= end_date)]
            hbs_size_sc_history['MANAGER_NAME'] = self.manager_name
            hbs_size_sc_history = hbs_size_sc_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'SIZE_SHIFT', 'SIZE_CEN']]
            hbs_size_sc_history_list.append(hbs_size_sc_history)

            hbs_size_sc_rank_history = FEDB().read_hbs_rank_history_given_code(fund_code)
            hbs_size_sc_rank_history = hbs_size_sc_rank_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_size': 'SIZE_SHIFT_RANK', 'c_level_size': 'SIZE_CEN_RANK'})
            hbs_size_sc_rank_history['REPORT_DATE'] = hbs_size_sc_rank_history['REPORT_DATE'].astype(str)
            hbs_size_sc_rank_history['REPORT_DATE'] = hbs_size_sc_rank_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_size_sc_rank_history = hbs_size_sc_rank_history[(hbs_size_sc_rank_history['REPORT_DATE'] > begin_date) & (hbs_size_sc_rank_history['REPORT_DATE'] <= end_date)]
            hbs_size_sc_rank_history['MANAGER_NAME'] = self.manager_name
            hbs_size_sc_rank_history = hbs_size_sc_rank_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'SIZE_SHIFT_RANK', 'SIZE_CEN_RANK']]
            hbs_size_sc_rank_history_list.append(hbs_size_sc_rank_history)

            hbs_size_history = FEDB().read_size_exposure_given_code(fund_code)
            hbs_size_history = hbs_size_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'size_type': 'SIZE', 'zjbl': 'SIZE_EXPOSURE'})
            hbs_size_history['REPORT_DATE'] = hbs_size_history['REPORT_DATE'].astype(str)
            hbs_size_history['REPORT_DATE'] = hbs_size_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_size_history = hbs_size_history[(hbs_size_history['REPORT_DATE'] > begin_date) & (hbs_size_history['REPORT_DATE'] <= end_date)]
            hbs_size_history['MANAGER_NAME'] = self.manager_name
            hbs_size_history = hbs_size_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'SIZE', 'SIZE_EXPOSURE']]
            hbs_size_history_list.append(hbs_size_history)

        hbs_size_sc_history = pd.concat(hbs_size_sc_history_list)
        hbs_size_sc_rank_history = pd.concat(hbs_size_sc_rank_history_list)
        hbs_size_history = pd.concat(hbs_size_history_list)
        hbs_size_history = hbs_size_history.pivot(index=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], columns='SIZE', values='SIZE_EXPOSURE').reset_index()
        hbs_size_history = hbs_size_sc_history.merge(hbs_size_sc_rank_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left').merge(hbs_size_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left')
        hbs_size_history = hbs_size_history.sort_values('REPORT_DATE')
        for col in list(hbs_size_history.columns[3:7]):
            hbs_size_history[col] = hbs_size_history[col].fillna(hbs_size_history[col].interpolate())

        change_dates = [date for date in self.change_dates if date >= min(hbs_size_history['REPORT_DATE']) and date <= max(hbs_size_history['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_size_history['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(1, 2, figsize=(12, 6))
        ax1, ax3 = ax[0], ax[1]
        ax2, ax4 = ax1.twinx(), ax3.twinx()
        ax1.plot(hbs_size_history['REPORT_DATE'].values, hbs_size_history['SIZE_SHIFT'].values, label='规模换手率', color=line_color_list[0])
        ax2.plot(hbs_size_history['REPORT_DATE'].values, hbs_size_history['SIZE_SHIFT_RANK'].values, label='规模换手率排名（右轴）', color=line_color_list[1])
        ax3.plot(hbs_size_history['REPORT_DATE'].values, hbs_size_history['SIZE_CEN'].values, label='规模集中度', color=line_color_list[0])
        ax4.plot(hbs_size_history['REPORT_DATE'].values, hbs_size_history['SIZE_CEN_RANK'].values, label='规模集中度排名（右轴）', color=line_color_list[1])
        h1, l1 = ax1.get_legend_handles_labels()
        h2, l2 = ax2.get_legend_handles_labels()
        h3, l3 = ax3.get_legend_handles_labels()
        h4, l4 = ax4.get_legend_handles_labels()
        ax1.legend(handles=h1 + h2, labels=l1 + l2, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax3.legend(handles=h3 + h4, labels=l3 + l4, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax1.set_xticks(np.arange(len(hbs_size_history)))
        ax1.set_xticklabels(labels=hbs_size_history['REPORT_DATE'], rotation=45)
        ax3.set_xticks(np.arange(len(hbs_size_history)))
        ax3.set_xticklabels(labels=hbs_size_history['REPORT_DATE'], rotation=45)
        ax1.set_ylim([0.0, 1.0])
        ax2.set_ylim([0.0, 1.0])
        ax3.set_ylim([0.0, 1.0])
        ax4.set_ylim([0.0, 1.0])
        ax1.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax3.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax1.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax2.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax3.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax4.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax1.set_title('规模换手率', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        ax3.set_title('规模集中度', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=False)
        plt.savefig('{0}{1}_size_sc.png'.format(self.file_path, self.manager_name))

        hbs_size_history_disp = pd.concat(hbs_size_history_list)
        hbs_size_history_disp = hbs_size_history_disp.sort_values('REPORT_DATE')
        change_dates = [date for date in self.change_dates if date >= min(hbs_size_history_disp['REPORT_DATE']) and date <= max(hbs_size_history_disp['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_size_history_disp['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.barplot(x='REPORT_DATE', y='SIZE_EXPOSURE', data=hbs_size_history_disp, hue='SIZE', hue_order=['大盘', '中盘', '小盘'], palette=[bar_color_list[0], bar_color_list[7], bar_color_list[14]])
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=3)
        plt.xticks(rotation=45)
        ax.set_xlabel('')
        ax.set_ylabel('')
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(hbs_size_history_disp['SIZE_EXPOSURE']), linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.title('规模暴露', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_size.png'.format(self.file_path, self.manager_name))

        size_shift_value = hbs_size_history['SIZE_SHIFT_RANK'].mean()
        size_cen_value = hbs_size_history['SIZE_CEN_RANK'].mean()
        size_shift_mark = '偏下' if size_shift_value <= 0.3 else '中等偏下' if size_shift_value > 0.3 and size_shift_value <= 0.5 else '中等偏上' if size_shift_value > 0.5 and size_shift_value <= 0.7 else '偏上'
        size_cen_mark = '偏下' if size_cen_value <= 0.3 else '中等偏下' if size_cen_value > 0.3 and size_cen_value <= 0.5 else '中等偏上' if size_cen_value > 0.5 and size_cen_value <= 0.7 else '偏上'
        size_mark = '配置型' if size_shift_value <= 0.5 and size_cen_value <= 0.5 else '专注型' if size_shift_value <= 0.5 and size_cen_value > 0.5 else '轮动型' if size_shift_value > 0.5 and size_cen_value <= 0.5 else '博弈型'
        size_type_index = hbs_size_history[['大盘', '中盘', '小盘']].mean()[hbs_size_history[['大盘', '中盘', '小盘']].mean() == hbs_size_history[['大盘', '中盘', '小盘']].mean().max()].index.tolist()
        size_type_mark = size_type_index[0] if len(size_type_index) == 1 else '、'.join(size_type_index)

        size_text = '在规模配置上，{0}的换手率在偏股型基金整体中处于{1}水平，集中度在偏股型基金整体中处于{2}水平，属于规模{3}基金经理，且从大中小盘平均绝对暴露看更偏{4}。'.format(self.manager_name, size_shift_mark, size_cen_mark, size_mark, size_type_mark)
        return size_text

    def get_manager_style(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_style_sc_history_list, hbs_style_sc_rank_history_list, hbs_style_history_list = [], [], []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_style_sc_history = FEDB().read_hbs_size_style_property_history_given_code(fund_code)
            hbs_style_sc_history = hbs_style_sc_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_style': 'STYLE_SHIFT', 'c_level_style': 'STYLE_CEN'})
            hbs_style_sc_history['REPORT_DATE'] = hbs_style_sc_history['REPORT_DATE'].astype(str)
            hbs_style_sc_history['REPORT_DATE'] = hbs_style_sc_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_style_sc_history = hbs_style_sc_history[(hbs_style_sc_history['REPORT_DATE'] > begin_date) & (hbs_style_sc_history['REPORT_DATE'] <= end_date)]
            hbs_style_sc_history['MANAGER_NAME'] = self.manager_name
            hbs_style_sc_history = hbs_style_sc_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'STYLE_SHIFT', 'STYLE_CEN']]
            hbs_style_sc_history_list.append(hbs_style_sc_history)

            hbs_style_sc_rank_history = FEDB().read_hbs_rank_history_given_code(fund_code)
            hbs_style_sc_rank_history = hbs_style_sc_rank_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_style': 'STYLE_SHIFT_RANK', 'c_level_style': 'STYLE_CEN_RANK'})
            hbs_style_sc_rank_history['REPORT_DATE'] = hbs_style_sc_rank_history['REPORT_DATE'].astype(str)
            hbs_style_sc_rank_history['REPORT_DATE'] = hbs_style_sc_rank_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_style_sc_rank_history = hbs_style_sc_rank_history[(hbs_style_sc_rank_history['REPORT_DATE'] > begin_date) & (hbs_style_sc_rank_history['REPORT_DATE'] <= end_date)]
            hbs_style_sc_rank_history['MANAGER_NAME'] = self.manager_name
            hbs_style_sc_rank_history = hbs_style_sc_rank_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'STYLE_SHIFT_RANK', 'STYLE_CEN_RANK']]
            hbs_style_sc_rank_history_list.append(hbs_style_sc_rank_history)

            hbs_style_history = FEDB().read_style_exposure_given_code(fund_code)
            hbs_style_history = hbs_style_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'style_type': 'STYLE', 'zjbl': 'STYLE_EXPOSURE'})
            hbs_style_history['REPORT_DATE'] = hbs_style_history['REPORT_DATE'].astype(str)
            hbs_style_history['REPORT_DATE'] = hbs_style_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_style_history = hbs_style_history[(hbs_style_history['REPORT_DATE'] > begin_date) & (hbs_style_history['REPORT_DATE'] <= end_date)]
            hbs_style_history['MANAGER_NAME'] = self.manager_name
            hbs_style_history = hbs_style_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'STYLE', 'STYLE_EXPOSURE']]
            hbs_style_history_list.append(hbs_style_history)

        hbs_style_sc_history = pd.concat(hbs_style_sc_history_list)
        hbs_style_sc_rank_history = pd.concat(hbs_style_sc_rank_history_list)
        hbs_style_history = pd.concat(hbs_style_history_list)
        hbs_style_history = hbs_style_history.pivot(index=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], columns='STYLE', values='STYLE_EXPOSURE').reset_index()
        hbs_style_history = hbs_style_sc_history.merge(hbs_style_sc_rank_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left').merge(hbs_style_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left')
        hbs_style_history = hbs_style_history.sort_values('REPORT_DATE')
        for col in list(hbs_style_history.columns[3:7]):
            hbs_style_history[col] = hbs_style_history[col].fillna(hbs_style_history[col].interpolate())

        change_dates = [date for date in self.change_dates if date >= min(hbs_style_history['REPORT_DATE']) and date <= max(hbs_style_history['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_style_history['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(1, 2, figsize=(12, 6))
        ax1, ax3 = ax[0], ax[1]
        ax2, ax4 = ax1.twinx(), ax3.twinx()
        ax1.plot(hbs_style_history['REPORT_DATE'].values, hbs_style_history['STYLE_SHIFT'].values, label='风格换手率', color=line_color_list[0])
        ax2.plot(hbs_style_history['REPORT_DATE'].values, hbs_style_history['STYLE_SHIFT_RANK'].values, label='风格换手率排名（右轴）', color=line_color_list[1])
        ax3.plot(hbs_style_history['REPORT_DATE'].values, hbs_style_history['STYLE_CEN'].values, label='风格集中度', color=line_color_list[0])
        ax4.plot(hbs_style_history['REPORT_DATE'].values, hbs_style_history['STYLE_CEN_RANK'].values, label='风格集中度排名（右轴）', color=line_color_list[1])
        h1, l1 = ax1.get_legend_handles_labels()
        h2, l2 = ax2.get_legend_handles_labels()
        h3, l3 = ax3.get_legend_handles_labels()
        h4, l4 = ax4.get_legend_handles_labels()
        ax1.legend(handles=h1 + h2, labels=l1 + l2, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax3.legend(handles=h3 + h4, labels=l3 + l4, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax1.set_xticks(np.arange(len(hbs_style_history)))
        ax1.set_xticklabels(labels=hbs_style_history['REPORT_DATE'], rotation=45)
        ax3.set_xticks(np.arange(len(hbs_style_history)))
        ax3.set_xticklabels(labels=hbs_style_history['REPORT_DATE'], rotation=45)
        ax1.set_ylim([0.0, 1.0])
        ax2.set_ylim([0.0, 1.0])
        ax3.set_ylim([0.0, 1.0])
        ax4.set_ylim([0.0, 1.0])
        ax1.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax3.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax1.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax2.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax3.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax4.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax1.set_title('风格换手率', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        ax3.set_title('风格集中度', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=False)
        plt.savefig('{0}{1}_style_sc.png'.format(self.file_path, self.manager_name))

        hbs_style_history_disp = pd.concat(hbs_style_history_list)
        hbs_style_history_disp = hbs_style_history_disp.sort_values('REPORT_DATE')
        change_dates = [date for date in self.change_dates if date >= min(hbs_style_history_disp['REPORT_DATE']) and date <= max(hbs_style_history_disp['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_style_history_disp['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.barplot(x='REPORT_DATE', y='STYLE_EXPOSURE', data=hbs_style_history_disp, hue='STYLE', hue_order=['成长', '价值'], palette=[bar_color_list[0], bar_color_list[7]])
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        plt.xticks(rotation=45)
        ax.set_xlabel('')
        ax.set_ylabel('')
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(hbs_style_history_disp['STYLE_EXPOSURE']), linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.title('风格暴露', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_style.png'.format(self.file_path, self.manager_name))

        style_shift_value = hbs_style_history['STYLE_SHIFT_RANK'].mean()
        style_cen_value = hbs_style_history['STYLE_CEN_RANK'].mean()
        style_shift_mark = '偏下' if style_shift_value <= 0.3 else '中等偏下' if style_shift_value > 0.3 and style_shift_value <= 0.5 else '中等偏上' if style_shift_value > 0.5 and style_shift_value <= 0.7 else '偏上'
        style_cen_mark = '偏下' if style_cen_value <= 0.3 else '中等偏下' if style_cen_value > 0.3 and style_cen_value <= 0.5 else '中等偏上' if style_cen_value > 0.5 and style_cen_value <= 0.7 else '偏上'
        style_mark = '配置型' if style_shift_value <= 0.5 and style_cen_value <= 0.5 else '专注型' if style_shift_value <= 0.5 and style_cen_value > 0.5 else '轮动型' if style_shift_value > 0.5 and style_cen_value <= 0.5 else '博弈型'
        style_type_index = hbs_style_history[['成长', '价值']].mean()[hbs_style_history[['成长', '价值']].mean() == hbs_style_history[['成长', '价值']].mean().max()].index.tolist()
        style_type_mark = style_type_index[0] if len(style_type_index) == 1 else '、'.join(style_type_index)

        style_text = '在风格配置上，{0}的换手率在偏股型基金整体中处于{1}水平，集中度在偏股型基金整体中处于{2}水平，属于风格{3}基金经理，且从成长价值平均绝对暴露看更偏{4}。'.format(self.manager_name, style_shift_mark, style_cen_mark, style_mark, style_type_mark)
        return style_text

    def get_manager_theme(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_theme_sc_history_list, hbs_theme_sc_rank_history_list, hbs_theme_history_list = [], [], []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_theme_sc_history = FEDB().read_hbs_theme_industry_property_history_given_code(fund_code)
            hbs_theme_sc_history = hbs_theme_sc_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_theme': 'THEME_SHIFT', 'c_level_theme': 'THEME_CEN'})
            hbs_theme_sc_history['FUND_CODE'] = hbs_theme_sc_history['FUND_CODE'].apply(lambda x: str(x).zfill(6))
            hbs_theme_sc_history['REPORT_DATE'] = hbs_theme_sc_history['REPORT_DATE'].astype(str)
            hbs_theme_sc_history['REPORT_DATE'] = hbs_theme_sc_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_theme_sc_history = hbs_theme_sc_history[(hbs_theme_sc_history['REPORT_DATE'] > begin_date) & (hbs_theme_sc_history['REPORT_DATE'] <= end_date)]
            hbs_theme_sc_history['MANAGER_NAME'] = self.manager_name
            hbs_theme_sc_history = hbs_theme_sc_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'THEME_SHIFT', 'THEME_CEN']]
            hbs_theme_sc_history_list.append(hbs_theme_sc_history)

            hbs_theme_sc_rank_history = FEDB().read_hbs_rank_history_given_code(fund_code)
            hbs_theme_sc_rank_history = hbs_theme_sc_rank_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_theme': 'THEME_SHIFT_RANK', 'c_level_theme': 'THEME_CEN_RANK'})
            hbs_theme_sc_rank_history['REPORT_DATE'] = hbs_theme_sc_rank_history['REPORT_DATE'].astype(str)
            hbs_theme_sc_rank_history['REPORT_DATE'] = hbs_theme_sc_rank_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_theme_sc_rank_history = hbs_theme_sc_rank_history[(hbs_theme_sc_rank_history['REPORT_DATE'] > begin_date) & (hbs_theme_sc_rank_history['REPORT_DATE'] <= end_date)]
            hbs_theme_sc_rank_history['MANAGER_NAME'] = self.manager_name
            hbs_theme_sc_rank_history = hbs_theme_sc_rank_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'THEME_SHIFT_RANK', 'THEME_CEN_RANK']]
            hbs_theme_sc_rank_history_list.append(hbs_theme_sc_rank_history)

            hbs_theme_history = HBDB().read_fund_theme_given_codes([fund_code])
            hbs_theme_history = hbs_theme_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zclb': 'IS_ZC', 'zblx': 'THEME', 'zgpszb': 'THEME_EXPOSURE', 'jsqbd': 'MV_IN_NA_DIFF', 'tlpj': 'HOMO_AVG'})
            hbs_theme_history['THEME'] = hbs_theme_history['THEME'].replace({'TMT': 'TMT', 'ZZ': '制造', 'ZQ': '周期', 'DJR': '大金融', 'XF': '消费', 'QT': '其他'})
            hbs_theme_history = hbs_theme_history[hbs_theme_history['IS_ZC'] == '1']
            hbs_theme_history['REPORT_DATE'] = hbs_theme_history['REPORT_DATE'].astype(str)
            hbs_theme_history['REPORT_DATE'] = hbs_theme_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_theme_history = hbs_theme_history[(hbs_theme_history['REPORT_DATE'] > begin_date) & (hbs_theme_history['REPORT_DATE'] <= end_date)]
            hbs_theme_history['MANAGER_NAME'] = self.manager_name
            hbs_theme_history = hbs_theme_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'THEME', 'THEME_EXPOSURE']]
            hbs_theme_history_list.append(hbs_theme_history)

        hbs_theme_sc_history = pd.concat(hbs_theme_sc_history_list)
        hbs_theme_sc_rank_history = pd.concat(hbs_theme_sc_rank_history_list)
        hbs_theme_history = pd.concat(hbs_theme_history_list)
        hbs_theme_history = hbs_theme_history.pivot(index=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], columns='THEME', values='THEME_EXPOSURE').reset_index()
        hbs_theme_history = hbs_theme_sc_history.merge(hbs_theme_sc_rank_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left').merge(hbs_theme_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left')
        hbs_theme_history = hbs_theme_history.sort_values('REPORT_DATE')
        for col in list(hbs_theme_history.columns[3:7]):
            hbs_theme_history[col] = hbs_theme_history[col].fillna(hbs_theme_history[col].interpolate())

        change_dates = [date for date in self.change_dates if date >= min(hbs_theme_history['REPORT_DATE']) and date <= max(hbs_theme_history['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_theme_history['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(1, 2, figsize=(12, 6))
        ax1, ax3 = ax[0], ax[1]
        ax2, ax4 = ax1.twinx(), ax3.twinx()
        ax1.plot(hbs_theme_history['REPORT_DATE'].values, hbs_theme_history['THEME_SHIFT'].values, label='主题换手率', color=line_color_list[0])
        ax2.plot(hbs_theme_history['REPORT_DATE'].values, hbs_theme_history['THEME_SHIFT_RANK'].values, label='主题换手率排名（右轴）', color=line_color_list[1])
        ax3.plot(hbs_theme_history['REPORT_DATE'].values, hbs_theme_history['THEME_CEN'].values, label='主题集中度', color=line_color_list[0])
        ax4.plot(hbs_theme_history['REPORT_DATE'].values, hbs_theme_history['THEME_CEN_RANK'].values, label='主题集中度排名（右轴）', color=line_color_list[1])
        h1, l1 = ax1.get_legend_handles_labels()
        h2, l2 = ax2.get_legend_handles_labels()
        h3, l3 = ax3.get_legend_handles_labels()
        h4, l4 = ax4.get_legend_handles_labels()
        ax1.legend(handles=h1 + h2, labels=l1 + l2, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax3.legend(handles=h3 + h4, labels=l3 + l4, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax1.set_xticks(np.arange(len(hbs_theme_history)))
        ax1.set_xticklabels(labels=hbs_theme_history['REPORT_DATE'], rotation=45)
        ax3.set_xticks(np.arange(len(hbs_theme_history)))
        ax3.set_xticklabels(labels=hbs_theme_history['REPORT_DATE'], rotation=45)
        ax1.set_ylim([0.0, 1.0])
        ax2.set_ylim([0.0, 1.0])
        ax3.set_ylim([0.0, 1.0])
        ax4.set_ylim([0.0, 1.0])
        ax1.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax3.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax1.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax2.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax3.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax4.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax1.set_title('主题换手率', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        ax3.set_title('主题集中度', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=False)
        plt.savefig('{0}{1}_theme_sc.png'.format(self.file_path, self.manager_name))

        theme_list = ['大金融', '消费', 'TMT', '制造', '周期']
        theme_list = [theme for theme in theme_list if theme in list(hbs_theme_history.columns)]
        hbs_theme_history_disp = hbs_theme_history[['REPORT_DATE'] + theme_list]
        hbs_theme_history_disp = hbs_theme_history_disp.sort_values('REPORT_DATE')
        hbs_theme_history_disp[theme_list] = hbs_theme_history_disp[theme_list].fillna(0.0)
        for i in range(1, len(theme_list)):
            hbs_theme_history_disp[theme_list[i]] = hbs_theme_history_disp[theme_list[i]] + hbs_theme_history_disp[theme_list[i - 1]]
        change_dates = [date for date in self.change_dates if date >= min(hbs_theme_history_disp['REPORT_DATE']) and date <= max(hbs_theme_history_disp['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_theme_history_disp['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        color_list = [bar_color_list[0], bar_color_list[1], bar_color_list[7], bar_color_list[8], bar_color_list[14]]
        for i in range(len(theme_list) - 1, -1, -1):
            sns.barplot(x='REPORT_DATE', y=theme_list[i], data=hbs_theme_history_disp, label=theme_list[i], color=color_list[i])
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=5)
        plt.xticks(rotation=45)
        ax.set_xlabel('')
        ax.set_ylabel('')
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(hbs_theme_history_disp[theme_list[-1]]), linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.title('主题暴露', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_theme.png'.format(self.file_path, self.manager_name))

        theme_shift_value = hbs_theme_history['THEME_SHIFT_RANK'].mean()
        theme_cen_value = hbs_theme_history['THEME_CEN_RANK'].mean()
        theme_shift_mark = '偏下' if theme_shift_value <= 0.3 else '中等偏下' if theme_shift_value > 0.3 and theme_shift_value <= 0.5 else '中等偏上' if theme_shift_value > 0.5 and theme_shift_value <= 0.7 else '偏上'
        theme_cen_mark = '偏下' if theme_cen_value <= 0.3 else '中等偏下' if theme_cen_value > 0.3 and theme_cen_value <= 0.5 else '中等偏上' if theme_cen_value > 0.5 and theme_cen_value <= 0.7 else '偏上'
        theme_mark = '配置型' if theme_shift_value <= 0.5 and theme_cen_value <= 0.5 else '专注型' if theme_shift_value <= 0.5 and theme_cen_value > 0.5 else '轮动型' if theme_shift_value > 0.5 and theme_cen_value <= 0.5 else '博弈型'
        theme_type_index = hbs_theme_history[theme_list].mean()[hbs_theme_history[theme_list].mean() == hbs_theme_history[theme_list].mean().max()].index.tolist()
        theme_type_mark = theme_type_index[0] if len(theme_type_index) == 1 else '、'.join(theme_type_index)

        theme_text = '在主题配置上，{0}的换手率在偏股型基金整体中处于{1}水平，集中度在偏股型基金整体中处于{2}水平，属于主题{3}基金经理，且从各主题平均绝对暴露看更偏{4}。'.format(self.manager_name, theme_shift_mark, theme_cen_mark, theme_mark, theme_type_mark)
        return theme_text

    def get_manager_industry(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_industry_sc_history_list, hbs_industry_sc_rank_history_list, hbs_industry_history_list = [], [], []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_industry_sc_history = FEDB().read_hbs_theme_industry_property_history_given_code(fund_code)
            hbs_industry_sc_history = hbs_industry_sc_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_ind': 'INDUSTRY_SHIFT', 'c_level_ind': 'INDUSTRY_CEN'})
            hbs_industry_sc_history['FUND_CODE'] = hbs_industry_sc_history['FUND_CODE'].apply(lambda x: str(x).zfill(6))
            hbs_industry_sc_history['REPORT_DATE'] = hbs_industry_sc_history['REPORT_DATE'].astype(str)
            hbs_industry_sc_history['REPORT_DATE'] = hbs_industry_sc_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_industry_sc_history = hbs_industry_sc_history[(hbs_industry_sc_history['REPORT_DATE'] > begin_date) & (hbs_industry_sc_history['REPORT_DATE'] <= end_date)]
            hbs_industry_sc_history['MANAGER_NAME'] = self.manager_name
            hbs_industry_sc_history = hbs_industry_sc_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'INDUSTRY_SHIFT', 'INDUSTRY_CEN']]
            hbs_industry_sc_history_list.append(hbs_industry_sc_history)

            hbs_industry_sc_rank_history = FEDB().read_hbs_rank_history_given_code(fund_code)
            hbs_industry_sc_rank_history = hbs_industry_sc_rank_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'shift_ratio_ind': 'INDUSTRY_SHIFT_RANK', 'c_level_ind': 'INDUSTRY_CEN_RANK'})
            hbs_industry_sc_rank_history['REPORT_DATE'] = hbs_industry_sc_rank_history['REPORT_DATE'].astype(str)
            hbs_industry_sc_rank_history['REPORT_DATE'] = hbs_industry_sc_rank_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_industry_sc_rank_history = hbs_industry_sc_rank_history[(hbs_industry_sc_rank_history['REPORT_DATE'] > begin_date) & (hbs_industry_sc_rank_history['REPORT_DATE'] <= end_date)]
            hbs_industry_sc_rank_history['MANAGER_NAME'] = self.manager_name
            hbs_industry_sc_rank_history = hbs_industry_sc_rank_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'INDUSTRY_SHIFT_RANK', 'INDUSTRY_CEN_RANK']]
            hbs_industry_sc_rank_history_list.append(hbs_industry_sc_rank_history)

            hbs_industry_history = HBDB().read_fund_holding_given_codes([fund_code])
            hbs_industry_history = hbs_industry_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zqdm': 'TICKER_SYMBOL', 'zqmc': 'SEC_SHORT_NAME', 'ccsz': 'HOLDING_MARKET_VALUE', 'ccsl': 'HOLDING_AMOUNT', 'zjbl': 'INDUSTRY_EXPOSURE'})
            hbs_industry_history = hbs_industry_history.sort_values(['FUND_CODE', 'REPORT_DATE', 'INDUSTRY_EXPOSURE'], ascending=[True, True, False]).groupby(['FUND_CODE', 'REPORT_DATE']).head(10)
            stock_industry = HBDB().read_stock_industry()
            stock_industry = stock_industry.rename(columns={'zqdm': 'TICKER_SYMBOL', 'flmc': 'INDUSTRY', 'fldm': 'INDUSTRY_ID', 'fljb': 'INDUSTRY_TYPE', 'hyhfbz': 'INDUSTRY_VERSION', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'sfyx': 'IS_NEW'})
            stock_industry = stock_industry[stock_industry['IS_NEW'] == 1]
            stock_industry = stock_industry[stock_industry['INDUSTRY_VERSION'] == 2]
            stock_industry = stock_industry[stock_industry['INDUSTRY_TYPE'] == '1']
            stock_industry = stock_industry.sort_values(['TICKER_SYMBOL', 'credt_etl']).drop_duplicates('TICKER_SYMBOL', keep='last')
            hbs_industry_history = hbs_industry_history.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY']], on=['TICKER_SYMBOL'], how='left')
            hbs_industry_history = hbs_industry_history[['FUND_CODE', 'REPORT_DATE', 'INDUSTRY', 'INDUSTRY_EXPOSURE']].groupby(['FUND_CODE', 'REPORT_DATE', 'INDUSTRY']).sum().reset_index()
            hbs_industry_history['REPORT_DATE'] = hbs_industry_history['REPORT_DATE'].astype(str)
            hbs_industry_history['REPORT_DATE'] = hbs_industry_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_industry_history = hbs_industry_history[(hbs_industry_history['REPORT_DATE'] > begin_date) & (hbs_industry_history['REPORT_DATE'] <= end_date)]
            hbs_industry_history['MANAGER_NAME'] = self.manager_name
            hbs_industry_history = hbs_industry_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'INDUSTRY', 'INDUSTRY_EXPOSURE']]
            hbs_industry_history_list.append(hbs_industry_history)

        hbs_industry_sc_history = pd.concat(hbs_industry_sc_history_list)
        hbs_industry_sc_rank_history = pd.concat(hbs_industry_sc_rank_history_list)
        hbs_industry_history = pd.concat(hbs_industry_history_list)
        hbs_industry_history = hbs_industry_history.pivot(index=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], columns='INDUSTRY', values='INDUSTRY_EXPOSURE').reset_index()
        hbs_industry_history = hbs_industry_sc_history.merge(hbs_industry_sc_rank_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left').merge(hbs_industry_history, on=['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE'], how='left')
        hbs_industry_history = hbs_industry_history.sort_values('REPORT_DATE')
        for col in list(hbs_industry_history.columns[3:7]):
            hbs_industry_history[col] = hbs_industry_history[col].fillna(hbs_industry_history[col].interpolate())

        change_dates = [date for date in self.change_dates if date >= min(hbs_industry_history['REPORT_DATE']) and date <= max(hbs_industry_history['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_industry_history['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(1, 2, figsize=(12, 6))
        ax1, ax3 = ax[0], ax[1]
        ax2, ax4 = ax1.twinx(), ax3.twinx()
        ax1.plot(hbs_industry_history['REPORT_DATE'].values, hbs_industry_history['INDUSTRY_SHIFT'].values, label='行业换手率', color=line_color_list[0])
        ax2.plot(hbs_industry_history['REPORT_DATE'].values, hbs_industry_history['INDUSTRY_SHIFT_RANK'].values, label='行业题换手率排名（右轴）', color=line_color_list[1])
        ax3.plot(hbs_industry_history['REPORT_DATE'].values, hbs_industry_history['INDUSTRY_CEN'].values, label='行业集中度', color=line_color_list[0])
        ax4.plot(hbs_industry_history['REPORT_DATE'].values, hbs_industry_history['INDUSTRY_CEN_RANK'].values, label='行业集中度排名（右轴）', color=line_color_list[1])
        h1, l1 = ax1.get_legend_handles_labels()
        h2, l2 = ax2.get_legend_handles_labels()
        h3, l3 = ax3.get_legend_handles_labels()
        h4, l4 = ax4.get_legend_handles_labels()
        ax1.legend(handles=h1 + h2, labels=l1 + l2, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax3.legend(handles=h3 + h4, labels=l3 + l4, loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax1.set_xticks(np.arange(len(hbs_industry_history)))
        ax1.set_xticklabels(labels=hbs_industry_history['REPORT_DATE'], rotation=45)
        ax3.set_xticks(np.arange(len(hbs_industry_history)))
        ax3.set_xticklabels(labels=hbs_industry_history['REPORT_DATE'], rotation=45)
        ax1.set_ylim([0.0, 1.0])
        ax2.set_ylim([0.0, 1.0])
        ax3.set_ylim([0.0, 1.0])
        ax4.set_ylim([0.0, 1.0])
        ax1.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax3.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax1.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax2.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax3.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax4.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax1.set_title('行业换手率', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        ax3.set_title('行业集中度', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=False)
        plt.savefig('{0}{1}_industry_sc.png'.format(self.file_path, self.manager_name))

        industry_list = ['电力设备', '商贸零售', '食品饮料', '农林牧渔', '有色金属', '银行', '非银金融', '汽车', '石油石化', '建筑材料', '房地产', '钢铁', '建筑装饰', '公用事业', '医药生物', '家用电器', '电子', '通信', '轻工制造', '机械设备', '传媒', '基础化工', '计算机', '煤炭', '交通运输', '纺织服装', '环保', '社会服务', '美容护理', '国防军工', '综合']
        industry_list = [industry for industry in industry_list if industry in list(hbs_industry_history.columns)]
        hbs_industry_history_disp = hbs_industry_history[['REPORT_DATE'] + industry_list]
        hbs_industry_history_disp = hbs_industry_history_disp.sort_values('REPORT_DATE')
        hbs_industry_history_disp[industry_list] = hbs_industry_history_disp[industry_list].fillna(0.0)
        for i in range(1, len(industry_list)):
            hbs_industry_history_disp[industry_list[i]] = hbs_industry_history_disp[industry_list[i]] + hbs_industry_history_disp[industry_list[i - 1]]
        change_dates = [date for date in self.change_dates if date >= min(hbs_industry_history_disp['REPORT_DATE']) and date <= max(hbs_industry_history_disp['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_industry_history_disp['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        for i in range(len(industry_list) - 1, -1, -1):
            sns.barplot(x='REPORT_DATE', y=industry_list[i], data=hbs_industry_history_disp, label=industry_list[i], color=(bar_color_list + bar_color_list)[i])
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.45), ncol=8)
        plt.xticks(rotation=45)
        ax.set_xlabel('')
        ax.set_ylabel('')
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(hbs_industry_history_disp[industry_list[-1]]), linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.title('行业暴露', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_industry.png'.format(self.file_path, self.manager_name))

        industry_shift_value = hbs_industry_history['INDUSTRY_SHIFT_RANK'].mean()
        industry_cen_value = hbs_industry_history['INDUSTRY_CEN_RANK'].mean()
        industry_shift_mark = '偏下' if industry_shift_value <= 0.3 else '中等偏下' if industry_shift_value > 0.3 and industry_shift_value <= 0.5 else '中等偏上' if industry_shift_value > 0.5 and industry_shift_value <= 0.7 else '偏上'
        industry_cen_mark = '偏下' if industry_cen_value <= 0.3 else '中等偏下' if industry_cen_value > 0.3 and industry_cen_value <= 0.5 else '中等偏上' if industry_cen_value > 0.5 and industry_cen_value <= 0.7 else '偏上'
        industry_mark = '配置型' if industry_shift_value <= 0.5 and industry_cen_value <= 0.5 else '专注型' if industry_shift_value <= 0.5 and industry_cen_value > 0.5 else '轮动型' if industry_shift_value > 0.5 and industry_cen_value <= 0.5 else '博弈型'
        industry_type_index = hbs_industry_history[industry_list].mean()[hbs_industry_history[industry_list].mean() == hbs_industry_history[industry_list].mean().max()].index.tolist()
        industry_type_mark = industry_type_index[0] if len(industry_type_index) == 1 else '、'.join(industry_type_index)

        industry_text = '在行业配置上，{0}的换手率在偏股型基金整体中处于{1}水平，集中度在偏股型基金整体中处于{2}水平，属于行业{3}基金经理，且在{4}行业上平均绝对暴露最大。'.format(self.manager_name, industry_shift_mark, industry_cen_mark, industry_mark, industry_type_mark)
        return industry_text

    def get_manager_turnover(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_turnover_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_turnover = HBDB().read_fund_turnover_given_code(fund_code)
            hbs_turnover = hbs_turnover.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'tjqj': 'TYPE', 'hsl': 'TURNOVER'})
            hbs_turnover['REPORT_DATE'] = hbs_turnover['REPORT_DATE'].astype(str)
            hbs_turnover = hbs_turnover.sort_values('REPORT_DATE')
            hbs_turnover = hbs_turnover[(hbs_turnover['REPORT_DATE'] > begin_date) & (hbs_turnover['REPORT_DATE'] <= end_date)]
            hbs_turnover = hbs_turnover[hbs_turnover['TYPE'].isin(['1', '3'])]
            hbs_turnover['TURNOVER'] = hbs_turnover['TURNOVER'] * 2.0
            hbs_turnover = hbs_turnover[['REPORT_DATE', 'TURNOVER']]
            hbs_turnover_list.append(hbs_turnover)

        hbs_turnover = pd.concat(hbs_turnover_list)
        change_dates = [date for date in self.change_dates if date >= min(hbs_turnover['REPORT_DATE']) and date <= max(hbs_turnover['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_turnover['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.barplot(ax=ax, x='REPORT_DATE', y='TURNOVER', data=hbs_turnover, palette=['#C94649'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(hbs_turnover['TURNOVER']), linestyles='dashed', color='#959595')
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.xlabel('')
        plt.ylabel('')
        plt.xticks(rotation=45)
        plt.title('换手率', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_turnover.png'.format(self.file_path, self.manager_name))

        fund_turnover_text = '在基金换手率方面，{0}管理期间单边年化换手倍数平均稳定在{1}倍左右。'.format(self.manager_name,  round(hbs_turnover['TURNOVER'].mean() / 100.0, 2))
        return fund_turnover_text

    def get_manager_cr(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_cr_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_cr = HBDB().read_fund_cr_given_code(fund_code)
            hbs_cr = hbs_cr.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'qsdzb': 'CR10', 'qsdzbtlpj': 'CR10_MEAN', 'qwdzb': 'CR5', 'qwdzbtlpj': 'CR5_MEAN'})
            hbs_cr['REPORT_DATE'] = hbs_cr['REPORT_DATE'].astype(str)
            hbs_cr = hbs_cr.sort_values('REPORT_DATE')
            hbs_cr = hbs_cr[(hbs_cr['REPORT_DATE'] > begin_date) & (hbs_cr['REPORT_DATE'] <= end_date)]
            hbs_cr = hbs_cr[['REPORT_DATE', 'CR5', 'CR5_MEAN', 'CR10', 'CR10_MEAN']]
            hbs_cr.columns = ['报告日期', '前五大股票占比（%）', '前五大股票占比同类平均（%）', '前十大股票占比（%）', '前十大股票占比同类平均（%）']
            hbs_cr_list.append(hbs_cr)

        hbs_cr = pd.concat(hbs_cr_list)
        cr5 = hbs_cr[['报告日期', '前五大股票占比（%）']].rename(columns={'前五大股票占比（%）': '占净资产比（%）'})
        cr5['前五大占比'] = '前五大股票占比'
        cr5_mean = hbs_cr[['报告日期', '前五大股票占比同类平均（%）']].rename(columns={'前五大股票占比同类平均（%）': '占净资产比（%）'})
        cr5_mean['前五大占比'] = '前五大股票占比同类平均'
        cr5_mean = pd.concat([cr5, cr5_mean]).reset_index().drop('index', axis=1)
        cr10 = hbs_cr[['报告日期', '前十大股票占比（%）']].rename(columns={'前十大股票占比（%）': '占净资产比（%）'})
        cr10['前十大占比'] = '前十大股票占比'
        cr10_mean = hbs_cr[['报告日期', '前十大股票占比同类平均（%）']].rename(columns={'前十大股票占比同类平均（%）': '占净资产比（%）'})
        cr10_mean['前十大占比'] = '前十大股票占比同类平均'
        cr10_mean = pd.concat([cr10, cr10_mean]).reset_index().drop('index', axis=1)

        change_dates = [date for date in self.change_dates if date >= min(cr5_mean['报告日期']) and date <= max(cr5_mean['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(cr5_mean['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.barplot(ax=ax, x='报告日期', y='占净资产比（%）', data=cr5_mean, hue='前五大占比', palette=['#C94649', '#8588B7'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(cr5_mean['占净资产比（%）']), linestyles='dashed', color='#959595')
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.xlabel('')
        plt.ylabel('')
        plt.xticks(rotation=45)
        ax.set_title('前五大重仓占净资产比', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_cr5.png'.format(self.file_path, self.manager_name))

        fig, ax = plt.subplots(figsize=(12, 6))
        sns.barplot(ax=ax, x='报告日期', y='占净资产比（%）', data=cr10_mean, hue='前十大占比', palette=['#C94649', '#8588B7'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(cr10_mean['占净资产比（%）']), linestyles='dashed', color='#959595')
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.xlabel('')
        plt.ylabel('')
        plt.xticks(rotation=45)
        ax.set_title('前十大重仓占净资产比', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_cr10.png'.format(self.file_path, self.manager_name))

        fund_cr5_mean = round(hbs_cr['前五大股票占比（%）'].mean(), 2)
        fund_cr5_mean_mean = round(hbs_cr['前五大股票占比同类平均（%）'].mean(), 2)
        fund_cr10_mean = round(hbs_cr['前十大股票占比（%）'].mean(), 2)
        fund_cr10_mean_mean = round(hbs_cr['前十大股票占比同类平均（%）'].mean(), 2)
        cr5_text = '低' if fund_cr5_mean < fund_cr5_mean_mean else '高'
        cr10_text = '低' if fund_cr10_mean < fund_cr10_mean_mean else '高'
        text = '，持股集中度低' if cr5_text == '低' and cr10_text == '低' else ''

        fund_cr_text = '在持股集中度方面，{0}管理期间前五大权重保持在{1}%左右，整体{2}于同类平均水平，前十大权重保持在{3}%左右，整体{4}于同类平均水平{5}。' .format(self.manager_name, fund_cr5_mean, cr5_text, fund_cr10_mean, cr10_text, text)
        return fund_cr_text
    
    def get_manager_huddle(self):
        fund_holding = pd.read_hdf('D:/Git/hbshare/hbshare/fe/xwq/data/manager_report/fund_holding.hdf', key='table')
        fund_holding = fund_holding.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zqdm': 'TICKER_SYMBOL', 'zqmc': 'SEC_SHORT_NAME', 'ccsz': 'HOLDING_MARKET_VALUE', 'ccsl': 'HOLDING_AMOUNT', 'zjbl': 'MV_IN_NA'})
        fund_holding['REPORT_DATE'] = fund_holding['REPORT_DATE'].astype(str)
        stock_name = fund_holding.sort_values(['TICKER_SYMBOL', 'REPORT_DATE'])[['TICKER_SYMBOL', 'SEC_SHORT_NAME']].drop_duplicates('TICKER_SYMBOL', keep='last')
        fund_holding_stock = fund_holding[['REPORT_DATE', 'TICKER_SYMBOL', 'HOLDING_MARKET_VALUE']].groupby(['REPORT_DATE', 'TICKER_SYMBOL']).sum().reset_index()
        fund_holding_stock_weight = fund_holding_stock[['REPORT_DATE', 'HOLDING_MARKET_VALUE']].groupby(['REPORT_DATE']).sum().reset_index().rename(columns={'HOLDING_MARKET_VALUE': 'TOTAL_HOLDING_MARKET_VALUE'})
        fund_holding_stock = fund_holding_stock.merge(fund_holding_stock_weight, on=['REPORT_DATE'], how='left')
        fund_holding_stock['RATIO'] = fund_holding_stock['HOLDING_MARKET_VALUE'] / fund_holding_stock['TOTAL_HOLDING_MARKET_VALUE']
        fund_holding_stock = fund_holding_stock.merge(stock_name, on=['TICKER_SYMBOL'], how='left')
        fund_holding_stock_list = []
        for date in sorted(fund_holding_stock['REPORT_DATE'].unique().tolist()):
            fund_holding_stock_date = fund_holding_stock[fund_holding_stock['REPORT_DATE'] == date]
            fund_holding_stock_date['SCORE'] = (fund_holding_stock_date['RATIO'] - fund_holding_stock_date['RATIO'].min()) / (fund_holding_stock_date['RATIO'].max() - fund_holding_stock_date['RATIO'].min())
            fund_holding_stock_list.append(fund_holding_stock_date)
        fund_holding_stock = pd.concat(fund_holding_stock_list)
        fund_holding_stock.to_hdf('D:/Git/hbshare/hbshare/fe/xwq/data/manager_report/stock_huddle.hdf', key='table', mode='w')
        fund_holding = fund_holding.sort_values(['FUND_CODE', 'REPORT_DATE', 'MV_IN_NA'], ascending=[True, True, False]).groupby(['FUND_CODE', 'REPORT_DATE']).head(10)
        fund_holding = fund_holding.merge(fund_holding_stock[['REPORT_DATE', 'TICKER_SYMBOL', 'SCORE']], on=['REPORT_DATE', 'TICKER_SYMBOL'], how='left')
        fund_holding = fund_holding.dropna(subset=['SCORE'])
        fund_holding['WEIGHT_SCORE'] = fund_holding['MV_IN_NA'] * fund_holding['SCORE']
        fund_huddle = fund_holding[['FUND_CODE', 'REPORT_DATE', 'WEIGHT_SCORE']].groupby(['FUND_CODE', 'REPORT_DATE']).sum().reset_index().rename(columns={'WEIGHT_SCORE': 'SCORE'})
        fund_huddle.to_hdf('D:/Git/hbshare/hbshare/fe/xwq/data/manager_report/fund_huddle.hdf', key='table', mode='w')

        fund_huddle = pd.read_hdf('D:/Git/hbshare/hbshare/fe/xwq/data/manager_report/fund_huddle.hdf', key='table')
        fund = self.fund[self.fund['FUND_TYPE_2ND'].isin(['普通股票型', '灵活配置型', '偏股混合型'])]
        fund = fund.dropna(subset=['BEGIN_DATE'])
        fund['BEGIN_DATE'] = fund['BEGIN_DATE'].astype(int).astype(str)
        fund = fund[fund['BEGIN_DATE'] < (datetime.strptime('20220630', '%Y%m%d') - timedelta(730)).strftime('%Y%m%d')]
        fund_huddle = fund_huddle[fund_huddle['FUND_CODE'].isin(fund['FUND_CODE'].unique().tolist())]
        fund_huddle_mean = fund_huddle[['REPORT_DATE', 'SCORE']].groupby('REPORT_DATE').mean().reset_index().rename(columns={'SCORE': 'SCORE_MEAN'})
        fund_huddle_mean['REPORT_DATE'] = fund_huddle_mean['REPORT_DATE'].astype(str)

        representative_fund = self.representative_fund.copy(deep=True)
        hbs_huddle_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_huddle = fund_huddle[fund_huddle['FUND_CODE'] == fund_code]
            hbs_huddle['REPORT_DATE'] = hbs_huddle['REPORT_DATE'].astype(str)
            hbs_huddle = hbs_huddle[(hbs_huddle['REPORT_DATE'] > begin_date) & (hbs_huddle['REPORT_DATE'] <= end_date)]
            hbs_huddle = hbs_huddle[['REPORT_DATE', 'SCORE']]
            hbs_huddle_list.append(hbs_huddle)

        hbs_huddle = pd.concat(hbs_huddle_list)
        hbs_huddle = hbs_huddle.merge(fund_huddle_mean, on=['REPORT_DATE'], how='left')
        huddle = hbs_huddle[['REPORT_DATE', 'SCORE']].rename(columns={'SCORE': 'HUDDLE'})
        huddle['TYPE'] = '抱团程度'
        huddle_mean = hbs_huddle[['REPORT_DATE', 'SCORE_MEAN']].rename(columns={'SCORE_MEAN': 'HUDDLE'})
        huddle_mean['TYPE'] = '抱团程度同类平均'
        huddle = pd.concat([huddle, huddle_mean]).reset_index().drop('index', axis=1)

        change_dates = [date for date in self.change_dates if date >= min(huddle['REPORT_DATE']) and date <= max(huddle['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(huddle['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.barplot(ax=ax, x='REPORT_DATE', y='HUDDLE', data=huddle, hue='TYPE', palette=['#C94649', '#8588B7'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(huddle['HUDDLE']), linestyles='dashed', color='#959595')
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.xlabel('')
        plt.ylabel('')
        plt.xticks(rotation=45)
        ax.set_title('抱团程度', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_huddle.png'.format(self.file_path, self.manager_name))

        huddle_mean = round(hbs_huddle['SCORE'].mean(), 2)
        huddle_mean_mean = round(hbs_huddle['SCORE_MEAN'].mean(), 2)
        huddle_text = '低' if huddle_mean < huddle_mean_mean else '高'

        huddle_text = '在持股热度方面，{0}管理期间持股的抱团程度平均保持在{1}%左右，整体{2}于同类平均水平。'.format(self.manager_name, huddle_mean, huddle_text)
        return huddle_text
        

    def get_manager_industry_stock(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_industry_stock_rank_history_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_industry_stock_rank_history = FEDB().read_hbs_rank_history_given_code(fund_code)
            hbs_industry_stock_rank_history = hbs_industry_stock_rank_history.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'c_level_ind': 'INDUSTRY_CEN_RANK', 'c_level_stock': 'STOCK_CEN_RANK'})
            hbs_industry_stock_rank_history['REPORT_DATE'] = hbs_industry_stock_rank_history['REPORT_DATE'].astype(str)
            hbs_industry_stock_rank_history['REPORT_DATE'] = hbs_industry_stock_rank_history['REPORT_DATE'].apply(lambda x: x[:6] + '31' if x[4:6] == '03' or x[4:6] == '12' else x[:6] + '30')
            hbs_industry_stock_rank_history = hbs_industry_stock_rank_history[(hbs_industry_stock_rank_history['REPORT_DATE'] > begin_date) & (hbs_industry_stock_rank_history['REPORT_DATE'] <= end_date)]
            hbs_industry_stock_rank_history['MANAGER_NAME'] = self.manager_name
            hbs_industry_stock_rank_history = hbs_industry_stock_rank_history[['MANAGER_NAME', 'FUND_CODE', 'REPORT_DATE', 'INDUSTRY_CEN_RANK', 'STOCK_CEN_RANK']]
            hbs_industry_stock_rank_history_list.append(hbs_industry_stock_rank_history)

        hbs_industry_stock_rank_history = pd.concat(hbs_industry_stock_rank_history_list)
        hbs_industry_stock_rank_history = hbs_industry_stock_rank_history.sort_values('REPORT_DATE')
        change_dates = [date for date in self.change_dates if date >= min(hbs_industry_stock_rank_history['REPORT_DATE']) and date <= max(hbs_industry_stock_rank_history['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_industry_stock_rank_history['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        ax.plot(hbs_industry_stock_rank_history['REPORT_DATE'].values, hbs_industry_stock_rank_history['INDUSTRY_CEN_RANK'].values, label='行业集中度排名', color=line_color_list[0])
        ax.plot(hbs_industry_stock_rank_history['REPORT_DATE'].values, hbs_industry_stock_rank_history['STOCK_CEN_RANK'].values, label='个股集中度排名', color=line_color_list[1])
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        ax.set_xticks(np.arange(len(hbs_industry_stock_rank_history)))
        ax.set_xticklabels(labels=hbs_industry_stock_rank_history['REPORT_DATE'], rotation=45)
        ax.set_ylim([0.0, 1.0])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        ax.set_title('行业/个股集中度情况', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_industry_stock.png'.format(self.file_path, self.manager_name))

        industry_cen_mark = '低于50%' if hbs_industry_stock_rank_history['INDUSTRY_CEN_RANK'].mean() < 0.5 else '高于70%' if hbs_industry_stock_rank_history['INDUSTRY_CEN_RANK'].mean() > 0.7 else '处于50%与70%之间'
        stock_cen_mark = '低于50%' if hbs_industry_stock_rank_history['STOCK_CEN_RANK'].mean() < 0.5 else '高于70%' if  hbs_industry_stock_rank_history['STOCK_CEN_RANK'].mean() > 0.7 else '处于50%与70%之间'
        industry_stock_cen_mark = '选股方法为自下而上' if stock_cen_mark == '高于70%' and industry_cen_mark == '低于50%' else '选股方法为自上而下' if stock_cen_mark == '低于50%' and industry_cen_mark == '高于70%' else '无明显的自上而下或者自下而上选股特征'
        industry_stock_text = '{0}管理期间，在行业上的平均集中度排名{1}，在个股上的平均集中度{2}，{3}。'.format(self.manager_name, industry_cen_mark, stock_cen_mark, industry_stock_cen_mark)
        return industry_stock_text

    def get_manager_left(self):
        representative_fund = self.representative_fund.copy(deep=True)
        representative_fund = representative_fund.replace(datetime.today().date().strftime('%Y%m%d'), '-')
        representative_fund = representative_fund.merge(self.fund[['FUND_CODE', 'FUND_SHORT_NAME']], on=['FUND_CODE'], how='left')
        representative_fund = representative_fund[['FUND_SHORT_NAME', 'FUND_CODE', 'BEGIN_DATE', 'END_DATE']]
        representative_fund = representative_fund.sort_values('BEGIN_DATE')
        hbs_left_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            fund_name = row['FUND_SHORT_NAME']

            hbs_left = FEDB().read_hbs_stock_trading_property_given_code(fund_code)
            hbs_left = hbs_left.rename(columns={'jjdm': '代表基金代码'})
            hbs_left = hbs_left[['代表基金代码', '左侧概率（出持仓前,半年线）_rank', '左侧程度（出持仓前,半年线）']]
            hbs_left['代表基金名称'] = fund_name
            hbs_left['代表基金左侧特征'] = '深度左侧' if (hbs_left['左侧概率（出持仓前,半年线）_rank'][0] > 0.5 and hbs_left['左侧程度（出持仓前,半年线）'][0] > 0.5) else '左侧特征' if (hbs_left['左侧概率（出持仓前,半年线）_rank'][0] > 0.5 and hbs_left['左侧程度（出持仓前,半年线）'][0] <= 0.5) else '无明显的左侧特征'
            hbs_left = hbs_left[['代表基金名称', '代表基金代码', '左侧概率（出持仓前,半年线）_rank', '左侧程度（出持仓前,半年线）', '代表基金左侧特征']]
            for col in ['左侧概率（出持仓前,半年线）_rank', '左侧程度（出持仓前,半年线）']:
                hbs_left[col].iloc[0] = '{0}%'.format(round(hbs_left[col].iloc[0] * 100.0, 2))
            hbs_left_list.append(hbs_left)

        hbs_left = pd.concat(hbs_left_list)
        hbs_left = hbs_left.rename(columns={'左侧概率（出持仓前,半年线）_rank': '左侧概率（出持仓前,半年线）排名', '左侧程度（出持仓前,半年线）': '左侧程度（出持仓前,半年线）排名'})
        hbs_left = hbs_left.T.reset_index().T.reset_index().drop('index', axis=1)
        hbs_left_text = '{0}管理期间，其代表基金左侧情况如下：'.format(self.manager_name)
        return hbs_left_text, hbs_left

    def get_manager_new(self):
        representative_fund = self.representative_fund.copy(deep=True)
        representative_fund = representative_fund.replace(datetime.today().date().strftime('%Y%m%d'), '-')
        representative_fund = representative_fund.merge(self.fund[['FUND_CODE', 'FUND_SHORT_NAME']], on=['FUND_CODE'], how='left')
        representative_fund = representative_fund[['FUND_SHORT_NAME', 'FUND_CODE', 'BEGIN_DATE', 'END_DATE']]
        representative_fund = representative_fund.sort_values('BEGIN_DATE')
        hbs_new_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            fund_name = row['FUND_SHORT_NAME']

            hbs_new = FEDB().read_hbs_stock_trading_property_given_code(fund_code)
            hbs_new = hbs_new.rename(columns={'jjdm': '代表基金代码'})
            hbs_new = hbs_new[['代表基金代码', '新股概率（出持仓前）_rank', '次新股概率（出持仓前）_rank']]
            hbs_new['代表基金名称'] = fund_name
            hbs_new['代表基金新股/次新股偏好'] = '偏好新股/次新股' if hbs_new['新股概率（出持仓前）_rank'][0] > 0.75 and hbs_new['次新股概率（出持仓前）_rank'][0] > 0.75 else '偏好新股' if hbs_new['新股概率（出持仓前）_rank'][0] > 0.75 and hbs_new['次新股概率（出持仓前）_rank'][0] <= 0.75 else '偏好次新股' if hbs_new['次新股概率（出持仓前）_rank'][0] <= 0.75 and hbs_new['次新股概率（出持仓前）_rank'][0] > 0.75 else '无明显的新股/次新股偏好'
            hbs_new = hbs_new[['代表基金名称', '代表基金代码', '新股概率（出持仓前）_rank', '次新股概率（出持仓前）_rank', '代表基金新股/次新股偏好']]
            for col in ['新股概率（出持仓前）_rank', '次新股概率（出持仓前）_rank']:
                hbs_new[col].iloc[0] = '{0}%'.format(round(hbs_new[col].iloc[0] * 100.0, 2))
            hbs_new_list.append(hbs_new)

        hbs_new = pd.concat(hbs_new_list)
        hbs_new = hbs_new.rename(columns={'新股概率（出持仓前）_rank': '偏好新股概率（出持仓前）排名', '次新股概率（出持仓前）_rank': '偏好次新股概率（出持仓前）排名'})
        hbs_new = hbs_new.T.reset_index().T.reset_index().drop('index', axis=1)
        hbs_new_text = '{0}管理期间，其代表基金新股/次新股偏好情况如下：'.format(self.manager_name)
        return hbs_new_text, hbs_new
    
    def get_manager_holder(self):
        representative_fund = self.representative_fund.copy(deep=True)
        hbs_holder_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            hbs_holder = HBDB().read_fund_holder_given_code(fund_code)
            hbs_holder = hbs_holder.rename( columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'ggrq': 'PUBLISH_DATE', 'jgcyfe': 'ORGAN_HOLDING', 'jgcybl': 'ORGAN_HOLDING_RATIO', 'grcyfe': 'PERSON_HOLDING', 'grcybl': 'PERSON_HOLDING_RATIO'})
            hbs_holder['REPORT_DATE'] = hbs_holder['REPORT_DATE'].astype(str)
            hbs_holder = hbs_holder[(hbs_holder['REPORT_DATE'] > begin_date) & (hbs_holder['REPORT_DATE'] <= end_date)]
            hbs_holder = hbs_holder[['REPORT_DATE', 'ORGAN_HOLDING_RATIO', 'PERSON_HOLDING_RATIO']]
            hbs_holder_list.append(hbs_holder)
            
        hbs_holder = pd.concat(hbs_holder_list)
        hbs_holder = hbs_holder.sort_values('REPORT_DATE')
        hbs_holder['PERSON_HOLDING_RATIO'] = hbs_holder['PERSON_HOLDING_RATIO'] + hbs_holder['ORGAN_HOLDING_RATIO']
        change_dates = [date for date in self.change_dates if date >= min(hbs_holder['REPORT_DATE']) and date <= max(hbs_holder['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(hbs_holder['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.barplot(ax=ax, x='REPORT_DATE', y='PERSON_HOLDING_RATIO', data=hbs_holder, label='个人持有比例', color='#8588B7')
        sns.barplot(ax=ax, x='REPORT_DATE', y='ORGAN_HOLDING_RATIO', data=hbs_holder, label='机构持有比例', color='#C94649')
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=100.0, linestyles='dashed', color='#959595')
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=2)
        plt.xticks(rotation=45)
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.xlabel('')
        plt.ylabel('')
        plt.title('机构/个人持有比例', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_holder.png'.format(self.file_path, self.manager_name))

        holder_mark = '高于' if hbs_holder['ORGAN_HOLDING_RATIO'].mean() > 50 else '低于'
        position_text = '从{0}代表基金的机构/个人持仓占比看，机构对基金经理的关注度{1}个人投资者。'.format(self.manager_name, holder_mark)
        return position_text

    def get_manager_performance(self):
        representative_fund = self.representative_fund.copy(deep=True)
        fund_nav_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            fund_nav = HBDB().read_fund_cumret_given_code(fund_code, begin_date, end_date)
            fund_nav = fund_nav.rename(columns={'JJDM': 'FUND_CODE', 'JZRQ': 'TRADE_DATE', 'HBCL': 'CUM_RET'})
            fund_nav['TRADE_DATE'] = fund_nav['TRADE_DATE'].astype(str)
            fund_nav = fund_nav.sort_values('TRADE_DATE')
            fund_nav['CUM_RET'] = 0.01 * fund_nav['CUM_RET']
            fund_nav['NAV'] = fund_nav['CUM_RET'] + 1
            fund_nav['NAV'] = fund_nav['NAV'] / fund_nav['NAV'].iloc[0]
            fund_nav['RET'] = fund_nav['NAV'].pct_change()
            fund_nav = fund_nav[(fund_nav['TRADE_DATE'] > begin_date) & (fund_nav['TRADE_DATE'] <= end_date)]
            fund_nav_list.append(fund_nav)

        fund_nav = pd.concat(fund_nav_list)
        fund_nav = fund_nav[['TRADE_DATE', 'RET']].set_index('TRADE_DATE')
        fund_nav.loc[representative_fund['BEGIN_DATE'].iloc[0]] = 0.0
        fund_nav = fund_nav.sort_index().reset_index()
        fund_nav['NAV'] = (fund_nav['RET'] + 1).cumprod()
        fund_nav['CUM_RET'] = fund_nav['NAV'] / fund_nav['NAV'].iloc[0] - 1.0
        fund_nav['IDX'] = range(len(fund_nav))
        fund_nav['DRAWDOWN'] = fund_nav['IDX'].apply(lambda x: fund_nav['NAV'].iloc[x] / max(fund_nav['NAV'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        fund_nav['MANAGER_NAME'] = self.manager_name
        fund_cumret = fund_nav.pivot(index='TRADE_DATE', columns='MANAGER_NAME', values='CUM_RET')
        fund_d = fund_nav.pivot(index='TRADE_DATE', columns='MANAGER_NAME', values='DRAWDOWN')

        stock_fund_index_name = '偏股混合型基金指数'
        stock_fund_index_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(representative_fund['BEGIN_DATE'].iloc[0], ['885001'])
        stock_fund_index_daily_k = stock_fund_index_daily_k[['zqmc', 'jyrq', 'spjg']].rename(columns={'zqmc': 'INDEX_NAME', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        stock_fund_index_daily_k['TRADE_DATE'] = stock_fund_index_daily_k['TRADE_DATE'].astype(str)
        stock_fund_index_daily_k = stock_fund_index_daily_k.sort_values('TRADE_DATE')
        stock_fund_index_daily_k = stock_fund_index_daily_k[(stock_fund_index_daily_k['TRADE_DATE'] > min(fund_cumret.index)) & (stock_fund_index_daily_k['TRADE_DATE'] <= max(fund_cumret.index))]
        stock_fund_index_daily_k['INDEX_NAME'] = stock_fund_index_name
        stock_fund_index_daily_k['CUM_RET'] = stock_fund_index_daily_k['CLOSE_INDEX'] / stock_fund_index_daily_k['CLOSE_INDEX'].iloc[0] - 1.0
        stock_fund_index_daily_k['NAV'] = stock_fund_index_daily_k['CLOSE_INDEX'] / stock_fund_index_daily_k['CLOSE_INDEX'].iloc[0]
        stock_fund_index_daily_k['IDX'] = range(len(stock_fund_index_daily_k))
        stock_fund_index_daily_k['DRAWDOWN'] = stock_fund_index_daily_k['IDX'].apply(lambda x: stock_fund_index_daily_k['CLOSE_INDEX'].iloc[x] / max(stock_fund_index_daily_k['CLOSE_INDEX'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        stock_fund_index_cumret = stock_fund_index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CUM_RET')
        stock_fund_index_d = stock_fund_index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='DRAWDOWN')

        index = '000300'
        index_name = '沪深300'
        index_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(representative_fund['BEGIN_DATE'].iloc[0], [index])
        index_daily_k = index_daily_k[['zqmc', 'jyrq', 'spjg']].rename(columns={'zqmc': 'INDEX_NAME', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        index_daily_k['TRADE_DATE'] = index_daily_k['TRADE_DATE'].astype(str)
        index_daily_k = index_daily_k.sort_values('TRADE_DATE')
        index_daily_k = index_daily_k[(index_daily_k['TRADE_DATE'] > min(fund_cumret.index)) & (index_daily_k['TRADE_DATE'] <= max(fund_cumret.index))]
        index_daily_k['CUM_RET'] = index_daily_k['CLOSE_INDEX'] / index_daily_k['CLOSE_INDEX'].iloc[0] - 1.0
        index_daily_k['NAV'] = index_daily_k['CLOSE_INDEX'] / index_daily_k['CLOSE_INDEX'].iloc[0]
        index_daily_k['IDX'] = range(len(index_daily_k))
        index_daily_k['DRAWDOWN'] = index_daily_k['IDX'].apply(lambda x: index_daily_k['CLOSE_INDEX'].iloc[x] / max(index_daily_k['CLOSE_INDEX'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        index_cumret = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CUM_RET')
        index_d = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='DRAWDOWN')

        fund_index_cumret = fund_cumret.merge(stock_fund_index_cumret, left_index=True, right_index=True, how='left').merge(index_cumret, left_index=True, right_index=True, how='left').reset_index()
        fund_index_cumret.iloc[0, 1:] = 0.0
        fund_index_cumret['TRADE_DATE'] = fund_index_cumret['TRADE_DATE'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_index_cumret['TRADE_DATE']) and date <= max(fund_index_cumret['TRADE_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_index_cumret['TRADE_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 8))
        sns.lineplot(x='TRADE_DATE', y=self.manager_name, data=fund_index_cumret, color='#F04950', label=self.manager_name)
        sns.lineplot(x='TRADE_DATE', y=stock_fund_index_name, data=fund_index_cumret, color='#6268A2', label=stock_fund_index_name)
        sns.lineplot(x='TRADE_DATE', y=index_name, data=fund_index_cumret, color='#959595', label=index_name)
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(max(fund_index_cumret[self.manager_name]), max(fund_index_cumret[index_name]), max(fund_index_cumret[stock_fund_index_name])), linestyles='dashed', color='#959595')
        ax.xaxis.set_major_locator(ticker.MultipleLocator(90))
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.xticks(rotation=45)
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=3)
        plt.xlabel('')
        plt.ylabel('')
        plt.title('累计收益率', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_return.png'.format(self.file_path, self.manager_name))

        fund_index_d = fund_d.merge(stock_fund_index_d, left_index=True, right_index=True, how='left').merge(index_d, left_index=True, right_index=True, how='left').reset_index()
        fund_index_d.iloc[0, 1:] = 0.0
        fund_index_d['TRADE_DATE'] = fund_index_d['TRADE_DATE'].astype(str)
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.lineplot(x='TRADE_DATE', y=self.manager_name, data=fund_index_d, color='#F04950', label=self.manager_name)
        sns.lineplot(x='TRADE_DATE', y=stock_fund_index_name, data=stock_fund_index_d, color='#6268A2', label=stock_fund_index_name)
        sns.lineplot(x='TRADE_DATE', y=index_name, data=fund_index_d, color='#959595', label=index_name)
        ax.vlines(x=change_dates_loc, ymin=min(min(fund_index_d[self.manager_name]), min(fund_index_d[index_name]), min(fund_index_d[stock_fund_index_name])), ymax=0.0, linestyles='dashed', color='#959595')
        ax.xaxis.set_major_locator(ticker.MultipleLocator(90))
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.xticks(rotation=45)
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=3)
        plt.xlabel('')
        plt.ylabel('')
        plt.title('回撤', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_drawdown.png'.format(self.file_path, self.manager_name))

        fund_nav = fund_nav.pivot(index='TRADE_DATE', columns='MANAGER_NAME', values='NAV')
        index_daily_k = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CLOSE_INDEX')
        stock_fund_index_daily_k = stock_fund_index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME',  values='CLOSE_INDEX')
        nav_df = fund_nav.merge(stock_fund_index_daily_k, left_index=True, right_index=True, how='left').merge(index_daily_k, left_index=True, right_index=True, how='left')
        nav_df = nav_df.replace(0.0, np.nan).sort_index().fillna(method='ffill')

        achievement = pd.DataFrame(index=['业绩统计', '年化收益率', '年化波动率', 'Sharpe比率', '最大回撤', 'Calmar比率', '投资胜率', '平均损益比'], columns=range(len(list(nav_df.columns))))
        for idx, col in enumerate(list(nav_df.columns)):
            col_nav_df = nav_df[col].dropna()
            col_ret_df = col_nav_df.pct_change().dropna()
            achievement.loc['业绩统计', idx] = col
            achievement.loc['年化收益率', idx] = (col_nav_df.iloc[-1] / col_nav_df.iloc[0]) ** (250.0 / len(col_nav_df)) - 1.0
            achievement.loc['年化波动率', idx] = np.std(col_ret_df, ddof=1) * np.sqrt(250.0)
            achievement.loc['Sharpe比率', idx] = (achievement.loc['年化收益率', idx] - 0.03) / achievement.loc['年化波动率', idx]
            achievement.loc['最大回撤', idx] = -1.0 * max(
                [(min(col_nav_df.iloc[i:]) / col_nav_df.iloc[i] - 1.0) * (-1.0) for i in range(len(col_nav_df))])
            achievement.loc['Calmar比率', idx] = achievement.loc['年化收益率', idx] / achievement.loc['最大回撤', idx]
            achievement.loc['投资胜率', idx] = len(col_ret_df[col_ret_df >= 0]) / float(len(col_ret_df))
            achievement.loc['平均损益比', idx] = col_ret_df[col_ret_df >= 0].mean() / col_ret_df[col_ret_df < 0].mean() * (-1.0)

            achievement.loc['年化收益率', idx] = '{0}%'.format(round(achievement.loc['年化收益率', idx] * 100.0, 2))
            achievement.loc['年化波动率', idx] = '{0}%'.format(round(achievement.loc['年化波动率', idx] * 100.0, 2))
            achievement.loc['Sharpe比率', idx] = '{0}'.format(round(achievement.loc['Sharpe比率', idx], 2))
            achievement.loc['最大回撤', idx] = '{0}%'.format(round(achievement.loc['最大回撤', idx] * 100.0, 2))
            achievement.loc['Calmar比率', idx] = '{0}'.format(round(achievement.loc['Calmar比率', idx], 2))
            achievement.loc['投资胜率', idx] = '{0}%'.format(round(achievement.loc['投资胜率', idx] * 100.0, 2))
            achievement.loc['平均损益比', idx] = '{0}'.format(round(achievement.loc['平均损益比', idx], 2))

        achievement = achievement.reset_index()
        return achievement

    def get_manager_performance_all(self):
        manager_fund = self.manager_fund.merge(self.fund[['FUND_CODE', 'FUND_FULL_NAME', 'FUND_SHORT_NAME', 'FUND_TYPE_2ND']], on=['FUND_CODE'], how='left')
        manager_fund = manager_fund[manager_fund['FUND_TYPE_2ND'].isin(['偏股混合型', '灵活配置型', '普通股票型'])]
        fund_nav_list = []
        for idx, row in manager_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            fund_nav = HBDB().read_fund_cumret_given_code(fund_code, begin_date, end_date)
            fund_nav = fund_nav.rename(columns={'JJDM': 'FUND_CODE', 'JZRQ': 'TRADE_DATE', 'HBCL': 'CUM_RET'})
            fund_nav['TRADE_DATE'] = fund_nav['TRADE_DATE'].astype(str)
            fund_nav = fund_nav.sort_values('TRADE_DATE')
            fund_nav['CUM_RET'] = 0.01 * fund_nav['CUM_RET']
            fund_nav['NAV'] = fund_nav['CUM_RET'] + 1
            fund_nav['NAV'] = fund_nav['NAV'] / fund_nav['NAV'].iloc[0]
            fund_nav = fund_nav[(fund_nav['TRADE_DATE'] >= begin_date) & (fund_nav['TRADE_DATE'] <= end_date)]
            fund_nav_list.append(fund_nav)

        fund_nav = pd.concat(fund_nav_list)
        fund_nav['TRADE_DATE'] = fund_nav['TRADE_DATE'].astype(str)
        fund_nav = fund_nav.drop_duplicates(['FUND_CODE', 'TRADE_DATE'])
        fund_nav = fund_nav.pivot(index='TRADE_DATE', columns='FUND_CODE', values='NAV').sort_index()
        fund_nav = fund_nav.interpolate().pct_change()
        fund_nav = fund_nav.unstack().reset_index()
        fund_nav.columns = ['FUND_CODE', 'TRADE_DATE', 'RET']
        fund_nav = fund_nav.dropna(subset=['RET'])

        fund_scale = self.manager_fund_scale.drop_duplicates(['FUND_CODE', 'REPORT_DATE'])
        fund_scale = fund_scale.pivot(index='REPORT_DATE', columns='FUND_CODE', values='AUM').sort_index()
        fund_scale = fund_scale.reindex(self.calendar_df['CALENDAR_DATE'].unique().tolist()).sort_index().fillna(method='ffill')
        fund_scale = fund_scale.unstack().reset_index()
        fund_scale.columns = ['FUND_CODE', 'TRADE_DATE', 'AUM']
        fund_scale = fund_scale.dropna(subset=['AUM'])

        fund_position = self.manager_fund_position.drop_duplicates(['FUND_CODE', 'REPORT_DATE'])
        fund_position = fund_position.pivot(index='REPORT_DATE', columns='FUND_CODE', values='STOCK_RATIO').sort_index()
        fund_position = fund_position.reindex(self.calendar_df['CALENDAR_DATE'].unique().tolist()).sort_index().fillna(method='ffill')
        fund_position = fund_position.unstack().reset_index()
        fund_position.columns = ['FUND_CODE', 'TRADE_DATE', 'STOCK_RATIO']
        fund_position = fund_position.dropna(subset=['STOCK_RATIO'])

        fund_nav = fund_nav.merge(fund_position, on=['FUND_CODE', 'TRADE_DATE'], how='left')
        fund_nav = fund_nav[fund_nav['STOCK_RATIO'] >= 60]
        fund_nav = fund_nav.merge(fund_scale, on=['FUND_CODE', 'TRADE_DATE'], how='left')
        fund_nav = fund_nav.sort_values(['TRADE_DATE', 'FUND_CODE'])
        fund_nav_aum = fund_nav[['TRADE_DATE', 'AUM']].groupby('TRADE_DATE').sum().reset_index().rename(columns={'AUM': 'TOTAL_AUM'})
        fund_nav = fund_nav.merge(fund_nav_aum, on=['TRADE_DATE'], how='left')
        fund_nav['WEIGHT_RET'] = fund_nav['RET'] * fund_nav['AUM'] / fund_nav['TOTAL_AUM']
        fund_nav = fund_nav[['TRADE_DATE', 'WEIGHT_RET']].groupby('TRADE_DATE').sum().reset_index().rename(columns={'WEIGHT_RET': 'RET'})
        fund_nav = fund_nav.set_index('TRADE_DATE')
        fund_nav = fund_nav.sort_index().reset_index()
        fund_nav['NAV'] = (fund_nav['RET'] + 1).cumprod()
        fund_nav['CUM_RET'] = fund_nav['NAV'] / fund_nav['NAV'].iloc[0] - 1.0
        fund_nav['IDX'] = range(len(fund_nav))
        fund_nav['DRAWDOWN'] = fund_nav['IDX'].apply(lambda x: fund_nav['NAV'].iloc[x] / max(fund_nav['NAV'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        fund_nav['MANAGER_NAME'] = self.manager_name
        fund_cumret = fund_nav.pivot(index='TRADE_DATE', columns='MANAGER_NAME', values='CUM_RET')
        fund_d = fund_nav.pivot(index='TRADE_DATE', columns='MANAGER_NAME', values='DRAWDOWN')

        stock_fund_index_name = '偏股混合型基金指数'
        stock_fund_index_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(manager_fund['BEGIN_DATE'].min(), ['885001'])
        stock_fund_index_daily_k = stock_fund_index_daily_k[['zqmc', 'jyrq', 'spjg']].rename(columns={'zqmc': 'INDEX_NAME', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        stock_fund_index_daily_k['TRADE_DATE'] = stock_fund_index_daily_k['TRADE_DATE'].astype(str)
        stock_fund_index_daily_k = stock_fund_index_daily_k.sort_values('TRADE_DATE')
        stock_fund_index_daily_k = stock_fund_index_daily_k[(stock_fund_index_daily_k['TRADE_DATE'] > min(fund_cumret.index)) & (stock_fund_index_daily_k['TRADE_DATE'] <= max(fund_cumret.index))]
        stock_fund_index_daily_k['INDEX_NAME'] = stock_fund_index_name
        stock_fund_index_daily_k['CUM_RET'] = stock_fund_index_daily_k['CLOSE_INDEX'] / stock_fund_index_daily_k['CLOSE_INDEX'].iloc[0] - 1.0
        stock_fund_index_daily_k['NAV'] = stock_fund_index_daily_k['CLOSE_INDEX'] / stock_fund_index_daily_k['CLOSE_INDEX'].iloc[0]
        stock_fund_index_daily_k['IDX'] = range(len(stock_fund_index_daily_k))
        stock_fund_index_daily_k['DRAWDOWN'] = stock_fund_index_daily_k['IDX'].apply(lambda x: stock_fund_index_daily_k['CLOSE_INDEX'].iloc[x] / max(stock_fund_index_daily_k['CLOSE_INDEX'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        stock_fund_index_cumret = stock_fund_index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CUM_RET')
        stock_fund_index_d = stock_fund_index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='DRAWDOWN')

        index = '000300'
        index_name = '沪深300'
        index_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(manager_fund['BEGIN_DATE'].min(), [index])
        index_daily_k = index_daily_k[['zqmc', 'jyrq', 'spjg']].rename(columns={'zqmc': 'INDEX_NAME', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        index_daily_k['TRADE_DATE'] = index_daily_k['TRADE_DATE'].astype(str)
        index_daily_k = index_daily_k.sort_values('TRADE_DATE')
        index_daily_k = index_daily_k[(index_daily_k['TRADE_DATE'] > min(fund_cumret.index)) & (index_daily_k['TRADE_DATE'] <= max(fund_cumret.index))]
        index_daily_k['CUM_RET'] = index_daily_k['CLOSE_INDEX'] / index_daily_k['CLOSE_INDEX'].iloc[0] - 1.0
        index_daily_k['NAV'] = index_daily_k['CLOSE_INDEX'] / index_daily_k['CLOSE_INDEX'].iloc[0]
        index_daily_k['IDX'] = range(len(index_daily_k))
        index_daily_k['DRAWDOWN'] = index_daily_k['IDX'].apply(lambda x: index_daily_k['CLOSE_INDEX'].iloc[x] / max(index_daily_k['CLOSE_INDEX'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        index_cumret = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CUM_RET')
        index_d = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='DRAWDOWN')

        fund_index_cumret = fund_cumret.merge(stock_fund_index_cumret, left_index=True, right_index=True, how='left').merge(index_cumret, left_index=True, right_index=True, how='left').reset_index()
        fund_index_cumret.iloc[0, 1:] = 0.0
        fund_index_cumret['TRADE_DATE'] = fund_index_cumret['TRADE_DATE'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_index_cumret['TRADE_DATE']) and date <= max(fund_index_cumret['TRADE_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_index_cumret['TRADE_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 8))
        sns.lineplot(x='TRADE_DATE', y=self.manager_name, data=fund_index_cumret, color='#F04950', label=self.manager_name)
        sns.lineplot(x='TRADE_DATE', y=stock_fund_index_name, data=fund_index_cumret, color='#6268A2', label=stock_fund_index_name)
        sns.lineplot(x='TRADE_DATE', y=index_name, data=fund_index_cumret, color='#959595', label=index_name)
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(max(fund_index_cumret[self.manager_name]), max(fund_index_cumret[index_name]), max(fund_index_cumret[stock_fund_index_name])), linestyles='dashed', color='#959595')
        ax.xaxis.set_major_locator(ticker.MultipleLocator(90))
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.xticks(rotation=45)
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=3)
        plt.xlabel('')
        plt.ylabel('')
        plt.title('累计收益率', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_return_all.png'.format(self.file_path, self.manager_name))

        fund_index_d = fund_d.merge(stock_fund_index_d, left_index=True, right_index=True, how='left').merge(index_d, left_index=True, right_index=True, how='left').reset_index()
        fund_index_d.iloc[0, 1:] = 0.0
        fund_index_d['TRADE_DATE'] = fund_index_d['TRADE_DATE'].astype(str)
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.lineplot(x='TRADE_DATE', y=self.manager_name, data=fund_index_d, color='#F04950', label=self.manager_name)
        sns.lineplot(x='TRADE_DATE', y=stock_fund_index_name, data=stock_fund_index_d, color='#6268A2', label=stock_fund_index_name)
        sns.lineplot(x='TRADE_DATE', y=index_name, data=fund_index_d, color='#959595', label=index_name)
        ax.vlines(x=change_dates_loc, ymin=min(min(fund_index_d[self.manager_name]), min(fund_index_d[index_name]), min(fund_index_d[stock_fund_index_name])), ymax=0.0, linestyles='dashed', color='#959595')
        ax.xaxis.set_major_locator(ticker.MultipleLocator(90))
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.xticks(rotation=45)
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=3)
        plt.xlabel('')
        plt.ylabel('')
        plt.title('回撤', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_drawdown_all.png'.format(self.file_path, self.manager_name))

        fund_nav = fund_nav.pivot(index='TRADE_DATE', columns='MANAGER_NAME', values='NAV')
        index_daily_k = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CLOSE_INDEX')
        stock_fund_index_daily_k = stock_fund_index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME',  values='CLOSE_INDEX')
        nav_df = fund_nav.merge(stock_fund_index_daily_k, left_index=True, right_index=True, how='left').merge(index_daily_k, left_index=True, right_index=True, how='left')
        nav_df = nav_df.replace(0.0, np.nan).sort_index().fillna(method='ffill')

        achievement = pd.DataFrame(index=['业绩统计', '年化收益率', '年化波动率', 'Sharpe比率', '最大回撤', 'Calmar比率', '投资胜率', '平均损益比'], columns=range(len(list(nav_df.columns))))
        for idx, col in enumerate(list(nav_df.columns)):
            col_nav_df = nav_df[col].dropna()
            col_ret_df = col_nav_df.pct_change().dropna()
            achievement.loc['业绩统计', idx] = col
            achievement.loc['年化收益率', idx] = (col_nav_df.iloc[-1] / col_nav_df.iloc[0]) ** (250.0 / len(col_nav_df)) - 1.0
            achievement.loc['年化波动率', idx] = np.std(col_ret_df, ddof=1) * np.sqrt(250.0)
            achievement.loc['Sharpe比率', idx] = (achievement.loc['年化收益率', idx] - 0.03) / achievement.loc['年化波动率', idx]
            achievement.loc['最大回撤', idx] = -1.0 * max([(min(col_nav_df.iloc[i:]) / col_nav_df.iloc[i] - 1.0) * (-1.0) for i in range(len(col_nav_df))])
            achievement.loc['Calmar比率', idx] = achievement.loc['年化收益率', idx] / achievement.loc['最大回撤', idx]
            achievement.loc['投资胜率', idx] = len(col_ret_df[col_ret_df >= 0]) / float(len(col_ret_df))
            achievement.loc['平均损益比', idx] = col_ret_df[col_ret_df >= 0].mean() / col_ret_df[col_ret_df < 0].mean() * (-1.0)

            achievement.loc['年化收益率', idx] = '{0}%'.format(round(achievement.loc['年化收益率', idx] * 100.0, 2))
            achievement.loc['年化波动率', idx] = '{0}%'.format(round(achievement.loc['年化波动率', idx] * 100.0, 2))
            achievement.loc['Sharpe比率', idx] = '{0}'.format(round(achievement.loc['Sharpe比率', idx], 2))
            achievement.loc['最大回撤', idx] = '{0}%'.format(round(achievement.loc['最大回撤', idx] * 100.0, 2))
            achievement.loc['Calmar比率', idx] = '{0}'.format(round(achievement.loc['Calmar比率', idx], 2))
            achievement.loc['投资胜率', idx] = '{0}%'.format(round(achievement.loc['投资胜率', idx] * 100.0, 2))
            achievement.loc['平均损益比', idx] = '{0}'.format(round(achievement.loc['平均损益比', idx], 2))

        achievement = achievement.reset_index()
        return achievement

    def get_manager_brison(self):
        representative_fund = self.representative_fund.copy(deep=True)
        attr_list = ['大类资产配置收益', '行业配置收益', '行业选股收益', '交易收益']
        fund_brison_list = []
        for idx, row in representative_fund.iterrows():
            fund_code = row['FUND_CODE']
            begin_date = row['BEGIN_DATE']
            end_date = row['END_DATE']

            fund_brison = HBDB().read_fund_brinson_attribution_given_code(fund_code)
            fund_brison = fund_brison.rename(columns={'jjdm': 'FUND_CODE', 'tjrq': 'REPORT_DATE', 'asset_allo': '大类资产配置收益', 'sector_allo': '行业配置收益', 'equity_selection': '行业选股收益', 'trading': '交易收益'})
            fund_brison['REPORT_DATE'] = fund_brison['REPORT_DATE'].astype(str)
            fund_brison['REPORT_DATE'] = fund_brison['REPORT_DATE'].apply(lambda x: str(x)[:6] + '30' if str(x)[4:6] == '06' else str(x)[:6] + '31')
            fund_brison = fund_brison.sort_values('REPORT_DATE')
            fund_brison = fund_brison[(fund_brison['REPORT_DATE'] > begin_date) & (fund_brison['REPORT_DATE'] <= end_date)]
            fund_brison['总超额收益'] = fund_brison['portfolio_return'] - fund_brison['benchmark_return']
            fund_brison = fund_brison[['REPORT_DATE', '总超额收益'] + attr_list]
            fund_brison_list.append(fund_brison)

        fund_brison = pd.concat(fund_brison_list)
        fund_brison_disp = fund_brison.copy(deep=True)
        for i in range(1, len(attr_list)):
            fund_brison_disp[attr_list[i]] = fund_brison_disp[attr_list[i]] + fund_brison_disp[attr_list[i - 1]]
        change_dates = [date for date in self.change_dates if date >= min(fund_brison_disp['REPORT_DATE']) and date <= max(fund_brison_disp['REPORT_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_brison_disp['REPORT_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(12, 6))
        color_list = [bar_color_list[0], bar_color_list[1], bar_color_list[7], bar_color_list[14]]
        for i in range(len(attr_list) - 1, -1, -1):
            sns.barplot(ax=ax, x='REPORT_DATE', y=attr_list[i], data=fund_brison_disp, label=attr_list[i], color=color_list[i])
        sns.lineplot(ax=ax, x='REPORT_DATE', y='总超额收益', data=fund_brison_disp.reset_index().drop('index', axis=1), label='总超额收益', color='#F04950')
        ax.vlines(x=change_dates_loc, ymin=min([min(fund_brison_disp[col]) for col in attr_list + ['总超额收益']]), ymax=max([max(fund_brison_disp[col]) for col in attr_list + ['总超额收益']]), linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.legend(loc=8, bbox_to_anchor=(0.5, -0.2), ncol=5)
        plt.xlabel('')
        plt.xticks(rotation=45)
        plt.ylabel('')
        plt.title('Brison归因', fontdict={'font': 'kaiti', 'weight': 'bold', 'size': 16})
        plt.tight_layout()
        sns.despine(left=False, bottom=False, top=True, right=True)
        plt.savefig('{0}{1}_brison.png'.format(self.file_path, self.manager_name))

        fund_brison_index = fund_brison[attr_list].mean()[fund_brison[attr_list].mean() == fund_brison[attr_list].mean().max()].index.tolist()
        fund_brison_mark = fund_brison_index[0] if len(fund_brison_index) == 1 else '、'.join(fund_brison_index)

        fund_brison_text = '{0}管理期间业绩的Brison归因结果如下，其中{1}贡献最大。'.format(self.manager_name, fund_brison_mark)
        return fund_brison_text

    def get_quantitative_report(self):
        """
        生成公募基金量化分析报告
        """
        document = Document()
        document.styles['Normal'].font.size = Pt(8)
        head = document.add_heading(text='', level=0)
        run_head = head.add_run('【{0}】基金经理定量分析报告'.format(self.manager_name))
        run_head.font.color.rgb = RGBColor(201, 70, 73)

        risk, statement1, statement2 = self.report_info()
        document.add_paragraph(text='风险提示').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=risk)
        document.add_paragraph(text='重要声明').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=statement1)
        document.add_paragraph(text=statement2)
        document.add_page_break()

        head_1 = document.add_heading(text='', level=1)
        run_head_1 = head_1.add_run('一. 基金经理概况')
        run_head_1.font.color.rgb = RGBColor(201, 70, 73)
        head_1_1 = document.add_heading(text='', level=2)
        run_head_1_1 = head_1_1.add_run('1. 简介')
        run_head_1_1.font.color.rgb = RGBColor(201, 70, 73)
        manager_info_text, manager_fund = self.get_manager_info()
        document.add_paragraph(text=manager_info_text)
        rowc, colc = manager_fund.shape[0], manager_fund.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(manager_fund.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_1_1 = document.add_heading(text='', level=2)
        run_head_1_1 = head_1_1.add_run('2. 管理产品规模')
        run_head_1_1.font.color.rgb = RGBColor(201, 70, 73)
        manager_scale_text = self.get_manager_scale()
        document.add_paragraph(text=manager_scale_text)
        document.add_picture('{0}{1}_scale.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2 = document.add_heading(text='', level=1)
        run_head_2 = head_2.add_run('二. 基金经理定量分析')
        run_head_2.font.color.rgb = RGBColor(201, 70, 73)
        repr_fund_text, repr_fund = self.get_repr_fund()
        document.add_paragraph(text=repr_fund_text)
        rowc, colc = repr_fund.shape[0], repr_fund.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(repr_fund.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_1 = document.add_heading(text='', level=2)
        run_head_2_1 = head_2_1.add_run('1. 仓位分析')
        run_head_2_1.font.color.rgb = RGBColor(201, 70, 73)
        manager_position_text = self.get_manager_position()
        document.add_paragraph(text=manager_position_text)
        document.add_picture('{0}{1}_position.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_2 = document.add_heading(text='', level=2)
        run_head_2_2 = head_2_2.add_run('2. 规模分析')
        run_head_2_2.font.color.rgb = RGBColor(201, 70, 73)
        manager_size_text = self.get_manager_size()
        document.add_paragraph(text=manager_size_text)
        document.add_picture('{0}{1}_size_sc.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_size.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_3 = document.add_heading(text='', level=2)
        run_head_2_3 = head_2_3.add_run('3. 风格分析')
        run_head_2_3.font.color.rgb = RGBColor(201, 70, 73)
        manager_style_text = self.get_manager_style()
        document.add_paragraph(text=manager_style_text)
        document.add_picture('{0}{1}_style_sc.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_style.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_4 = document.add_heading(text='', level=2)
        run_head_2_4 = head_2_4.add_run('4. 主题分析')
        run_head_2_4.font.color.rgb = RGBColor(201, 70, 73)
        manager_theme_text = self.get_manager_theme()
        document.add_paragraph(text=manager_theme_text)
        document.add_picture('{0}{1}_theme_sc.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_theme.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_5 = document.add_heading(text='', level=2)
        run_head_2_5 = head_2_5.add_run('5. 行业分析')
        run_head_2_5.font.color.rgb = RGBColor(201, 70, 73)
        manager_industry_text = self.get_manager_industry()
        document.add_paragraph(text=manager_industry_text)
        document.add_picture('{0}{1}_industry_sc.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_industry.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_6 = document.add_heading(text='', level=2)
        run_head_2_6 = head_2_6.add_run('6. 持股分析')
        run_head_2_6.font.color.rgb = RGBColor(201, 70, 73)
        manager_turnover_text = self.get_manager_turnover()
        document.add_paragraph(text=manager_turnover_text)
        document.add_picture('{0}{1}_turnover.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        manager_cr_text = self.get_manager_cr()
        document.add_paragraph(text=manager_cr_text)
        document.add_picture('{0}{1}_cr5.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_cr10.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        manager_huddle_text = self.get_manager_huddle()
        document.add_paragraph(text=manager_huddle_text)
        document.add_picture('{0}{1}_huddle.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_7 = document.add_heading(text='', level=2)
        run_head_2_7 = head_2_7.add_run('7. 选股方法')
        run_head_2_7.font.color.rgb = RGBColor(201, 70, 73)
        manager_industry_stock_text = self.get_manager_industry_stock()
        document.add_paragraph(text=manager_industry_stock_text)
        document.add_picture('{0}{1}_industry_stock.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        manager_left_text, manager_left = self.get_manager_left()
        document.add_paragraph(text=manager_left_text)
        rowc, colc = manager_left.shape[0], manager_left.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(manager_left.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        manager_new_text, manager_new = self.get_manager_new()
        document.add_paragraph(text=manager_new_text)
        rowc, colc = manager_new.shape[0], manager_new.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(manager_new.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_2_8 = document.add_heading(text='', level=2)
        run_head_2_8 = head_2_8.add_run('8. 机构关注度')
        run_head_2_8.font.color.rgb = RGBColor(201, 70, 73)
        manager_holder_text = self.get_manager_holder()
        document.add_paragraph(text=manager_holder_text)
        document.add_picture('{0}{1}_holder.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_3 = document.add_heading(text='', level=1)
        run_head_3 = head_3.add_run('三. 基金经理业绩')
        run_head_3.font.color.rgb = RGBColor(201, 70, 73)
        head_3_1 = document.add_heading(text='', level=2)
        run_head_3_1 = head_3_1.add_run('1. 业绩走势')
        run_head_3_1.font.color.rgb = RGBColor(201, 70, 73)
        manager_performance = self.get_manager_performance()
        document.add_paragraph(text='该基金经理代表产品业绩收益和回撤情况统计如下：', style=None)
        document.add_picture('{0}{1}_return.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_drawdown.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text='相关统计指标如下：')
        rowc, colc = manager_performance.shape[0], manager_performance.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(manager_performance.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        manager_performance_all = self.get_manager_performance_all()
        document.add_paragraph(text='将该基金经理所管理的偏股混合型、灵活配置型、普通股票型基金业绩按照规模进行加权，构造该基金经理的整体业绩表现，则收益和回撤情况如下：', style=None)
        document.add_picture('{0}{1}_return_all.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_drawdown_all.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text='相关统计指标如下：')
        rowc, colc = manager_performance_all.shape[0], manager_performance_all.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(manager_performance_all.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_3_2 = document.add_heading(text='', level=2)
        run_head_3_2 = head_3_2.add_run('2. 业绩归因')
        run_head_3_2.font.color.rgb = RGBColor(201, 70, 73)
        manager_brison_text = self.get_manager_brison()
        document.add_paragraph(text=manager_brison_text, style=None)
        document.add_picture('{0}{1}_brison.png'.format(self.file_path, self.manager_name), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_4 = document.add_heading(text='', level=1)
        run_head_4 = head_4.add_run('四. 总结')
        run_head_4.font.color.rgb = RGBColor(201, 70, 73)
        fund_summary = '1. 管理产品规模：{0}'.format(manager_scale_text.split('变化情况如下，')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '2. 仓位：{0}{1}'.format(self.manager_name, manager_position_text.split('对各类资产的配置情况如下，其中')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '3. 规模：{0}属于{1}'.format(self.manager_name, manager_size_text.split('属于')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '4. 风格：{0}属于{1}'.format(self.manager_name, manager_style_text.split('属于')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '5. 主题：{0}属于{1}'.format(self.manager_name, manager_theme_text.split('属于')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '6. 行业：{0}属于{1}'.format(self.manager_name, manager_industry_text.split('属于')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '7. 持股：{0}{1}{2}'.format(manager_turnover_text.split('换手率方面，')[1].replace('。', '，'), manager_cr_text.split('管理期间')[1].replace('。', '，'), manager_huddle_text.split('管理期间')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '8. 选股方法：{0}选股{1}{2}，{3}。'.format(self.manager_name, manager_industry_stock_text.split('，')[3].replace('。', '，'), manager_left.iloc[1, 4], manager_new.iloc[1, 4]) if len(self.representative_fund) == 1 else '8. 选股方法：{0}选股{1}代表基金{2}的左侧特征分别为{3}，新股/次新股偏好分别为{4}。'.format(self.manager_name, manager_industry_stock_text.split('，')[3].replace('。', '，'), '、'.join(manager_left.iloc[1:, 0].tolist()), '、'.join(manager_left.iloc[1:, 4].tolist()).replace('特征', ''), '、'.join(manager_new.iloc[1:, 4].tolist()))
        document.add_paragraph(text=fund_summary)
        fund_summary = '9. 机构关注度：{0}'.format(manager_holder_text.split('持仓占比看，')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '10. 业绩归因：Brison业绩归因中{0}'.format(manager_brison_text.split('其中')[1])
        document.add_paragraph(text=fund_summary)

        for paragraph in document.paragraphs:
            if paragraph.style.name == 'Normal' and paragraph.text != '':
                paragraph.paragraph_format.first_line_indent = paragraph.style.font.size * 2

        document.save('{0}【{1}】基金经理定量分析报告.docx'.format(self.file_path, self.manager_name))

if __name__ == "__main__":
    # 确定基金经理代码
    fund_manager = HBDB().read_fund_manager_given_code('519002')
    fund_manager = fund_manager.rename(columns={'jjdm': 'FUND_CODE', 'rydm': 'MANAGER_CODE', 'ryxm': 'MANAGER_NAME', 'ryzw': 'POSITION', 'ryzt': 'STATUS', 'bdrq': 'CHANGE_DATE', 'rzrq': 'BEGIN_DATE', 'lrrq': 'END_DATE', 'lryy': 'REASON'})

    file_path = 'D:/Git/hbshare/hbshare/fe/xwq/data/manager_report/'
    manager_code = '30569470'
    fund_code = None
    QuantitativeReport(file_path, manager_code).get_quantitative_report()