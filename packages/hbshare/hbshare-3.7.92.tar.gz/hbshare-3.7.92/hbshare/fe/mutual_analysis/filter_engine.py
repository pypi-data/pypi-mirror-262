import datetime
import pandas as pd
import numpy as np
from hbshare.fe.XZ import db_engine
from hbshare.fe.XZ import functionality
from hbshare.fe.mutual_analysis import pool_analysis as pla
from sklearn import  preprocessing as pp




util=functionality.Untils()
hbdb=db_engine.HBDB()
localdb=db_engine.PrvFunDB().engine

jjdm_list = util.get_mutual_stock_funds(datetime.datetime.today().strftime('%Y%m%d'))
jjdm_con = util.list_sql_condition(jjdm_list)
sql = "select jjjc,jjdm from st_fund.t_st_gm_jjxx where jjdm in ({0})".format(jjdm_con)
jjdm_name_map = hbdb.db2df(sql, db='funduser')

class Equilibrium:

    @staticmethod
    def ind_equ():

        latest_asofdate=\
            pd.read_sql("select max(asofdate) as asofdate from hbs_industry_property_new ",con=localdb)['asofdate'][0]

        sql="SELECT jjdm,cen_ind_1,asofdate from hbs_industry_property_new where asofdate='{0}' order by cen_ind_1 ".format(latest_asofdate)
        ind_equ=pd.read_sql(sql,con=localdb)

        return ind_equ

    @staticmethod
    def style_equ():

        latest_asofdate = \
        pd.read_sql("select max(asofdate) as asofdate from hbs_style_property ", con=localdb)['asofdate'][0]

        sql="""SELECT jjdm,cen_lv_rank,asofdate 
        from hbs_style_property where asofdate='{0}' order by cen_lv_rank 
        """.format(latest_asofdate)
        style_equ = pd.read_sql(sql, con=localdb)

        return  style_equ

    @staticmethod
    def size_equ():

        latest_asofdate = \
        pd.read_sql("select max(asofdate) as asofdate from hbs_size_property ", con=localdb)['asofdate'][0]

        sql = """SELECT jjdm,cen_lv_rank,asofdate 
        from hbs_size_property where asofdate='{0}' order by cen_lv_rank 
        """.format(latest_asofdate)

        size_equ = pd.read_sql(sql, con=localdb)

        return  size_equ

    @staticmethod
    def pepbroe_equ():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_holding_property ", con=localdb)[
                'asofdate'][0]

        sql="""
        select jjdm,PE_rank,PB_rank,ROE_rank,股息率_rank,asofdate 
        from hbs_holding_property where asofdate='{0}'
        """.format(latest_asofdate)

        pepbroe_equ=pd.read_sql(sql,con=localdb)
        pepbroe_equ[
            ['PE_rank','PB_rank','ROE_rank','股息率_rank']
        ]=abs(pepbroe_equ[['PE_rank','PB_rank','ROE_rank','股息率_rank']]-0.5)

        pepbroe_equ=pepbroe_equ.sort_values('PE_rank')

        return pepbroe_equ

    @staticmethod
    def nav_equ():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from nav_ret_bias ", con=localdb)[
                'asofdate'][0]

        sql="""
        select jjdm,mean_rank_monthly,std_rank_monthly,mean_rank_yearly,std_rank_yearly,asofdate
        from nav_ret_bias where asofdate='{0}' and mean_rank_yearly is not null  order by mean_rank_monthly
        """.format(latest_asofdate)

        nav_equ=pd.read_sql(sql,con=localdb)

        return  nav_equ

    def get_equilibrium(self,threshield,show_num=300):

        # method_1
        ind_equ = self.ind_equ()
        style_equ = self.style_equ()
        size_equ = self.size_equ()
        pepbroe_equ = self.pepbroe_equ()
        nav_equ = self.nav_equ()

        #rename columns :
        ind_equ=pd.merge(ind_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').\
            rename(columns={'cen_ind_1':'行业集中度'})

        style_equ=pd.merge(style_equ,
                         jjdm_name_map,
                         how='left', on='jjdm')\
            .rename(columns={'cen_lv_rank': '风格集中度'})

        size_equ=pd.merge(size_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').\
            rename(columns={'cen_lv_rank': '规模集中度'})

        pepbroe_equ=pd.merge(pepbroe_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').rename(columns={'PE_rank': 'pe偏离度',
                                    'PB_rank': 'pb偏离度',
                                    'ROE_rank': 'roe偏离度',
                                    '股息率_rank': '股息率偏离度'})

        nav_equ=pd.merge(nav_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').rename(columns={'mean_rank_monthly': '净值偏离度（月）',
                                'std_rank_monthly': '净值偏离度方差（月）',
                                'mean_rank_yearly': '净值偏离度（年）',
                                'std_rank_yearly': '净值偏离度方差（年）',
                                })

        # method_2
        joint_rank = pd.merge(ind_equ, style_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, size_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, pepbroe_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, nav_equ, how='inner', on=['jjdm','jjjc'])

        col_list = ['行业集中度', '风格集中度',
                    '规模集中度', 'pe偏离度',
                    '净值偏离度（月）']

        joint_rank['平均集中度'] = joint_rank[col_list].mean(axis=1)
        joint_rank = joint_rank.sort_values('平均集中度')

        # method_3
        joint_result = joint_rank[(joint_rank[col_list] <= threshield).prod(axis=1) == 1][['jjdm','jjjc','asofdate'] + col_list]

        joint_rank=joint_rank[['jjdm','jjjc','平均集中度','asofdate']+col_list]


        #format change
        ind_equ['行业集中度'] = ind_equ['行业集中度'].map("{:.2%}".format)
        style_equ['风格集中度'] = style_equ['风格集中度'].map("{:.2%}".format)
        size_equ['规模集中度'] = size_equ['规模集中度'].map("{:.2%}".format)

        for col in ['pe偏离度','pb偏离度','roe偏离度','股息率偏离度']:
            pepbroe_equ[col] = pepbroe_equ[col].map("{:.2%}".format)
        for col in ['净值偏离度（月）', '净值偏离度方差（月）', '净值偏离度（年）', '净值偏离度方差（年）']:
            nav_equ[col] = nav_equ[col].map("{:.2%}".format)
        for col in col_list:
            joint_rank[col]=joint_rank[col].map("{:.2%}".format)
            joint_result[col]=joint_result[col].map("{:.2%}".format)
        joint_rank['平均集中度']=joint_rank['平均集中度'].map("{:.2%}".format)

        return ind_equ[0:show_num],style_equ[0:show_num],\
               size_equ[0:show_num],pepbroe_equ[0:show_num],nav_equ[0:show_num],\
               joint_rank[0:show_num],joint_result[0:show_num]

class Leftside:

    @staticmethod
    def stock_left():

        latest_asofdate=\
            pd.read_sql("select max(asofdate) as asofdate from hbs_stock_trading_property ",
                        con=localdb)['asofdate'][0]

        sql="""
        SELECT jjdm,`左侧概率（出持仓前,半年线）_rank`,`左侧概率（出持仓前,年线）_rank`,asofdate 
        from hbs_stock_trading_property where asofdate='{0}'  """\
            .format(latest_asofdate)
        stock_left=pd.read_sql(sql,con=localdb)

        stock_left['stock_left_rank']=stock_left[['左侧概率（出持仓前,半年线）_rank','左侧概率（出持仓前,年线）_rank']]\
            .max(axis=1)

        stock_left=stock_left.sort_values('stock_left_rank',ascending=False)

        return stock_left

    @staticmethod
    def ind_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_industry_shift_property_new "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_industry_shift_property_new where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        ind_left=pd.read_sql(sql,con=localdb)

        ind_left=ind_left.groupby('jjdm').max('Total_rank')
        ind_left.reset_index(inplace=True)
        ind_left['asofdate']=latest_asofdate

        ind_left.sort_values('Total_rank',ascending=False,inplace=True)

        ind_left.rename(columns={'Total_rank':'ind_rank'},inplace=True)

        return ind_left

    @staticmethod
    def value_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_shift_property_value "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_shift_property_value where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        value_left=pd.read_sql(sql,con=localdb)

        value_left=value_left.groupby('jjdm').max('Total_rank')
        value_left.reset_index(inplace=True)
        value_left['asofdate']=latest_asofdate

        value_left.sort_values('Total_rank',ascending=False,inplace=True)
        value_left.rename(columns={'Total_rank': 'value_rank'}, inplace=True)

        return value_left

    @staticmethod
    def size_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_shift_property_size "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_shift_property_size where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        size_left=pd.read_sql(sql,con=localdb)

        size_left=size_left.groupby('jjdm').max('Total_rank')
        size_left.reset_index(inplace=True)
        size_left['asofdate']=latest_asofdate

        size_left.sort_values('Total_rank',ascending=False,inplace=True)
        size_left.rename(columns={'Total_rank': 'size_rank'}, inplace=True)


        return size_left

    def get_left(self,threshield,show_num=300):

        # method_1
        stock_left = self.stock_left()
        ind_left = self.ind_left()
        value_left = self.value_left()
        size_left = self.size_left()

        #rename columns
        stock_left=stock_left[['jjdm','stock_left_rank','asofdate']]

        stock_left=pd.merge(stock_left,jjdm_name_map,how='left',on='jjdm')\
            .rename(columns={'stock_left_rank':'个股交易左侧率'})
        ind_left = pd.merge(ind_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'ind_rank':'行业切换左侧率'})
        value_left = pd.merge(value_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'value_rank':'风格切换左侧率'})
        size_left = pd.merge(size_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'size_rank':'规模切换左侧率'})


        # method_2
        joint_rank = pd.merge(stock_left, ind_left, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, value_left, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, size_left, how='inner', on=['jjdm','jjjc'])
        col_list = ['个股交易左侧率', '行业切换左侧率', '风格切换左侧率', '规模切换左侧率']
        joint_rank['平均左侧率'] = joint_rank[col_list].mean(axis=1)

        # method_3
        joint_restult = joint_rank[(joint_rank[col_list] >= (1 - threshield)).prod(axis=1) == 1][['jjdm','jjjc'] + col_list]

        joint_rank = joint_rank.sort_values('平均左侧率', ascending=False)[['jjdm','jjjc', '平均左侧率', 'asofdate_x']+col_list]


        #re format
        stock_left['个股交易左侧率']=stock_left['个股交易左侧率'].map("{:.2%}".format)
        ind_left['行业切换左侧率'] = ind_left['行业切换左侧率'].map("{:.2%}".format)
        value_left['风格切换左侧率'] = value_left['风格切换左侧率'].map("{:.2%}".format)
        size_left['规模切换左侧率'] = size_left['规模切换左侧率'].map("{:.2%}".format)
        for col in col_list:
            joint_rank[col] = joint_rank[col].map("{:.2%}".format)
            joint_restult[col] = joint_restult[col].map("{:.2%}".format)

        joint_rank['平均左侧率']=joint_rank['平均左侧率'].map("{:.2%}".format)

        return stock_left[0:show_num],ind_left[0:show_num],value_left[0:show_num],size_left[0:show_num],joint_rank[0:show_num],joint_restult[0:show_num]

class Size:
    @staticmethod
    def size_property(fre):
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from hbs_size_property".format(fre), con=localdb)['asofdate'][0]
        sql = "select jjdm, shift_lv, cen_lv, shift_lv_rank, cen_lv_rank,大盘, 中盘, 小盘 ,大盘_rank, 中盘_rank, 小盘_rank, asofdate from hbs_size_property where asofdate='{0}'".format( latest_asofdate)
        size_property = pd.read_sql(sql, con=localdb)
        return size_property

    def get_size(self, fre, show_num=200, shift_ratio_threshold=0.5, centralization_threshold=0.5):
        size_property = self.size_property(fre)
        size_property.columns = ['jjdm', '换手率', '集中度', '换手率排名', '集中度排名','大盘(绝对值）', '中盘(绝对值）', '小盘(绝对值）' ,'大盘', '中盘', '小盘', 'asofdate']
        size_property = size_property.merge(jjdm_name_map, on=['jjdm'], how='left')

        size = size_property[(size_property['换手率排名'] < shift_ratio_threshold) & (size_property['集中度排名'] > centralization_threshold)]
        big_size = size[(size['大盘'] > size['中盘']) & (size['大盘'] > size['小盘'])].sort_values('大盘', ascending=False)
        medium_size = size[(size['中盘'] > size['大盘']) & (size['中盘'] > size['小盘'])].sort_values('中盘', ascending=False)
        small_size = size[(size['小盘'] > size['大盘']) & (size['小盘'] > size['中盘'])].sort_values('小盘', ascending=False)

        big_size = big_size[['jjdm', 'jjjc', 'asofdate', '大盘']]
        medium_size = medium_size[['jjdm', 'jjjc', 'asofdate', '中盘']]
        small_size = small_size[['jjdm', 'jjjc', 'asofdate', '小盘']]
        big_size['大盘'] = big_size['大盘'].map("{:.2%}".format)
        medium_size['中盘'] = medium_size['中盘'].map("{:.2%}".format)
        small_size['小盘'] = small_size['小盘'].map("{:.2%}".format)
        return big_size[0: show_num], medium_size[0: show_num], small_size[0: show_num]

class Value:
    @staticmethod
    def value_property():
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from jjpic_value_p_hbs ", con=localdb)['asofdate'][0]
        sql = "SELECT jjdm,`集中度排名(持仓)`,`换手率排名(持仓)`,`成长绝对暴露(持仓)`,`价值绝对暴露(持仓)`,asofdate from jjpic_value_p_hbs where  asofdate='{0}'".format(latest_asofdate)
        value_property = pd.read_sql(sql, con=localdb)
        return value_property

    @staticmethod
    def holding_property():
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from hbs_holding_property", con=localdb)['asofdate'][0]
        sql = "select jjdm, PE_rank, PB_rank, PE_REL_rank, PB_REL_rank, ROE_rank, 股息率_rank, asofdate from hbs_holding_property where asofdate='{0}'".format(latest_asofdate)
        holding_property = pd.read_sql(sql, con=localdb)
        return holding_property

    def get_value(self, show_num=200, shift_ratio_threshold=0.5, centralization_threshold=0.5):

        value_property = self.value_property()
        value_property.columns = ['jjdm', '集中度排名','换手率排名','成长', '价值','asofdate']

        #filter the industry and theme fund

        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from jjpic_theme_p ", con=localdb)['asofdate'][0]
        sql = "SELECT * from jjpic_theme_p where  asofdate='{0}'".format(latest_asofdate)
        jj_theme = pd.read_sql(sql, con=localdb)
        jj_theme['max_theme_w']=jj_theme[[ '大金融', '消费', 'TMT', '周期','制造']].max(axis=1)
        jj_theme=jj_theme[jj_theme['max_theme_w']<=0.7][['jjdm']]

        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from jjpic_industry_detail_1 ", con=localdb)['asofdate'][0]
        sql = "SELECT jjdm,max(`占持仓比例(时序均值)`) as max_ind_w from jjpic_industry_detail_1 where asofdate='{0}' GROUP BY jjdm ".format(latest_asofdate)
        jj_max_ind_w = pd.read_sql(sql, con=localdb)
        jj_max_ind_w=jj_max_ind_w[jj_max_ind_w['max_ind_w']<=0.5][['jjdm']]

        value_property=pd.merge(value_property,jj_theme,how='inner',on='jjdm')
        value_property = pd.merge(value_property, jj_max_ind_w, how='inner', on='jjdm')

        value_property = value_property.merge(jjdm_name_map, on=['jjdm'], how='left')

        #value = value_property[(value_property['换手率排名'] < shift_ratio_threshold) & (value_property['集中度排名'] > centralization_threshold)]
        growth = value_property[value_property['成长'] > value_property['价值']].sort_values('成长', ascending=False)
        value = value_property[value_property['价值'] > value_property['成长']].sort_values('价值', ascending=False)
        growth=growth[(growth['换手率排名']<growth['换手率排名'].quantile(shift_ratio_threshold))]
        value = value[(value['换手率排名'] < value['换手率排名'].quantile(shift_ratio_threshold))]


        growth = growth[['jjdm', 'jjjc', 'asofdate', '成长']]
        value = value[['jjdm', 'jjjc', 'asofdate', '价值']]
        growth['成长'] = growth['成长'].map("{:.2%}".format)
        value['价值'] = value['价值'].map("{:.2%}".format)

        holding_property = self.holding_property()
        holding_property.columns = ['jjdm', 'PE排名', 'PB排名', 'PE相对行业均值排名', 'PB相对行业均值排名', 'ROE排名', '股息率排名', 'asofdate']
        value_holding_property = value.merge(holding_property.drop('asofdate', axis=1), on=['jjdm'], how='left')
        absolute_pe_value = value_holding_property.sort_values('PE排名')
        absolute_pb_value = value_holding_property.sort_values('PB排名')
        relative_pe_value = value_holding_property.sort_values('PE相对行业均值排名')
        relative_pb_value = value_holding_property.sort_values('PB相对行业均值排名')
        dividend_value = value_holding_property.sort_values('股息率排名', ascending=False)
        reverse_value = value_holding_property.sort_values('ROE排名')
        high_quality_value = value_holding_property.sort_values('ROE排名', ascending=False)

        absolute_pe_value = absolute_pe_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PE排名']]
        absolute_pb_value = absolute_pb_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PB排名']]
        relative_pe_value = relative_pe_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PE相对行业均值排名']]
        relative_pb_value = relative_pb_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PB相对行业均值排名']]
        dividend_value = dividend_value[['jjdm', 'jjjc', 'asofdate', '价值', '股息率排名']]
        reverse_value = reverse_value[['jjdm', 'jjjc', 'asofdate', '价值', 'ROE排名']]
        high_quality_value = high_quality_value[['jjdm', 'jjjc', 'asofdate', '价值', 'ROE排名']]

        # absolute_pe_value['PE排名'] = absolute_pe_value['PE排名'].map("{:.2%}".format)
        # absolute_pb_value['PB排名'] = absolute_pb_value['PB排名'].map("{:.2%}".format)
        # relative_pe_value['PE相对行业均值排名'] = relative_pe_value['PE相对行业均值排名'].map("{:.2%}".format)
        # relative_pb_value['PB相对行业均值排名'] = relative_pb_value['PB相对行业均值排名'].map("{:.2%}".format)
        # dividend_value['股息率排名'] = dividend_value['股息率排名'].map("{:.2%}".format)
        # reverse_value['ROE排名'] = reverse_value['ROE排名'].map("{:.2%}".format)
        # high_quality_value['ROE排名'] = high_quality_value['ROE排名'].map("{:.2%}".format)
        return growth[0: show_num], value[0: show_num]
            # , absolute_pe_value[0: show_num], absolute_pb_value[0: show_num], \
            #    relative_pe_value[0: show_num], relative_pb_value[0: show_num], dividend_value[0: show_num], \
            #    reverse_value[0: show_num], high_quality_value[0: show_num]

class Industry_rotating:

    def get_industry_rotating(self,shift_count_rank_thresheld=0.5
                              ,shift_ret_thresheld=0.5,shift_winpro_thresheld=0.5,left_trade_ratio=0.5,pool_size=30):

        max_asofdate=pd.read_sql("select max(asofdate) as max_date from jjpic_industry_sp",con=localdb)['max_date'][0]

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='切换次数' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_count_rank_thresheld,max_asofdate)
        jjdm_list1=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='胜率（直到下次切换）' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_winpro_thresheld,max_asofdate)
        jjdm_list2=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='胜率（下季度）' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_winpro_thresheld,max_asofdate)
        jjdm_list22=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        max_asofdate2 = pd.read_sql("select max(asofdate) as max_date from jjpic_industry_p", con=localdb)['max_date'][0]
        sql="SELECT jjdm from jjpic_industry_p where `一级行业类型` in ('轮动','博弈') and asofdate={0}"\
            .format(max_asofdate2)
        jjdm_list3=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='持有平均收益' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_ret_thresheld,max_asofdate)
        jjdm_list4=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='下季平均收益' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_ret_thresheld,max_asofdate)
        jjdm_list44=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='左侧比率' and Total_rank>={0} and asofdate='{1}'"\
            .format(left_trade_ratio,max_asofdate)
        jjdm_list5=pd.read_sql(sql,con=localdb)['jjdm'].tolist()


        jjdm_list=set(jjdm_list1).intersection(set(jjdm_list22))
        jjdm_list=jjdm_list.intersection(set(jjdm_list3))
        jjdm_list = jjdm_list.intersection(set(jjdm_list44))
        jjdm_list=list(jjdm_list.intersection(set(jjdm_list5)))

        sql="select jjdm,avg(Total_rank) as new_factor from jjpic_industry_sp where (`项目名`='下季平均收益' or `项目名`='胜率（下季度）') and jjdm in ({0}) and asofdate='{1}' group by jjdm order by new_factor DESC "\
            .format(util.list_sql_condition(jjdm_list),max_asofdate)
        pool_funds=pd.read_sql(sql,con=localdb).iloc[0:pool_size]['jjdm'].tolist()

        # sql = "select jjdm,info_ratio from factor_infor_ratio where jjdm in ({0}) and date>='{1}' and date <='{2}' and zblb='{3}' and info_ratio!=99999 order by info_ratio DESC" \
        #     .format(util.list_sql_condition(jjdm_list)
        #             , max_asofdate[0:6] + '01', max_asofdate[0:6] + '31'
        #             , '2101')
        # pool_funds = pd.read_sql(sql, con=localdb).iloc[0:pool_size]['jjdm'].tolist()

        sql = "select tjyf,avg(hb1y) from st_fund.t_st_gm_yhb where jjdm in ({0}) and hb1y!=99999 and tjyf>'202112' and tjyf<='202208' group by tjyf " \
            .format(util.list_sql_condition(pool_funds))
        navdf = hbdb.db2df(sql, db='funduser').set_index('tjyf')
        navdf=navdf/100+1
        print((navdf.cumprod()['avg(hb1y)'].iloc[-1]-1)*100)


        return  pool_funds

class Positive_equilibrium:

    def get_nav_equilibrium_pool(self,asofdate,poo_size=30):

        jjdm_list_0=util.get_bmk_funds_list(asofdate)

        sql="""SELECT jjdm,mean_rank_yearly,mean_rank_monthly from nav_ret_bias where mean_yearly is not null 
        and mean_rank_yearly<=0.1 and mean_rank_monthly<=0.1 and asofdate='{0}' and jjdm in ({1}) ORDER BY mean_yearly
        """.format(asofdate[0:6],util.list_sql_condition(jjdm_list_0))
        jjdm_list=pd.read_sql(sql,con=localdb)

        jjdm_list=jjdm_list.iloc[0:poo_size]['jjdm'].tolist()

        return jjdm_list

class Report:


    @staticmethod
    def get_silly_summary(doc,equ_list,equ_name_list,fold,
                          paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num):

        from PIL import Image
        from docx.shared import Inches, Pt


        def style_analysis(df,colname):
            style_style=''
            for style in df.index:
                style_style+=str(style)+"占比"+str(round(df.loc[style][colname],2))+"%,"

            return  style_style

        pic_name_list = ['池列表.png',
                         '行业风格分布.png', '1级行业持仓分布_占持仓比例.png',
                         '1级行业持仓分布_占持仓比例排名.png',
                         '2级行业持仓分布_占持仓比例.png',
                         '2级行业持仓分布_占持仓比例排名.png',
                         '3级行业持仓分布_占持仓比例.png',
                         '3级行业持仓分布_占持仓比例排名.png',
                         '风格类型分布.png',
                         '风格专注型分布.png',
                         '成长价值持仓占比.png',
                         '成长价值持仓占比排名.png',
                         '规模风格类型分布.png',
                         '专注型风格分布.png',
                         '个股持仓特征.png',
                         '大中小盘持仓占比.png',
                         '大中小盘持仓占比排名.png',
                         '个股风格类型分布A.png',
                         '个股风格类型分布B.png',
                         '左侧类型分布.png',
                         '新股偏好分布.png',
                         ]

        for i in range(len(equ_list)):
            equ=equ_list[i]
            name=equ_name_list[i]

            doc.add_paragraph("")
            paragraphs_count += 1

            p=doc.add_paragraph("{2}.{3} {0}：对{1}排名前30的基金进行了统计画像，统计结果如下：".format(name, name,title_num,i+1))
            p.style.font.size=Pt(10)
            p.paragraph_format.first_line_indent=p.style.font.size*2
            p.paragraph_format.line_spacing = Pt(20)  # 行距
            string_paragraphs_list.append(paragraphs_count)
            paragraphs_count+=1


            industry_style, class_list, style_type_dis, style_incline_dis, \
            style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
            stock_style_a, stock_style_b, stock_left, stock_new, stock_fin = pla.pool_picturing(equ, returndata=True,save_local_file=True)

            for pic in pic_name_list:
                file_path=fold+"\\"+pic
                try:
                    doc.add_picture(file_path, width=Inches(3), height=Inches(2.8))
                    pic_paragraphs_list.append(paragraphs_count)
                    paragraphs_count += 1
                except Exception as e:
                    pic_temp = Image.open(file_path)
                    pic_temp.save(pic_temp)
                    doc.add_picture(file_path, width=Inches(4.5), height=Inches(2.5))

            ind_style=style_analysis(industry_style,'num')
            style_style=style_analysis(style_type_dis,'风格类型')
            style_weight=style_analysis(style_weight_dis,'成长价值权重分布')
            size_style=style_analysis(style_type_dis2,'规模风格类型')
            size_weight=style_analysis(style_weight_dis2,'成长价值权重分布')
            left_property=style_analysis(stock_left,'左侧类型分布')
            new_stock_property=style_analysis(stock_new,'新股偏好分布')
            stock_hld_property=style_analysis(stock_fin,'个股持仓特征')

            summary="总结来看：从行业上看，{0}池中".format(name)+ind_style+\
                    "一级行业中，{0}占比较高".format(util.list_sql_condition(class_list[0].index[0:3].tolist()))+\
                    "从风格上看，{0}池中".format(name)+style_style+"具体的持仓来看"+style_weight+"。"+\
                    "从规模上看，{0}池中".format(name) + size_style + "具体的持仓来看" + size_weight + "。"+\
                    "从左侧属性上看，{0}池中".format(name) + left_property +"。"+\
                    "从新股属性上看，{0}池中".format(name) + new_stock_property + "。" +\
                    "从个股持仓属性上看，{0}池中".format(name) + stock_hld_property

            doc.add_paragraph("")
            paragraphs_count += 1

            p=doc.add_paragraph(summary)
            p.style.font.size = Pt(10)
            p.paragraph_format.first_line_indent=p.style.font.size*2
            p.paragraph_format.line_spacing = Pt(20)  # 行距
            string_paragraphs_list.append(paragraphs_count)
            paragraphs_count+=1

        return doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list

    def filter_pool_report(self):

        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        #save png 2 doc

        equclass=Equilibrium()
        ind_equ, style_equ, size_equ,pepbroe_equ, nav_equ,\
        joint_rank,joint_result=equclass.get_equilibrium(threshield=0.3,show_num=30)
        equ_list=[joint_rank,joint_result]
        equ_name_list=['均衡平均','均衡交集']

        leftclass=Leftside()
        stock_left, ind_left, value_left,size_left,\
        joint_rank, joint_restult=leftclass.get_left(threshield=0.4,show_num=30)
        left_list=[joint_rank, joint_restult]
        left_name_list = ['左侧平均', '左侧交集']

        sizeclass = Size()
        size_list=sizeclass.get_size(fre='M', show_num=30,
                                                               shift_ratio_threshold=0.5,
                                                               centralization_threshold=0.5)
        size_name_list = ['大盘', '中盘', '小盘']

        valueclass = Value()
        value_list=valueclass.get_value(fre='M', show_num=30,
                                        shift_ratio_threshold=0.5, centralization_threshold=0.5)
        value_list=value_list[0:2]
        value_name_list = ['成长', '价值']

        doc = Document()
        fold=r"E:\GitFolder\hbshare\fe\mutual_analysis"

        pic_paragraphs_list=[]
        string_paragraphs_list=[]
        paragraphs_count=0

        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc,equ_list,equ_name_list,fold
                               ,paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=2)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, left_list, left_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=3)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, size_list, size_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=4)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, value_list, value_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=5)

        for j in pic_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for j in string_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


        doc.save(r"E:\GitFolder\hbshare\fe\mutual_analysis\筛选报告\筛选池画像报告.docx")


        print('')

class Advance_man:

    @staticmethod
    def advance_man_from3ind(asofdate,asofdate_2,dir,theme_name):

        advance_man_3ind_cluster=pd.read_excel(dir,sheet_name=theme_name)[['板块','三级行业']]

        target_ind=advance_man_3ind_cluster['三级行业'].tolist()


        sql="SELECT distinct(jsrq) as jsrq  from hbs_industry_class1_exp where jjdm='000001' ORDER BY jsrq DESC "
        jsrq_list=pd.read_sql(sql,con=localdb)['jsrq']
        jsrq_list=(jsrq_list[jsrq_list<=asofdate[0:6]+'31'].iloc[0:8]).tolist()

        bmk = util.get_bmk_funds_list_filter(asofdate_2, 2)


        sql="SELECT jjdm,sjxymc,sum(zjbl) as zjbl from hbs_industry_class3_exp where sjxymc in ({0}) and jsrq in ({1}) and jjdm in ({2}) group by jjdm,sjxymc "\
            .format(util.list_sql_condition(target_ind),util.list_sql_condition(jsrq_list),util.list_sql_condition(bmk))
        adv_man=pd.read_sql(sql,con=localdb)

        sql="SELECT jjdm,count(distinct(jsrq)) as count from hbs_style_exp where jsrq in ({})  GROUP BY jjdm "\
            .format(util.list_sql_condition(jsrq_list))
        jsrq_cout=pd.read_sql(sql,con=localdb)
        adv_man=pd.merge(adv_man,jsrq_cout,how='left',on='jjdm')
        adv_man['zjbl']=adv_man['zjbl']/adv_man['count']
        adv_man.drop('count',axis=1,inplace=True)

        ind3=adv_man.groupby('jjdm').sum().rename(columns={'zjbl':'目标三级行业占比'})

        adv_man=adv_man.pivot_table('zjbl','jjdm','sjxymc').fillna(0)

        adv_man.columns=[x+'_w' for x in adv_man.columns.tolist()]
        ind3=pd.merge(ind3,adv_man
                      ,how='left',on='jjdm')

        ind3=ind3[ind3['目标三级行业占比']>0].sort_values('目标三级行业占比',ascending=False)

        return  ind3

    @staticmethod
    def advance_man_wording():

        key_words=['先进制造','高端制造','智能制造','国产替代','专精特新','新能源','5G','芯片','半导体']
        jjname=pd.DataFrame(data=util.get_mutual_stock_funds('20211231'),columns=['jjdm'])

        sql="select jjdm,jjmc from st_fund.t_st_gm_jjxx "
        jjname=pd.merge(jjname,hbdb.db2df(sql,db='funduser'),how='left',on='jjdm')

        orconditon=[False]*len(jjname)
        for word in key_words:
            orconditon=(orconditon)|(jjname['jjmc'].str.contains(word))

        return  jjname.loc[orconditon]
    @staticmethod
    def advance_man_fromticker(jjdm_list):
        from hbshare.fe.mutual_analysis import holdind_based as hb

        hld=hb.read_hld_fromdb(jjdm_list,start_date='20211231',end_date='20211231')

        ticker_list=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\指数成份汇总.xlsx",sheet_name='成份汇总')
        ticker_list['Wind代码']=ticker_list['Wind代码'].astype(str).str.replace("'", "")
        hld=pd.merge(hld,ticker_list,how='left',left_on='zqdm',right_on='Wind代码')
        hld=hld[hld['Wind代码'].notnull()]
        hld=hld.groupby('jjdm').sum()
        hld['zjbl']=hld['zjbl'].rank(method='min')/len(hld)
        hld.rename(columns={'zjbl':'zjbl_ticker'})
        return hld.reset_index()
    @staticmethod
    def advance_ols(jjdm_list):
        from hbshare.fe.mutual_analysis import  nav_based as nb
        import statsmodels.api as sm

        index_range=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\先进制造指数的区间.xlsx")
        index_range['start date'] = index_range['start date'].astype(str).str.replace("-", "")
        index_range['end date'] = index_range['end date'].astype(str).str.replace("-", "")

        index_ret=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\指数成份汇总.xlsx",
                                  sheet_name='走势').drop('881001.WI',axis=1)
        index_ret[['930820.CSI','931167.CSI','866003.WI']]=\
            index_ret[['930820.CSI','931167.CSI','866003.WI']].pct_change()
        index_ret['Date']=index_ret['Date'].astype(str).str.replace("-","")
        index_ret=index_ret.iloc[2:]
        index_ret.set_index('Date',inplace=True)
        index_ret = index_ret.sort_index()
        index_ret=index_ret*100

        result=[[],[],[],[],[],[],[],[],[]]
        result2=[[],[],[],[],[],[],[],[],[]]
        result3 = [[], [], [], [], [], [], [], [], []]
        para_df=pd.DataFrame(data=jjdm_list,columns=['jjdm'])

        for jjdm in jjdm_list:
            print(jjdm)
            jj_ret = nb.get_jj_daily_ret([jjdm])
            jj_ret.index=jj_ret.index.astype(str)
            olsdf = pd.merge(index_ret,jj_ret, how='left', left_index=True, right_index=True).fillna(0)

            # laster_3_year = (datetime.datetime.strptime(str(olsdf.index[-1]),
            #                                                                               '%Y%m%d') - datetime.timedelta(
            #     days=365*3)).strftime('%Y%m%d')

            for j in range(3):
                ind=['930820.CSI','931167.CSI','866003.WI'][j]

                #for 正超额
                time_range=index_range[(index_range['index name']==ind)&(index_range['type']=='正超额')]
                tempolsdf=[]
                for i in range(len(time_range)):
                    tempolsdf.append(olsdf.loc[time_range['start date'].iloc[i]:time_range['end date'].iloc[i]])
                tempolsdf=pd.concat(tempolsdf,axis=0)
                #keep only daily ret part,remove weekly return part
                tempolsdf['index_num'] = tempolsdf.reset_index().index.to_list()
                tempolsdf['con']=1
                new_start_date=tempolsdf[tempolsdf[jjdm]==0][tempolsdf[tempolsdf[jjdm]==0]['index_num'].diff()==1]
                if(len(new_start_date)>0):
                    new_start_date=new_start_date.index[-1]
                    tempolsdf=tempolsdf.loc[new_start_date:]
                ols_result=sm.OLS(tempolsdf[ind].values,
                              tempolsdf[[jjdm,'con']].values).fit()
                para =ols_result.params.tolist()
                rsq_adj=ols_result.rsquared_adj
                result[j*3].append(para[0])
                result[(j * 3)+1].append(para[1])
                result[(j * 3)+2].append(rsq_adj)

                #for 负超额
                time_range=index_range[(index_range['index name']==ind)&(index_range['type']=='负超额')]
                tempolsdf=[]
                for i in range(len(time_range)):
                    tempolsdf.append(olsdf.loc[time_range['start date'].iloc[i]:time_range['end date'].iloc[i]])
                tempolsdf=pd.concat(tempolsdf,axis=0)
                #keep only daily ret part,remove weekly return part
                tempolsdf['index_num'] = tempolsdf.reset_index().index.to_list()
                tempolsdf['con']=1
                new_start_date=tempolsdf[tempolsdf[jjdm]==0][tempolsdf[tempolsdf[jjdm]==0]['index_num'].diff()==1]
                if(len(new_start_date)>0):
                    new_start_date=new_start_date.index[-1]
                    tempolsdf=tempolsdf.loc[new_start_date:]
                ols_result=sm.OLS(tempolsdf[ind].values,
                              tempolsdf[[jjdm,'con']].values).fit()
                para =ols_result.params.tolist()
                rsq_adj=ols_result.rsquared_adj
                result3[j*3].append(para[0])
                result3[(j * 3)+1].append(para[1])
                result3[(j * 3)+2].append(rsq_adj)

                # for 负超额+正超额
                time_range = index_range[(index_range['index name'] == ind) ]
                tempolsdf = []
                time_range=time_range.sort_values('start date')
                for i in range(len(time_range)):
                    tempolsdf.append(olsdf.loc[time_range['start date'].iloc[i]:time_range['end date'].iloc[i]])
                tempolsdf = pd.concat(tempolsdf, axis=0)
                tempolsdf.drop_duplicates(ignore_index=False, inplace=True)
                # keep only daily ret part,remove weekly return part
                tempolsdf['index_num'] = tempolsdf.reset_index().index.to_list()
                tempolsdf['con'] = 1
                new_start_date=tempolsdf[tempolsdf[jjdm]==0][tempolsdf[tempolsdf[jjdm]==0]['index_num'].diff()==1]
                if(len(new_start_date)>0):
                    new_start_date=new_start_date.index[-1]
                    tempolsdf=tempolsdf.loc[new_start_date:]

                ols_result=sm.OLS(tempolsdf[ind].values,
                              tempolsdf[[jjdm,'con']].values).fit()
                para =ols_result.params.tolist()
                rsq_adj=ols_result.rsquared_adj
                result2[j*3].append(para[0])
                result2[(j * 3)+1].append(para[1])
                result2[(j * 3)+2].append(rsq_adj)

        for j in range(3):
            ind = ['930820.CSI', '931167.CSI', '866003.WI'][j]

            para_df[ind+'_positive_beta']=result[(j*3)+0]
            para_df[ind + '_positive_alpha'] = result[(j * 3) + 1]
            para_df[ind + '_positive_rsquared_adj'] = result[(j * 3) + 2]

            para_df[ind+'_negative_beta']=result3[(j*3)+0]
            para_df[ind + '_negative_alpha'] = result3[(j * 3) + 1]
            para_df[ind + '_negative_rsquared_adj'] = result3[(j * 3) + 2]

            para_df[ind+'_positive&negative_beta']=result2[(j*3)+0]
            para_df[ind + '_positive&negative_alpha'] = result2[(j * 3) + 1]
            para_df[ind + '_positive&negative_rsquared_adj'] = result2[(j * 3) + 2]


        return  para_df
    @staticmethod
    def advance_contribution(asofdate,dir,theme_name):


        #read ticker contribution from local
        # sql="SELECT * from hbs_ticker_contribution where asofdate='{}'".format(asofdate)
        # ticker_contribution=pd.read_sql(sql,con=localdb)
        # #only take top 20 remove last 5
        # ticker_contribution['rank']=ticker_contribution.groupby('jjdm').rank(ascending=False)
        # ticker_contribution=ticker_contribution[ticker_contribution['rank']<=20].drop('rank',axis=1)


        # total_contribution = ticker_contribution.groupby('jjdm').sum()['contribution'].to_frame('total_con')
        # ticker_contribution=pd.merge(ticker_contribution,total_contribution,how='left',on='jjdm')
        # ticker_contribution['contribution_ratio']=ticker_contribution['contribution']/ticker_contribution['total_con']
        # ticker_contribution=ticker_contribution.groupby('jjdm').sum()['contribution_ratio'].to_frame('ticker_contribution_ratio')



        advance_man_3ind_cluster=pd.read_excel(dir,sheet_name=theme_name)[['板块','三级行业']]

        #read 3rd industry contribution from local
        sql="SELECT jjdm,industry_name,contribution from hbs_industry_contribution where industry_lv='sjxymc' and   asofdate='{0}' and industry_name in ({1}) "\
            .format(asofdate,util.list_sql_condition(advance_man_3ind_cluster['三级行业'].unique().tolist()))
        industry_contribution=pd.read_sql(sql,con=localdb)

        industry_total_con=industry_contribution.groupby('jjdm').sum().rename(columns={'contribution':'total_con'})
        industry_contribution=industry_contribution.pivot_table('contribution','jjdm','industry_name').fillna(0)
        industry_contribution.columns=[x+'_con' for x in industry_contribution.columns.tolist()]
        industry_contribution=pd.merge(industry_contribution
                                       ,industry_total_con,how='left',on='jjdm')

        # industry_contribution['contribution_ratio']=industry_contribution['contribution']/industry_contribution['total_con']
        # output_df=ticker_contribution.copy()
        #
        # for sector in advance_man_3ind_cluster['板块'].unique():
        #
        #     target_ind = advance_man_3ind_cluster[advance_man_3ind_cluster['板块']==sector]['三级行业'].tolist()
        #     target_ind=pd.DataFrame(data=target_ind,columns=['industry_name'])
        #     target_ind['flag']=1
        #     industry_contribution_temp=pd.merge(industry_contribution,target_ind,how='left',on='industry_name')
        #     industry_contribution_temp=industry_contribution_temp[industry_contribution_temp['flag'].notnull()]
        #
        #     industry_contribution_temp=industry_contribution_temp.groupby('jjdm').sum()['contribution_ratio'].to_frame('{}_contribution_ratio'.format(sector))
        #     output_df=pd.merge(output_df,industry_contribution_temp,how='left',on='jjdm')



        return  industry_contribution

    def get_advance_man(self,asofdate,industry_rang_dir,theme_name,filter_pool=False):
        # print(asofdate)
        # jjdm_list = util.get_mutual_stock_funds(asofdate)
        asofdate_2=(datetime.datetime.strptime(asofdate, '%Y%m%d') + datetime.timedelta(days=90)).strftime('%Y%m%d')
        adv_man_contribution=self.advance_contribution(asofdate,industry_rang_dir,theme_name)
        adv_man_3ind = self.advance_man_from3ind(asofdate,asofdate_2,industry_rang_dir,theme_name)
        # adv_man_ols=self.advance_ols(jjdm_list)
        # adv_man_ticker=self.advance_man_fromticker(jjdm_list)
        # adv_man_words=self.advance_man_wording()
        #
        # adv_man=pd.merge(adv_man_3ind,adv_man_words
        #                  ,how='left',on='jjdm')
        # adv_man=pd.merge(adv_man,adv_man_ticker
        #                  ,how='left',on='jjdm')
        # adv_man=pd.merge(adv_man,adv_man_ols
        #                  ,how='left',on='jjdm')
        adv_man=pd.merge(adv_man_3ind,adv_man_contribution
                         ,how='left',on='jjdm')

        if(filter_pool):
            pool=adv_man.copy().reset_index(drop=False)
            bmk=util.get_bmk_funds_list_filter(asofdate_2,2)
            pool=pool[pool['jjdm'].isin(bmk)]
            pool=pool[pool['total_con'].notnull()]
            weight_threshold=pool['目标三级行业占比'].nlargest(100).min()
            # con_threshold = pool['total_con'].nlargest(100).min()
            if(theme_name in ['HB消费','HBTMT','HB制造','HB医药']):
                weight_min=45
            else:
                weight_min=35

            pool=pool[((pool['目标三级行业占比']>=weight_threshold)
                       |(pool['目标三级行业占比']>=80))
                      &(pool['目标三级行业占比']>=weight_min)]
            ind3_ret_rank=pd.read_excel(r"E:\GitFolder\docs\三级行业近三年收益排名.xlsx")
            ind3_ret_rank['证券简称']=[x.replace('(申万)','') for x in ind3_ret_rank['证券简称']]
            advance_man_3ind_cluster = pd.read_excel(industry_rang_dir,sheet_name=theme_name)[['板块','三级行业']]
            ind3_ret_rank=pd.merge(ind3_ret_rank,advance_man_3ind_cluster
                                   ,how='inner',left_on='证券简称',right_on='三级行业')
            top20_ind3=[x+'_w' for x in
                        ind3_ret_rank.sort_values(int(asofdate),ascending=False)[0:round(len(ind3_ret_rank)*0.2)]['证券简称']]
            pool['top20_w']=pool[top20_ind3].sum(axis=1).fillna(0)
            pool['top20_w']=pool['top20_w']/pool['目标三级行业占比']

            pool['top20_w'] = pp.scale(pool['top20_w'])
            # pool['top20_w']=pool['top20_w'].rank(method='min')

            # for col in advance_man_3ind_cluster['三级行业'].tolist():
            #     pool[col+'_unit_con']=pool[col+'_con']/pool[col+'_w']
            # pool['avg_unit_con']=pool[[x+'_unit_con' for x in advance_man_3ind_cluster['三级行业'].tolist()]].mean(axis=1)

            pool['avg_unit_con']=pool['total_con']/pool['目标三级行业占比']
            pool['avg_unit_con'] = pp.scale(pool['avg_unit_con'])
            # pool['avg_unit_con'] = pool['avg_unit_con'].rank(method='min')
            pool['new_rank']=pool['avg_unit_con']*0.25+pool['top20_w']*0.75
            # pool=pool.sort_values('new_rank',ascending=False)[0:20][['jjdm']]
            pool['asofdate']=asofdate

            # theme_name="长江"+theme_name

            pool['theme']=theme_name
            pool.rename(columns={'目标三级行业占比':'主题持仓占比',
                                 'total_con':'主题贡献占比',
                                 'avg_unit_con':'单位贡献排名',
                                 'top20_w':'行业配置排名',
                                 'new_rank':'综合因子'},inplace=True)

            localdb.execute("delete from theme_pool where asofdate='{0}' and theme='{1}'"
                            .format(asofdate,theme_name))
            pool[['jjdm','主题持仓占比','主题贡献占比','单位贡献排名','行业配置排名','综合因子','theme','asofdate']].to_sql('theme_pool',con=localdb,index=False,if_exists='append')

        else:
            adv_man.to_excel('先进制造_基于三级行业列表.xlsx')

            print('Done')

    @staticmethod
    def save_all_market_funds2db(filter_date):


        sql="select max(asofdate) as max_date from jjpic_industry_p where asofdate<='{0}' "\
            .format(filter_date)
        max_date1=pd.read_sql(sql,con=localdb)['max_date'].iloc[0]

        f1="select distinct(jjdm) from jjpic_industry_p where `一级行业集中度`<=0.5 and asofdate='{0}'"\
            .format(max_date1)
        f2="select distinct(jjdm) from jjpic_theme_p where `主题集中度`<=0.5  and `大金融`<=0. and `TMT`<=0.3 and `消费`<=0.3 and `制造`<=0.3 and  `周期`<=0.3 and asofdate='{0}'"\
            .format(max_date1)

        jj_pool=\
            pd.merge(pd.read_sql(f1,con=localdb),
                     pd.read_sql(f2,con=localdb),
                     how='inner',on='jjdm')['jjdm'].tolist()

        #get pool style and size  pic
        sql="select max(asofdate) as max_date from jjpic_size_p_hbs where asofdate<='{0}' "\
            .format(filter_date)
        max_date2=pd.read_sql(sql,con=localdb)['max_date'].iloc[0]

        sql="select * from jjpic_size_p_hbs where jjdm in ({0}) and asofdate='{1}' "\
            .format(util.list_sql_condition(jj_pool),max_date2)
        size_pic=pd.read_sql(sql,con=localdb)

        sql="select * from jjpic_value_p_hbs where jjdm in ({0}) and asofdate='{1}' "\
            .format(util.list_sql_condition(jj_pool),max_date2)
        style_pic=pd.read_sql(sql,con=localdb)


        #big_g
        big_g=pd.merge(size_pic[size_pic['规模偏好']=='大盘'][['jjdm']],
                       style_pic[style_pic['风格偏好']=='成长'][['jjdm']],how='inner',on='jjdm')
        big_g['theme']='大盘成长'


        #big_v
        big_v=pd.merge(size_pic[size_pic['规模偏好']=='大盘'][['jjdm']],
                       style_pic[style_pic['风格偏好']=='价值'][['jjdm']],how='inner',on='jjdm')
        big_v['theme']='大盘价值'


        #small_g
        small_g=pd.merge(size_pic[size_pic['规模偏好']=='中小盘'][['jjdm']],
                       style_pic[style_pic['风格偏好']=='成长'][['jjdm']],how='inner',on='jjdm')
        small_g['theme']='中小盘成长'


        #small_v
        small_v=pd.merge(size_pic[size_pic['规模偏好']=='中小盘'][['jjdm']],
                       style_pic[style_pic['风格偏好']=='价值'][['jjdm']],how='inner',on='jjdm')
        small_v['theme']='中小盘价值'


        #balance
        balance=pd.merge(size_pic[size_pic['规模偏好']=='均衡'][['jjdm']],
                       style_pic[style_pic['风格偏好']=='均衡'][['jjdm']],how='inner',on='jjdm')
        balance['theme']='均衡均衡'

        for df in [big_g,big_v,small_g,small_v,balance]:
            df['asofdate'] = max_date2
            if(len(df)>0):
                # delete already exist data
                localdb.execute("delete from theme_pool where theme='{0}' and asofdate='{1}' "
                                .format(df['theme'].iloc[0],max_date2))
                df.to_sql('theme_pool',con=localdb,index=False,if_exists='append')


        print('data inserted Done')

class Public_theme_and_all_market_fund:

    def __init__(self):

        theme_col = ['大金融', '消费', 'TMT', '周期', '制造', '医药']
        theme_map = dict(zip(theme_col,
                             [['银行', '非银金融', '房地产'],
                              ['食品饮料', '家用电器', '社会服务', '农林牧渔', '商贸零售', '美容护理','纺织服饰'],
                              ['通信', '计算机', '电子', '传媒'],
                              ['钢铁', '有色金属', '建筑装饰', '建筑材料', '基础化工', '石油石化', '煤炭'],
                              ['交通运输', '机械设备', '汽车',  '轻工制造', '电力设备', '环保','公用事业','国防军工'],
                              ['医药生物']
                              ]
                             ))
        lista = []
        listb = []
        for theme in theme_col:
            for col in theme_map[theme]:
                lista.append(col)
                listb.append(theme)
        self.ind2thememap = pd.DataFrame()
        self.ind2thememap['industry_name'] = lista
        self.ind2thememap['theme'] = listb

    def save_theme_and_allmarket_fund2db(self,start_date,end_date):

        def get_ind_dis(jjdm_list,start_date,end_date):

            sql = "select jsrq,jjdm,flmc,zgpbl from st_fund.t_st_gm_jjhyzhyszb where jjdm in ({0}) and jsrq>='{1}' and jsrq<='{2}' and hyhfbz='2' and zclb='2' " \
                .format(util.list_sql_condition(jjdm_list),start_date, end_date)
            ind_dis = hbdb.db2df(sql, db='funduser')
            ind_dis = pd.merge(ind_dis, self.ind2thememap, how='left', left_on='flmc', right_on='industry_name')
            ind_dis = ind_dis.groupby(['jjdm', 'jsrq', 'theme']).sum()['zgpbl'].reset_index()

            return ind_dis

        jjdm_list,jjdm_list2 = util.get_potient_theme_funds_pool(end_date, 2)

        # ind_dis=pd.read_pickle(r"E:\GitFolder\hbshare\fe\mutual_analysis\行业数据")
        ind_dis = get_ind_dis(jjdm_list, start_date, end_date)
        # ind_dis = pd.merge(ind_dis, self.ind2thememap, how='left', left_on='flmc', right_on='industry_name')
        # ind_dis = ind_dis.groupby(['jjdm', 'jsrq', 'theme']).sum()['zgpbl'].reset_index()
        # ind_dis2 = pd.read_pickle(r"E:\GitFolder\hbshare\fe\mutual_analysis\行业数据")
        # ind_dis2 = pd.merge(ind_dis2, self.ind2thememap, how='left', left_on='flmc', right_on='industry_name')
        # ind_dis2 = ind_dis2.groupby(['jjdm', 'jsrq', 'theme']).sum()['zgpbl'].reset_index()
        # ind_dis=get_ind_dis(jjdm_list,start_date,end_date)
        ind_dis2 = get_ind_dis(jjdm_list2,start_date,end_date)
        print('industry weight data load done ')
        date_list = ind_dis['jsrq'].sort_values().unique().tolist()


        def save_ext_stable_fund2db(end_date,jjdm_list11):

            year_list=[]
            for i in range(8):
                year_list.append(util._shift_date(str(end_date-i*10000)))
            #get fund ret
            sql="select jjdm,jzrq,fqdwjz from st_fund.t_st_gm_rhb where jjdm in ({0}) and jzrq in ({1}) and fqdwjz!=99999"\
                .format(util.list_sql_condition(jjdm_list11),util.list_sql_condition(year_list))
            fund_year_ret=hbdb.db2df(sql,db='funduser')
            fund_year_ret['ret'] = fund_year_ret.groupby('jjdm')['fqdwjz'].pct_change()
            fund_year_ret=fund_year_ret[fund_year_ret['ret'].notnull()]
            sql="select zqdm,spjg,jyrq from st_market.t_st_zs_hq where zqdm='930950' and jyrq in ({0})"\
                .format(util.list_sql_condition(year_list))
            bmk_yearly_ret=hbdb.db2df(sql,db='alluser')
            bmk_yearly_ret['ret']=bmk_yearly_ret['spjg'].pct_change()
            fund_year_ret=pd.merge(fund_year_ret,bmk_yearly_ret,how='left',left_on='jzrq',right_on='jyrq')

            sql="select jjdm,hb1n as ret,tjnf from st_fund.t_st_gm_nhb where jjdm in ({0}) and tjnf >='{1}' and tjnf<='{2}' and hb1n!=99999"\
                .format(util.list_sql_condition(jjdm_list11),year_list[-1][0:4],year_list[0][0:4])
            fund_year_ret = hbdb.db2df(sql, db='funduser')
            sql="select zbnp as ret ,jjdm from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb='2998' and zbnp!=99999"\
                .format(util.list_sql_condition(jjdm_list11),util._shift_date(str(end_date)))
            temp=hbdb.db2df(sql, db='funduser')
            temp['tjnf']=year_list[0][0:4]
            fund_year_ret=pd.concat([fund_year_ret,temp],axis=0)


            sql="select hb1n as ret,tjnf from st_market.t_st_zs_nhb where zqdm='930950' and tjnf >='{0}' and tjnf<='{1}'"\
                .format(year_list[-1][0:4],year_list[0][0:4])
            bmk_yearly_ret=hbdb.db2df(sql,db='alluser')
            sql="select zbnp as ret from st_market.t_st_zs_rqjhb where zqdm='930950' and jyrq='{0}' and zblb='2998' and zbnp!=99999"\
                .format(util._shift_date(str(end_date)))
            temp=hbdb.db2df(sql,db='alluser')
            temp['tjnf'] = year_list[0][0:4]
            bmk_yearly_ret = pd.concat([bmk_yearly_ret, temp], axis=0)
            fund_year_ret=pd.merge(fund_year_ret,bmk_yearly_ret,how='left',on='tjnf')
            fund_year_ret = fund_year_ret.sort_values(['jjdm', 'tjnf']).reset_index(drop=False)

            fund_year_ret['winning_flag']=(fund_year_ret['ret_x']-fund_year_ret['ret_y'])>0
            fund_year_ret['3_year_winning_pro']=(fund_year_ret.groupby('jjdm').rolling(3)['winning_flag'].sum()/3).values
            fund_year_ret['5_year_winning_pro'] =( fund_year_ret.groupby('jjdm').rolling(5)['winning_flag'].sum()/5).values
            fund_year_ret['7_year_winning_pro'] =( fund_year_ret.groupby('jjdm').rolling(7)['winning_flag'].sum()/7).values
            fund_year_ret.drop_duplicates('jjdm', keep='last',inplace=True)
            ext_ret_list=\
                fund_year_ret[(fund_year_ret['3_year_winning_pro']==1)|(fund_year_ret['5_year_winning_pro']>=0.8)|(fund_year_ret['7_year_winning_pro']==5/7)]['jjdm'].tolist()
            ext_ret_list=pd.DataFrame(data=ext_ret_list,columns=['jjdm'])
            ext_ret_list['asofdate']=end_date
            localdb.execute("delete from extra_table_funds where asofdate='{}'".format(end_date))
            ext_ret_list.to_sql('extra_table_funds',con=localdb,if_exists='append',index=False)

        def get_temp_ind(date1,date2,ind_dis,shift_threshold=None):

            temp_ind = \
                ind_dis[(ind_dis['jsrq'] >= date1) \
                        & (ind_dis['jsrq'] <= date2)]


            temp_ind['theme_change']=\
                temp_ind.sort_values(['jjdm','jsrq']).groupby(['jjdm', 'theme']).diff().abs()['zgpbl']
            temp_shift=temp_ind.groupby('jjdm')['theme_change'].sum()
            if(shift_threshold is not None):
                temp_shift=temp_shift[temp_shift<=shift_threshold]
            else:
                temp_shift=temp_shift[temp_shift.rank()/len(temp_shift)<=0.5]

            shift_threshold=temp_shift.max()
            theme_count=temp_ind.groupby(['jjdm', 'theme'])['jsrq'].count()
            temp_ind =pd.merge(temp_ind.groupby(['jjdm', 'theme'])['zgpbl'].sum()/4,
                               temp_ind.groupby(['jjdm', 'theme'])['zgpbl'].min()
                               ,how='left',on=['jjdm', 'theme']).reset_index()

            temp_ind =pd.merge(temp_ind,theme_count
                               ,how='left',on=['jjdm', 'theme']).reset_index()
            temp_ind.loc[temp_ind['jsrq']<4,'zgpbl_y']=0
            temp_ind.drop('jsrq',axis=1,inplace=True)

            temp_ind['rank']=\
                temp_ind.groupby('jjdm').rank(ascending=False)['zgpbl_x']
            f1=temp_ind[temp_ind['rank'] <= 2]
            # f1=f1.sort_values(['jjdm','zgpbl_x'],ascending=True)
            # f1['pct_change'] = f1.groupby('jjdm')['zgpbl_x'].pct_change().abs()
            f1['pct_change'] = f1.sort_values(['zgpbl_x'], ascending=False).groupby('jjdm')[
                'zgpbl_x'].pct_change().abs()
            f1 = f1[(f1['rank']==1)&((f1['pct_change'] > 0.3)|(f1['pct_change'].isnull()))]['jjdm'].tolist()

            temp_ind = temp_ind[temp_ind['zgpbl_x'] >= 10]

            return  temp_ind,f1,temp_shift,shift_threshold

        def filter_and_save2db(temp_ind,theme_threshold,temp_shift,f1,date2,table_name):

            temp_ind = pd.merge(temp_ind, theme_threshold, how='left', on='theme')
            temp_ind1 = temp_ind[(temp_ind['zgpbl_x_x'] >= temp_ind['zgpbl_x_y'])
                                &(temp_ind['zgpbl_y'] >= (temp_ind['zgpbl_x_y']/(1.5/1.2)))]
            temp_ind2 = temp_ind[(temp_ind['zgpbl_x_x'] >= temp_ind['zgpbl_x_y'])]

            temp_ind1=temp_ind1[(temp_ind1['jjdm'].isin(temp_shift.index.tolist()))&(temp_ind1['jjdm'].isin(f1))]
            temp_ind1=temp_ind1[temp_ind1['rank'] == 1]
            temp_ind2 = temp_ind2[~temp_ind2['jjdm'].isin(temp_ind1['jjdm'].tolist())]
            temp_ind2['theme']=[x+'+' for x in temp_ind2['theme']]
            temp_ind2=temp_ind2.groupby('jjdm')['theme'].sum().reset_index()

            temp_ind1=pd.concat([temp_ind1,temp_ind2],axis=0)
            temp_ind1['asofdate'] = date2
            temp_ind1.rename(columns={"zgpbl_x_x":"平均权重","zgpbl_y":"最低权重","zgpbl_x_y":"阈值"},inplace=True)
            localdb.execute("delete from {0} where asofdate='{1}'".format(table_name,date2))
            temp_ind1[['jjdm', 'theme', 'asofdate', '平均权重','最低权重','阈值']].to_sql(table_name, con=localdb,
                                                                      index=False, if_exists='append')

        for i in range(3, len(date_list)):
            date1 = date_list[i - 3]
            date2 = date_list[i]

            jjdm_list11, jjdm_list22 = util.get_potient_theme_funds_pool(str(date2), 2)
            ind_dis11=ind_dis[ind_dis['jjdm'].isin(jjdm_list11)]
            ind_dis22 = ind_dis2[ind_dis2['jjdm'].isin(jjdm_list22)]

            temp_ind, f1, temp_shift,shift_threshold=get_temp_ind(date1,date2,ind_dis11)
            temp_ind2, f12, temp_shift2,shift_threshold = get_temp_ind(date1, date2, ind_dis22,shift_threshold)

            theme_threshold = temp_ind.groupby('theme').mean()['zgpbl_x'] * 1.5

            filter_and_save2db(temp_ind, theme_threshold, temp_shift, f1, date2, 'public_theme_pool_history')
            filter_and_save2db(temp_ind2, theme_threshold, temp_shift2, f12, date2, 'public_negative_index_theme_pool_history')
            save_ext_stable_fund2db(date2,jjdm_list11)
            print('filter result updated done for {}'.format(date2))

    @staticmethod
    def get_public_advance_theme_pools(end_date):

        latest_asofdate=pd.read_sql("select max(asofdate) as asofdate from public_theme_pool_history where asofdate<='{}'"
                                    .format(end_date),con=localdb)['asofdate'].iloc[0]
        #read raw theme pool
        sql="select jjdm,theme from  public_theme_pool_history where asofdate='{0}'"\
            .format(latest_asofdate)
        theme_pool=pd.read_sql(sql,con=localdb)

        sql="select jjdm,zbnp from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb='2202'"\
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),end_date)
        fund_lj_ret=hbdb.db2df(sql,db='funduser')

        sql="select jjdm,zbnp from st_fund.t_st_gm_rqjzdhc where jjdm in ({0}) and jzrq='{1}' and zblb='2202'"\
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),end_date)
        fund_draw_back=hbdb.db2df(sql,db='funduser')

        sql="select jjdm,tjyf,hb1y from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' "\
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),int(end_date[0:6])-200,end_date[0:6])
        fund_month_ret = hbdb.db2df(sql, db='funduser').pivot_table('hb1y','jjdm','tjyf')
        month_list=fund_month_ret.columns.tolist()

        sql = "select jjdm,ryxm as  基金经理 ,rzrq as 任职日期 from st_fund.t_st_gm_jjjl where jjdm in ({0}) and ryzt=-1 and ryzw='基金经理' and rzrq<={1}" \
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),int(end_date)-20000)
        manager_info = hbdb.db2df(sql, db='funduser')
        manager_info = manager_info.sort_values(['jjdm', '任职日期']).drop_duplicates('jjdm', keep='first')

        theme_pool=pd.merge(theme_pool,fund_lj_ret,how='left',on='jjdm')
        theme_pool = pd.merge(theme_pool, fund_draw_back, how='left', on='jjdm')
        theme_pool = pd.merge(theme_pool, fund_month_ret, how='left', on='jjdm')
        theme_pool = pd.merge(theme_pool, manager_info, how='left', on='jjdm')


        print('start filting')
        advance_theme_pool=[]
        for theme in theme_pool['theme'].unique():
            temp_theme_pool=theme_pool[theme_pool['theme']==theme]
            temp_theme_pool['zbnp_x']=temp_theme_pool['zbnp_x'].rank(ascending=False)/len(temp_theme_pool)
            temp_theme_pool['zbnp_y'] = temp_theme_pool['zbnp_y'].rank(ascending=False)/len(temp_theme_pool)
            for month in month_list:
                temp_theme_pool[month]=temp_theme_pool[month]>temp_theme_pool[month].mean()

            temp_theme_pool['winning']=temp_theme_pool[month_list].sum(axis=1)/len(month_list)
            temp_theme_pool=temp_theme_pool[(temp_theme_pool['zbnp_x']<0.5)
                            &(temp_theme_pool['zbnp_y']<0.3)
                            &(temp_theme_pool['winning']>=0.5)
                            &(temp_theme_pool['基金经理'].notnull())]
            advance_theme_pool.append(temp_theme_pool[['theme','jjdm']])

        advance_theme_pool=pd.concat(advance_theme_pool,axis=0)
        advance_theme_pool['asofdate']=end_date
        localdb.execute("delete from public_advance_theme_pool_history where asofdate='{}'".format(end_date))
        advance_theme_pool.to_sql('public_advance_theme_pool_history',index=False,if_exists='append',con=localdb)
        print('done')

class Agent_theme_fund:

    def __init__(self):

        theme_col = ['大金融', '消费', 'TMT', '周期上游', '制造', '医药']
        theme_map = dict(zip(theme_col,
                             [['银行', '非银金融', '房地产'],
                              ['食品饮料', '家用电器', '社会服务', '农林牧渔', '商贸零售', '美容护理','纺织服饰'],
                              ['通信', '计算机', '电子', '传媒'],
                              ['有色金属','石油石化','煤炭'],
                              ['交通运输', '机械设备', '汽车',  '轻工制造', '电力设备', '环保','公用事业','国防军工'],
                              ['医药生物']
                              ]
                             ))

        lista = []
        listb = []
        for theme in theme_col:
            for col in theme_map[theme]:
                lista.append(col)
                listb.append(theme)
        self.ind2thememap = pd.DataFrame()
        self.ind2thememap['industry_name'] = lista
        self.ind2thememap['theme'] = listb

    def save_theme_and_allmarket_fund2db(self,start_date,end_date):

        def get_ind_dis(jjdm_list,start_date,end_date):

            sql = "select jsrq,jjdm,flmc,zgpbl from st_fund.t_st_gm_jjhyzhyszb where jjdm in ({0}) and jsrq>='{1}' and jsrq<='{2}' and hyhfbz='2' and zclb='2' " \
                .format(util.list_sql_condition(jjdm_list),start_date, end_date)
            ind_dis = hbdb.db2df(sql, db='funduser')
            ind_dis = pd.merge(ind_dis, self.ind2thememap, how='left', left_on='flmc', right_on='industry_name')
            ind_dis = ind_dis.groupby(['jjdm', 'jsrq', 'theme']).sum()['zgpbl'].reset_index()
            return ind_dis

        jjdm_list,jjdm_list2 = util.get_potient_theme_funds_pool(end_date, 2)
        ind_dis = get_ind_dis(jjdm_list,start_date,end_date)
        # ind_dis=pd.read_pickle(r"E:\GitFolder\hbshare\fe\mutual_analysis\行业数据")
        ind_dis = pd.merge(ind_dis, self.ind2thememap, how='left', left_on='flmc', right_on='industry_name')
        ind_dis = ind_dis.groupby(['jjdm', 'jsrq', 'theme']).sum()['zgpbl'].reset_index()
        # ind_dis2 = pd.read_pickle(r"E:\GitFolder\hbshare\fe\mutual_analysis\行业数据")
        # ind_dis2 = pd.merge(ind_dis2, self.ind2thememap, how='left', left_on='flmc', right_on='industry_name')
        # ind_dis2 = ind_dis2.groupby(['jjdm', 'jsrq', 'theme']).sum()['zgpbl'].reset_index()
        # ind_dis=get_ind_dis(jjdm_list,start_date,end_date)
        ind_dis2 = get_ind_dis(jjdm_list2,start_date,end_date)
        print('industry weight data load done ')
        date_list = ind_dis['jsrq'].sort_values().unique().tolist()

        def get_temp_ind(date1,date2,ind_dis,shift_threshold=None):

            temp_ind = \
                ind_dis[(ind_dis['jsrq'] >= date1) \
                        & (ind_dis['jsrq'] <= date2)]


            temp_ind['theme_change']=\
                temp_ind.sort_values(['jjdm','jsrq']).groupby(['jjdm', 'theme']).diff().abs()['zgpbl']
            temp_shift=temp_ind.groupby('jjdm')['theme_change'].sum()
            if(shift_threshold is not None):
                temp_shift=temp_shift[temp_shift<=shift_threshold]
            else:
                temp_shift=temp_shift[temp_shift.rank()/len(temp_shift)<=0.5]

            shift_threshold=temp_shift.max()
            theme_count=temp_ind.groupby(['jjdm', 'theme'])['jsrq'].count()
            temp_ind =pd.merge(temp_ind.groupby(['jjdm', 'theme'])['zgpbl'].sum()/4,
                               temp_ind.groupby(['jjdm', 'theme'])['zgpbl'].min()
                               ,how='left',on=['jjdm', 'theme']).reset_index()

            temp_ind =pd.merge(temp_ind,theme_count
                               ,how='left',on=['jjdm', 'theme']).reset_index()
            temp_ind.loc[temp_ind['jsrq']<4,'zgpbl_y']=0
            temp_ind.drop('jsrq',axis=1,inplace=True)

            temp_ind['rank']=\
                temp_ind.groupby('jjdm').rank(ascending=False)['zgpbl_x']
            f1=temp_ind[temp_ind['rank'] <= 2]
            f1['pct_change'] = f1.sort_values('zgpbl_x',ascending=False).groupby('jjdm')['zgpbl_x'].pct_change().abs()
            f1 = f1[(f1['rank']==1)&((f1['pct_change'] > 0.3)|(f1['pct_change'].isnull()))]['jjdm'].tolist()

            temp_ind = temp_ind[temp_ind['zgpbl_x'] >= 10]

            return  temp_ind,f1,temp_shift,shift_threshold

        def filter_and_save2db(temp_ind,theme_threshold,temp_shift,f1,date2,table_name):

            temp_ind = pd.merge(temp_ind, theme_threshold, how='left', on='theme')
            temp_ind1 = temp_ind[(temp_ind['zgpbl_x_x'] >= temp_ind['zgpbl_x_y'])
                                &(temp_ind['zgpbl_y'] >= (temp_ind['zgpbl_x_y']/(1.5/1.2)))]
            temp_ind2 = temp_ind[(temp_ind['zgpbl_x_x'] >= temp_ind['zgpbl_x_y'])]

            temp_ind1=temp_ind1[(temp_ind1['jjdm'].isin(temp_shift.index.tolist()))&(temp_ind1['jjdm'].isin(f1))]
            temp_ind1=temp_ind1[temp_ind1['rank'] == 1]
            temp_ind2 = temp_ind2[~temp_ind2['jjdm'].isin(temp_ind1['jjdm'].tolist())]
            temp_ind2['theme']=[x+'+' for x in temp_ind2['theme']]
            temp_ind2=temp_ind2.groupby('jjdm')['theme'].sum().reset_index()

            temp_ind1=pd.concat([temp_ind1,temp_ind2],axis=0)
            temp_ind1['asofdate'] = date2
            temp_ind1.rename(columns={"zgpbl_x_x":"平均权重","zgpbl_y":"最低权重","zgpbl_x_y":"阈值"},inplace=True)
            localdb.execute("delete from {0} where asofdate='{1}'".format(table_name,date2))
            temp_ind1[['jjdm', 'theme', 'asofdate', '平均权重','最低权重','阈值']].to_sql(table_name, con=localdb,
                                                                      index=False, if_exists='append')

        for i in range(3, len(date_list)):
            date1 = date_list[i - 3]
            date2 = date_list[i]

            jjdm_list11, jjdm_list22 = util.get_potient_theme_funds_pool(str(date2), 2)
            ind_dis11=ind_dis[ind_dis['jjdm'].isin(jjdm_list11)]
            ind_dis22 = ind_dis2[ind_dis2['jjdm'].isin(jjdm_list22)]

            temp_ind, f1, temp_shift,shift_threshold=get_temp_ind(date1,date2,ind_dis11)
            temp_ind2, f12, temp_shift2,shift_threshold = get_temp_ind(date1, date2, ind_dis22,shift_threshold)

            theme_threshold = temp_ind.groupby('theme').mean()['zgpbl_x'] * 1.5

            filter_and_save2db(temp_ind, theme_threshold, temp_shift, f1, date2, 'agency_theme_pool_history')
            filter_and_save2db(temp_ind2, theme_threshold, temp_shift2, f12, date2, 'agency_negative_index_theme_pool_history')
            # save_ext_stable_fund2db(int(util._shift_date(str(date2))))
            print('filter result updated done for {}'.format(date2))

    @staticmethod
    def get_public_advance_theme_pools(end_date):

        latest_asofdate=pd.read_sql("select max(asofdate) as asofdate from agency_theme_pool_history where asofdate<='{}'"
                                    .format(end_date),con=localdb)['asofdate'].iloc[0]
        #read raw theme pool
        sql="select jjdm,theme from  agency_theme_pool_history where asofdate='{0}'"\
            .format(latest_asofdate)
        theme_pool=pd.read_sql(sql,con=localdb)

        sql="select jjdm,zbnp from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb='2202'"\
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),end_date)
        fund_lj_ret=hbdb.db2df(sql,db='funduser')

        sql="select jjdm,zbnp from st_fund.t_st_gm_rqjzdhc where jjdm in ({0}) and jzrq='{1}' and zblb='2202'"\
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),end_date)
        fund_draw_back=hbdb.db2df(sql,db='funduser')

        sql="select jjdm,tjyf,hb1y from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' "\
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),int(end_date[0:6])-200,end_date[0:6])
        fund_month_ret = hbdb.db2df(sql, db='funduser').pivot_table('hb1y','jjdm','tjyf')
        month_list=fund_month_ret.columns.tolist()

        sql = "select jjdm,ryxm as  基金经理 ,rzrq as 任职日期 from st_fund.t_st_gm_jjjl where jjdm in ({0}) and ryzt=-1 and ryzw='基金经理' and rzrq<={1}" \
            .format(util.list_sql_condition(theme_pool['jjdm'].tolist()),int(end_date)-20000)
        manager_info = hbdb.db2df(sql, db='funduser')
        manager_info = manager_info.sort_values(['jjdm', '任职日期']).drop_duplicates('jjdm', keep='first')

        theme_pool=pd.merge(theme_pool,fund_lj_ret,how='left',on='jjdm')
        theme_pool = pd.merge(theme_pool, fund_draw_back, how='left', on='jjdm')
        theme_pool = pd.merge(theme_pool, fund_month_ret, how='left', on='jjdm')
        theme_pool = pd.merge(theme_pool, manager_info, how='left', on='jjdm')


        print('start filting')
        advance_theme_pool=[]
        for theme in theme_pool['theme'].unique():
            temp_theme_pool=theme_pool[theme_pool['theme']==theme]
            temp_theme_pool['zbnp_x']=temp_theme_pool['zbnp_x'].rank(ascending=False)/len(temp_theme_pool)
            temp_theme_pool['zbnp_y'] = temp_theme_pool['zbnp_y'].rank(ascending=False)/len(temp_theme_pool)
            for month in month_list:
                temp_theme_pool[month]=temp_theme_pool[month]>temp_theme_pool[month].mean()

            temp_theme_pool['winning']=temp_theme_pool[month_list].sum(axis=1)/len(month_list)
            temp_theme_pool=temp_theme_pool[(temp_theme_pool['zbnp_x']<0.5)
                            &(temp_theme_pool['zbnp_y']<0.3)
                            &(temp_theme_pool['winning']>=0.5)
                            &(temp_theme_pool['基金经理'].notnull())]
            advance_theme_pool.append(temp_theme_pool[['theme','jjdm']])

        advance_theme_pool=pd.concat(advance_theme_pool,axis=0)
        advance_theme_pool['asofdate']=end_date
        #localdb.execute("delete from agency_advance_theme_pool_history where asofdate='{}'".format(end_date))
        advance_theme_pool.to_sql('agency_advance_theme_pool_history',index=False,if_exists='append',con=localdb)
        print('done')

class Similarity_filter:

    @staticmethod
    def theme_similarity(target_jjdm,ext1,ext2):

        maxasofdate=pd.read_sql("select max(asofdate) as max from jjpic{0}_theme_p".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_theme_p where asofdate='{1}' and jjdm='{2}' "\
            .format(ext1,maxasofdate,target_jjdm)
        target_pic=pd.read_sql(sql,con=localdb).fillna(0)

        maxasofdate=pd.read_sql("select max(asofdate) as max from jjpic{0}_theme_p".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_theme_p where asofdate='{1}' "\
            .format(ext2,maxasofdate)
        theme_pic=pd.read_sql(sql,con=localdb).fillna(0)

        theme_weight_difference=((theme_pic['大金融']-target_pic['大金融'].values[0]).abs()\
                                +(theme_pic['消费']-target_pic['消费'].values[0]).abs()+\
                                (theme_pic['TMT']-target_pic['TMT'].values[0]).abs()\
                                +(theme_pic['周期']-target_pic['周期'].values[0]).abs()\
                                +(theme_pic['制造']-target_pic['制造'].values[0]).abs())/2

        theme_difference=((0.1*((theme_pic['主题集中度']-target_pic['主题集中度'].values[0]).abs().rank(method='min')/len(theme_pic))\
                         +0.1*((theme_pic['主题换手率']-target_pic['主题换手率'].values[0]).abs().rank(method='min')/len(theme_pic))\
                         +(theme_weight_difference.rank(method='min')/len(theme_pic))*0.8)/1).to_frame('theme_similarity')
        theme_difference['jjdm']=theme_pic['jjdm']
        theme_difference['theme_similarity']=theme_difference['theme_similarity'].rank(method='min',ascending=False)/len(theme_difference)
        theme_difference['theme_weight_difference']=theme_weight_difference.values

        return  theme_difference

    @staticmethod
    def industry_similarity(target_jjdm,ext1,ext2):

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_p".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_p where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_industry_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_p".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_p where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        industry_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = \
        pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_detail_1".format(ext1), con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_detail_1 where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_industry_detail_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = \
        pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_detail_1".format(ext2), con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_detail_1 where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        industry_detail_pic = pd.read_sql(sql, con=localdb).fillna(0)

        industry_weight_difference=pd.merge(industry_detail_pic,
                                            target_industry_detail_pic,
                                            how='outer',on='行业名称').fillna(0)
        industry_weight_difference['weight_difference']=(industry_weight_difference['占持仓比例(时序均值)_x']-industry_weight_difference['占持仓比例(时序均值)_y']).abs()
        industry_weight_difference=industry_weight_difference.groupby('jjdm_x').sum()['weight_difference']/2

        industry_difference=pd.Series(((0.1*((industry_pic['一级行业集中度']-target_industry_pic['一级行业集中度'].values[0]).abs().rank(method='min')/len(industry_pic)).values
                             +0.1*((industry_pic['一级行业换手率']-target_industry_pic['一级行业换手率'].values[0]).abs().rank(method='min')/len(industry_pic)).values
                              +0.8*(industry_weight_difference.rank(method='min') / len(industry_weight_difference)).values)/1)).to_frame('industry_similarity')
        industry_difference['jjdm']=industry_pic['jjdm']
        industry_difference['industry_similarity'] = industry_difference['industry_similarity'].rank(method='min', ascending=False)/len(industry_difference)
        industry_difference['industry_weight_difference']=industry_weight_difference.values
        return  industry_difference

    @staticmethod
    def style_similarity(target_jjdm,ext1,ext2):


        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_value_p_hbs".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_value_p_hbs where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_style_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_value_p_hbs".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_value_p_hbs where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        style_pic = pd.read_sql(sql, con=localdb).fillna(0)

        style_weight_difference = ((style_pic['成长绝对暴露(持仓)'] - target_style_pic['成长绝对暴露(持仓)'].values[0]).abs() \
                                   + (style_pic['价值绝对暴露(持仓)'] - target_style_pic['价值绝对暴露(持仓)'].values[0]).abs() )/2
        style_difference = ((0.1*((style_pic['集中度(持仓)'] - target_style_pic['集中度(持仓)'].values[0]).abs().rank(method='min')/len(style_pic)) \
                            + 0.1*((style_pic['换手率(持仓)'] - target_style_pic['换手率(持仓)'].values[0]).abs().rank(method='min')/len(style_pic)) \
                            + 0.8*(style_weight_difference.rank(method='min')/len(style_pic))) / 1).to_frame('style_similarity')

        style_difference['jjdm']=style_pic['jjdm']
        style_difference['style_similarity']=style_difference['style_similarity'].rank(method='min',ascending=False)/len(style_difference)
        style_difference['style_weight_difference']=style_weight_difference.values

        return style_difference

    @staticmethod
    def size_similarity(target_jjdm,ext1,ext2):

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_size_p_hbs".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_size_p_hbs where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_size_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_size_p_hbs".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_size_p_hbs where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        size_pic = pd.read_sql(sql, con=localdb).fillna(0)

        size_weight_difference = ((size_pic['大盘绝对暴露(持仓)'] - target_size_pic['大盘绝对暴露(持仓)'].values[0]).abs()
                                  + (size_pic['中盘绝对暴露(持仓)'] - target_size_pic['中盘绝对暴露(持仓)'].values[0]).abs()+
                                  (size_pic['小盘绝对暴露(持仓)'] - target_size_pic['小盘绝对暴露(持仓)'].values[0]).abs() )/2
        size_difference = ((0.1*((size_pic['集中度(持仓)'] - target_size_pic['集中度(持仓)'].values[0]).abs().rank(method='min')/len(size_pic)) \
                            + 0.1*((size_pic['换手率(持仓)'] - target_size_pic['换手率(持仓)'].values[0]).abs().rank(method='min')/len(size_pic)) \
                            + 0.8*(size_weight_difference.rank(method='min')/len(size_pic))) / 1).to_frame('size_similarity')
        size_difference['jjdm']=size_pic['jjdm']
        size_difference['size_similarity']=size_difference['size_similarity'].rank(method='min',ascending=False)/len(size_difference)
        size_difference['size_weight_difference']=size_weight_difference.values

        return size_difference

    def get_similarity(self,target_jjdm,mu_target_type=True,mu_pool_type=True,shown_num=30):

        if(mu_target_type==True):
            ext1=''
        else:
            ext1='_prv'

        if(mu_pool_type==True):
            ext2=''
        else:
            ext2='_prv'


        theme_difference=self.theme_similarity(target_jjdm,ext1,ext2)
        industry_difference=self.industry_similarity(target_jjdm,ext1,ext2)
        style_differene=self.style_similarity(target_jjdm,ext1,ext2)
        size_difference=self.size_similarity(target_jjdm,ext1,ext2)

        similarity=pd.merge(pd.merge(theme_difference,industry_difference,how='inner',on='jjdm'),
                            pd.merge(style_differene,size_difference,how='inner',on='jjdm'),how='inner',on='jjdm')

        similarity['total_similarity']=(1.5*similarity['theme_similarity']+
                                        0.5*similarity['industry_similarity']+
                                        0.5*similarity['style_similarity']+
                                        0.5*similarity['size_similarity'])/3
        similarity=similarity.sort_values('total_similarity',ascending=False).iloc[0:shown_num]

        jj_base_info=hbdb.db2df("select jjdm,jjjc from st_fund.t_st_gm_jjxx where jjdm in ({})"
                                .format(util.list_sql_condition(similarity['jjdm'].tolist())),db='funduser')
        similarity=pd.merge(similarity,jj_base_info,how='left',on='jjdm')

        return similarity[['jjdm','jjjc','total_similarity','theme_similarity','theme_weight_difference',
                           'industry_similarity','industry_weight_difference','style_similarity','style_weight_difference',
                           'size_similarity','size_weight_difference']]

class future_star_funds:

    @staticmethod
    def get_fund_return(jjdm_list,jzrq,zblb):

        sql="select jjdm,zbnp from st_fund.t_st_gm_rqjhb where jjdm in {0} and jzrq='{1}' and zblb='{2}'"\
            .format(tuple(jjdm_list),jzrq,zblb)
        fund_ret=hbdb.db2df(sql,db='funduser')

        return fund_ret

    @staticmethod
    def get_manager_infor(jjdm_list,jzrq):

        sql="select jjdm,rzrq,rydm,ryxm from st_fund.t_st_gm_jjjl where jjdm in {0} and rzrq<='{1}' and lrrq>='{2}' and ryzw='基金经理'"\
            .format(tuple(jjdm_list),jzrq,jzrq)
        manager_infor=hbdb.db2df(sql,db='funduser')

        manager_infor=manager_infor.sort_values('rzrq').drop_duplicates('rydm', keep='first')

        return  manager_infor

    @staticmethod
    def get_funds_size(jjdm_list,jsrq):

        sql="select max(jsrq) as jsrq from st_fund.t_st_gm_zcpz where jsrq<='{0}' and jjdm in {1} and jjzzc>0"\
            .format(jsrq,tuple(jjdm_list))
        max_jsrq=hbdb.db2df(sql,db='funduser')['jsrq'].iloc[0]

        sql="select jjdm,jjzzc/100000000 as jjzzc from st_fund.t_st_gm_zcpz where jjdm in {0} and jsrq='{1}'"\
            .format(tuple(jjdm_list),max_jsrq)
        fund_size=hbdb.db2df(sql,db='funduser')

        return  fund_size

    def filter(self,jzrq,
               return_threshield=0.7,
               size_threshield=50,
               time_threshield=365):

        jjdm_list=util.get_stock_funds_pool(jzrq)

        fund_ret=self.get_fund_return(jjdm_list,jzrq,'2201')

        manager_infor=self.get_manager_infor(jjdm_list,jzrq)

        fund_size=self.get_funds_size(jjdm_list,jzrq)

        #filter by funds return
        fund_ret['rank']=fund_ret['zbnp'].rank(pct=True)

        f1=fund_ret[fund_ret['rank'] >= return_threshield]['jjdm'].tolist()

        #filter by size_threshild

        f2=fund_size[fund_size['jjzzc'] <= size_threshield]['jjdm'].tolist()

        #filter by manager career lifetime

        earliest_time=util.str_date_shift(jzrq,time_threshield,'-')
        f3=manager_infor[manager_infor['rzrq']>=int(earliest_time)]['rydm'].tolist()

        final_f=list(set(f1).intersection(set(f2)))

        final_f=\
            manager_infor[(manager_infor['jjdm'].isin(final_f))]

        final_f=pd.merge(final_f,fund_ret[['jjdm','zbnp']],how='left',on='jjdm')
        final_f = pd.merge(final_f, fund_size[['jjdm', 'jjzzc']], how='left', on='jjdm')

        final_f=pd.merge(final_f,final_f.groupby('rydm')['jjzzc'].sum().to_frame('totalmv')
                         ,how='left',on='rydm')

        final_f['w']=final_f['jjzzc']/final_f['totalmv']
        final_f['zbnp']=final_f['zbnp']*final_f['w']
        final_f=pd.merge(final_f,final_f.groupby('rydm')['zbnp'].sum().to_frame('avg_ret')
                         ,how='left',on='rydm').drop_duplicates('rydm').drop(['jjzzc', 'w','zbnp'],axis=1)

        final_f=final_f[final_f['rydm'].isin(f3)]

        final_f.columns=['任职日期', 'rydm', '基金代码', '人员姓名', '基金规模(亿)', '区间收益']

        final_f.to_excel(r"E:\GitFolder\docs\基金筛选\黑马基金\{}.xlsx".format(jzrq),index=False)


def factor_test(asofdate1,asofdate2,theme_name,pool_size,save_to_db=False):

    if(theme_name=='超额稳定'):
        orginal_pool=pd.read_sql("SELECT * from extra_table_funds where asofdate >='{0}' and asofdate<='{1}' "
                                 .format(pd.Series(asofdate1).min(),pd.Series(asofdate1).max()),con=localdb)
        if_theme_funds=False
    elif(theme_name in ['成长','价值']):
        orginal_pool=pd.read_sql("SELECT jjdm,asofdate from jjpic_value_p_hbs where asofdate >='{0}' and asofdate<='{1}' "
                                 .format(pd.Series(asofdate1).min(),pd.Series(asofdate1).max()),con=localdb)
        if_theme_funds = False
    else:
        orginal_pool=pd.read_sql("SELECT * from public_theme_pool_history where theme='{0}' and asofdate >='{1}' and asofdate<='{2}' "
                                 .format(theme_name,pd.Series(asofdate1).min(),pd.Series(asofdate1).max()),con=localdb)
        if_theme_funds = True

    # get the theme and industry pic factor
    sql="SELECT * from jjpic_theme_p where asofdate >='{0}' and asofdate<='{2}'  and jjdm in ({1})"\
        .format(pd.Series(asofdate1).min(),
                util.list_sql_condition(orginal_pool['jjdm'].tolist())
                ,pd.Series(asofdate2).max())
    theme_pic=pd.read_sql(sql,con=localdb)

    theme_pic['max_theme']=theme_pic[['大金融', '消费', 'TMT', '周期', '制造']].max(axis=1)
    sql="select jjdm,`一级行业集中度`,`一级行业换手率`,`二级行业集中度`,`二级行业换手率`,`三级行业集中度`,`三级行业换手率`,`龙头占比(时序均值)`,asofdate  from jjpic_industry_p where asofdate>='{0}' and jjdm in ({1}) "\
        .format(pd.Series(asofdate1).min(),util.list_sql_condition(orginal_pool['jjdm'].tolist()))
    ind_pic=pd.read_sql(sql,con=localdb)
    ind_pic['asofdate']=ind_pic['asofdate'].astype(str)

    theme_pic=pd.merge(theme_pic,ind_pic,how='inner',on=['jjdm','asofdate'])

    # get the stock factor
    sql="SELECT * from jjpic_stock_p where asofdate >='{0}'  and jjdm in ({1})"\
        .format(pd.Series(asofdate1).min(),util.list_sql_condition(orginal_pool['jjdm'].tolist()))
    stock_pic=pd.read_sql(sql,con=localdb)
    sql="SELECT * from jjpic_stock_tp where asofdate >='{0}'  and jjdm in ({1})"\
        .format(pd.Series(asofdate1).min(),util.list_sql_condition(orginal_pool['jjdm'].tolist()))
    stock_tp_pic=pd.read_sql(sql,con=localdb)

    asofdate_map2=dict(zip(stock_pic['asofdate'].sort_values().unique().tolist()
                          ,orginal_pool['asofdate'].sort_values().unique().tolist()))
    stock_pic['asofdate']=[asofdate_map2[x] for x in stock_pic['asofdate']]
    asofdate_map3=dict(zip(stock_tp_pic['asofdate'].sort_values().unique().tolist()
                          ,orginal_pool['asofdate'].sort_values().unique().tolist()))
    stock_tp_pic['asofdate']=[asofdate_map3[x] for x in stock_tp_pic['asofdate']]


    stock_pic=pd.merge(stock_pic[['jjdm','asofdate','个股集中度','持股数量','PE', 'PB', 'ROE', '股息率','PCF', '毛利率', '净利率','净利增速']],
                       stock_tp_pic[['jjdm','asofdate','换手率','左侧概率(出重仓前,半年线)','买入 ROE(出重仓前)_rank',
                                     '买入 ROE(出持仓前)_rank', '卖出 ROE(出重仓前)_rank', '卖出 ROE(出持仓前)_rank',
                                     '买入 净利润增速(出重仓前)_rank', '买入 净利润增速(出持仓前)_rank', '卖出 净利润增速(出重仓前)_rank',
                                     '卖出 净利润增速(出持仓前)_rank', '买入 股息率(出重仓前)_rank', '买入 股息率(出持仓前)_rank',
                                     '卖出 股息率(出重仓前)_rank', '卖出 股息率(出持仓前)_rank', '买入 PE(出重仓前)_rank', '买入 PE(出持仓前)_rank',
                                     '卖出 PE(出重仓前)_rank', '卖出 PE(出持仓前)_rank', '买入 PB(出重仓前)_rank', '买入 PB(出持仓前)_rank',
                                     '卖出 PB(出重仓前)_rank', '卖出 PB(出持仓前)_rank', '买入 毛利率(出重仓前)_rank', '买入 毛利率(出持仓前)_rank',
                                     '卖出 毛利率(出重仓前)_rank', '卖出 毛利率(出持仓前)_rank', '买入 净利率(出重仓前)_rank', '买入 净利率(出持仓前)_rank',
                                     '卖出 净利率(出重仓前)_rank', '卖出 净利率(出持仓前)_rank', '买入 PEG(出重仓前)_rank', '买入 PEG(出持仓前)_rank',
                                     '卖出 PEG(出重仓前)_rank', '卖出 PEG(出持仓前)_rank', '买入 PCF(出重仓前)_rank', '买入 PCF(出持仓前)_rank',
                                     '卖出 PCF(出重仓前)_rank', '卖出 PCF(出持仓前)_rank', '买入 ROE-PB比例(出重仓前)_rank', '买入 ROE-PB比例(出持仓前)_rank',
                                     '卖出 ROE-PB比例(出重仓前)_rank', '卖出 ROE-PB比例(出持仓前)_rank', '买入 PEG比例(出重仓前)_rank', '买入 PEG比例(出持仓前)_rank',
                                     '卖出 PEG比例(出重仓前)_rank', '卖出 PEG比例(出持仓前)_rank','逆向买入比例(出持仓前)', '逆向买入比例(出持仓前)_rank', 'baotuan_ratio']],
                       how='inner',on=['jjdm','asofdate'])

    #get thr brinson result
    sql="select jjdm,tjrq as asofdate,sector_allo,equity_selection,trading from st_fund.r_st_hold_excess_attr_df where tjrq>='{0}' and jjdm in ({1})"\
        .format(pd.Series(asofdate1).min(),util.list_sql_condition(orginal_pool['jjdm'].tolist()))
    brinson_factor=hbdb.db2df(sql,db='funduser')


    #get the fund size
    sql="select jjdm,jsrq,jjzzc from st_fund.t_st_gm_zcpz where jjdm in ({0}) and jsrq >='{1}' and jsrq<='{2}'"\
        .format(util.list_sql_condition(orginal_pool['jjdm'].tolist()),pd.Series(asofdate1).min(),pd.Series(asofdate1).max())
    jjzzc=hbdb.db2df(sql,db='funduser')
    stock_pic['ym']=stock_pic['asofdate'].astype(str).str[0:6]
    jjzzc['ym'] = jjzzc['jsrq'].astype(str).str[0:6]
    stock_pic=pd.merge(stock_pic,jjzzc,how='left',on=['jjdm','ym'])


    if(if_theme_funds):
        #get the theme infor ratio

        sql="select jzrq,jjdm,hbdr from st_fund.t_st_gm_rhb where jzrq<='{0}' and jjdm in ({1}) and jzrq>='{2}' "\
            .format(pd.Series(asofdate2).max(),util.list_sql_condition(orginal_pool['jjdm'].tolist()),pd.Series(asofdate2).min())
        jj_daily_ret=hbdb.db2df(sql,db='funduser')

        sql="select jzrq,{0} from theme_daily_ret where jzrq<='{1}'"\
            .format(theme_name,pd.Series(asofdate1).max())
        theme_daily_ret=pd.read_sql(sql,con=localdb)

        rolling_window=120
        jj_daily_ret=pd.merge(jj_daily_ret,theme_daily_ret,how='inner',on='jzrq').sort_values(['jjdm','jzrq'])
        jj_daily_ret['ext_ret'] = jj_daily_ret['hbdr']/100 - jj_daily_ret[theme_name.replace('HB', '')]
        jj_daily_ret['ir_theme'] = (jj_daily_ret.groupby('jjdm')[['ext_ret']].rolling(rolling_window).mean() / jj_daily_ret.groupby('jjdm')[
            ['ext_ret']].rolling(rolling_window).std())['ext_ret'].tolist()

    else:
        jj_daily_ret=pd.DataFrame(columns=['jjdm','ir_theme','jzrq'])
        jj_daily_ret['jzrq']=19900101

    #standarlize all the factor data
    def scale_by_group(arr):

       return pd.Series(data=pp.scale(arr),index=arr.index)

    for col in ['主题集中度', '主题换手率','一级行业集中度', '一级行业换手率', '二级行业集中度',
       '二级行业换手率', '三级行业集中度', '三级行业换手率', '龙头占比(时序均值)']:
        theme_pic[col] = theme_pic.groupby('asofdate')[col].apply(scale_by_group)

    for col in ['个股集中度','持股数量','PE', 'PB', 'ROE', '股息率','PCF', '毛利率', '净利率','净利增速','换手率','左侧概率(出重仓前,半年线)','jjzzc','买入 ROE(出重仓前)_rank',
                                     '买入 ROE(出持仓前)_rank', '卖出 ROE(出重仓前)_rank', '卖出 ROE(出持仓前)_rank',
                                     '买入 净利润增速(出重仓前)_rank', '买入 净利润增速(出持仓前)_rank', '卖出 净利润增速(出重仓前)_rank',
                                     '卖出 净利润增速(出持仓前)_rank', '买入 股息率(出重仓前)_rank', '买入 股息率(出持仓前)_rank',
                                     '卖出 股息率(出重仓前)_rank', '卖出 股息率(出持仓前)_rank', '买入 PE(出重仓前)_rank', '买入 PE(出持仓前)_rank',
                                     '卖出 PE(出重仓前)_rank', '卖出 PE(出持仓前)_rank', '买入 PB(出重仓前)_rank', '买入 PB(出持仓前)_rank',
                                     '卖出 PB(出重仓前)_rank', '卖出 PB(出持仓前)_rank', '买入 毛利率(出重仓前)_rank', '买入 毛利率(出持仓前)_rank',
                                     '卖出 毛利率(出重仓前)_rank', '卖出 毛利率(出持仓前)_rank', '买入 净利率(出重仓前)_rank', '买入 净利率(出持仓前)_rank',
                                     '卖出 净利率(出重仓前)_rank', '卖出 净利率(出持仓前)_rank', '买入 PEG(出重仓前)_rank', '买入 PEG(出持仓前)_rank',
                                     '卖出 PEG(出重仓前)_rank', '卖出 PEG(出持仓前)_rank', '买入 PCF(出重仓前)_rank', '买入 PCF(出持仓前)_rank',
                                     '卖出 PCF(出重仓前)_rank', '卖出 PCF(出持仓前)_rank', '买入 ROE-PB比例(出重仓前)_rank', '买入 ROE-PB比例(出持仓前)_rank',
                                     '卖出 ROE-PB比例(出重仓前)_rank', '卖出 ROE-PB比例(出持仓前)_rank', '买入 PEG比例(出重仓前)_rank', '买入 PEG比例(出持仓前)_rank',
                                     '卖出 PEG比例(出重仓前)_rank', '卖出 PEG比例(出持仓前)_rank','逆向买入比例(出持仓前)', '逆向买入比例(出持仓前)_rank', 'baotuan_ratio']:
        stock_pic[col] = stock_pic.groupby('asofdate')[col].apply(scale_by_group)
    for col in ['trading','equity_selection','sector_allo']:
        brinson_factor[col]=brinson_factor.groupby('asofdate')[col].apply(scale_by_group).values


    asofdate_map=dict(zip(theme_pic['asofdate'].sort_values().unique().tolist()
                          ,orginal_pool['asofdate'].sort_values().unique().tolist()))
    theme_pic['asofdate']=[asofdate_map[x] for x in theme_pic['asofdate']]



    orginal_pool=pd.merge(orginal_pool,
                          theme_pic[['jjdm','asofdate','主题集中度', '主题换手率','一级行业集中度', '一级行业换手率', '二级行业集中度',
       '二级行业换手率', '三级行业集中度', '三级行业换手率', '龙头占比(时序均值)']],
                          how='left',on=['jjdm','asofdate'])

    orginal_pool=pd.merge(orginal_pool,
                          stock_pic[['jjdm','asofdate','个股集中度','持股数量','PE',
                                     'PB', 'ROE', '股息率','PCF', '毛利率', '净利率','净利增速','换手率','左侧概率(出重仓前,半年线)','jjzzc','买入 ROE(出重仓前)_rank',
                                     '买入 ROE(出持仓前)_rank', '卖出 ROE(出重仓前)_rank', '卖出 ROE(出持仓前)_rank',
                                     '买入 净利润增速(出重仓前)_rank', '买入 净利润增速(出持仓前)_rank', '卖出 净利润增速(出重仓前)_rank',
                                     '卖出 净利润增速(出持仓前)_rank', '买入 股息率(出重仓前)_rank', '买入 股息率(出持仓前)_rank',
                                     '卖出 股息率(出重仓前)_rank', '卖出 股息率(出持仓前)_rank', '买入 PE(出重仓前)_rank', '买入 PE(出持仓前)_rank',
                                     '卖出 PE(出重仓前)_rank', '卖出 PE(出持仓前)_rank', '买入 PB(出重仓前)_rank', '买入 PB(出持仓前)_rank',
                                     '卖出 PB(出重仓前)_rank', '卖出 PB(出持仓前)_rank', '买入 毛利率(出重仓前)_rank', '买入 毛利率(出持仓前)_rank',
                                     '卖出 毛利率(出重仓前)_rank', '卖出 毛利率(出持仓前)_rank', '买入 净利率(出重仓前)_rank', '买入 净利率(出持仓前)_rank',
                                     '卖出 净利率(出重仓前)_rank', '卖出 净利率(出持仓前)_rank', '买入 PEG(出重仓前)_rank', '买入 PEG(出持仓前)_rank',
                                     '卖出 PEG(出重仓前)_rank', '卖出 PEG(出持仓前)_rank', '买入 PCF(出重仓前)_rank', '买入 PCF(出持仓前)_rank',
                                     '卖出 PCF(出重仓前)_rank', '卖出 PCF(出持仓前)_rank', '买入 ROE-PB比例(出重仓前)_rank', '买入 ROE-PB比例(出持仓前)_rank',
                                     '卖出 ROE-PB比例(出重仓前)_rank', '卖出 ROE-PB比例(出持仓前)_rank', '买入 PEG比例(出重仓前)_rank', '买入 PEG比例(出持仓前)_rank',
                                     '卖出 PEG比例(出重仓前)_rank', '卖出 PEG比例(出持仓前)_rank','逆向买入比例(出持仓前)', '逆向买入比例(出持仓前)_rank', 'baotuan_ratio']],
                          how='left',on=['jjdm','asofdate'])
    brinson_factor['asofdate']=brinson_factor['asofdate'].astype(int)
    orginal_pool['asofdate']=orginal_pool['asofdate'].astype(int)
    orginal_pool=pd.merge(orginal_pool,
                          brinson_factor,
                          how='left',on=['jjdm','asofdate'])

    orginal_pool['asofdate']=orginal_pool['asofdate'].astype(str)

    factor_list=['主题集中度', '主题换手率','一级行业集中度', '一级行业换手率', '二级行业集中度',
       '二级行业换手率', '三级行业集中度', '三级行业换手率', '龙头占比(时序均值)','ir','ir_theme'
            ,'个股集中度','持股数量','PE', 'PB', 'ROE', '股息率','PCF', '毛利率', '净利率',  '净利增速','换手率','左侧概率(出重仓前,半年线)','trading','equity_selection','sector_allo','jjzzc','买入 ROE(出重仓前)_rank',
                                     '买入 ROE(出持仓前)_rank', '卖出 ROE(出重仓前)_rank', '卖出 ROE(出持仓前)_rank',
                                     '买入 净利润增速(出重仓前)_rank', '买入 净利润增速(出持仓前)_rank', '卖出 净利润增速(出重仓前)_rank',
                                     '卖出 净利润增速(出持仓前)_rank', '买入 股息率(出重仓前)_rank', '买入 股息率(出持仓前)_rank',
                                     '卖出 股息率(出重仓前)_rank', '卖出 股息率(出持仓前)_rank', '买入 PE(出重仓前)_rank', '买入 PE(出持仓前)_rank',
                                     '卖出 PE(出重仓前)_rank', '卖出 PE(出持仓前)_rank', '买入 PB(出重仓前)_rank', '买入 PB(出持仓前)_rank',
                                     '卖出 PB(出重仓前)_rank', '卖出 PB(出持仓前)_rank', '买入 毛利率(出重仓前)_rank', '买入 毛利率(出持仓前)_rank',
                                     '卖出 毛利率(出重仓前)_rank', '卖出 毛利率(出持仓前)_rank', '买入 净利率(出重仓前)_rank', '买入 净利率(出持仓前)_rank',
                                     '卖出 净利率(出重仓前)_rank', '卖出 净利率(出持仓前)_rank', '买入 PEG(出重仓前)_rank', '买入 PEG(出持仓前)_rank',
                                     '卖出 PEG(出重仓前)_rank', '卖出 PEG(出持仓前)_rank', '买入 PCF(出重仓前)_rank', '买入 PCF(出持仓前)_rank',
                                     '卖出 PCF(出重仓前)_rank', '卖出 PCF(出持仓前)_rank', '买入 ROE-PB比例(出重仓前)_rank', '买入 ROE-PB比例(出持仓前)_rank',
                                     '卖出 ROE-PB比例(出重仓前)_rank', '卖出 ROE-PB比例(出持仓前)_rank', '买入 PEG比例(出重仓前)_rank', '买入 PEG比例(出持仓前)_rank',
                                     '卖出 PEG比例(出重仓前)_rank', '卖出 PEG比例(出持仓前)_rank','逆向买入比例(出持仓前)', '逆向买入比例(出持仓前)_rank', 'baotuan_ratio']
    ic_histor=pd.DataFrame(index=[asofdate2],columns=factor_list)
    ex_ret_histor=pd.DataFrame(index=[asofdate2],columns=factor_list)
    ic_pvalue_histor=pd.DataFrame(columns=factor_list)
    corr=pd.DataFrame(columns=factor_list,index=factor_list,
                      data=np.zeros(len(factor_list)*len(factor_list)).reshape(len(factor_list),len(factor_list)))

    asofdate1=orginal_pool['asofdate'].unique().tolist()
    asofdate2.sort()

    for i in range(len(asofdate2)):

        date=asofdate2[i]

        if(date[4:6]=='03'):
            end_date=date[0:4]+'09'
        else:
            end_date=str(int(date[0:4])+1)+'03'

        initial_pool=orginal_pool[orginal_pool['asofdate']==asofdate1[i]]

        sql = "select avg(hb1y) as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>='{1}' and tjyf<='{2}' and hb1y!=99999 group by tjyf" \
            .format(util.list_sql_condition(initial_pool['jjdm'].tolist())
                    , date[0:6],end_date)
        pool_ret=hbdb.db2df(sql,db='funduser')
        pool_ret['pool_ret'] = (pool_ret['pool_ret'] / 100 + 1)
        pool_ret['pool_ret'] = pool_ret['pool_ret'].cumprod()


        sql="select max(tjrq) as date from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq<={1}  "\
            .format(util.list_sql_condition(initial_pool['jjdm'].tolist()),date)
        infor_rate_date=hbdb.db2df(sql,db='funduser')['date'][0]
        sql = "select jjdm,zbnp as ir from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq={1} and zblb='2201' and zbnp!=99999 ".format(
            util.list_sql_condition(initial_pool['jjdm'].tolist()),infor_rate_date)
        info_ratio=hbdb.db2df(sql,db='funduser')

        initial_pool=pd.merge(initial_pool,info_ratio,how='left',on='jjdm')
        initial_pool = pd.merge(initial_pool,
                                jj_daily_ret[jj_daily_ret['jzrq']==jj_daily_ret[jj_daily_ret['jzrq']<=int(date)]['jzrq'].max()][['jjdm','ir_theme']],
                                how='left', on='jjdm')
        initial_pool['ir']=pp.scale(initial_pool['ir'])
        initial_pool['ir_theme'] = pp.scale(initial_pool['ir_theme'])
        #+initial_pool['trading']
        #get the combined factor
        if(theme_name=='HB医药'):
            if(pool_size==3 or pool_size==5):
                #top5,top3
                initial_pool['final_score'] = \
                    (initial_pool['sector_allo']+initial_pool['主题换手率']+initial_pool['ir']
                     -initial_pool['持股数量']-initial_pool['主题集中度']+initial_pool['换手率'])
            #top20
            else:
                initial_pool['final_score'] = \
                    (-initial_pool['持股数量']+initial_pool['二级行业换手率']+initial_pool['sector_allo']+initial_pool['ir']
                     )
        elif(theme_name=='HB消费'):
            if (pool_size == 3 ):
            #top3
                initial_pool['final_score'] = \
                    initial_pool['三级行业集中度']
            elif(pool_size == 5 ):
            #top5
                initial_pool['final_score'] = \
                    initial_pool['龙头占比(时序均值)']-initial_pool['单位贡献排名']
            else:
            #top20
                initial_pool['final_score'] = \
                    initial_pool['一级行业集中度']-initial_pool['持股数量']+initial_pool['左侧概率(出重仓前,半年线)']
        elif(theme_name=='HBTMT'):
            if(pool_size==5 or pool_size==3):
            #top5
                # initial_pool['final_score'] = \
                #     -initial_pool['主题集中度']-initial_pool['单位贡献排名']+\
                #     initial_pool['三级行业换手率']-initial_pool['一级行业集中度']
                # initial_pool['final_score'] = \
                #     initial_pool['三级行业换手率']+initial_pool['PB']
                initial_pool['final_score'] = \
                    initial_pool['ir_theme'] - initial_pool['jjzzc']-initial_pool['持股数量']
            else:
            # top20
                initial_pool['final_score'] = \
                    initial_pool['ir_theme']-initial_pool['jjzzc']-initial_pool['持股数量'] +initial_pool['股息率']
        elif (theme_name == 'HB大金融'):
            if(pool_size==3):
                #top3
                initial_pool['final_score'] = \
                    (initial_pool['换手率'] )
            elif(pool_size==5):
                initial_pool['final_score'] = \
                    (-initial_pool['个股集中度'] + initial_pool['ir_theme']-initial_pool['jjzzc']+initial_pool['换手率'])
            else:
            # top5 and top20
                initial_pool['final_score'] = \
                    (-initial_pool['龙头占比(时序均值)']+initial_pool['换手率'])
        elif (theme_name == 'HB周期'):
            if(pool_size==3):
            #top3
                initial_pool['final_score'] = \
                    (initial_pool['ir']-initial_pool['龙头占比(时序均值)'])
            elif(pool_size==5):
            #top5
                initial_pool['final_score'] = \
                    (initial_pool['ir'])
            else:
            #top20
                initial_pool['final_score'] = \
                    (initial_pool['ir']-initial_pool['行业配置排名']   )
        elif (theme_name == 'HB必选消费'):
            initial_pool['final_score'] = \
                (initial_pool['个股集中度'] +
                 initial_pool['左侧概率(出重仓前,半年线)']
                 )
        elif (theme_name == 'HB新制造'):
            initial_pool['final_score'] = \
                (initial_pool['PE'] -
                 initial_pool['持股数量']+initial_pool['trading']
                 )
        elif (theme_name == 'HB资源'):
            initial_pool['final_score'] = \
                (
                 initial_pool['trading']+initial_pool['ir']
                 )
        elif (theme_name == 'HB制造'):
            if(pool_size==3):
                #top3
                initial_pool['final_score'] = \
                    (initial_pool['一级行业集中度']
                     )
            elif(pool_size==5):
                #top5
                initial_pool['final_score'] = \
                    (-initial_pool['ROE']+initial_pool['trading']-initial_pool['持股数量']

                     )
            else:
                #top20
                initial_pool['final_score'] = \
                    (initial_pool['trading']-initial_pool['ROE']
                     )
        elif (theme_name == 'HB材料'):
            initial_pool['final_score'] = \
                (initial_pool['ir'] -
                 initial_pool['个股集中度']
                 )
        elif (theme_name == 'TMT'):
            initial_pool['final_score'] = \
                (
                 -initial_pool['baotuan_ratio']+initial_pool['ir_theme']+initial_pool['一级行业换手率']+initial_pool['卖出 PEG比例(出持仓前)_rank']
                 )
        elif (theme_name == '制造'):
            initial_pool['final_score'] = \
                (
                 -initial_pool['jjzzc']+initial_pool['ir_theme']-initial_pool['持股数量']-initial_pool['买入 PB(出重仓前)_rank']
                 )
        elif (theme_name == '消费'):
            initial_pool['final_score'] = \
                (
                 -initial_pool['卖出 PEG(出重仓前)_rank']+initial_pool['ir_theme']+initial_pool['左侧概率(出重仓前,半年线)']
                 )
        elif (theme_name == '医药'):
            initial_pool['final_score'] = \
                (
                 -initial_pool['买入 PE(出重仓前)_rank']+initial_pool['ir_theme']+initial_pool['卖出 净利润增速(出持仓前)_rank']
                 )
        elif (theme_name == '超额稳定'):
            initial_pool['final_score'] = \
                (
                 initial_pool['ir']
                 )
        elif (theme_name == '价值'):
            initial_pool['final_score'] = \
                (
                 initial_pool['ir']
                 )
        elif (theme_name == '成长'):
            initial_pool['final_score'] = \
                (
                 initial_pool['ir']
                 )


        corr+=initial_pool[factor_list].corr()

        sql = "select jjdm,hb1y as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>='{1}' and tjyf<='{2}' and hb1y!=99999 " \
            .format(util.list_sql_condition(initial_pool['jjdm'].tolist())
                    , date[0:6], end_date)
        test_ret1 = hbdb.db2df(sql, db='funduser')
        test_ret1['pool_ret'] = (test_ret1['pool_ret'] / 100 + 1)
        test_ret1['pool_ret'] = test_ret1.groupby('jjdm').cumprod()['pool_ret']
        test_ret1.drop_duplicates('jjdm', keep='last', inplace=True)
        initial_pool = pd.merge(initial_pool, test_ret1[['jjdm', 'pool_ret']], how='left', on='jjdm')

        for factor in ['final_score']:
        # for factor in factor_list:

            initial_pool[factor+'_rank']=initial_pool[factor].rank()
            initial_pool['pool_ret'+'_rank']=initial_pool['pool_ret'].rank()
            ic=initial_pool[[factor+'_rank','pool_ret'+ '_rank']].corr().iloc[0]['pool_ret_rank']
            ic_histor.loc[date,factor]=ic

            initial_pool=initial_pool.sort_values(factor,ascending=False)
            test=initial_pool[0:pool_size]

            sql = "select avg(hb1y) as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>='{1}' and tjyf<='{2}' and hb1y!=99999 group by tjyf" \
                .format(util.list_sql_condition(test['jjdm'].tolist())
                        , date[0:6],end_date)
            test_ret=hbdb.db2df(sql,db='funduser')
            test_ret['pool_ret'] = (test_ret['pool_ret'] / 100 + 1)
            test_ret['pool_ret'] = test_ret['pool_ret'].cumprod()

            ex_ret_histor.loc[date,factor]=test_ret['pool_ret'].iloc[-1]-pool_ret['pool_ret'].iloc[-1]
            if(factor=='final_score'):

                print("date:{0},IC:{1} 超额:{2}".format(date,
                                                      ic,
                                                      test_ret['pool_ret'].iloc[-1]-pool_ret['pool_ret'].iloc[-1]))
                # print(pool_ret['pool_ret'].iloc[-1])

                # print(util.list_sql_condition(test['jjdm'].tolist()))
                if(save_to_db):
                    # localdb.execute("delete from advance_public_theme_pool_history where asofdate='{0}' and pool_name='{1}' "
                    #                 .format(asofdate1[i],theme_name))
                    inputdata=test[['jjdm']]
                    inputdata['asofdate']=asofdate1[i][0:6]
                    inputdata['pool_name'] = theme_name
                    inputdata.to_sql('advance_public_theme_pool_history',con=localdb,index=False,if_exists='append')


    corr=corr/len(asofdate2)
    from scipy.stats import  ttest_ind
    for col in ic_histor.columns:
        ic_pvalue_histor[col]=[ttest_ind(ic_histor[col].tolist(),np.zeros(len(ic_histor))).pvalue]
    ic_histor['theme_name']=theme_name
    ex_ret_histor['theme_name'] = theme_name
    corr['theme_name']=theme_name
    ic_pvalue_histor['theme_name']=theme_name

    return  ic_histor,ex_ret_histor,corr,ic_pvalue_histor

def all_market_funds_factor_test(asofdate1,asofdate2,theme_name,pool_size,save_to_db=False):

    orginal_pool=pd.read_sql("SELECT * from theme_pool where theme='{0}' and asofdate in ({1}) "
                             .format(theme_name,util.list_sql_condition(asofdate1)),con=localdb)
    asofdate_map=dict(zip(asofdate1,asofdate2))
    asofdate1=orginal_pool['asofdate'].unique().tolist()
    asofdate2=[asofdate_map[x] for x in asofdate1]


    # get the style and size pic
    sql = "select * from jjpic_size_p_hbs where jjdm in ({0}) and asofdate>='{1}' and asofdate<='{2}' " \
        .format(util.list_sql_condition(orginal_pool['jjdm'].tolist()), pd.Series(asofdate1).min(),pd.Series(asofdate2).max())
    size_pic = pd.read_sql(sql, con=localdb)

    sql = "select * from jjpic_value_p_hbs where jjdm in ({0}) and asofdate>='{1}' and asofdate<='{2}' " \
        .format(util.list_sql_condition(orginal_pool['jjdm'].tolist()), pd.Series(asofdate1).min(),pd.Series(asofdate2).max())
    style_pic = pd.read_sql(sql, con=localdb)

    style_pic = pd.merge(style_pic, size_pic.drop(['基金简称'],axis=1), how='inner', on=['jjdm', 'asofdate'])


    # get the stock factor
    sql="SELECT * from jjpic_stock_p where asofdate >='{0}'  and jjdm in ({1})"\
        .format(pd.Series(asofdate1).min(),util.list_sql_condition(orginal_pool['jjdm'].tolist()))
    stock_pic=pd.read_sql(sql,con=localdb)
    sql="SELECT * from jjpic_stock_tp where asofdate >='{0}'  and jjdm in ({1})"\
        .format(pd.Series(asofdate1).min(),util.list_sql_condition(orginal_pool['jjdm'].tolist()))
    stock_tp_pic=pd.read_sql(sql,con=localdb)

    stock_pic=pd.merge(stock_pic[['jjdm','asofdate','个股集中度','持股数量','PE', 'PB', 'ROE', '股息率']],
                       stock_tp_pic[['jjdm','asofdate','换手率','左侧概率(出重仓前,半年线)']],
                       how='inner',on=['jjdm','asofdate'])

    #get thr brinson result
    sql="select jjdm,tjrq as asofdate,sector_allo,equity_selection,trading from st_fund.r_st_hold_excess_attr_df where tjrq>='{0}' and jjdm in ({1})"\
        .format(pd.Series(asofdate1).min(),util.list_sql_condition(orginal_pool['jjdm'].tolist()))
    brinson_factor=hbdb.db2df(sql,db='funduser')


    #get the fund size
    sql="select jjdm,jsrq,jjzzc from st_fund.t_st_gm_zcpz where jjdm in ({0}) and jsrq >='{1}' and jsrq<='{2}'"\
        .format(util.list_sql_condition(orginal_pool['jjdm'].tolist()),pd.Series(asofdate1).min(),pd.Series(asofdate1).max())
    jjzzc=hbdb.db2df(sql,db='funduser')
    stock_pic['ym']=stock_pic['asofdate'].str[0:6]
    jjzzc['ym'] = jjzzc['jsrq'].astype(str).str[0:6]
    stock_pic=pd.merge(stock_pic,jjzzc,how='left',on=['jjdm','ym'])

    #get the theme infor ratio

    sql="select jzrq,jjdm,hbdr from st_fund.t_st_gm_rhb where jzrq<='{0}' and jjdm in ({1}) and jzrq>='{2}' "\
        .format(pd.Series(asofdate2).max(),util.list_sql_condition(orginal_pool['jjdm'].tolist())
                ,str(int(pd.Series(asofdate2).min()[0:4])-1)+pd.Series(asofdate2).min()[4:])
    jj_daily_ret=hbdb.db2df(sql,db='funduser')


    index_name_map=dict(zip(['大盘成长','中小盘成长','大盘价值','中小盘价值','均衡均衡'],
                            ['399372','399376','399373','399377','885001']))
    sql = "select jyrq as jzrq,hbdr from st_market.t_st_zs_rhb where jyrq<={0} and jyrq>='{1}' and zqdm='{2}'" \
        .format(pd.Series(asofdate2).max(),
                str(int(pd.Series(asofdate2).min()[0:4])-1)+pd.Series(asofdate2).min()[4:],
                index_name_map[theme_name])
    index_ret = hbdb.db2df(sql, db='alluser')

    rolling_window=120
    jj_daily_ret=pd.merge(jj_daily_ret,index_ret,how='inner',on='jzrq').sort_values(['jjdm','jzrq'])
    jj_daily_ret['ext_ret'] = jj_daily_ret['hbdr_x'] - jj_daily_ret['hbdr_y']

    jj_daily_ret['ir_style'] = (jj_daily_ret.groupby('jjdm')[['ext_ret']].rolling(rolling_window).mean() / jj_daily_ret.groupby('jjdm')[
        ['ext_ret']].rolling(rolling_window).std())['ext_ret'].tolist()


    #standarlize all the factor data
    def scale_by_group(arr):

       return pd.Series(data=pp.scale(arr),index=arr.index)

    for col in ['集中度(持仓)_x', '换手率(持仓)_x',
       '成长绝对暴露(持仓)', '价值绝对暴露(持仓)', '集中度排名(持仓)_x', '换手率排名(持仓)_x', '成长暴露排名(持仓)',
       '价值暴露排名(持仓)', '集中度(持仓)_y',
       '换手率(持仓)_y', '大盘绝对暴露(持仓)', '中小盘绝对暴露(持仓)', '集中度排名(持仓)_y', '换手率排名(持仓)_y',
       '大盘暴露排名(持仓)', '中小盘暴露排名(持仓)']:
        style_pic[col] = style_pic.groupby('asofdate')[col].apply(scale_by_group)

    for col in ['个股集中度','持股数量','PE', 'PB', 'ROE', '股息率','换手率','左侧概率(出重仓前,半年线)','jjzzc']:
        stock_pic[col] = stock_pic.groupby('asofdate')[col].apply(scale_by_group)
    for col in ['trading','equity_selection','sector_allo']:
        brinson_factor[col]=brinson_factor.groupby('asofdate')[col].apply(scale_by_group).values


    # asofdate_map=dict(zip(theme_pic['asofdate'].sort_values().unique().tolist()
    #                       ,orginal_pool['asofdate'].sort_values().unique().tolist()))
    # theme_pic['asofdate']=[asofdate_map[x] for x in theme_pic['asofdate']]

    orginal_pool=pd.merge(orginal_pool,
                          style_pic[['jjdm','asofdate','集中度(持仓)_x', '换手率(持仓)_x',
       '成长绝对暴露(持仓)', '价值绝对暴露(持仓)', '集中度排名(持仓)_x', '换手率排名(持仓)_x', '成长暴露排名(持仓)',
       '价值暴露排名(持仓)', '集中度(持仓)_y',
       '换手率(持仓)_y', '大盘绝对暴露(持仓)', '中小盘绝对暴露(持仓)', '集中度排名(持仓)_y', '换手率排名(持仓)_y',
       '大盘暴露排名(持仓)', '中小盘暴露排名(持仓)']],
                          how='left',on=['jjdm','asofdate'])
    orginal_pool=pd.merge(orginal_pool,
                          stock_pic[['jjdm','asofdate','个股集中度','持股数量','PE',
                                     'PB', 'ROE', '股息率','换手率','左侧概率(出重仓前,半年线)','jjzzc']],
                          how='left',on=['jjdm','asofdate'])
    orginal_pool=pd.merge(orginal_pool,
                          brinson_factor,
                          how='left',on=['jjdm','asofdate'])



    factor_list=['集中度(持仓)_x', '换手率(持仓)_x', '成长绝对暴露(持仓)', '价值绝对暴露(持仓)',
       '集中度排名(持仓)_x', '换手率排名(持仓)_x', '成长暴露排名(持仓)', '价值暴露排名(持仓)', '集中度(持仓)_y',
       '换手率(持仓)_y', '大盘绝对暴露(持仓)', '中小盘绝对暴露(持仓)', '集中度排名(持仓)_y', '换手率排名(持仓)_y',
       '大盘暴露排名(持仓)', '中小盘暴露排名(持仓)','ir','ir_style'
            ,'个股集中度','持股数量','PE', 'PB', 'ROE', '股息率','换手率','左侧概率(出重仓前,半年线)','trading','equity_selection','sector_allo','jjzzc']
    ic_histor=pd.DataFrame(index=[asofdate2],columns=factor_list)
    ex_ret_histor=pd.DataFrame(index=[asofdate2],columns=factor_list)
    ic_pvalue_histor=pd.DataFrame(columns=factor_list)
    corr=pd.DataFrame(columns=factor_list,index=factor_list,
                      data=np.zeros(len(factor_list)*len(factor_list)).reshape(len(factor_list),len(factor_list)))



    for i in range(len(asofdate2)):

        date=asofdate2[i]

        if(date[4:6]=='03'):
            end_date=date[0:4]+'09'
        else:
            end_date=str(int(date[0:4])+1)+'03'

        initial_pool=orginal_pool[orginal_pool['asofdate']==asofdate1[i]]

        sql = "select avg(hb1y) as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>='{1}' and tjyf<='{2}' and hb1y!=99999 group by tjyf" \
            .format(util.list_sql_condition(initial_pool['jjdm'].tolist())
                    , date[0:6],end_date)
        pool_ret=hbdb.db2df(sql,db='funduser')
        pool_ret['pool_ret'] = (pool_ret['pool_ret'] / 100 + 1)
        pool_ret['pool_ret'] = pool_ret['pool_ret'].cumprod()


        sql="select max(tjrq) as date from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq<={1}  "\
            .format(util.list_sql_condition(initial_pool['jjdm'].tolist()),date)
        infor_rate_date=hbdb.db2df(sql,db='funduser')['date'][0]
        sql = "select jjdm,zbnp as ir from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq={1} and zblb='2201' and zbnp!=99999 ".format(
            util.list_sql_condition(initial_pool['jjdm'].tolist()),infor_rate_date)
        info_ratio=hbdb.db2df(sql,db='funduser')

        initial_pool=pd.merge(initial_pool,info_ratio,how='left',on='jjdm')
        initial_pool = pd.merge(initial_pool,
                                jj_daily_ret[jj_daily_ret['jzrq']==jj_daily_ret[jj_daily_ret['jzrq']<=int(date)]['jzrq'].max()][['jjdm','ir_style']],
                                how='left', on='jjdm')
        initial_pool['ir']=pp.scale(initial_pool['ir'])
        initial_pool['ir_style'] = pp.scale(initial_pool['ir_style'])
        #+initial_pool['trading']
        #get the combined factor
        if(theme_name=='中小盘成长'):
            if( pool_size==5):
                #top5,top3
                initial_pool['final_score'] = \
                    (initial_pool['换手率']-initial_pool['equity_selection'])

        elif(theme_name=='中小盘价值'):
            if(pool_size==5):
                #top5,top3
                initial_pool['final_score'] = \
                    (initial_pool['ir_style']-initial_pool['ROE']+initial_pool['PE'])

        elif(theme_name=='均衡均衡'):
            if(pool_size==5):
                #top5,top3
                initial_pool['final_score'] = \
                    (initial_pool['ir_style']-initial_pool['jjzzc'])

        elif(theme_name=='大盘价值'):
            if(pool_size==5):
                #top5,top3
                initial_pool['final_score'] = \
                    (initial_pool['ir_style']-initial_pool['集中度(持仓)_y'])

        elif(theme_name=='大盘成长'):
            if(pool_size==5):
                #top5,top3
                initial_pool['final_score'] = \
                    initial_pool['ir_style']


        corr+=initial_pool[factor_list].corr()

        sql = "select jjdm,hb1y as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>='{1}' and tjyf<='{2}' and hb1y!=99999 " \
            .format(util.list_sql_condition(initial_pool['jjdm'].tolist())
                    , date[0:6], end_date)
        test_ret1 = hbdb.db2df(sql, db='funduser')
        test_ret1['pool_ret'] = (test_ret1['pool_ret'] / 100 + 1)
        test_ret1['pool_ret'] = test_ret1.groupby('jjdm').cumprod()['pool_ret']
        test_ret1.drop_duplicates('jjdm', keep='last', inplace=True)
        initial_pool = pd.merge(initial_pool, test_ret1[['jjdm', 'pool_ret']], how='left', on='jjdm')

        for factor in ['final_score']:
        #for factor in factor_list:

            initial_pool[factor+'_rank']=initial_pool[factor].rank()
            initial_pool['pool_ret'+'_rank']=initial_pool['pool_ret'].rank()
            ic=initial_pool[[factor+'_rank','pool_ret'+ '_rank']].corr().iloc[0]['pool_ret_rank']
            ic_histor.loc[date,factor]=ic

            initial_pool=initial_pool.sort_values(factor,ascending=False)
            test=initial_pool[0:pool_size]

            sql = "select avg(hb1y) as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>='{1}' and tjyf<='{2}' and hb1y!=99999 group by tjyf" \
                .format(util.list_sql_condition(test['jjdm'].tolist())
                        , date[0:6],end_date)
            test_ret=hbdb.db2df(sql,db='funduser')
            test_ret['pool_ret'] = (test_ret['pool_ret'] / 100 + 1)
            test_ret['pool_ret'] = test_ret['pool_ret'].cumprod()

            ex_ret_histor.loc[date,factor]=test_ret['pool_ret'].iloc[-1]-pool_ret['pool_ret'].iloc[-1]
            if(factor=='final_score'):

                import warnings
                warnings.filterwarnings('ignore')
                print("date:{0},IC:{1} 超额:{2}".format(date,
                                                      ic,
                                                      test_ret['pool_ret'].iloc[-1]-pool_ret['pool_ret'].iloc[-1]))
                # print(pool_ret['pool_ret'].iloc[-1])

                # print(util.list_sql_condition(test['jjdm'].tolist()))
                if(save_to_db):
                    localdb.execute("delete from pool_history where asofdate='{0}' and pool_name='{1}' "
                                    .format(asofdate1[i],theme_name))
                    inputdata=test[['jjdm']]
                    inputdata['asofdate']=asofdate1[i]
                    inputdata['pool_name'] = theme_name
                    inputdata.to_sql('pool_history',con=localdb,index=False,if_exists='append')


    corr=corr/len(asofdate2)
    from scipy.stats import  ttest_ind
    for col in ic_histor.columns:
        ic_pvalue_histor[col]=[ttest_ind(ic_histor[col].tolist(),np.zeros(len(ic_histor))).pvalue]
    ic_histor['theme_name']=theme_name
    ex_ret_histor['theme_name'] = theme_name
    corr['theme_name']=theme_name
    ic_pvalue_histor['theme_name']=theme_name

    return  ic_histor,ex_ret_histor,corr,ic_pvalue_histor

def get_the_theme_pool(asofdate1,asofdate2,theme_name,pool_size):


    initial_pool=pd.read_sql("SELECT * from theme_pool where theme='{0}' and asofdate='{1}' ORDER BY `综合因子` DESC"
                             .format(theme_name,asofdate1),con=localdb)

    if(asofdate2[4:6]=='03'):
        end_date=asofdate2[0:4]+'09'
    else:
        end_date=str(int(asofdate2[0:4])+1)+'03'

    sql = "select avg(hb1y) as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' and hb1y!=99999 group by tjyf" \
        .format(util.list_sql_condition(initial_pool['jjdm'].tolist())
                , asofdate2[0:6],end_date)
    pool_ret=hbdb.db2df(sql,db='funduser')
    pool_ret['pool_ret'] = (pool_ret['pool_ret'] / 100 + 1)
    pool_ret['pool_ret'] = pool_ret['pool_ret'].cumprod()

    # filter out fund that in the past 3 years rank at least twice at the down 30%
    # sql = "select jjdm,zbnp as info_ratio,jzrq,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq>='{1}' and jzrq <='{2}' and zblb in ('2998') and zbnp!=99999 " \
    #     .format(util.list_sql_condition(initial_pool['jjdm'].tolist()), asofdate2[0:6] + '01', asofdate2, '2203')
    # target_df = hbdb.db2df(sql, db='funduser').sort_values('jzrq').drop_duplicates(['jjdm', 'zblb'], keep='last')[
    #     ['jjdm', 'info_ratio', 'zblb']]
    #
    # sql = "select jjdm,hb1n as info_ratio,tjnf as zblb from st_fund.t_st_gm_nhb where jjdm in ({0}) and tjnf>='{1}' and tjnf<='{2}' and hb1n!=99999 " \
    #     .format(util.list_sql_condition(initial_pool['jjdm'].tolist()), int(asofdate2[0:4])-2, asofdate2[0:4])
    # target_df=pd.concat([target_df,hbdb.db2df(sql, db='funduser')],axis=0)
    #
    # target_df['rank'] = target_df.groupby('zblb').rank()
    # target_df=pd.merge(target_df,target_df.groupby('zblb').count()['jjdm'].to_frame('count'),
    #                    how='left',on='zblb')
    # target_df['rank']=target_df['rank']/target_df['count']
    # target_df.loc[target_df['rank']<0.3,'rank']=0
    # target_df.loc[target_df['rank'] >= 0.3, 'rank'] = 1
    # new_jjdm_list=(target_df.groupby('jjdm').sum()[target_df.groupby('jjdm').sum()['rank']>=3]).index.tolist()
    # initial_pool=initial_pool[initial_pool['jjdm'].isin(new_jjdm_list)]


    sql = "SELECT max(asofdate) as maxdate from jjpic_theme_p "
    max_date=pd.read_sql(sql,con=localdb)['maxdate'][0]

    sql="SELECT * from jjpic_theme_p where asofdate='{0}' and jjdm in ({1})"\
        .format(max_date,util.list_sql_condition(initial_pool['jjdm'].tolist()))
    theme_pic=pd.read_sql(sql,con=localdb)
    theme_pic['max_theme']=theme_pic[['大金融', '消费', 'TMT', '周期', '制造']].max(axis=1)
    sql="select jjdm,`一级行业集中度`,`三级行业类型`  from jjpic_industry_p where asofdate='{0}' and jjdm in ({1}) "\
        .format(max_date,util.list_sql_condition(initial_pool['jjdm'].tolist()))
    ind_pic=pd.read_sql(sql,con=localdb)
    theme_pic=pd.merge(theme_pic,ind_pic,how='inner',on='jjdm')
    theme_pic['theme_stable_rank'] =pp.scale(theme_pic['主题换手率'])
    theme_pic['一级行业集中度'] = pp.scale(theme_pic['一级行业集中度'])

    sql="select max(tjrq) as date from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq<={1}  "\
        .format(util.list_sql_condition(initial_pool['jjdm'].tolist()),asofdate2)
    infor_rate_date=hbdb.db2df(sql,db='funduser')['date'][0]
    sql = "select jjdm,zbnp as ir from st_fund.t_st_gm_zqjxxbl where jjdm in ({0}) and tjrq={1} and zblb='2201' ".format(
        util.list_sql_condition(initial_pool['jjdm'].tolist()),infor_rate_date)
    info_ratio=hbdb.db2df(sql,db='funduser')

    initial_pool=\
        pd.merge(initial_pool,theme_pic[['jjdm','theme_stable_rank','一级行业集中度']],how='left',on='jjdm')
    initial_pool=pd.merge(initial_pool,info_ratio,how='left',on='jjdm')
    initial_pool['ir']=pp.scale(initial_pool['ir'])


    initial_pool['final_socre']=\
        initial_pool['综合因子']*0.7+initial_pool['theme_stable_rank']*0.2+initial_pool['一级行业集中度']*0.1

    sql = "select jjdm,hb1y as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' and hb1y!=99999 " \
        .format(util.list_sql_condition(initial_pool['jjdm'].tolist())
                , asofdate2[0:6], end_date)
    test_ret1 = hbdb.db2df(sql, db='funduser')
    test_ret1['pool_ret'] = (test_ret1['pool_ret'] / 100 + 1)
    test_ret1['pool_ret'] = test_ret1.groupby('jjdm').cumprod()['pool_ret']
    test_ret1.drop_duplicates('jjdm', keep='last', inplace=True)
    initial_pool = pd.merge(initial_pool, test_ret1[['jjdm', 'pool_ret']], how='left', on='jjdm')

    for factor in ['theme_stable_rank','一级行业集中度','单位贡献排名','行业配置排名','综合因子','ir']:

        initial_pool[factor+'_rank']=initial_pool[factor].rank()
        initial_pool['pool_ret'+'_rank']=initial_pool['pool_ret'].rank()
        print('the ic of {0} is {1}'.format(factor,
                                            initial_pool[[factor+'_rank','pool_ret'+ '_rank']].corr().iloc[0]['pool_ret_rank']))

        test=initial_pool.sort_values(factor,ascending=False)[0:pool_size]


        sql = "select avg(hb1y) as pool_ret,tjyf from st_fund.t_st_gm_yhb where jjdm in ({0}) and tjyf>'{1}' and tjyf<='{2}' and hb1y!=99999 group by tjyf" \
            .format(util.list_sql_condition(test['jjdm'].tolist())
                    , asofdate2[0:6],end_date)
        test_ret=hbdb.db2df(sql,db='funduser')
        test_ret['pool_ret'] = (test_ret['pool_ret'] / 100 + 1)
        test_ret['pool_ret'] = test_ret['pool_ret'].cumprod()

        print("the size of the pool is {2} pool_ret is {0} and selected pool ret is {1}".format(pool_ret['pool_ret'].iloc[-1]
                                                                    ,test_ret['pool_ret'].iloc[-1],len(test)))


    # localdb.execute("delete from pool_history where asofdate='{0}' and pool_name='{1}'"
    #                 .format(test['asofdate'].iloc[0],theme_name))
    # test['pool_name']=theme_name
    # test[['pool_name','jjdm','asofdate']].to_sql('pool_history',index=False,con=localdb,if_exists='append')

def save_the_hbtheme_index2db(start_date,end_date):


    theme_col = ['大金融', '消费', 'TMT', '周期', '制造','医药']
    theme_map = dict(zip(theme_col,
                              [['银行', '非银金融', '房地产'],
                               ['食品饮料', '家用电器', '社会服务', '农林牧渔', '商贸零售', '美容护理'],
                               ['通信', '计算机', '电子', '传媒', '国防军工'],
                               ['钢铁', '有色金属', '建筑装饰', '建筑材料', '基础化工', '石油石化', '煤炭', '环保',
                                '公用事业'],
                               ['交通运输', '机械设备', '汽车', '纺织服饰', '轻工制造', '电力设备'],[ '医药生物',]
                               ]
                              ))

    lista = []
    listb = []
    for theme in theme_col:
        for col in theme_map[theme]:
            lista.append(col)
            listb.append(theme)
    ind2thememap = pd.DataFrame()
    ind2thememap['industry_name'] = lista
    ind2thememap['theme'] = listb

    new_col_name = ['农林牧渔', '基础化工', '钢铁', '有色金属', '电子', '家用电器', '食品饮料',
                    '纺织服饰', '轻工制造', '医药生物', '公用事业', '交通运输', '房地产', '商贸零售',
                    '社会服务', '综合', '建筑材料', '建筑装饰', '电力设备', '国防军工', '计算机', '传媒',
                    '通信', '银行', '非银金融', '汽车', '机械设备', '煤炭', '石油石化', '环保', '美容护理']
    index_code_list = ['801010', '801030', '801040', '801050', '801080', '801110',
                            '801120', '801130', '801140', '801150', '801160', '801170',
                            '801180', '801200', '801210', '801230', '801710', '801720',
                            '801730', '801740', '801750', '801760', '801770', '801780',
                            '801790', '801880', '801890', '801950', '801960', '801970', '801980']
    index_code_map = dict(zip(index_code_list, new_col_name))


    jjdm_list=util.get_stock_funds_pool(end_date)


    #read the mutual fund industry alloocation inside each theme
    sql="select flmc,jsrq,avg(zzjbl) as zzjbl from st_fund.t_st_gm_jjhyzhyszb where hyhfbz='2' and zclb='2' and jjdm in ({0}) and jsrq>='{1}' and jsrq<='{2}' group by jsrq,flmc "\
        .format(util.list_sql_condition(jjdm_list),start_date,end_date)
    fund_industry_allocation=hbdb.db2df(sql,db='funduser')

    fund_industry_allocation=pd.merge(fund_industry_allocation,ind2thememap,
                                      how='left',left_on='flmc',right_on='industry_name').pivot_table('zzjbl','jsrq','flmc')

    new_inde=[]
    for date in fund_industry_allocation.index.tolist():
        if('1231' in str(date)):
            new_inde.append( str(int(str(date)[0:4])+1)+"0401" )
        else:
            new_inde.append(str(int(str(date)[0:4])) + "0901")
    fund_industry_allocation.index=new_inde

    # get the industry nav

    sql = "select zqdm,jyrq,spjg from st_market.t_st_zs_hqql where zqdm in ({0}) and jyrq>='{1}' and jyrq<='{2}'  " \
        .format(util.list_sql_condition(index_code_list),start_date,end_date)
    navdf = hbdb.db2df(sql, db='alluser').pivot_table('spjg','jyrq','zqdm')
    navdf=navdf.pct_change()

    navdf.rename(columns=index_code_map,inplace=True)

    final_df=[]

    for i in range(len(fund_industry_allocation)-1):

        sdate=fund_industry_allocation.index[i]
        edate=fund_industry_allocation.index[i+1]

        if(i>=3):
            tempallocation = fund_industry_allocation.loc[fund_industry_allocation.index[i-2]:sdate].mean(axis=0).to_frame('w').fillna(0)
        else:
            tempallocation=fund_industry_allocation.loc[:sdate].mean(axis=0).to_frame('w').fillna(0)

        tempnav=navdf.loc[(navdf.index>=int(sdate))&(navdf.index<int(edate))]

        tempdf=\
            pd.DataFrame(columns=theme_col,index=tempnav.index.tolist())

        for theme in theme_col:

            theme_ret=pd.Series(index=tempnav.index.tolist()).fillna(0)
            total_w=\
                tempallocation.loc[theme_map[theme]].sum().iloc[0]

            for col in theme_map[theme]:
                theme_ret+=tempnav[col] * tempallocation.loc[col].iloc[0]

            theme_ret=theme_ret/total_w
            tempdf[theme]=theme_ret

        final_df.append(tempdf)

    final_df=pd.concat(final_df,axis=0)

    #localdb.execute("delete from theme_daily_ret ")
    final_df.reset_index(drop=False).rename(columns={'index':'jzrq'}).to_sql('theme_daily_ret',con=localdb
                    ,if_exists='append',index=False)

class recomand_pool_mothly_report:

    @staticmethod
    def get_the_pool_data(date,pool_list):

        sql = "SELECT distinct(asofdate) as asofdate from pool_history where pool_name in ({0}) and asofdate<='{1}' order by asofdate"\
            .format(util.list_sql_condition(pool_list),date)
        date_list=pd.read_sql(sql,con=localdb)['asofdate'].tolist()

        date_list=date_list[-4:]

        date_map=dict(zip(['20221230','20220630','20211231','20210630','20201231','20200630','20191231','20190628','20181228','20180629','20171229'],
                          ['20230331','20220930','20220330','20210930','20210330','20200930','20200330','20190930','20190330','20180930','20180331']))

        sql="SELECT * from pool_history where pool_name in ({0}) and asofdate in ({1}) "\
            .format(util.list_sql_condition(pool_list),util.list_sql_condition(date_list))
        data=pd.read_sql(sql,con=localdb)

        data['asofdate']=[ date_map[x] for x in data['asofdate']]

        date_list=data['asofdate'].sort_values(ascending=False).unique().tolist()
        max_asofdate=date_list[0]

        outputdict=[]

        for asofdate in date_list:
            # print(asofdate)
            tempdata= data[data['asofdate']==asofdate]
            temp_list=tempdata['jjdm'].tolist()

            sql="select jjdm,zbnp/100 as zbnp,zblb from st_fund.t_st_gm_rqjhb where jjdm in ({0}) and jzrq='{1}' and zblb in ('2007','2101','2103','2106','2998')"\
                .format(util.list_sql_condition(temp_list),date)
            ret_data=hbdb.db2df(sql,db='funduser').pivot_table('zbnp','jjdm','zblb')

            ret_data.rename(columns={'2007':'最近一周',
                                     '2101': '最近一月',
                                     '2103':'最近三月',
                                     '2106': '最近六月',
                                     '2998': '今年以来'},inplace=True)


            sql="select max(tjrq) as tjrq from st_fund.t_st_gm_zqjbdl where tjrq>='{0}' and  tjrq<='{1}' "\
                .format((datetime.datetime.strptime(date, '%Y%m%d') - datetime.timedelta(days=30)).strftime('%Y%m%d')
                        ,date)
            weekend_data=hbdb.db2df(sql,db='funduser')['tjrq'].iloc[0]

            sql = "select jjdm,nhzbnp/100 as '近一年波动率' from st_fund.t_st_gm_zqjbdl where jjdm in ({0}) and tjrq='{1}' and  zblb='2201' "\
                .format(util.list_sql_condition(temp_list),weekend_data)
            vol=hbdb.db2df(sql,db='funduser')

            sql = "select jjdm,nhzbnp as '近一年夏普' from st_fund.t_st_gm_zqjxpbl where jjdm in ({0}) and tjrq='{1}' and  zblb='2201' "\
                .format(util.list_sql_condition(temp_list),weekend_data)
            sharp=hbdb.db2df(sql,db='funduser')

            sql = "select jjdm,zbnp/100 as '近一年最大回撤' from st_fund.t_st_gm_rqjzdhc where jjdm in ({0}) and jzrq='{1}' and  zblb='2201' "\
                .format(util.list_sql_condition(temp_list),date)
            max_draw_back=hbdb.db2df(sql,db='funduser')


            if(asofdate!=max_asofdate):

                end_date=date_list[date_list.index(asofdate)-1]
            else:
                end_date=date

            sql = "select jjdm,jzrq,ljjz from st_fund.t_st_gm_jjjz where jjdm in ({0}) and jzrq in ('{1}','{2}')" \
                .format(util.list_sql_condition(temp_list), asofdate, end_date)
            pool_ret = hbdb.db2df(sql, db='funduser').pivot_table('ljjz', 'jjdm', 'jzrq')
            pool_ret['入池以来'] = pool_ret[int(end_date)] / pool_ret[int(asofdate)] - 1

            ret_data=pd.merge(ret_data,vol,how='left',on='jjdm')
            ret_data = pd.merge(ret_data, sharp, how='left', on='jjdm')
            ret_data = pd.merge(ret_data, max_draw_back, how='left', on='jjdm')
            ret_data = pd.merge(ret_data, pool_ret[['入池以来']], how='left', on='jjdm')
            ret_data=pd.merge(ret_data,tempdata[['jjdm','pool_name']],how='left',on='jjdm')
            ret_data['asofdate']=asofdate
            outputdict.append(ret_data)

        outputdict=pd.concat(outputdict,axis=0)
        sql="select jjdm,jjjc as '基金简称' from st_fund.t_st_gm_jjxx where jjdm in ({0})"\
            .format(util.list_sql_condition(outputdict['jjdm'].unique().tolist()))
        outputdict=pd.merge(outputdict,hbdb.db2df(sql,db='funduser'))


        return outputdict

    def update_pool_report_file(self,date,pool_list):

        data=\
            self.get_the_pool_data(date,pool_list)

        asofdate_list=data['asofdate'].sort_values(ascending=False).unique().tolist()
        import xlwings as xw
        filename = "E:\GitFolder\docs\个基研究\主题基金研究\主题基金池统计.xlsx"
        app = xw.App(visible=False)
        wb = app.books.open(filename)
        sheet_list=data['pool_name'].unique().tolist()

        for i in range(len(sheet_list)):
            pool_name=sheet_list[i]
            ws = wb.sheets[sheet_list[i]+"池历史"]
            x_loc=1
            for date in asofdate_list:
                tempdata=data[(data['asofdate']==date)&(data['pool_name']==pool_name)]
                tempdata.rename(columns={'asofdate':'入池时间'},inplace=True)
                tempdata['jjdm']=["'"+str(x) for x in tempdata['jjdm']]
                ws["A"+str(x_loc)].options(pd.DataFrame,
                                           header=1,
                                           index=False,
                                           expand='table').value = \
                    tempdata[['入池时间','jjdm','基金简称','最近一周', '最近一月', '最近三月', '最近六月','入池以来', '今年以来', '近一年波动率', '近一年夏普',
       '近一年最大回撤' ]]
                x_loc+=22

            print('{} done'.format(pool_name))

        wb.save(filename)
        wb.close()
        app.quit()


if __name__ == '__main__':

    # future_star_funds().filter('20231229',0.8,20,365*2)

    # ptm=Public_theme_and_all_market_fund()
    # ind_dis=pd.DataFrame(columns=['flmc'])
    # theme_dis=pd.DataFrame(columns=['theme'])
    # for i in range(8):
    #
    #     jjdm_list=util.get_885001_funds(str(2022-i)+'1231')
    #     sql = "select flmc,sum(zgpbl)/{4} as '{3}' from st_fund.t_st_gm_jjhyzhyszb where jjdm in ({0}) and jsrq>='{1}' and jsrq<='{2}' and hyhfbz='2' and zclb='2' group by flmc " \
    #         .format(util.list_sql_condition(jjdm_list), str(2022-i)+'0701', str(2022-i)+'1231',str(2022-i)+'1231',len(jjdm_list))
    #     temp_ind_dis = hbdb.db2df(sql, db='funduser')
    #     ind_dis = pd.merge(ind_dis,temp_ind_dis,how='outer',on='flmc')
    #     theme_dis=pd.merge(theme_dis,pd.merge(temp_ind_dis, ptm.ind2thememap, how='left', left_on='flmc', right_on='industry_name').groupby([ 'theme']).sum()[str(2022-i)+'1231'].reset_index(),
    #                        how='outer',on='theme')
    #
    #     jjdm_list=util.get_885001_funds(str(2022-i)+'0630')
    #     sql = "select flmc,sum(zgpbl)/{4} as '{3}' from st_fund.t_st_gm_jjhyzhyszb where jjdm in ({0}) and jsrq>='{1}' and jsrq<='{2}' and hyhfbz='2' and zclb='2' group by flmc " \
    #         .format(util.list_sql_condition(jjdm_list), str(2022-i)+'0101', str(2022-i)+'0630',str(2022-i)+'0630',len(jjdm_list))
    #     temp_ind_dis = hbdb.db2df(sql, db='funduser')
    #     ind_dis = pd.merge(ind_dis,temp_ind_dis,how='outer',on='flmc')
    #     theme_dis=pd.merge(theme_dis,pd.merge(temp_ind_dis, ptm.ind2thememap, how='left', left_on='flmc', right_on='industry_name').groupby([ 'theme']).sum()[str(2022-i)+'0630'].reset_index(),
    #                        how='outer',on='theme')


    # ind_dis.to_excel('885001_行业.xlsx')
    # theme_dis.to_excel('885001_主题.xlsx')
    # jjdm_list=util.get_all_mutual_stock_funds('20230630')
    # sql = "select jsrq,jjdm,flmc,zgpbl from st_fund.t_st_gm_jjhyzhyszb where jjdm in ({0}) and jsrq>='{1}' and jsrq<='{2}' and hyhfbz='2' and zclb='2' " \
    #     .format(util.list_sql_condition(jjdm_list), '20150101', '20221231')
    # ind_dis = hbdb.db2df(sql, db='funduser')
    # ind_dis.to_pickle('行业数据')

    #
    # ptm=Public_theme_and_all_market_fund()
    # ptm.save_theme_and_allmarket_fund2db('20210701', '20230630')
    #ptm.get_public_advance_theme_pools('20230831')
    #
    # atf=Agent_theme_fund()
    # atf.save_theme_and_allmarket_fund2db('20210101','20221231')
    # atf.get_public_advance_theme_pools('20230731')



    #save_the_hbtheme_index2db('20161231','20230531')

    # jjdm_list=util.get_stock_funds_pool('20221230')
    # sql="select * from jjpic_theme_p where jjdm in ({0}) and asofdate='20230331'"\
    #     .format(util.list_sql_condition(jjdm_list))
    # sql="SELECT zjbl/know_weight as zsbl from hbs_industry_class1_exp where jsrq='20221230' and jjdm in ({}) and yjxymc='医药生物'"\
    #     .format(util.list_sql_condition(jjdm_list))
    # data=pd.read_sql(sql,con=localdb)
    # pr=recomand_pool_mothly_report()
    # pr.update_pool_report_file('20230531',
    #                                ['HB周期','HBTMT','HB医药','HB大金融', 'HB消费', 'HB制造'])

    # plot=functionality.Plot(1000,500)
    #





    #
    # ir=Industry_rotating()
    # industry_roation=ir.get_industry_rotating(shift_count_rank_thresheld=0.6
    #                           ,shift_ret_thresheld=0.6,shift_winpro_thresheld=0.6
    #                                           ,left_trade_ratio=0.4,pool_size=20)

    #sf=Similarity_filter()
    #similarity=sf.get_similarity('006624')
    #similarity = sf.get_similarity('S66391',mu_pool_type=True,mu_target_type=False)
    #similarity.to_excel('similarity_of_{}.xlsx'.format('006624'),index=False)
    #
    # am=Advance_man()
    #am.save_all_market_funds2db('20230301')
    # am.save_all_market_funds2db('20220901')
    # am.save_all_market_funds2db('20220301')
    # am.save_all_market_funds2db('20210901')
    # am.save_all_market_funds2db('20210301')
    # am.save_all_market_funds2db('20200901')
    # am.save_all_market_funds2db('20200301')
    # am.save_all_market_funds2db('20190901')
    # am.save_all_market_funds2db('20190301')
    # am.save_all_market_funds2db('20180901')
    # am.save_all_market_funds2db('20180301')



    ic_data=[]
    ext_ret=[]
    corr_list=[]
    ic_pvalue_data=[]
    # for theme_name in ['HB消费','HB医药','HB可选消费','HB必选消费','HB制造','HB新制造','HBTMT','HB大金融','HB周期'
    #     ,'HB资源','HB材料','HB基础设施']:
    # for theme_name in ['HB周期']:
    #
    for theme_name in ['TMT','医药', '消费', '制造','成长','价值','超额稳定']:

    # for theme_name in ['超额稳定']:
        ic_data = []
        ext_ret = []
        corr_list = []
        ic_pvalue_data = []

        print(theme_name)
        ic_histor,ex_ret_histor,corr,ic_pvalue_histor=factor_test(['20230630','20221231','20220630','20211231','20210630','20201231','20200630','20191231','20190630','20181231','20180630','20171201',],
                    ['20230930','20230331','20220930','20220330','20210930','20210330','20200930','20200330','20190930','20190330','20180930','20180331'],
                    theme_name=theme_name, pool_size=20,save_to_db=True)
        ic_data.append(ic_histor)
        ext_ret.append(ex_ret_histor)
        corr_list.append(corr)
        ic_pvalue_data.append(ic_pvalue_histor)

        pd.concat(ic_data).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\IC数据.xlsx'.format(theme_name))
        pd.concat(ext_ret).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\超额数据.xlsx'.format(theme_name))
        pd.concat(corr_list).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\因子相关性.xlsx'.format(theme_name))
        pd.concat(ic_pvalue_data).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\ICP值.xlsx'.format(theme_name))

    # for theme_name in ['TMT','医药', '消费', '制造']:
    #
    #
    #     ic_data = []
    #     ext_ret = []
    #     corr_list = []
    #     ic_pvalue_data = []
    #     print(theme_name)
    #
    #     ic_histor,ex_ret_histor,corr,ic_pvalue_histor=factor_test(['20230630','20221231','20220630','20211231','20210630','20201231','20200630','20191231','20190630','20181231','20180630','20171201',],
    #                 ['20230930','20230331','20220930','20220330','20210930','20210330','20200930','20200330','20190930','20190330','20180930','20180331'],
    #                 theme_name=theme_name, pool_size=20,save_to_db=False)
    #
    #     ic_data.append(ic_histor)
    #     ext_ret.append(ex_ret_histor)
    #     corr_list.append(corr)
    #     ic_pvalue_data.append(ic_pvalue_histor)
    #
    #
    #     pd.concat(ic_data).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\IC数据.xlsx'.format(theme_name))
    #     pd.concat(ext_ret).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\超额数据.xlsx'.format(theme_name))
    #     pd.concat(corr_list).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\因子相关性.xlsx'.format(theme_name))
    #     pd.concat(ic_pvalue_data).to_excel('E:\GitFolder\docs\基金筛选\主题基金\{}\ICP值.xlsx'.format(theme_name))


    # get_the_theme_pool('20211231', '20220331', theme_name='先进制造', pool_size=10)
    # get_the_theme_pool('20211231', '20220331', theme_name='能源化工', pool_size=10)
    # get_the_theme_pool('20211231', '20220331', theme_name='医药生物', pool_size=10)
    # get_the_theme_pool('20211231', '20220331', theme_name='消费', pool_size=10)
    # get_the_theme_pool('20211231', '20220331', theme_name='金融地产', pool_size=10)

    # pe=Positive_equilibrium()
    # pe_pool=[]
    # for asofdate in ['20151231', '20160630', '20161231', '20170630', '20171231', '20180630'
    #     , '20181231', '20190630', '20191231'
    #     , '20200630', '20201231', '20210630', '20211231', '20220630']:
    #
    #     tempdf=pd.DataFrame()
    #     tempdf['jjdm']=pe.get_nav_equilibrium_pool(asofdate)
    #     tempdf['asofdate']=asofdate
    #     pe_pool.append(tempdf)
    # pd.concat(pe_pool,axis=0).to_excel('被动均衡（净值）池.xlsx')

    # sql="select ticker,category,vcg_score,trade_date,type from st_fund.r_st_gp_equity_score where trade_date>='20200630'    "
    # test=hbdb.db2df(sql,db='funduser')
    # test=test.sort_values(['ticker','trade_date','type'])
    # test=test.drop_duplicates(['ticker','trade_date'],keep='first')
    # test['size_type']=test['category'].str[0:2]
    # test['style_type']='成长'
    # test.loc[test['vcg_score']<=150,'style_type'] = '价值'
    # test['category']=test['size_type']+test['style_type']
    # reprot_date=test['trade_date'].unique().tolist()

    #
    # outdf=[]
    #
    # for i in range(len(reprot_date)):
    #
    #     if(i!=len(reprot_date)-1):
    #         t0 = reprot_date[i]
    #         t1 = reprot_date[i + 1]
    #         sql = """ select TDATE,SYMBOL,VATURNOVER from finchina.CHDQUOTE
    #         where TDATE>{0} and TDATE<={1} and SYMBOL='000001'
    #         """.format(t0,t1)
    #         date_list=hbdb.db2df(sql, db='readonly').drop('ROW_ID', axis=1)['TDATE'].tolist()
    #     else:
    #         t0 = reprot_date[i]
    #         sql = """ select TDATE,SYMBOL,VATURNOVER from finchina.CHDQUOTE
    #         where TDATE>{0}  and SYMBOL='000001'
    #         """.format(t0)
    #         date_list=hbdb.db2df(sql, db='readonly').drop('ROW_ID', axis=1)['TDATE'].tolist()
    #
    #     for date in date_list:
    #         print(date)
    #         sql = """ select TDATE,SYMBOL,VATURNOVER from finchina.CHDQUOTE
    #         where TDATE={0}
    #         """.format(date)
    #         cjl = hbdb.db2df(sql, db='readonly').drop('ROW_ID', axis=1)
    #         cjl=pd.merge(cjl,test[test['trade_date']==t0]
    #                      ,how='left',left_on='SYMBOL',right_on='ticker')
    #         outdf.append(cjl.groupby(['category', 'TDATE']).sum()['VATURNOVER'].reset_index())
    #
    # outdf=pd.concat(outdf,axis=0)
    # outdf=outdf.pivot_table('VATURNOVER','TDATE','category')
    # cols=outdf.columns.tolist()
    # outdf['total']=outdf.sum(axis=1)
    # for col in cols:
    #     outdf[col]=outdf[col]/outdf['total']
    #
    # outdf.to_excel('成交量分布.xlsx')
    # plot= functionality.Plot(1200,600)
    # plot.plotly_area(outdf.drop('total',axis=1),'成交量分布')

    #
    # val=Value()
    # growth_pool,value_pool=val.get_value(show_num=100, shift_ratio_threshold=0.5)
    # growth_pool.to_excel(r"E:\GitFolder\docs\金工基金池\成长池.xlsx",index=False)
    # value_pool.to_excel(r"E:\GitFolder\docs\金工基金池\价值池.xlsx",index=False)

    #
    #
    # sql = "select jjdm from st_fund.t_st_gm_jjxx where cpfl='2' and ejfl='37' "
    # jjdm_list = hbdb.db2df(sql, db='funduser')['jjdm'].tolist()
    #
    # jjdm_list=['004917','004959','005001']
    #
    # sql="select jjdm,qsrq,jsrq,qjsgfe,wjshfe from st_fund.t_st_gm_febd where jjdm in ({}) and bblb in ('3','5','6','7') "\
    #     .format(util.list_sql_condition(jjdm_list))
    # test=hbdb.db2df(sql,db='funduser')
    # test=test.sort_values(['jjdm','qsrq'])
    # test=test[(test['qjsgfe'].notnull())
    #           |(test['wjshfe'].notnull())]
    # date_list=test['jsrq'].unique().tolist()+test['qsrq'].unique().tolist()
    # map_list=[ util._shift_date(x) for x in date_list]
    # date_map=dict(zip(date_list,map_list))
    # test['qsrq']=[date_map[x] for x in test['qsrq']]
    # test['jsrq']=[date_map[x] for x in test['jsrq']]
    #
    # #get jj nav
    # sql="select jjdm,jjjz,jzrq from st_fund.t_st_gm_jjjz where jjdm in ({0}) and jzrq>='{1}' and jzrq<='{2}'"\
    #     .format(util.list_sql_condition(test['jjdm'].unique().tolist())
    #             ,min(map_list),max(map_list))
    # jj_nav=hbdb.db2df(sql,db='funduser')
    # jj_nav['jzrq']=jj_nav['jzrq'].astype(str)
    #
    # #get index nav
    # sql = "select jyrq,spjg as `偏股混指数` from st_market.t_st_zs_hqql where zqdm='885001' and jyrq>='{0}' and jyrq<='{1}'  " \
    #     .format(min(map_list),max(map_list))
    # index_nav = hbdb.db2df(sql, db='alluser')
    # index_nav['jyrq']=index_nav['jyrq'].astype(str)
    #
    #
    # test=pd.merge(jj_nav,test,how='right',
    #               left_on=['jjdm','jzrq'],right_on=['jjdm','qsrq'])
    # test=pd.merge(jj_nav,test,how='right',
    #               left_on=['jjdm','jzrq'],right_on=['jjdm','jsrq'])
    # test['avg_jjjz']=(test['jjjz_x']+test['jjjz_y'])/2
    # test['申购']=test['qjsgfe']*test['avg_jjjz']
    # test['赎回']=test['wjshfe']*test['avg_jjjz']*-100

    # total_buy_sell=test.groupby('jsrq').sum()[['申购','赎回']]
    # total_buy_sell=pd.merge(index_nav,total_buy_sell,how='left',
    #               left_on='jyrq',right_on='jsrq')
    # total_buy_sell['偏股混指数']=total_buy_sell['偏股混指数']/total_buy_sell['偏股混指数'].iloc[0]
    # total_buy_sell.set_index('jyrq',inplace=True)
    # total_buy_sell.to_excel('偏股混申赎.xlsx')
    # plot = functionality.Plot(1000, 500)
    # data,layout=plot.plotly_line_and_bar(total_buy_sell[['偏股混指数']]
    #                                      ,total_buy_sell[['申购','赎回']],'偏股混申赎','净值','金额',10)
    # plot.plot_render(data,layout)

    # test=pd.merge(jj_nav,
    #               test[['jjdm','jsrq','申购','赎回']],how='left',
    #               left_on=['jjdm','jzrq'],right_on=['jjdm','jsrq']).rename(columns={'jjjz':'基金净值'})
    #
    #
    # for jjdm in test['jjdm'].unique():
    #
    #     test_sample=test[test['jjdm']==jjdm]
    #     test_sample=pd.merge(test_sample,index_nav,how='left',left_on='jzrq',right_on='jyrq').drop('jyrq',axis=1)
    #     test_sample.set_index('jzrq',inplace=True)
    #     test_sample['偏股混指数']=test_sample['偏股混指数']/test_sample['偏股混指数'].iloc[0]
    #     test_sample.to_excel(jjdm+'.xlsx')
    #
    #     plot = functionality.Plot(1000, 500)
    #     data,layout=plot.plotly_line_and_bar(test_sample[['基金净值','偏股混指数']]
    #                                          ,test_sample[['申购','赎回']],jjdm,'净值','金额',10)
    #     plot.plot_render(data,layout)



    print('')
    #


    # rp=Report()
    # rp.filter_pool_report()

    #advance_man()


    #
    # equclass=Equilibrium()
    # ind_equ, style_equ,size_equ, \
    # pepbroe_equ, nav_equ,joint_rank, \
    # joint_result=equclass.get_equilibrium(threshield=0.3,show_num=100)
    #
    #
    #
    # plot=functionality.Plot(800,800)
    #
    # plot.plotly_table(ind_equ, 800, 'asdf')
    # pla.pool_picturing(ind_equ)
    #
    # plot.plotly_table(style_equ, 800, 'asdf')
    # pla.pool_picturing(style_equ)
    #
    # plot.plotly_table(size_equ, 800, 'asdf')
    # pla.pool_picturing(size_equ)
    #
    # plot.plotly_table(pepbroe_equ, 800, 'asdf')
    # pla.pool_picturing(pepbroe_equ)
    #
    # plot.plotly_table(nav_equ, 800, 'asdf')
    # pla.pool_picturing(nav_equ)
    #
    # plot.plotly_table(joint_result,800,'asdf')
    # pla.pool_picturing(joint_result)
    #
    # plot.plotly_table(joint_rank, 800, 'asdf')
    # pla.pool_picturing(joint_rank)

    #

    # leftclass=Leftside()
    #
    # stock_left, ind_left, value_left, \
    # size_left, joint_rank, joint_restult=leftclass.get_left(threshield=0.4,show_num=100)
    #
    #
    # plot=functionality.Plot(800,800)
    #
    # plot.plotly_table(stock_left,800,'asdf')
    # pla.pool_picturing(stock_left)
    #
    # plot.plotly_table(ind_left, 800, 'asdf')
    # pla.pool_picturing(ind_left)
    #
    # plot.plotly_table(value_left, 800, 'asdf')
    # pla.pool_picturing(value_left)
    #
    # plot.plotly_table(size_left, 800, 'asdf')
    # pla.pool_picturing(size_left)
    #
    # plot.plotly_table(joint_rank, 800, 'asdf')
    # pla.pool_picturing(joint_rank)

    # plot.plotly_table(joint_restult, 800, 'asdf')
    # pla.pool_picturing(joint_restult)

    # sizeclass = Size()
    # big_size, medium_size, small_size = sizeclass.get_size(fre='M', show_num=100,
    #                                                        shift_ratio_threshold=0.5,
    #                                                        centralization_threshold=0.5)

    #plot = functionality.Plot(800, 800)
    # plot.plotly_table(big_size, 800, 'asdf')
    # pla.pool_picturing(big_size)
    #
    # plot.plotly_table(medium_size, 800, 'asdf')
    # pla.pool_picturing(medium_size)
    #
    # plot.plotly_table(small_size, 800, 'asdf')
    # pla.pool_picturing(small_size)


    # valueclass = Value()
    # growth, value=valueclass.get_value(fre='M', show_num=20, shift_ratio_threshold=0.5, centralization_threshold=0.5)
    #     , absolute_pe_value, absolute_pb_value, relative_pe_value, relative_pb_value, \
    # dividend_value, reverse_value, high_quality_value \


    # plot = functionality.Plot(800, 800)
    # plot.plotly_table(growth, 800, 'asdf')
    # pla.pool_picturing(growth)
    #
    # plot.plotly_table(value, 800, 'asdf')
    # pla.pool_picturing(value)
    #
    # plot.plotly_table(absolute_pe_value, 800, 'asdf')
    # pla.pool_picturing(absolute_pe_value)
    #
    # plot.plotly_table(absolute_pb_value, 800, 'asdf')
    # pla.pool_picturing(absolute_pb_value)
    #
    # plot.plotly_table(relative_pe_value, 800, 'asdf')
    # pla.pool_picturing(relative_pe_value)
    #
    # plot.plotly_table(relative_pb_value, 800, 'asdf')
    # pla.pool_picturing(relative_pb_value)
    #
    # plot.plotly_table(dividend_value, 800, 'asdf')
    # pla.pool_picturing(dividend_value)
    #
    # plot.plotly_table(reverse_value, 800, 'asdf')
    # pla.pool_picturing(reverse_value)
    #
    # plot.plotly_table(high_quality_value, 800, 'asdf')
    # pla.pool_picturing(high_quality_value)
    #
    # print("Done")


