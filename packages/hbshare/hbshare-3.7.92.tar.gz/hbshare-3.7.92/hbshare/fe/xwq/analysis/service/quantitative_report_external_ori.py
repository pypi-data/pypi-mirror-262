# -*- coding: utf-8 -*-

from datetime import datetime, timedelta
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from hbshare.fe.mutual_analysis import holdind_based as hlb
from hbshare.fe.mutual_analysis.holding_compare_analysis import HoldingCompare
from hbshare.fe.xwq.analysis.orm.fedb import FEDB
from hbshare.fe.xwq.analysis.orm.hbdb import HBDB
import docx
import os
import numpy as np
import pandas as pd
from matplotlib.ticker import FuncFormatter
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import seaborn as sns
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False
sns.set_style('white', {'font.sans-serif': ['simhei', 'Arial']})
line_color_list = ['#F04950', '#6268A2', '#959595', '#333335', '#EE703F', '#7E4A9B', '#8A662C',
                  '#44488E', '#BA67E9', '#3FAEEE']
bar_color_list = ['#C94649', '#EEB2B4', '#E1777A', '#D57C56', '#E39A79', '#DB8A66', '#E5B88C',
                  '#8588B7', '#B4B6D1', '#55598D', '#628497', '#A9C6CB', '#866EA9', '#B79BC7',
                  '#7D7D7E', '#CACACA', '#A7A7A8', '#606063', '#C4C4C4', '#99999B', '#B7B7B7']


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def to_percent(temp, position):
    return '%1.0f'%(temp) + '%'

def to_100percent(temp, position):
    return '%1.0f'%(temp * 100) + '%'

def get_change_dates_loc(change_dates, all_dates):
    change_dates_loc = []
    for change_date in change_dates:
        for idx, date in enumerate(all_dates[:-1]):
            if change_date >= all_dates[idx] and change_date <= all_dates[idx + 1]:
                loc = idx + (datetime.strptime(change_date, '%Y%m%d') - datetime.strptime(all_dates[idx], '%Y%m%d')) / (datetime.strptime(all_dates[idx + 1], '%Y%m%d') - datetime.strptime(all_dates[idx], '%Y%m%d'))
                change_dates_loc.append(loc)
                break
    return change_dates_loc


class QuantitativeReport:
    def __init__(self, file_path, target_fund_code, compare_fund_code_list=[]):
        self.file_path = file_path + target_fund_code + '/'
        if not os.path.exists(self.file_path):
            os.makedirs(self.file_path)
        self.target_fund_code = target_fund_code
        self.compare_fund_code_list = compare_fund_code_list
        self.stock_fund = HBDB().read_stock_fund_info()
        self.stock_fund = self.stock_fund.rename(columns={'jjdm': 'FUND_CODE', 'jjmc': 'FUND_FULL_NAME', 'jjjc': 'FUND_SHORT_NAME', 'clrq': 'BEGIN_DATE', 'zzrq': 'END_DATE', 'ejfl': 'FUND_TYPE', 'kffb': 'OPEN_CLOSE'})
        self.stock_fund['FUND_TYPE'] = self.stock_fund['FUND_TYPE'].replace({'13': '普通股票型', '35': '灵活配置型', '37': '偏股混合型'})
        self.target_fund = self.stock_fund[self.stock_fund['FUND_CODE'] == self.target_fund_code]
        self.compare_fund = self.stock_fund[self.stock_fund['FUND_CODE'].isin(compare_fund_code_list)]
        self.target_fund_name = self.target_fund['FUND_SHORT_NAME'].values[0]
        self.compare_fund_name_dic = self.compare_fund.set_index('FUND_CODE')['FUND_SHORT_NAME'].to_dict()
        self.compare_fund_name_dic[self.target_fund_code] = self.target_fund_name
        self.target_fund_manager = HBDB().read_fund_manager_given_code(self.target_fund_code)
        self.target_fund_manager = self.target_fund_manager.rename(columns={'jjdm': 'FUND_CODE', 'rydm': 'MANAGER_CODE', 'ryxm': 'MANAGER_NAME', 'ryzw': 'POSITION', 'bdrq': 'CHANGE_DATE', 'ryzt': 'STATUS', 'rzrq': 'BEGIN_DATE', 'lrrq': 'END_DATE', 'lryy': 'REASON'})
        self.target_fund_manager = self.target_fund_manager[self.target_fund_manager['POSITION'] == '基金经理']
        self.target_fund_manager['BEGIN_DATE'] = self.target_fund_manager['BEGIN_DATE'].astype(str)
        self.target_fund_manager['END_DATE'] = self.target_fund_manager['END_DATE'].astype(str)
        self.target_fund_manager['END_DATE'] = self.target_fund_manager['END_DATE'].replace('29991231', '-')
        self.change_dates = [date for date in self.target_fund_manager['BEGIN_DATE'].unique().tolist() + self.target_fund_manager['END_DATE'].unique().tolist() if date != '-']
        self.change_dates = sorted(list(set(self.change_dates)))[1:]
        self.fund = HBDB().read_fund_info()
        self.fund = self.fund.rename(columns={'jjdm': 'FUND_CODE', 'jjmc': 'FUND_FULL_NAME', 'jjjc': 'FUND_SHORT_NAME', 'clrq': 'BEGIN_DATE', 'zzrq': 'END_DATE', 'jjfl': 'FUND_TYPE_1ST', 'ejfl': 'FUND_TYPE_2ND', 'kffb': 'OPEN_CLOSE'})
        self.fund['FUND_TYPE_1ST'] = self.fund['FUND_TYPE_1ST'].replace({'f': 'FOF型', '1': '股票型', '2': '债券型', '3': '混合型', '4': '另类投资型', '7': '货币型', '9': 'QDII型'})
        self.fund['FUND_TYPE_2ND'] = self.fund['FUND_TYPE_2ND'].replace(
            {'f3': '债券型FOF', 'f4': '货币型FOF', 'f1': '股票型FOF', 'f2': '混合型FOF', 'f5': '另类投资FOF',
             '13': '普通股票型', '14': '股票型', '15': '增强指数型', '16': '被动指数型',
             '21': '被动指数型债券', '22': '短期纯债型', '23': '混合债券型一级', '24': '混合债券型二级', '25': '增强指数型债券', '26': '债券型',
             '27': '中长期纯债型', '28': '可转换债券型',
             '34': '平衡混合型', '35': '灵活配置型', '36': '混合型', '37': '偏股混合型', '38': '偏债混合型',
             '41': '股票多空', '42': '商品型', '43': 'REITs',
             '91': '股票型QDII', '93': '混合型QDII', '94': '债券型QDII', '95': '另类型QDII'})

    def report_info(self):
        """
        风险提示及重要声明
        """
        risk = '投资有风险。基金的过往业绩并不预示其未来表现。基金管理人管理的其他基金的业绩并不构成基金业绩表现的保证。相关数据仅供参考，不构成投资建议。投资人请详阅基金合同等法律文件，了解产品风险收益特征，根据自身资产状况、风险承受能力审慎决策，独立承担投资风险。本资料仅为宣传用品，本机构及工作人员不存在直接或间接主动推介相关产品的行为。'
        statement1 = '本文件中的信息基于已公开的信息、数据及尽调访谈等，好买基金或好买基金研究中心（以下简称“本公司”）对这些信息的及时性、准确性及完整性不做任何保证，也不保证所包含的信息不会发生变更。文件中的内容仅供参考，不代表任何确定性的判断。本文件及其内容均不构成投资建议，也没有考虑个别客户特殊的投资目标、财务状况或需要。获得本文件的机构或个人据此做出投资决策，应自行承担投资风险。'
        statement2 = '本文件版权为本公司所有。未经本公司书面许可，任何机构或个人不得以翻版、复制、 发表、引用或再次分发他人等任何形式侵犯本公司版权。本文件中的信息均为保密信息，未经本公司事先同意，不得以任何目的，复制或传播本文本中所含信息，亦不可向任何第三方披露。'
        statement3 = '本文件系为好买基金备制并仅以非公开方式提交给符合监管规定之合格投资者。'
        return risk, statement1, statement2, statement3

    def get_fund_info(self):
        """
        生成基金简介数据
        """
        fund_full_name = self.target_fund['FUND_FULL_NAME'].values[0]
        establish_date = str(self.target_fund['BEGIN_DATE'].values[0])
        establish_date_cn = '{0}年{1}月{2}日'.format(int(establish_date[:4]), int(establish_date[4:6]), int(establish_date[6:]))
        fund_type = self.target_fund['FUND_TYPE'].values[0]

        company = HBDB().read_fund_company_given_code(self.target_fund_code)
        company = company.rename(columns={'JJDM': 'FUND_CODE', 'JGDM': 'COMPANY_CODE', 'JGMC': 'COMPANY_NAME'})
        company_name = company['COMPANY_NAME'].values[0]

        manager = HBDB().read_fund_manager_given_code(self.target_fund_code)
        manager = manager.rename(columns={'jjdm': 'FUND_CODE', 'rydm': 'MANAGER_CODE', 'ryxm': 'MANAGER_NAME', 'ryzw': 'POSITION', 'bdrq': 'CHANGE_DATE', 'ryzt': 'STATUS', 'rzrq': 'BEGIN_DATE', 'lrrq': 'END_DATE', 'lryy': 'REASON'})
        manager = manager[(manager['STATUS'] == '-1') & (manager['POSITION'] == '基金经理')].sort_values('BEGIN_DATE')
        manager['BEGIN_DATE'] = manager['BEGIN_DATE'].astype(int)
        manager['BEGIN_DATE_CN'] = manager['BEGIN_DATE'].apply(lambda x: '{0}年{1}月{2}日'.format(int(str(x)[:4]), int(str(x)[4:6]), int(str(x)[6:])))
        manager_name = manager['MANAGER_NAME'].values[0] if len(manager) == 1 else '、'.join(manager['MANAGER_NAME'].tolist())
        begin_date = '自{}'.format(manager['BEGIN_DATE_CN'].values[0]) if len(manager) == 1 else '分别自{}'.format('、'.join(manager['BEGIN_DATE_CN'].tolist()))

        scale = HBDB().read_fund_scale_given_code(self.target_fund_code)
        scale = scale.rename(columns={'JJDM': 'FUND_CODE', 'BBLB1': 'REPORT_TYPE', 'JSRQ': 'REPORT_DATE', 'GGRQ': 'PUBLISH_DATE', 'ZCJZ': 'EQUITY'})
        scale = scale[scale['REPORT_TYPE'] == 13]
        scale = scale.sort_values(['REPORT_DATE', 'PUBLISH_DATE']).drop_duplicates('REPORT_DATE', keep='last')
        scale_num = round(scale['EQUITY'].values[-1] / 100000000.0, 2)

        target = HBDB().read_fund_target_given_code(self.target_fund_code)
        target = target.rename(columns={'JJDM': 'FUND_CODE', 'TZMB': 'TARGET'})
        target_description = target['TARGET'].values[0].replace(' ', '').replace('\r\n', '')

        fund_info_text = '{0}（{1}），全称为{2}，是由{3}于{4}发行的{5}基金，最新资产规模为{6}亿，基金经理为{7}，{8}担任此基金的基金经理至今。{9}'\
            .format(self.target_fund_name, self.target_fund_code, fund_full_name, company_name, establish_date_cn, fund_type, scale_num, manager_name, begin_date, target_description)
        return fund_info_text

    def get_fund_company_info(self):
        """
        生成基金公司简介数据
        """
        company = HBDB().read_fund_company_given_code(self.target_fund_code)
        company = company.rename(columns={'JJDM': 'FUND_CODE', 'JGDM': 'COMPANY_CODE', 'JGMC': 'COMPANY_NAME'})
        company_name = company['COMPANY_NAME'].values[0]
        company_code = company['COMPANY_CODE'].values[0]

        company_fund = HBDB().read_company_fund_given_code(company_code)
        company_fund = company_fund.rename(columns={'JGDM': 'COMPANY_CODE', 'JGMC': 'COMPANY_NAME', 'JJDM': 'FUND_CODE'})
        company_fund_scale_list = []
        for code in company_fund['FUND_CODE'].unique().tolist():
            company_fund_scale_code = HBDB().read_fund_scale_given_code(code)
            company_fund_scale_list.append(company_fund_scale_code)
        company_fund_scale = pd.concat(company_fund_scale_list)
        company_fund_scale = company_fund_scale.rename(columns={'JJDM': 'FUND_CODE', 'BBLB1': 'REPORT_TYPE', 'JSRQ': 'REPORT_DATE', 'GGRQ': 'PUBLISH_DATE', 'ZCJZ': 'EQUITY'})
        company_fund_scale = company_fund_scale[company_fund_scale['REPORT_TYPE'] == 13]
        company_fund_scale = company_fund_scale.sort_values(['FUND_CODE', 'REPORT_DATE']).drop_duplicates(['FUND_CODE'], keep='last')
        company_fund_scale['EQUITY'] = company_fund_scale['EQUITY'] / 100000000.0
        company_fund_scale = company_fund_scale[['FUND_CODE', 'REPORT_DATE', 'EQUITY']].merge(self.fund, on=['FUND_CODE'], how='left')
        company_fund_scale = company_fund_scale[company_fund_scale['END_DATE'].isnull()]
        company_fund_scale.loc[company_fund_scale['FUND_TYPE_2ND'].isnull(), 'FUND_TYPE_2ND'] = company_fund_scale.loc[company_fund_scale['FUND_TYPE_2ND'].isnull(), 'FUND_TYPE_1ST']
        company_fund_count_1st = company_fund_scale.drop_duplicates('FUND_FULL_NAME')[['FUND_TYPE_1ST', 'FUND_CODE']].groupby('FUND_TYPE_1ST').count().rename(columns={'FUND_CODE': '数量'})
        company_fund_count_2nd = company_fund_scale.drop_duplicates('FUND_FULL_NAME')[['FUND_TYPE_2ND', 'FUND_CODE']].groupby('FUND_TYPE_2ND').count().rename(columns={'FUND_CODE': '数量'})
        company_fund_count_1st['数量占比'] = company_fund_count_1st['数量'] / float(company_fund_count_1st['数量'].sum())
        company_fund_count_2nd['数量占比'] = company_fund_count_2nd['数量'] / float(company_fund_count_2nd['数量'].sum())
        company_fund_count_1st['数量占比'] = company_fund_count_1st['数量占比'].apply(lambda x: '{0}%'.format(round(x * 100, 2)))
        company_fund_count_2nd['数量占比'] = company_fund_count_2nd['数量占比'].apply(lambda x: '{0}%'.format(round(x * 100, 2)))
        company_fund_scale_1st = company_fund_scale[['FUND_TYPE_1ST', 'EQUITY']].groupby('FUND_TYPE_1ST').sum().rename(columns={'EQUITY': '规模（亿元）'})
        company_fund_scale_2nd = company_fund_scale[['FUND_TYPE_2ND', 'EQUITY']].groupby('FUND_TYPE_2ND').sum().rename(columns={'EQUITY': '规模（亿元）'})
        company_fund_scale_1st['规模（亿元）'] = company_fund_scale_1st['规模（亿元）'].apply(lambda x: round(x, 2))
        company_fund_scale_2nd['规模（亿元）'] = company_fund_scale_2nd['规模（亿元）'].apply(lambda x: round(x, 2))
        company_fund_scale_1st['规模占比'] = company_fund_scale_1st['规模（亿元）'] / float(company_fund_scale_1st['规模（亿元）'].sum())
        company_fund_scale_2nd['规模占比'] = company_fund_scale_2nd['规模（亿元）'] / float(company_fund_scale_2nd['规模（亿元）'].sum())
        company_fund_scale_1st['规模占比'] = company_fund_scale_1st['规模占比'].apply(lambda x: '{0}%'.format(round(x * 100, 2)))
        company_fund_scale_2nd['规模占比'] = company_fund_scale_2nd['规模占比'].apply(lambda x: '{0}%'.format(round(x * 100, 2)))
        company_fund_1st = pd.concat([company_fund_count_1st, company_fund_scale_1st], axis=1).sort_values('规模（亿元）', ascending=False)
        company_fund_2nd = pd.concat([company_fund_count_2nd, company_fund_scale_2nd], axis=1).sort_values('规模（亿元）', ascending=False)
        company_fund_1st = company_fund_1st.T.reset_index().T.reset_index()
        company_fund_2nd = company_fund_2nd.T.reset_index().T.reset_index()
        company_fund_1st.iloc[0, 0] = '一级分类'
        company_fund_2nd.iloc[0, 0] = '二级分类'

        fund_company_info_text = '该产品所在基金管理公司{0}旗下公募基金数量和规模分布如下：'.format(company_name)
        return fund_company_info_text, company_fund_1st, company_fund_2nd

    def get_fund_scale(self):
        """
        生成基金规模数据
        """
        scale = HBDB().read_fund_scale_given_code(self.target_fund_code)
        scale = scale.rename(columns={'JJDM': 'FUND_CODE', 'BBLB1': 'REPORT_TYPE', 'JSRQ': 'REPORT_DATE', 'GGRQ': 'PUBLISH_DATE', 'ZCJZ': 'EQUITY'})
        scale = scale[scale['REPORT_TYPE'] == 13]
        scale = scale.sort_values(['REPORT_DATE', 'PUBLISH_DATE']).drop_duplicates('REPORT_DATE', keep='last')
        scale['EQUITY'] = scale['EQUITY'] / 100000000.0
        scale = scale[['REPORT_DATE', 'EQUITY']]

        share = HBDB().read_fund_share_given_code(self.target_fund_code)
        share = share.rename(columns={'jjdm': 'FUND_CODE', 'bblb': 'REPORT_TYPE', 'jsrq': 'END_DATE', 'plrq': 'PUBLISH_DATE', 'qsrq': 'BEGIN_DATE', 'qcfe': 'BEGIN_SHARES', 'qjsgfe': 'PURCHASE_SHARES', 'cfzjfe': 'SPLIT_SHARES', 'wjshfe': 'REDEEM_SHARES', 'qmfe': 'END_SHARES'})
        share = share[share['REPORT_TYPE'].isin(['3', '5', '6', '7'])]
        share = share.sort_values(['END_DATE', 'PUBLISH_DATE']).drop_duplicates('END_DATE', keep='last')
        share[['END_SHARES', 'PURCHASE_SHARES', 'REDEEM_SHARES']] = share[['END_SHARES', 'PURCHASE_SHARES', 'REDEEM_SHARES']] / 100000000.0
        share = share[['END_DATE', 'END_SHARES', 'PURCHASE_SHARES', 'REDEEM_SHARES']].rename(columns={'END_DATE': 'REPORT_DATE'})

        fund_scale = scale.merge(share, on=['REPORT_DATE'], how='left')
        fund_scale = fund_scale.sort_values('REPORT_DATE').iloc[-12:]
        fund_scale.columns = ['报告日期', '净资产（亿元）', '份额（亿份）', '申购额（亿份）', '赎回额（亿份）']

        scale = fund_scale[['报告日期', '净资产（亿元）']].rename(columns={'净资产（亿元）': '规模（亿元）/份额（亿份）'})
        scale['规模/份额'] = '规模'
        share = fund_scale[['报告日期', '份额（亿份）']].rename(columns={'份额（亿份）': '规模（亿元）/份额（亿份）'})
        share['规模/份额'] = '份额'
        scale_share = pd.concat([scale, share]).reset_index().drop('index', axis=1)
        purchase = fund_scale[['报告日期', '申购额（亿份）']].rename(columns={'申购额（亿份）': '申赎额（亿份）'})
        purchase['申赎额'] = '申购额'
        redeem = fund_scale[['报告日期', '赎回额（亿份）']].rename(columns={'赎回额（亿份）': '申赎额（亿份）'})
        redeem['申赎额'] = '赎回额'
        purchase_redeem = pd.concat([purchase, redeem]).reset_index().drop('index', axis=1)

        scale_share['报告日期'] = scale_share['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(scale_share['报告日期']) and date <= max(scale_share['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(scale_share['报告日期'].unique().tolist()))
        fig, ax1 = plt.subplots(figsize=(6, 3))
        ax2 = ax1.twinx()
        sns.lineplot(ax=ax1, x='报告日期', y='规模（亿元）/份额（亿份）', data=scale_share, hue='规模/份额', palette=['#F04950', '#6268A2'])
        sns.barplot(ax=ax2, x='报告日期', y='申赎额（亿份）', data=purchase_redeem, hue='申赎额', palette=['#C94649', '#8588B7'])
        # handles_ax1, labels_ax1 = ax1.get_legend_handles_labels()
        # ax1.legend(handles=handles_ax1, labels=labels_ax1, loc=2)
        # handles_ax2, labels_ax2 = ax2.get_legend_handles_labels()
        # ax2.legend(handles=handles_ax2, labels=labels_ax2, loc=1)
        ax1.vlines(x=change_dates_loc, ymin=0.0, ymax=max(scale_share['规模（亿元）/份额（亿份）']), linestyles='dashed', color='#959595')
        ax1.legend(loc=2)
        ax2.legend(loc=1)
        ax1.set_xticklabels(labels=fund_scale['报告日期'], rotation=90)
        plt.tight_layout()
        plt.savefig('{0}{1}_scale.png'.format(self.file_path, self.target_fund_code))

        lastet_scale = round(fund_scale['净资产（亿元）'].iloc[-1], 2)
        fund_scale_text = '该产品规模变动如下，截至最新报告期，基金最新规模为{0}亿元。'.format(lastet_scale)

        compare_scale = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_scale_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_scale_code = HBDB().read_fund_scale_given_code(compare_fund_code)
                compare_scale_list.append(compare_scale_code)
            compare_scale = pd.concat(compare_scale_list)
            compare_scale = compare_scale.rename(columns={'JJDM': 'FUND_CODE', 'BBLB1': 'REPORT_TYPE', 'JSRQ': 'REPORT_DATE', 'GGRQ': 'PUBLISH_DATE', 'ZCJZ': 'EQUITY'})
            compare_scale = compare_scale[compare_scale['REPORT_TYPE'] == 13]
            compare_scale = compare_scale.sort_values(['FUND_CODE', 'REPORT_DATE', 'PUBLISH_DATE']).drop_duplicates(['FUND_CODE', 'REPORT_DATE'], keep='last')
            compare_scale = compare_scale[['FUND_CODE', 'REPORT_DATE', 'EQUITY']]
            compare_scale['EQUITY'] = compare_scale['EQUITY'].apply(lambda x: round(x / 100000000.0, 2))
            compare_scale = compare_scale.pivot(index='REPORT_DATE', columns='FUND_CODE', values='EQUITY')
            compare_scale = compare_scale.dropna(subset=[self.target_fund_code], axis=0).sort_index().iloc[-12:]
            compare_scale = compare_scale[[self.target_fund_code] + self.compare_fund_code_list]
            compare_scale.columns = [self.compare_fund_name_dic[col] for col in compare_scale.columns]
            compare_scale = compare_scale.replace(np.nan, '-')
            compare_scale = compare_scale.T.reset_index().T.reset_index()
            compare_scale.iloc[0, 0] = '规模（亿元）'
        return fund_scale_text, compare_scale

    def get_fund_holder(self):
        """
        生成基金持有人结构数据
        """
        fund_holder = HBDB().read_fund_holder_given_code(self.target_fund_code)
        fund_holder = fund_holder.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'ggrq': 'PUBLISH_DATE', 'jgcyfe': 'ORGAN_HOLDING', 'jgcybl': 'ORGAN_HOLDING_RATIO', 'grcyfe': 'PERSON_HOLDING', 'grcybl': 'PERSON_HOLDING_RATIO'})

        fund_holder = fund_holder.sort_values('REPORT_DATE').iloc[-12:]
        fund_holder = fund_holder[['REPORT_DATE', 'ORGAN_HOLDING_RATIO', 'PERSON_HOLDING_RATIO']]
        fund_holder.columns = ['报告日期', '机构持有比例', '个人持有比例']

        organ_person = fund_holder.copy(deep=True)
        organ_person['个人持有比例'] = organ_person['个人持有比例'] + organ_person['机构持有比例']

        organ_person['报告日期'] = organ_person['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(organ_person['报告日期']) and date <= max(organ_person['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(organ_person['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(ax=ax, x='报告日期', y='个人持有比例', data=organ_person, label='个人持有比例', color='#8588B7')
        sns.barplot(ax=ax, x='报告日期', y='机构持有比例', data=organ_person, label='机构持有比例', color='#C94649')
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=100.0, linestyles='dashed', color='#959595')
        plt.legend(loc=2)
        plt.xticks(rotation=90)
        plt.ylabel('持有比例（%）')
        plt.tight_layout()
        plt.savefig('{0}{1}_holder.png'.format(self.file_path, self.target_fund_code))

        latest_holder = fund_holder.iloc[-1]
        (larger_name, larger_ratio) = ('机构持有比例', round(latest_holder['机构持有比例'], 2)) if latest_holder['机构持有比例'] > latest_holder['个人持有比例'] else ('个人持有比例', round(latest_holder['个人持有比例'], 2))
        (smaller_name, smaller_ratio) = ('机构持有比例', round(latest_holder['机构持有比例'], 2)) if latest_holder['机构持有比例'] < latest_holder['个人持有比例'] else ('个人持有比例', round(latest_holder['个人持有比例'], 2))
        fund_scale_text = '该产品持有人结构如下，从最新持有人结构看，{0}大于{1}，占比分别为{2}%、{3}%。'.format(larger_name, smaller_name, larger_ratio, smaller_ratio)

        compare_holding = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_holding_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_holding_code = HBDB().read_fund_holder_given_code(compare_fund_code)
                compare_holding_list.append(compare_holding_code)
            compare_holding = pd.concat(compare_holding_list)
            compare_holding = compare_holding.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'ggrq': 'PUBLISH_DATE', 'jgcyfe': 'ORGAN_HOLDING', 'jgcybl': 'ORGAN_HOLDING_RATIO', 'grcyfe': 'PERSON_HOLDING', 'grcybl': 'PERSON_HOLDING_RATIO'})
            compare_holding['ORGAN_HOLDING_RATIO'] = compare_holding['ORGAN_HOLDING_RATIO'].fillna(0.0)
            compare_holding['PERSON_HOLDING_RATIO'] = compare_holding['PERSON_HOLDING_RATIO'].fillna(0.0)
            compare_holding = compare_holding.sort_values(['FUND_CODE', 'REPORT_DATE', 'PUBLISH_DATE']).drop_duplicates(['FUND_CODE', 'REPORT_DATE'], keep='last')
            compare_holding = compare_holding[['FUND_CODE', 'REPORT_DATE', 'ORGAN_HOLDING_RATIO']]
            compare_holding['ORGAN_HOLDING_RATIO'] = compare_holding['ORGAN_HOLDING_RATIO'].apply(lambda x: '{0}%'.format(round(x, 2)))
            compare_holding = compare_holding.pivot(index='REPORT_DATE', columns='FUND_CODE', values='ORGAN_HOLDING_RATIO')
            compare_holding = compare_holding.dropna(subset=[self.target_fund_code], axis=0).sort_index().iloc[-12:]
            compare_holding = compare_holding[[self.target_fund_code] + self.compare_fund_code_list]
            compare_holding.columns = [self.compare_fund_name_dic[col] for col in compare_holding.columns]
            compare_holding = compare_holding.replace(np.nan, '-')
            compare_holding = compare_holding.T.reset_index().T.reset_index()
            compare_holding.iloc[0, 0] = '机构持有比例'
        return fund_scale_text, compare_holding

    def get_fund_manager_info(self):
        """
        生成基金经理简介数据
        """
        manager = HBDB().read_fund_manager_given_code(self.target_fund_code)
        manager = manager.rename(columns={'jjdm': 'FUND_CODE', 'rydm': 'MANAGER_CODE', 'ryxm': 'MANAGER_NAME', 'ryzw': 'POSITION', 'bdrq': 'CHANGE_DATE', 'ryzt': 'STATUS', 'rzrq': 'BEGIN_DATE', 'lrrq': 'END_DATE', 'lryy': 'REASON'})
        manager_all = manager.copy(deep=True)
        manager = manager[(manager['STATUS'] == '-1') & (manager['POSITION'] == '基金经理')].sort_values('BEGIN_DATE')
        manager['BEGIN_DATE'] = manager['BEGIN_DATE'].astype(int)
        manager['BEGIN_DATE_CN'] = manager['BEGIN_DATE'].apply(lambda x: '{0}年{1}月{2}日'.format(int(str(x)[:4]), int(str(x)[4:6]), int(str(x)[6:])))
        manager_name = manager['MANAGER_NAME'].values[0] if len(manager) == 1 else '、'.join(manager['MANAGER_NAME'].tolist())
        begin_date = '自{}'.format(manager['BEGIN_DATE_CN'].values[0]) if len(manager) == 1 else '分别自{}'.format('、'.join(manager['BEGIN_DATE_CN'].tolist()))

        three_years_ago = (datetime.today() - timedelta(365 * 3)).strftime('%Y%m%d')
        change = '未经历' if max(manager['BEGIN_DATE'].unique().tolist()) <= int(three_years_ago) else '经历了'

        fund_manager_info_text = '{0}目前的基金经理为{1}，{2}担任此基金的基金经理至今，在定量分析统计期间（近三年）{3}基金经理变更。'\
            .format(self.target_fund_name, manager_name, begin_date, change)
        for manager_code in manager['MANAGER_CODE'].unique().tolist():
            manager_name = manager[manager['MANAGER_CODE'] == manager_code]['MANAGER_NAME'].values[0]

            manager_product = HBDB().read_fund_manager_product_given_code(manager_code)
            manager_product = manager_product.rename(columns={'rydm': 'MANAGER_CODE', 'jsrq': 'REPORT_DATE', 'zgcpsl': 'FUND_NUM', 'zgcpgm': 'FUND_SCALE', 'dbjj': 'FUND_REPRESENT'})
            manager_product = manager_product.sort_values('REPORT_DATE')
            manager_product_scale = round(manager_product['FUND_SCALE'].values[-1] / 100000000.0, 2)
            manager_product_repr = '，根据基金经理任职时长和基金净资产，其代表基金即为本产品' if manager_product['FUND_REPRESENT'].values[-1] == self.target_fund_code else ''

            manager_funds = HBDB().read_fund_given_manager_code(manager_code)
            manager_funds = manager_funds.rename(columns={'jjdm': 'FUND_CODE', 'rydm': 'MANAGER_CODE', 'ryxm': 'MANAGER_NAME', 'ryzw': 'POSITION', 'bdrq': 'CHANGE_DATE', 'ryzt': 'STATUS', 'rzrq': 'BEGIN_DATE', 'lrrq': 'END_DATE', 'lryy': 'REASON'})
            manager_funds = manager_funds[manager_funds['STATUS'] == '-1']
            manager_funds = self.fund[self.fund['FUND_CODE'].isin(manager_funds['FUND_CODE'].unique().tolist())]
            manager_product_num = len(manager_funds.drop_duplicates('FUND_FULL_NAME'))
            funds_scale_list = []
            for code in manager_funds['FUND_CODE'].unique().tolist():
                funds_scale_list.append(HBDB().read_fund_scale_given_code(code))
            funds_scale = pd.concat(funds_scale_list)
            funds_scale = funds_scale.rename(columns={'JJDM': 'FUND_CODE', 'BBLB1': 'REPORT_TYPE', 'JSRQ': 'REPORT_DATE', 'GGRQ': 'PUBLISH_DATE', 'ZCJZ': 'EQUITY'})
            funds_scale = funds_scale[funds_scale['REPORT_TYPE'] == 13]
            funds_scale = funds_scale.sort_values(['FUND_CODE', 'REPORT_DATE', 'PUBLISH_DATE']).drop_duplicates(['FUND_CODE'], keep='last')
            manager_funds = manager_funds.merge(funds_scale[['FUND_CODE', 'EQUITY']], on=['FUND_CODE'], how='left')
            manager_funds = manager_funds[['FUND_TYPE_2ND', 'EQUITY']].groupby('FUND_TYPE_2ND').sum().reset_index().sort_values('EQUITY', ascending=False)
            manager_funds['EQUITY'] = manager_funds['EQUITY'].apply(lambda x: str(round(x / 100000000.0, 2)))
            scale_type = '、'.join(manager_funds['FUND_TYPE_2ND'].unique().tolist())
            scale_value = '、'.join(manager_funds['EQUITY'].unique().tolist())
            manager_scale_repr = '{0}基金管理规模分别为{1}亿元'.format(scale_type, scale_value) if len(manager_funds) > 1 else '{0}基金管理规模为{1}亿元'.format(scale_type, scale_value)

            fund_manager_achievement = HBDB().read_fund_manager_achievement_given_code(manager_code, self.target_fund_code)
            fund_manager_achievement = fund_manager_achievement.rename(columns={'rydm': 'MANAGER_CODE', 'jjdm': 'FUND_CODE', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'rqts': 'MANAGE_DAYS', 'hbrq': 'RETURN', 'rqnhhb': 'ANNUAL_RETURN', 'zdyl': 'MAX_RETURN', 'zdhc': 'MAX_DRAWDOWN', 'nhxpbl': 'ANNUAL_SHARPE_RATIO', 'kmbl': 'CALMAR_RATIO'})
            annual_return = round(fund_manager_achievement['ANNUAL_RETURN'].values[0], 2)
            max_drawdown = round(fund_manager_achievement['MAX_DRAWDOWN'].values[0], 2)
            annual_sharpe_ratio = round(fund_manager_achievement['ANNUAL_SHARPE_RATIO'].values[0], 2)
            calmar_ratio = round(fund_manager_achievement['CALMAR_RATIO'].values[0], 2)

            fund_manager_info_text += '{0}目前在管产品数量为{1}只，在管基金规模为{2}亿元，其中{3}{4}。管理期内{5}年化收益为{6}%， 最大回撤为{7}%，年化夏普比率为{8}，卡玛比率为{9}。'\
                .format(manager_name, manager_product_num, manager_product_scale, manager_scale_repr, manager_product_repr, self.target_fund_name, annual_return, max_drawdown, annual_sharpe_ratio, calmar_ratio)

        fund_manager_info_text += '该产品自成立以来，基金经理变更情况如下：'
        manager_all = manager_all[manager_all['POSITION'] == '基金经理']
        manager_all['BEGIN_DATE'] = manager_all['BEGIN_DATE'].astype(str)
        manager_all['END_DATE'] = manager_all['END_DATE'].astype(str)
        manager_all['END_DATE'] = manager_all['END_DATE'].replace('29991231', '-')
        manager_all['REASON'] = manager_all['REASON'].replace(np.nan, '-')
        manager_all = manager_all.sort_values(['BEGIN_DATE', 'END_DATE'])
        manager_all = manager_all[['MANAGER_NAME', 'BEGIN_DATE', 'END_DATE', 'REASON']]
        manager_all.columns = ['基金经理', '任职日期', '离任日期', '离任原因']
        manager_all_ori = manager_all.copy(deep=True)
        manager_all = manager_all.T.reset_index().T.reset_index().drop('index', axis=1)
        return fund_manager_info_text, manager_all_ori, manager_all

    def get_fund_manager_method(self):
        """
        生成基金经理投资理念数据
        """
        # 换手率与集中度
        fund_size = FEDB().read_hbs_size_property_given_code(self.target_fund_code)
        fund_size = fund_size[['jjdm', 'asofdate', 'shift_lv_rank', 'cen_lv_rank']]
        fund_size = fund_size.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'shift_lv_rank': 'SIZE_TURNOVER', 'cen_lv_rank': 'SIZE_CONCENTRATION'})
        fund_size['FUND_NAME'] = self.target_fund_name
        size_shift = fund_size[['FUND_NAME', 'FUND_CODE', 'SIZE_TURNOVER']]
        size_cen = fund_size[['FUND_NAME', 'FUND_CODE', 'SIZE_CONCENTRATION']]

        fund_style = FEDB().read_hbs_style_property_given_code(self.target_fund_code)
        fund_style = fund_style[['jjdm', 'asofdate', 'shift_lv_rank', 'cen_lv_rank']]
        fund_style = fund_style.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'shift_lv_rank': 'STYLE_TURNOVER', 'cen_lv_rank': 'STYLE_CONCENTRATION'})
        fund_style['FUND_NAME'] = self.target_fund_name
        style_shift = fund_style[['FUND_NAME', 'FUND_CODE', 'STYLE_TURNOVER']]
        style_cen = fund_style[['FUND_NAME', 'FUND_CODE', 'STYLE_CONCENTRATION']]

        fund_theme = FEDB().read_hbs_theme_industry_property_given_code(self.target_fund_code)
        fund_theme = fund_theme[['jjdm', 'asofdate', 'ratio_theme_rank', 'cen_theme_rank']]
        fund_theme = fund_theme.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'ratio_theme_rank': 'THEME_TURNOVER', 'cen_theme_rank': 'THEME_CONCENTRATION'})
        fund_theme['FUND_NAME'] = self.target_fund_name
        theme_shift = fund_theme[['FUND_NAME', 'FUND_CODE', 'THEME_TURNOVER']]
        theme_cen = fund_theme[['FUND_NAME', 'FUND_CODE', 'THEME_CONCENTRATION']]

        fund_industry = FEDB().read_hbs_theme_industry_property_given_code(self.target_fund_code)
        fund_industry = fund_industry[['jjdm', 'asofdate', 'ratio_ind_rank', 'cen_ind_rank']]
        fund_industry = fund_industry.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'ratio_ind_rank': 'INDUSTRY_TURNOVER', 'cen_ind_rank': 'INDUSTRY_CONCENTRATION'})
        fund_industry['FUND_NAME'] = self.target_fund_name
        industry_shift = fund_industry[['FUND_NAME', 'FUND_CODE', 'INDUSTRY_TURNOVER']]
        industry_cen = fund_industry[['FUND_NAME', 'FUND_CODE', 'INDUSTRY_CONCENTRATION']]

        shift_disp = size_shift.merge(style_shift, on=['FUND_NAME', 'FUND_CODE'], how='left').merge(theme_shift, on=['FUND_NAME', 'FUND_CODE'], how='left').merge(industry_shift, on=['FUND_NAME', 'FUND_CODE'], how='left')
        shift_disp['TURNOVER'] = shift_disp.set_index(['FUND_NAME', 'FUND_CODE']).iloc[0].mean()
        shift_mark = '偏低' if shift_disp['TURNOVER'].iloc[0] < 0.4 else '偏高' if shift_disp['TURNOVER'].iloc[0] > 0.6 else '中等'
        for col in list(shift_disp.columns)[2:]:
            shift_disp[col].iloc[0] = '{0}%'.format(round(shift_disp[col].iloc[0] * 100.0, 2))
        shift_data = '、'.join(shift_disp.set_index(['FUND_NAME', 'FUND_CODE']).iloc[0].tolist()[:4])
        shift_disp.columns = ['基金名称', '基金代码', '规模换手率水平', '风格换手率水平', '主题换手率水平', '行业换手率水平', '整体换手率水平']
        shift_disp = shift_disp.T.reset_index().T.reset_index().drop('index', axis=1)
        cen_disp = size_cen.merge(style_cen, on=['FUND_NAME', 'FUND_CODE'], how='left').merge(theme_cen, on=['FUND_NAME', 'FUND_CODE'], how='left').merge(industry_cen, on=['FUND_NAME', 'FUND_CODE'], how='left')
        cen_disp['CONCENTRATION'] = cen_disp.set_index(['FUND_NAME', 'FUND_CODE']).iloc[0].mean()
        cen_mark = '偏低' if cen_disp['CONCENTRATION'].iloc[0] < 0.4 else '偏高' if cen_disp['CONCENTRATION'].iloc[0] > 0.6 else '中等'
        for col in list(cen_disp.columns)[2:]:
            cen_disp[col].iloc[0] = '{0}%'.format(round(cen_disp[col].iloc[0] * 100.0, 2))
        cen_data = '、'.join(cen_disp.set_index(['FUND_NAME', 'FUND_CODE']).iloc[0].tolist()[:4])
        cen_disp.columns = ['基金名称', '基金代码', '规模集中度水平', '风格集中度水平', '主题集中度水平', '行业集中度水平', '整体集中度水平']
        cen_disp = cen_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        hbs_rank_history = FEDB().read_hbs_rank_history_given_code(self.target_fund_code)
        hbs_rank_history = hbs_rank_history.sort_values('jsrq').iloc[-12:]
        shift_ts = hbs_rank_history[['jjdm', 'jsrq', 'shift_ratio_size', 'shift_ratio_style', 'shift_ratio_theme', 'shift_ratio_ind']]
        shift_ts = shift_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': '报告日期', 'shift_ratio_size': '规模换手率水平', 'shift_ratio_style': '风格换手率水平', 'shift_ratio_theme': '主题换手率水平', 'shift_ratio_ind': '行业换手率水平'})
        cen_ts = hbs_rank_history[['jjdm', 'jsrq', 'c_level_size', 'c_level_style', 'c_level_theme', 'c_level_ind']]
        cen_ts = cen_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': '报告日期', 'c_level_size': '规模集中度水平', 'c_level_style': '风格集中度水平', 'c_level_theme': '主题集中度水平', 'c_level_ind': '行业集中度水平'})

        size_style_abs_history = FEDB().read_hbs_size_style_property_history_given_code(self.target_fund_code)
        theme_ind_abs_history = FEDB().read_hbs_theme_industry_property_history_given_code(int(self.target_fund_code))
        theme_ind_abs_history['jjdm'] = theme_ind_abs_history['jjdm'].apply(lambda x: str(x).zfill(6))
        stock_abs_history = FEDB().read_hbs_holding_property_history_given_code(self.target_fund_code)
        hbs_abs_history = size_style_abs_history.merge(theme_ind_abs_history, on=['jjdm', 'jsrq'], how='left').merge(stock_abs_history, on=['jjdm', 'jsrq'], how='left')
        hbs_abs_history = hbs_abs_history.sort_values('jsrq').iloc[-12:]
        shift_abs_ts = hbs_abs_history[['jjdm', 'jsrq', 'shift_ratio_size', 'shift_ratio_style', 'shift_ratio_theme', 'shift_ratio_ind']]
        shift_abs_ts = shift_abs_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': '报告日期', 'shift_ratio_size': '规模换手率', 'shift_ratio_style': '风格换手率', 'shift_ratio_theme': '主题换手率', 'shift_ratio_ind': '行业换手率'})
        cen_abs_ts = hbs_abs_history[['jjdm', 'jsrq', 'c_level_size', 'c_level_style', 'c_level_theme', 'c_level_ind']]
        cen_abs_ts = cen_abs_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': '报告日期', 'c_level_size': '规模集中度', 'c_level_style': '风格集中度', 'c_level_theme': '主题集中度', 'c_level_ind': '行业集中度'})

        shift_ts['报告日期'] = shift_ts['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(shift_ts['报告日期']) and date <= max(shift_ts['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(shift_ts['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(2, 2, figsize=(6, 6))
        pic_list = [['规模换手率水平', '风格换手率水平'], ['主题换手率水平', '行业换手率水平']]
        for i in range(2):
            for j in range(2):
                axij = sns.lineplot(ax=ax[i][j], x='报告日期', y=pic_list[i][j], data=shift_ts, color='#F04950')
                axij2 = sns.lineplot(ax=ax[i][j].twinx(), x='报告日期', y=pic_list[i][j][:5], data=shift_abs_ts, color='#6268A2')
                axij.set_xticklabels(labels=shift_ts['报告日期'], rotation=90)
                axij.set_ylim([0.0, 1.0])
                axij.hlines(y=0.5, xmin=0.0, xmax=11.0, linestyles='dashed', color='#959595')
                axij.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
                axij.yaxis.set_major_formatter(FuncFormatter(to_100percent))
                axij2.yaxis.set_major_formatter(FuncFormatter(to_100percent))
                axij.set_ylabel(pic_list[i][j] + '（红）')
                axij2.set_ylabel(pic_list[i][j][:5] + '（蓝）')
        plt.tight_layout()
        plt.savefig('{0}{1}_shift.png'.format(self.file_path, self.target_fund_code))

        fig, ax = plt.subplots(2, 2, figsize=(6, 6))
        pic_list = [['规模集中度水平', '风格集中度水平'], ['主题集中度水平', '行业集中度水平']]
        for i in range(2):
            for j in range(2):
                axij = sns.lineplot(ax=ax[i][j], x='报告日期', y=pic_list[i][j], data=cen_ts, color='#F04950')
                axij2 = sns.lineplot(ax=ax[i][j].twinx(), x='报告日期', y=pic_list[i][j][:5], data=cen_abs_ts, color='#6268A2')
                axij.set_xticklabels(labels=cen_ts['报告日期'], rotation=90)
                axij.set_ylim([0.0, 1.0])
                axij.hlines(y=0.5, xmin=0.0, xmax=11.0, linestyles='dashed', color='#959595')
                axij.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
                axij.yaxis.set_major_formatter(FuncFormatter(to_100percent))
                axij2.yaxis.set_major_formatter(FuncFormatter(to_100percent))
                axij.set_ylabel(pic_list[i][j] + '（红）')
                axij2.set_ylabel(pic_list[i][j][:5] + '（蓝）')
        plt.tight_layout()
        plt.savefig('{0}{1}_cen.png'.format(self.file_path, self.target_fund_code))

        stock_cen = FEDB().read_hbs_holding_property_given_code(self.target_fund_code)
        stock_cen = stock_cen.rename(columns={'jjdm': 'FUND_CODE', '个股集中度': 'STOCK_CONCENTRATION'})
        industry_stock_cen_disp = industry_cen.merge(stock_cen[['FUND_CODE', 'STOCK_CONCENTRATION']], on=['FUND_CODE'], how='left')
        industry_cen_mark = '低于50%' if industry_stock_cen_disp['INDUSTRY_CONCENTRATION'].iloc[0] < 0.5 else '高于70%' if industry_stock_cen_disp['INDUSTRY_CONCENTRATION'].iloc[0] > 0.7 else '处于50%与70%之间'
        stock_cen_mark = '低于50%' if industry_stock_cen_disp['STOCK_CONCENTRATION'].iloc[0] < 0.5 else '高于70%' if industry_stock_cen_disp['STOCK_CONCENTRATION'].iloc[0] > 0.7 else '处于50%与70%之间'
        select_mark = '选股方法为自下而上' if stock_cen_mark == '高于70%' and industry_cen_mark == '低于50%' else '选股方法为自上而下' if stock_cen_mark == '低于50%' and industry_cen_mark == '高于70%' else '无明显的自上而下或者自下而上选股特征'
        for col in list(industry_stock_cen_disp.columns)[2:]:
            industry_stock_cen_disp[col].iloc[0] = '{0}%'.format(round(industry_stock_cen_disp[col].iloc[0] * 100.0, 2))
        industry_stock_cen_disp.columns = ['基金名称', '基金代码', '行业集中度水平', '个股集中度水平']
        industry_stock_cen_disp = industry_stock_cen_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        industry_stock_cen_ts = hbs_rank_history[['jjdm', 'jsrq', 'c_level_ind', 'c_level_stock']]
        industry_stock_cen_ts = industry_stock_cen_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': '报告日期', 'c_level_ind': '行业集中度水平', 'c_level_stock': '个股集中度水平'})

        fig, ax = plt.subplots(figsize=(6, 3))
        sns.lineplot(x='报告日期', y='行业集中度水平', data=industry_stock_cen_ts, color='#F04950', label='行业集中度水平')
        sns.lineplot(x='报告日期', y='个股集中度水平', data=industry_stock_cen_ts, color='#6268A2', label='个股集中度水平')
        ax.set_xticklabels(labels=industry_stock_cen_ts['报告日期'], rotation=90)
        ax.set_ylim([0.0, 1.0])
        ax.hlines(y=0.5, xmin=0.0, xmax=11.0, linestyles='dashed', color='#959595')
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=1.0, linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.ylabel('集中度水平')
        plt.legend(loc=2)
        plt.tight_layout()
        plt.savefig('{0}{1}_indu_stock_cen.png'.format(self.file_path, self.target_fund_code))

        holding_period = FEDB().read_hbs_stock_trading_property_given_code(self.target_fund_code)
        holding_period = holding_period[['jjdm', '平均持有时间（出持仓前）', '出全仓前平均收益率', '平均持有时间（出持仓前）_rank', '出全仓前平均收益率_rank']]
        holding_period = holding_period.rename(columns={'jjdm': 'FUND_CODE'})
        holding_period['FUND_NAME'] = self.target_fund_name
        holding_days = '{0}'.format(round(holding_period['平均持有时间（出持仓前）'].values[0], 0))
        holding_quarters = '{0}'.format(round(float(holding_days) / 91.0, 1))
        holding_rets = '{0}%'.format(round(holding_period['出全仓前平均收益率'].values[0] * 100, 2))
        holding_days_rank = '{0}%'.format(round(holding_period['平均持有时间（出持仓前）_rank'].values[0] * 100, 2))
        holding_rets_rank = '{0}%'.format(round(holding_period['出全仓前平均收益率_rank'].values[0] * 100, 2))
        holding_days_mark = '偏低' if holding_period['平均持有时间（出持仓前）_rank'].iloc[0] < 0.4 else '偏高' if holding_period['平均持有时间（出持仓前）_rank'].iloc[0] > 0.6 else '中等'
        holding_rets_mark = '偏低' if holding_period['出全仓前平均收益率_rank'].iloc[0] < 0.4 else '偏高' if holding_period['出全仓前平均收益率_rank'].iloc[0] > 0.6 else '中等'
        holding_period_disp = holding_period[['FUND_NAME', 'FUND_CODE', '平均持有时间（出持仓前）', '出全仓前平均收益率', '平均持有时间（出持仓前）_rank', '出全仓前平均收益率_rank']]
        for col in list(holding_period_disp.columns)[3:]:
            holding_period_disp[col].iloc[0] = '{0}%'.format(round(holding_period[col].iloc[0] * 100.0, 2))
        holding_period_disp.columns = ['基金名称', '基金代码', '平均持有时间', '持有期内平均收益率', '平均持有时间水平', '持有期内平均收益率水平']
        holding_period_disp = holding_period_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        left_disp = FEDB().read_hbs_stock_trading_property_given_code(self.target_fund_code)
        left_disp = left_disp[['jjdm', '左侧概率（出持仓前,半年线）_rank', '左侧程度（出持仓前,半年线）']]
        left_disp = left_disp.rename(columns={'jjdm': 'FUND_CODE'})
        left_disp['FUND_NAME'] = self.target_fund_name
        left_mark = '该产品交易具有深度左侧特征' if (left_disp['左侧概率（出持仓前,半年线）_rank'][0] > 0.5 and left_disp['左侧程度（出持仓前,半年线）'][0] > 0.5) \
                    else '该产品交易具有左侧特征' if (left_disp['左侧概率（出持仓前,半年线）_rank'][0] > 0.5 and left_disp['左侧程度（出持仓前,半年线）'][0] <= 0.5) \
                    else '无明显的左侧特征'
        left_prob_data = '{0}%'.format(round(left_disp['左侧概率（出持仓前,半年线）_rank'].values[0] * 100, 2))
        left_level_data = '{0}%'.format(round(left_disp['左侧程度（出持仓前,半年线）'].values[0] * 100, 2))
        left_disp = left_disp[['FUND_NAME', 'FUND_CODE', '左侧概率（出持仓前,半年线）_rank', '左侧程度（出持仓前,半年线）']]
        for col in list(left_disp.columns)[2:]:
            left_disp[col].iloc[0] = '{0}%'.format(round(left_disp[col].iloc[0] * 100.0, 2))
        left_disp.columns = ['基金名称', '基金代码', '左侧概率水平', '左侧程度水平']
        left_disp = left_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        new_disp = FEDB().read_hbs_stock_trading_property_given_code(self.target_fund_code)
        new_disp = new_disp[['jjdm', '新股概率（出持仓前）_rank', '次新股概率（出持仓前）_rank']]
        new_disp = new_disp.rename(columns={'jjdm': 'FUND_CODE'})
        new_disp['FUND_NAME'] = self.target_fund_name
        new_mark = '该产品选股偏好新股/次新股' if new_disp['新股概率（出持仓前）_rank'][0] > 0.75 and new_disp['次新股概率（出持仓前）_rank'][0] > 0.75 else '该产品选股偏好新股' if new_disp['新股概率（出持仓前）_rank'][0] > 0.75 and new_disp['次新股概率（出持仓前）_rank'][0] <= 0.75 else '该产品选股偏好次新股' if new_disp['次新股概率（出持仓前）_rank'][0] <= 0.75 and new_disp['次新股概率（出持仓前）_rank'][0] > 0.75 else '该产品选股无明显的新股/次新股偏好'
        new_data = '{0}%'.format(round(new_disp['新股概率（出持仓前）_rank'].values[0] * 100, 2))
        sub_new_data = '{0}%'.format(round(new_disp['次新股概率（出持仓前）_rank'].values[0] * 100, 2))
        new_disp = new_disp[['FUND_NAME', 'FUND_CODE', '新股概率（出持仓前）_rank', '次新股概率（出持仓前）_rank']]
        for col in list(new_disp.columns)[2:]:
            new_disp[col].iloc[0] = '{0}%'.format(round(new_disp[col].iloc[0] * 100.0, 2))
        new_disp.columns = ['基金名称', '基金代码', '新股概率水平', '次新股概率水平']
        new_disp = new_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        fund_manager_method_text = '该产品近三年投资逻辑如下：'
        fund_manager_method_shift_text = '该产品整体换手率在所有偏股型基金中处于{0}水平，规模、风格、主题、行业换手率排名水平分别为{1}。'.format(shift_mark, shift_data)
        fund_manager_method_cen_text = '该产品整体集中度在所有偏股型基金中处于{0}水平，规模、风格、主题、行业集中度排名水平分别为{1}。'.format(cen_mark, cen_data)
        fund_manager_method_select_text = '该产品行业集中度{0}，个股集中度{1}，{2}。'.format(industry_cen_mark, stock_cen_mark, select_mark)
        fund_manager_method_holding_period_text = '该产品持仓平均时间为{0}天（大约{1}个季度），排名水平为{2}，在所有偏股型基金中处于{3}水平，持有期内平均收益率为{4}，排名水平为{5}，在所有偏股型基金中处于{6}水平。'.format(holding_days, holding_quarters, holding_days_rank, holding_days_mark, holding_rets, holding_rets_rank, holding_rets_mark)
        fund_manager_method_left_text = '该产品左侧概率水平排名{0}，左侧程度水平排名{1}，{2}。'.format(left_prob_data, left_level_data, left_mark)
        fund_manager_method_new_text = '该产品新股概率水平排名{0}，次新股概率水平排名{1}，{2}。'.format(new_data, sub_new_data, new_mark)

        compare_shift, compare_cen, compare_industry_stock_cen, compare_holding_period, compare_left, compare_new = pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_shift_list, compare_cen_list, compare_industry_stock_cen_list, compare_holding_period_list, compare_left_list, compare_new_list = [], [], [], [], [], []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                fund_size_code = FEDB().read_hbs_size_property_given_code(compare_fund_code)
                fund_size_code = fund_size_code.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'shift_lv': 'SIZE_TURNOVER_ABS', 'cen_lv': 'SIZE_CONCENTRATION_ABS', 'shift_lv_rank': 'SIZE_TURNOVER', 'cen_lv_rank': 'SIZE_CONCENTRATION'})
                fund_size_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                size_shift_code = fund_size_code[['FUND_NAME', 'SIZE_TURNOVER_ABS', 'SIZE_TURNOVER']].set_index('FUND_NAME')
                size_cen_code = fund_size_code[['FUND_NAME', 'SIZE_CONCENTRATION_ABS', 'SIZE_CONCENTRATION']].set_index('FUND_NAME')
                size_shift_code.columns = ['规模换手率', '规模换手率排名']
                size_cen_code.columns = ['规模集中度', '规模集中度排名']

                fund_style_code = FEDB().read_hbs_style_property_given_code(compare_fund_code)
                fund_style_code = fund_style_code.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'shift_lv': 'STYLE_TURNOVER_ABS', 'cen_lv': 'STYLE_CONCENTRATION_ABS', 'shift_lv_rank': 'STYLE_TURNOVER', 'cen_lv_rank': 'STYLE_CONCENTRATION'})
                fund_style_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                style_shift_code = fund_style_code[['FUND_NAME', 'STYLE_TURNOVER_ABS', 'STYLE_TURNOVER']].set_index('FUND_NAME')
                style_cen_code = fund_style_code[['FUND_NAME', 'STYLE_CONCENTRATION_ABS', 'STYLE_CONCENTRATION']].set_index('FUND_NAME')
                style_shift_code.columns = ['风格换手率', '风格换手率排名']
                style_cen_code.columns = ['风格集中度', '风格集中度排名']

                fund_theme_code = FEDB().read_hbs_theme_industry_property_given_code(compare_fund_code)
                fund_theme_code = fund_theme_code.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'ratio_theme': 'THEME_TURNOVER_ABS', 'cen_theme': 'THEME_CONCENTRATION_ABS', 'ratio_theme_rank': 'THEME_TURNOVER', 'cen_theme_rank': 'THEME_CONCENTRATION'})
                fund_theme_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                theme_shift_code = fund_theme_code[['FUND_NAME', 'THEME_TURNOVER_ABS', 'THEME_TURNOVER']].set_index('FUND_NAME')
                theme_cen_code = fund_theme_code[['FUND_NAME', 'THEME_CONCENTRATION_ABS', 'THEME_CONCENTRATION']].set_index('FUND_NAME')
                theme_shift_code.columns = ['主题换手率', '主题换手率排名']
                theme_cen_code.columns = ['主题集中度', '主题集中度排名']

                fund_industry_code = FEDB().read_hbs_theme_industry_property_given_code(compare_fund_code)
                fund_industry_code = fund_industry_code.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE', 'ratio_ind': 'INDUSTRY_TURNOVER_ABS', 'cen_ind': 'INDUSTRY_CONCENTRATION_ABS', 'ratio_ind_rank': 'INDUSTRY_TURNOVER', 'cen_ind_rank': 'INDUSTRY_CONCENTRATION'})
                fund_industry_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                industry_shift_code = fund_industry_code[['FUND_NAME', 'INDUSTRY_TURNOVER_ABS', 'INDUSTRY_TURNOVER']].set_index('FUND_NAME')
                industry_cen_code = fund_industry_code[['FUND_NAME', 'INDUSTRY_CONCENTRATION_ABS', 'INDUSTRY_CONCENTRATION']].set_index('FUND_NAME')
                industry_shift_code.columns = ['行业换手率', '行业换手率排名']
                industry_cen_code.columns = ['行业集中度', '行业集中度排名']

                stock_cen_code = FEDB().read_hbs_holding_property_given_code(compare_fund_code)
                stock_cen_code = stock_cen_code[['jjdm', '个股集中度']]
                stock_cen_code = stock_cen_code.rename(columns={'jjdm': 'FUND_CODE', '个股集中度': '个股集中度排名'})
                stock_cen_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                stock_cen_code = stock_cen_code[['FUND_NAME', '个股集中度排名']].set_index('FUND_NAME')
                industry_stock_cen_code = industry_cen_code[['行业集中度排名']].merge(stock_cen_code, left_index=True, right_index=True, how='left')
                industry_stock_cen_code['MARK'] = industry_stock_cen_code.apply(lambda x: '自下而上' if x['行业集中度排名'] < 0.5 and x['个股集中度排名'] > 0.7 else '自上而下' if x['行业集中度排名'] > 0.7 and x['个股集中度排名'] < 0.5 else '无明显的自上而下或者自下而上选股特征', axis=1)
                industry_stock_cen_code.columns = ['行业集中度排名', '个股集中度排名', '选股方法']

                holding_period_code = FEDB().read_hbs_stock_trading_property_given_code(compare_fund_code)
                holding_period_code = holding_period_code[['jjdm', '平均持有时间（出持仓前）', '出全仓前平均收益率', '平均持有时间（出持仓前）_rank', '出全仓前平均收益率_rank']]
                holding_period_code = holding_period_code.rename(columns={'jjdm': 'FUND_CODE'})
                holding_period_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                holding_period_code = holding_period_code[['FUND_NAME', '平均持有时间（出持仓前）', '出全仓前平均收益率', '平均持有时间（出持仓前）_rank', '出全仓前平均收益率_rank']].set_index('FUND_NAME')
                holding_period_code.columns = ['平均持有时间', '持有期内平均收益率', '平均持有时间排名', '持有期内平均收益率排名']

                left_code = FEDB().read_hbs_stock_trading_property_given_code(compare_fund_code)
                left_code = left_code[['jjdm', '左侧概率（出持仓前,半年线）_rank', '左侧程度（出持仓前,半年线）']]
                left_code = left_code.rename(columns={'jjdm': 'FUND_CODE'})
                left_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                left_code['MARK'] = left_code.apply(lambda x: '深度左侧' if x['左侧概率（出持仓前,半年线）_rank'] > 0.5 and x['左侧程度（出持仓前,半年线）'] > 0.5 else '左侧' if x['左侧概率（出持仓前,半年线）_rank'] > 0.5 and x['左侧程度（出持仓前,半年线）'] <= 0.5 else '无明显的左侧特征', axis=1)
                left_code = left_code[['FUND_NAME', '左侧概率（出持仓前,半年线）_rank', '左侧程度（出持仓前,半年线）', 'MARK']].set_index('FUND_NAME')
                left_code.columns = ['左侧概率排名', '左侧程度排名', '左侧特征']

                new_code = FEDB().read_hbs_stock_trading_property_given_code(compare_fund_code)
                new_code = new_code[['jjdm', '新股概率（出持仓前）_rank', '次新股概率（出持仓前）_rank']]
                new_code = new_code.rename(columns={'jjdm': 'FUND_CODE'})
                new_code['FUND_NAME'] = self.compare_fund_name_dic[compare_fund_code]
                new_code['MARK'] = new_code.apply(lambda x: '偏好新股/次新股' if x['新股概率（出持仓前）_rank'] > 0.75 and x['次新股概率（出持仓前）_rank'] > 0.75 else '偏好新股' if x['新股概率（出持仓前）_rank'] > 0.75 and x['次新股概率（出持仓前）_rank'] <= 0.75 else '偏好次新股' if x['新股概率（出持仓前）_rank'] <= 0.75 and x['次新股概率（出持仓前）_rank'] > 0.75 else '无明显的新股/次新股偏好', axis=1)
                new_code = new_code[['FUND_NAME', '新股概率（出持仓前）_rank', '次新股概率（出持仓前）_rank', 'MARK']].set_index('FUND_NAME')
                new_code.columns = ['新股概率排名', '次新股概率排名', '新股/次新股特征']

                shift_code = pd.concat([size_shift_code, style_shift_code, theme_shift_code, industry_shift_code], axis=1)
                shift_code = shift_code[['规模换手率', '风格换手率', '主题换手率', '行业换手率', '规模换手率排名', '风格换手率排名', '主题换手率排名', '行业换手率排名']]
                cen_code = pd.concat([size_cen_code, style_cen_code, theme_cen_code, industry_cen_code], axis=1)
                cen_code = cen_code[['规模集中度', '风格集中度', '主题集中度', '行业集中度', '规模集中度排名', '风格集中度排名', '主题集中度排名', '行业集中度排名']]

                compare_shift_list.append(shift_code)
                compare_cen_list.append(cen_code)
                compare_industry_stock_cen_list.append(industry_stock_cen_code)
                compare_holding_period_list.append(holding_period_code)
                compare_left_list.append(left_code)
                compare_new_list.append(new_code)
            compare_shift = pd.concat(compare_shift_list)
            compare_cen= pd.concat(compare_cen_list)
            compare_industry_stock_cen = pd.concat(compare_industry_stock_cen_list)
            compare_holding_period = pd.concat(compare_holding_period_list)
            compare_left = pd.concat(compare_left_list)
            compare_new = pd.concat(compare_new_list)
            for col in list(compare_shift.columns):
                compare_shift[col] = compare_shift[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)) if not np.isnan(x) else '-')
            for col in list(compare_cen.columns):
                compare_cen[col] = compare_cen[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)) if not np.isnan(x) else '-')
            for col in list(compare_industry_stock_cen.columns)[:-1]:
                compare_industry_stock_cen[col] = compare_industry_stock_cen[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)) if not np.isnan(x) else '-')
            for col in list(compare_holding_period.columns)[1:]:
                compare_holding_period[col] = compare_holding_period[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)) if not np.isnan(x) else '-')
            for col in list(compare_left.columns)[:-1]:
                compare_left[col] = compare_left[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)) if not np.isnan(x) else '-')
            for col in list(compare_new.columns)[:-1]:
                compare_new[col] = compare_new[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)) if not np.isnan(x) else '-')
            compare_shift = compare_shift.T.reset_index().T.reset_index()
            compare_cen = compare_cen.T.reset_index().T.reset_index()
            compare_industry_stock_cen = compare_industry_stock_cen.T.reset_index().T.reset_index()
            compare_holding_period = compare_holding_period.T.reset_index().T.reset_index()
            compare_left = compare_left.T.reset_index().T.reset_index()
            compare_new = compare_new.T.reset_index().T.reset_index()
            compare_shift.iloc[0, 0] = ''
            compare_cen.iloc[0, 0] = ''
            compare_industry_stock_cen.iloc[0, 0] = ''
            compare_holding_period.iloc[0, 0] = ''
            compare_left.iloc[0, 0] = ''
            compare_new.iloc[0, 0] = ''
        return fund_manager_method_text, fund_manager_method_shift_text, shift_disp, fund_manager_method_cen_text, cen_disp, \
               fund_manager_method_select_text, industry_stock_cen_disp, fund_manager_method_holding_period_text, holding_period_disp, \
               fund_manager_method_left_text, left_disp, fund_manager_method_new_text, new_disp,\
               compare_shift, compare_cen, compare_industry_stock_cen, compare_holding_period, compare_left, compare_new

    def get_fund_position(self):
        """
        生成基金仓位数据
        """
        fund_position = HBDB().read_fund_position_given_code(self.target_fund_code)
        fund_position = fund_position.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'jjzzc': 'TOTAL', 'gptzsz': 'STOCK', 'zqzsz': 'BOND', 'jjtzszhj': 'FUND', 'hbzjsz': 'CASH'})
        fund_position['REPORT_DATE'] = fund_position['REPORT_DATE'].astype(str)
        fund_position['STOCK_RATIO'] = fund_position['STOCK'].fillna(0.0) / fund_position['TOTAL'] * 100.0
        fund_position['BOND_RATIO'] = fund_position['BOND'].fillna(0.0) / fund_position['TOTAL'] * 100.0
        fund_position['FUND_RATIO'] = fund_position['FUND'].fillna(0.0) / fund_position['TOTAL'] * 100.0
        fund_position['CASH_RATIO'] = fund_position['CASH'].fillna(0.0) / fund_position['TOTAL'] * 100.0
        fund_position['OTHER_RATIO'] = 100.0 - fund_position[['STOCK_RATIO', 'BOND_RATIO', 'FUND_RATIO', 'CASH_RATIO']].sum(axis=1)

        fund_position = fund_position.sort_values('REPORT_DATE').reset_index().drop('index', axis=1).iloc[-12:]
        fund_position = fund_position[['REPORT_DATE', 'STOCK_RATIO', 'BOND_RATIO', 'FUND_RATIO', 'CASH_RATIO', 'OTHER_RATIO']]
        fund_position.columns = ['报告日期', '股票', '债券', '基金', '现金', '其他']
        fund_position = fund_position.set_index('报告日期')

        change_dates = [date for date in self.change_dates if date >= min(fund_position.index) and date <= max(fund_position.index)]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_position.index.unique().tolist()))
        fig, ax = plt.subplots(figsize=(6, 3))
        fund_position.plot.area(ax=ax, stacked=True, color=['#D55659', '#E1777A', '#8588B7', '#626697', '#7D7D7E'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=100.0, linestyles='dashed', color='#959595')
        plt.legend(loc=2)
        plt.ylabel('占总资产比（%）')
        plt.tight_layout()
        plt.savefig('{0}{1}_position.png'.format(self.file_path, self.target_fund_code))

        fund_position_stock_mean = round(fund_position['股票'].mean(), 2)
        latest_fund_position = fund_position.iloc[-1].sort_values(ascending=False)
        fund_position_text = '该产品资产配置情况如下，其中股票资产近三年平均配置比例为{0}%，最新一期股票、债券、基金、现金、其他资产配置比例分别为{1}%、{2}%、{3}%、{4}%、{5}%。'\
            .format(fund_position_stock_mean, round(latest_fund_position['股票'].mean(), 2), round(latest_fund_position['债券'].mean(), 2), round(latest_fund_position['基金'].mean(), 2), round(latest_fund_position['现金'].mean(), 2), round(latest_fund_position['其他'].mean(), 2))

        compare_position = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_position_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_position_code = HBDB().read_fund_position_given_code(compare_fund_code)
                compare_position_list.append(compare_position_code)
            compare_position = pd.concat(compare_position_list)
            compare_position = compare_position.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'jjzzc': 'TOTAL', 'gptzsz': 'STOCK', 'zqzsz': 'BOND', 'jjtzszhj': 'FUND', 'hbzjsz': 'CASH'})
            compare_position['STOCK_RATIO'] = compare_position['STOCK'].fillna(0.0) / compare_position['TOTAL']
            compare_position = compare_position[['FUND_CODE', 'REPORT_DATE', 'STOCK_RATIO']]
            compare_position['STOCK_RATIO'] = compare_position['STOCK_RATIO'].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
            compare_position = compare_position.pivot(index='REPORT_DATE', columns='FUND_CODE', values='STOCK_RATIO')
            compare_position = compare_position.dropna(subset=[self.target_fund_code], axis=0).sort_index().iloc[-12:]
            compare_position = compare_position[[self.target_fund_code] + self.compare_fund_code_list]
            compare_position.columns = [self.compare_fund_name_dic[col] for col in compare_position.columns]
            compare_position = compare_position.replace(np.nan, '-')
            compare_position = compare_position.T.reset_index().T.reset_index()
            compare_position.iloc[0, 0] = '股票仓位占比'
        return fund_position_text, compare_position

    def get_fund_size(self):
        """
        生成基金规模数据
        """
        size_list = ['大盘', '中盘', '小盘']
        size_rank_list = ['大盘_rank', '中盘_rank', '小盘_rank']
        fund_size = FEDB().read_hbs_size_property_given_code(self.target_fund_code)
        fund_size = fund_size.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE'})

        fund_size_shift = round(fund_size['shift_lv_rank'].values[0] * 100.0, 2)
        fund_size_cen = round(fund_size['cen_lv_rank'].values[0] * 100.0, 2)
        size_type = '配置型' if fund_size_shift <= 50.0 and fund_size_cen <= 50.0 else '专注型' if fund_size_shift <= 50.0 and fund_size_cen > 50.0 else '轮动型' if fund_size_shift > 50.0 and fund_size_cen <= 50.0 else '博弈型'

        largest_size_df = fund_size[['FUND_CODE'] + size_list].set_index('FUND_CODE').T.sort_values(self.target_fund_code, ascending=False)
        largest_size_data = largest_size_df[self.target_fund_code].iloc[0]
        largest_size = '、'.join(list(largest_size_df[largest_size_df[self.target_fund_code] == largest_size_data].index))
        largest_size_data = round(largest_size_data * 100.0, 2)
        largest_size_rank_df = fund_size[['FUND_CODE'] + size_rank_list].set_index('FUND_CODE').T.sort_values(self.target_fund_code, ascending=False)
        largest_size_rank_df.index = map(lambda x: x.split('_')[0], largest_size_rank_df.index)
        largest_size_rank_data = largest_size_rank_df[self.target_fund_code].iloc[0]
        largest_size_rank = '、'.join(list(largest_size_rank_df[largest_size_rank_df[self.target_fund_code] == largest_size_rank_data].index))
        largest_size_rank_data = round(largest_size_rank_data * 100.0, 2)
        fund_size_disp = fund_size.copy(deep=True)
        fund_size_disp['FUND_NAME'] = self.target_fund_name
        for col in ['shift_lv', 'cen_lv'] + size_list + ['shift_lv_rank', 'cen_lv_rank'] + size_rank_list:
            fund_size_disp[col].iloc[0] = '{0}%'.format(round(fund_size_disp[col].iloc[0] * 100.0, 2))
        fund_size_disp = fund_size_disp[['FUND_NAME', 'FUND_CODE'] + ['shift_lv', 'cen_lv'] + size_list + ['shift_lv_rank', 'cen_lv_rank'] + size_rank_list]
        fund_size_disp.columns = ['基金名称', '基金代码'] + ['规模换手率', '规模集中度'] + size_list + ['规模换手率排名', '规模集中度排名'] + [size_rank.replace('_rank', '排名') for size_rank in size_rank_list]
        fund_size_disp = fund_size_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        fund_size_ts = FEDB().read_size_exposure_given_code(self.target_fund_code)
        fund_size_ts = fund_size_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'size_type': 'SIZE_TYPE', 'zjbl': 'MV_IN_NA'})
        fund_size_ts = fund_size_ts[fund_size_ts['REPORT_DATE'].isin(sorted(fund_size_ts['REPORT_DATE'].unique().tolist())[-12:])]
        fund_size_ts = fund_size_ts[['REPORT_DATE', 'SIZE_TYPE', 'MV_IN_NA']]
        fund_size_ts.columns = ['报告日期', '规模类型', '规模暴露水平（%）']

        fund_size_ts['报告日期'] = fund_size_ts['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_size_ts['报告日期']) and date <= max(fund_size_ts['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_size_ts['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(ax=ax, x='报告日期', y='规模暴露水平（%）', data=fund_size_ts, hue='规模类型', hue_order=size_list, palette=['#C94649', '#8588B7', '#7D7D7E'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(fund_size_ts['规模暴露水平（%）']), linestyles='dashed', color='#959595')
        ax.set_xticklabels(labels=sorted(fund_size_ts['报告日期'].unique().tolist()), rotation=90)
        plt.legend(loc=2)
        plt.tight_layout()
        plt.savefig('{0}{1}_size.png'.format(self.file_path, self.target_fund_code))

        fund_size_text = '该产品在统计区间内（近三年）规模换手率在所有偏股型基金池中排名{0}%，规模集中度在所有偏股型基金池中排名{1}%，属于规模{2}基金。从规模绝对暴露看，其在{3}上暴露最大，绝对暴露度为{4}%，从规模相对暴露看，其在{5}上暴露最大，相对暴露度为{6}%。' \
            .format(fund_size_shift, fund_size_cen, size_type, largest_size, largest_size_data, largest_size_rank, largest_size_rank_data)

        compare_size = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_size_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_size_code = FEDB().read_hbs_size_property_given_code(compare_fund_code)
                compare_size_list.append(compare_size_code)
            compare_size = pd.concat(compare_size_list)
            compare_size = compare_size.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE'})
            for col in size_list + size_rank_list:
                compare_size[col] = compare_size[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
            compare_size['FUND_NAME'] = compare_size['FUND_CODE'].apply(lambda x: self.compare_fund_name_dic[x])
            compare_size = compare_size[['FUND_NAME'] + size_list + size_rank_list].set_index('FUND_NAME')
            compare_size = compare_size.loc[[self.compare_fund_name_dic[code] for code in [self.target_fund_code] + self.compare_fund_code_list]]
            compare_size.columns = ['大盘绝对暴露', '中盘绝对暴露', '小盘绝对暴露', '大盘暴露排名', '中盘暴露排名', '小盘暴露排名']
            compare_size = compare_size.replace(np.nan, '0.00%')
            compare_size = compare_size.T.reset_index().T.reset_index()
            compare_size.iloc[0, 0] = ''
        return fund_size_disp, fund_size_text, compare_size

    def get_fund_style(self):
        """
        生成基金风格数据
        """
        style_list = ['成长', '价值']
        style_rank_list = ['成长_rank', '价值_rank']
        fund_style = FEDB().read_hbs_style_property_given_code(self.target_fund_code)
        fund_style = fund_style.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE'})

        fund_style_shift = round(fund_style['shift_lv_rank'].values[0] * 100.0, 2)
        fund_style_cen = round(fund_style['cen_lv_rank'].values[0] * 100.0, 2)
        style_type = '配置型' if fund_style_shift <= 50.0 and fund_style_cen <= 50.0 else '专注型' if fund_style_shift <= 50.0 and fund_style_cen > 50.0 else '轮动型' if fund_style_shift > 50.0 and fund_style_cen <= 50.0 else '博弈型'

        largest_style_df = fund_style[['FUND_CODE'] + style_list].set_index('FUND_CODE').T.sort_values(self.target_fund_code, ascending=False)
        largest_style_data = largest_style_df[self.target_fund_code].iloc[0]
        largest_style = '、'.join(list(largest_style_df[largest_style_df[self.target_fund_code] == largest_style_data].index))
        largest_style_data = round(largest_style_data * 100.0, 2)
        largest_style_rank_df = fund_style[['FUND_CODE'] + style_rank_list].set_index('FUND_CODE').T.sort_values(self.target_fund_code, ascending=False)
        largest_style_rank_df.index = map(lambda x: x.split('_')[0], largest_style_rank_df.index)
        largest_style_rank_data = largest_style_rank_df[self.target_fund_code].iloc[0]
        largest_style_rank = '、'.join(list(largest_style_rank_df[largest_style_rank_df[self.target_fund_code] == largest_style_rank_data].index))
        largest_style_rank_data = round(largest_style_rank_data * 100.0, 2)
        fund_style_disp = fund_style.copy(deep=True)
        fund_style_disp['FUND_NAME'] = self.target_fund_name
        for col in ['shift_lv', 'cen_lv'] + style_list + ['shift_lv_rank', 'cen_lv_rank'] + style_rank_list:
            fund_style_disp[col].iloc[0] = '{0}%'.format(round(fund_style_disp[col].iloc[0] * 100.0, 2))
        fund_style_disp = fund_style_disp[['FUND_NAME', 'FUND_CODE'] + ['shift_lv', 'cen_lv'] + style_list + ['shift_lv_rank', 'cen_lv_rank'] + style_rank_list]
        fund_style_disp.columns = ['基金名称', '基金代码'] + ['风格换手率', '风格集中度'] + style_list + ['风格换手率排名', '风格集中度排名'] + [style_rank.replace('_rank', '排名') for style_rank in style_rank_list]
        fund_style_disp = fund_style_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        fund_style_ts = FEDB().read_style_exposure_given_code(self.target_fund_code)
        fund_style_ts = fund_style_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'style_type': 'STYLE_TYPE', 'zjbl': 'MV_IN_NA'})
        fund_style_ts = fund_style_ts[fund_style_ts['REPORT_DATE'].isin(sorted(fund_style_ts['REPORT_DATE'].unique().tolist())[-12:])]
        fund_style_ts = fund_style_ts[['REPORT_DATE', 'STYLE_TYPE', 'MV_IN_NA']]
        fund_style_ts.columns = ['报告日期', '风格类型', '风格暴露水平（%）']

        fund_style_ts['报告日期'] = fund_style_ts['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_style_ts['报告日期']) and date <= max(fund_style_ts['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_style_ts['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(ax=ax, x='报告日期', y='风格暴露水平（%）', data=fund_style_ts, hue='风格类型', hue_order=style_list, palette=['#C94649', '#8588B7'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(fund_style_ts['风格暴露水平（%）']), linestyles='dashed', color='#959595')
        ax.set_xticklabels(labels=sorted(fund_style_ts['报告日期'].unique().tolist()), rotation=90)
        plt.legend(loc=2)
        plt.tight_layout()
        plt.savefig('{0}{1}_style.png'.format(self.file_path, self.target_fund_code))

        fund_style_text = '该产品在统计区间内（近三年）风格换手率在所有偏股型基金池中排名{0}%，风格集中度在所有偏股型基金池中排名{1}%，属于风格{2}基金。从风格绝对暴露看，其在{3}上暴露最大，绝对暴露度为{4}%，从风格相对暴露看，其在{5}上暴露最大，相对暴露度为{6}%。' \
            .format(fund_style_shift, fund_style_cen, style_type, largest_style, largest_style_data, largest_style_rank, largest_style_rank_data)

        compare_style = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_style_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_style_code = FEDB().read_hbs_style_property_given_code(compare_fund_code)
                compare_style_list.append(compare_style_code)
            compare_style = pd.concat(compare_style_list)
            compare_style = compare_style.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE'})
            for col in style_list + style_rank_list:
                compare_style[col] = compare_style[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
            compare_style['FUND_NAME'] = compare_style['FUND_CODE'].apply(lambda x: self.compare_fund_name_dic[x])
            compare_style = compare_style[['FUND_NAME'] + style_list + style_rank_list].set_index('FUND_NAME')
            compare_style = compare_style.loc[[self.compare_fund_name_dic[code] for code in [self.target_fund_code] + self.compare_fund_code_list]]
            compare_style.columns = ['成长绝对暴露', '价值绝对暴露', '成长暴露排名', '价值暴露排名']
            compare_style = compare_style.replace(np.nan, '0.00%')
            compare_style = compare_style.T.reset_index().T.reset_index()
            compare_style.iloc[0, 0] = ''
        return fund_style_disp, fund_style_text, compare_style

    def get_fund_theme(self):
        """
        生成基金主题数据
        """
        theme_list = ['大金融', '消费', 'TMT', '制造', '周期']
        fund_theme_shift_cen = FEDB().read_hbs_theme_industry_property_given_code(self.target_fund_code)
        fund_theme = FEDB().read_hbs_theme_property_given_code(self.target_fund_code)
        fund_theme = fund_theme_shift_cen.drop('asofdate', axis=1).merge(fund_theme.drop('asofdate', axis=1), on=['jjdm'], how='left')
        fund_theme = fund_theme[['jjdm', 'ratio_theme', 'cen_theme'] + theme_list + ['ratio_theme_rank', 'cen_theme_rank']]
        fund_theme = fund_theme.rename(columns={'jjdm': 'FUND_CODE', 'ratio_theme': 'THEME_TURNOVER', 'cen_theme': 'THEME_CONCENTRATION', 'ratio_theme_rank': 'THEME_TURNOVER_RANK', 'cen_theme_rank': 'THEME_CONCENTRATION_RANK'})

        fund_theme_ratio = round(fund_theme['THEME_TURNOVER_RANK'].values[0] * 100.0, 2)
        fund_theme_cen = round(fund_theme['THEME_CONCENTRATION_RANK'].values[0] * 100.0, 2)
        theme_type = '配置型' if fund_theme_ratio <= 50.0 and fund_theme_cen <= 50.0 else '专注型' if fund_theme_ratio <= 50.0and fund_theme_cen > 50.0 else '轮动型' if fund_theme_ratio > 50.0 and fund_theme_cen <= 50.0 else '博弈型'

        fund_theme_disp = fund_theme.set_index('FUND_CODE')[theme_list].T.sort_values(self.target_fund_code, ascending=False)
        fund_theme_rank = list(fund_theme_disp.index)
        fund_theme_data_rank = fund_theme_disp[self.target_fund_code].unique().tolist()
        fund_theme_data_rank = ['{0}%'.format(round(data * 100.0, 2)) for data in fund_theme_data_rank]
        fund_theme_disp = fund_theme_disp.T.reset_index()
        fund_theme_disp['FUND_NAME'] = self.target_fund_name
        fund_theme_disp = fund_theme_disp.merge(fund_theme[['FUND_CODE', 'THEME_TURNOVER', 'THEME_CONCENTRATION', 'THEME_TURNOVER_RANK', 'THEME_CONCENTRATION_RANK']], on=['FUND_CODE'], how='left')
        for col in ['THEME_TURNOVER', 'THEME_CONCENTRATION'] + fund_theme_rank + ['THEME_TURNOVER_RANK', 'THEME_CONCENTRATION_RANK']:
            fund_theme_disp[col].iloc[0] = '{0}%'.format(round(fund_theme_disp[col].iloc[0] * 100.0, 2))
        fund_theme_disp = fund_theme_disp[['FUND_NAME', 'FUND_CODE', 'THEME_TURNOVER', 'THEME_CONCENTRATION'] + fund_theme_rank + ['THEME_TURNOVER_RANK', 'THEME_CONCENTRATION_RANK']]
        fund_theme_disp.columns = ['基金名称', '基金代码', '主题换手率', '主题集中度'] + fund_theme_rank + ['主题换手率排名', '主题集中度排名']
        fund_theme_disp = fund_theme_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        fund_theme_ts = HBDB().read_fund_theme_given_codes([self.target_fund_code])
        fund_theme_ts = fund_theme_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zclb': 'IS_ZC', 'zblx': 'THEME', 'zgpszb': 'MV_IN_NA', 'jsqbd': 'MV_IN_NA_DIFF', 'tlpj': 'HOMO_AVG'})
        fund_theme_ts['THEME'] = fund_theme_ts['THEME'].replace({'TMT': 'TMT', 'ZZ': '制造', 'ZQ': '周期', 'DJR': '大金融', 'XF': '消费', 'QT': '其他'})
        fund_theme_ts = fund_theme_ts[fund_theme_ts['IS_ZC'] == '1']
        fund_theme_ts = fund_theme_ts.pivot(index='REPORT_DATE', columns='THEME', values='MV_IN_NA').reset_index().fillna(0.0)
        theme_list = [theme for theme in theme_list if theme in list(fund_theme_ts.columns)]
        for i in range(1, len(theme_list)):
            fund_theme_ts[theme_list[i]] = fund_theme_ts[theme_list[i]] + fund_theme_ts[theme_list[i - 1]]
        fund_theme_ts = fund_theme_ts.iloc[-12:]
        fund_theme_ts = fund_theme_ts[['REPORT_DATE'] + theme_list]
        fund_theme_ts.columns = ['报告日期'] + theme_list

        fund_theme_ts['报告日期'] = fund_theme_ts['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_theme_ts['报告日期']) and date <= max(fund_theme_ts['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_theme_ts['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(8, 4))
        color_list = [bar_color_list[0], bar_color_list[1], bar_color_list[7], bar_color_list[8], bar_color_list[14]]
        for i in range(len(theme_list) - 1, -1, -1):
            sns.barplot(ax=ax, x='报告日期', y=theme_list[i], data=fund_theme_ts, label=theme_list[i], color=color_list[i])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(fund_theme_ts[theme_list[-1]]), linestyles='dashed', color='#959595')
        plt.legend(loc='upper right', bbox_to_anchor=(1.2, 1.02))
        plt.xticks(rotation=90)
        plt.ylabel('主题占比（%）')
        plt.tight_layout()
        plt.savefig('{0}{1}_theme.png'.format(self.file_path, self.target_fund_code))

        fund_theme_text = '该产品在统计区间内（近三年）主题换手率在所有偏股型基金池中排名{0}%，主题集中度在所有偏股型基金池中排名{1}%，属于主题{2}基金，其主题暴露排序为{3}，相应的暴露占比为{4}。' \
            .format(fund_theme_ratio, fund_theme_cen, theme_type, '、'.join(fund_theme_rank), '、'.join(fund_theme_data_rank))

        compare_theme = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_theme_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_theme_code = FEDB().read_hbs_theme_property_given_code(compare_fund_code)
                compare_theme_list.append(compare_theme_code)
            compare_theme = pd.concat(compare_theme_list)
            compare_theme = compare_theme.rename(columns={'jjdm': 'FUND_CODE', 'asofdate': 'ASOFDATE'})
            for col in theme_list:
                compare_theme[col] = compare_theme[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
            compare_theme['FUND_NAME'] = compare_theme['FUND_CODE'].apply(lambda x: self.compare_fund_name_dic[x])
            compare_theme = compare_theme[['FUND_NAME'] +theme_list].set_index('FUND_NAME')
            compare_theme = compare_theme.loc[[self.compare_fund_name_dic[code] for code in [self.target_fund_code] + self.compare_fund_code_list]]
            compare_theme = compare_theme.replace(np.nan, '0.00%')
            compare_theme = compare_theme.T.reset_index().T.reset_index()
            compare_theme.iloc[0, 0] = ''
        return fund_theme_disp, fund_theme_text, compare_theme

    def get_fund_industry(self):
        """
        生成基金行业数据
        """
        fund_industry_shift_cen = FEDB().read_hbs_theme_industry_property_given_code(self.target_fund_code)
        fund_industry = FEDB().read_hbs_industry_property_given_code(self.target_fund_code)
        fund_industry = fund_industry.pivot(index='jjdm', columns='yjxymc', values='zsbl_mean')
        industry_list = list(fund_industry.columns)
        fund_industry = fund_industry_shift_cen.drop('asofdate', axis=1).merge(fund_industry.reset_index(), on=['jjdm'], how='left')
        fund_industry = fund_industry[['jjdm', 'ratio_ind', 'cen_ind'] + industry_list + ['ratio_ind_rank', 'cen_ind_rank']]
        fund_industry = fund_industry.rename(columns={'jjdm': 'FUND_CODE', 'ratio_ind': 'INDUSTRY_TURNOVER', 'cen_ind': 'INDUSTRY_CONCENTRATION', 'ratio_ind_rank': 'INDUSTRY_TURNOVER_RANK', 'cen_ind_rank': 'INDUSTRY_CONCENTRATION_RANK'})

        fund_industry_ratio = round(fund_industry['INDUSTRY_TURNOVER_RANK'].values[0] * 100.0, 2)
        fund_industry_cen = round(fund_industry['INDUSTRY_CONCENTRATION_RANK'].values[0] * 100.0, 2)
        industry_type = '配置型' if fund_industry_ratio <= 50.0 and fund_industry_cen <= 50.0 else '专注型' if fund_industry_ratio <= 50.0and fund_industry_cen > 50.0 else '轮动型' if fund_industry_ratio > 50.0 and fund_industry_cen <= 50.0 else '博弈型'

        fund_industry_disp = fund_industry.set_index('FUND_CODE')[industry_list].T.sort_values(self.target_fund_code, ascending=False)
        fund_industry_rank = list(fund_industry_disp.index)[:5]
        fund_industry_data_rank = fund_industry_disp[self.target_fund_code].unique().tolist()[:5]
        fund_industry_data_rank = ['{0}%'.format(round(data * 100.0, 2)) for data in fund_industry_data_rank]
        fund_industry_disp = fund_industry_disp.T.reset_index()
        fund_industry_disp['FUND_NAME'] = self.target_fund_name
        fund_industry_disp = fund_industry_disp.merge(fund_industry[['FUND_CODE', 'INDUSTRY_TURNOVER', 'INDUSTRY_CONCENTRATION', 'INDUSTRY_TURNOVER_RANK', 'INDUSTRY_CONCENTRATION_RANK']], on=['FUND_CODE'], how='left')
        for col in ['INDUSTRY_TURNOVER', 'INDUSTRY_CONCENTRATION'] + fund_industry_rank + ['INDUSTRY_TURNOVER_RANK', 'INDUSTRY_CONCENTRATION_RANK']:
            fund_industry_disp[col].iloc[0] = '{0}%'.format(round(fund_industry_disp[col].iloc[0] * 100.0, 2))
        fund_industry_disp = fund_industry_disp[['FUND_NAME', 'FUND_CODE', 'INDUSTRY_TURNOVER', 'INDUSTRY_CONCENTRATION'] + fund_industry_rank + ['INDUSTRY_TURNOVER_RANK', 'INDUSTRY_CONCENTRATION_RANK']]
        fund_industry_disp.columns = ['基金名称', '基金代码', '行业换手率', '行业集中度'] + fund_industry_rank + ['行业换手率排名', '行业集中度排名']
        fund_industry_disp = fund_industry_disp.T.reset_index().T.reset_index().drop('index', axis=1)

        industry_new_list = ['电力设备', '商贸零售', '食品饮料', '农林牧渔', '有色金属', '银行', '非银金融', '汽车', '石油石化', '建筑材料', '房地产', '钢铁', '建筑装饰', '公用事业', '医药生物', '家用电器', '电子', '通信', '轻工制造', '机械设备', '传媒', '基础化工', '计算机', '煤炭', '交通运输', '纺织服装', '环保', '社会服务', '美容护理', '国防军工', '综合']
        fund_holding = HBDB().read_fund_holding_given_codes([self.target_fund_code])
        fund_holding = fund_holding.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zqdm': 'TICKER_SYMBOL', 'zqmc': 'SEC_SHORT_NAME', 'ccsz': 'HOLDING_MARKET_VALUE', 'ccsl': 'HOLDING_AMOUNT', 'zjbl': 'MV_IN_NA'})
        fund_holding = fund_holding.sort_values(['FUND_CODE', 'REPORT_DATE', 'MV_IN_NA'], ascending=[True, True, False]).groupby(['FUND_CODE', 'REPORT_DATE']).head(10)
        stock_industry = HBDB().read_stock_industry()
        stock_industry = stock_industry.rename(columns={'zqdm': 'TICKER_SYMBOL', 'flmc': 'INDUSTRY_NAME', 'fldm': 'INDUSTRY_ID', 'fljb': 'INDUSTRY_TYPE', 'hyhfbz': 'INDUSTRY_VERSION', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'sfyx': 'IS_NEW'})
        stock_industry = stock_industry[stock_industry['IS_NEW'] == 1]
        stock_industry = stock_industry[stock_industry['INDUSTRY_VERSION'] == 2]
        stock_industry = stock_industry[stock_industry['INDUSTRY_TYPE'] == '1']
        stock_industry = stock_industry.sort_values(['TICKER_SYMBOL', 'credt_etl']).drop_duplicates('TICKER_SYMBOL', keep='last')
        fund_holding = fund_holding.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='left')
        fund_industry_ts = fund_holding[['REPORT_DATE', 'INDUSTRY_NAME', 'MV_IN_NA']].groupby(['REPORT_DATE', 'INDUSTRY_NAME']).sum().reset_index()
        fund_industry_ts = fund_industry_ts.pivot(index='REPORT_DATE', columns='INDUSTRY_NAME', values='MV_IN_NA').reset_index().fillna(0.0)
        industry_ts_list = [indu for indu in industry_new_list if indu in list(fund_industry_ts.columns)[1:]]
        for i in range(1, len(industry_ts_list)):
            fund_industry_ts[industry_ts_list[i]] = fund_industry_ts[industry_ts_list[i]] + fund_industry_ts[industry_ts_list[i - 1]]
        fund_industry_ts = fund_industry_ts.iloc[-12:]
        fund_industry_ts = fund_industry_ts[['REPORT_DATE'] + industry_ts_list]
        fund_industry_ts.columns = ['报告日期'] + industry_ts_list


        # fund_industry_ts = HBDB().read_fund_industry_given_codes([fund_code])
        # fund_industry_ts = fund_industry_ts.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zclb': 'IS_ZC', 'hyhfbz': 'INDUSTRY_VERSION', 'fldm': 'INDUSTRY_ID', 'flmc': 'INDUSTRY_NAME', 'zzjbl': 'MV_IN_NA', 'hyzjzbjbd': 'MV_IN_NA_DIFF', 'hyzjzbltlpj': 'HOMO_AVG'})
        # fund_industry_ts = fund_industry_ts[fund_industry_ts['INDUSTRY_VERSION'] == '2']
        # fund_industry_ts = fund_industry_ts[fund_industry_ts['IS_ZC'] == '1']
        # fund_industry_ts = fund_industry_ts.pivot(index='REPORT_DATE', columns='INDUSTRY_NAME', values='MV_IN_NA').reset_index().fillna(0.0)
        # industry_ts_list = list(fund_industry_ts.columns)[1:]
        # for i in range(1, len(industry_ts_list)):
        #     fund_industry_ts[industry_ts_list[i]] = fund_industry_ts[industry_ts_list[i]] + fund_industry_ts[industry_ts_list[i - 1]]
        # fund_industry_ts = fund_industry_ts.iloc[-12:]
        # fund_industry_ts = fund_industry_ts[['REPORT_DATE'] + industry_ts_list]
        # fund_industry_ts.columns = ['报告日期'] + industry_ts_list

        fund_industry_ts['报告日期'] = fund_industry_ts['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_industry_ts['报告日期']) and date <= max(fund_industry_ts['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_industry_ts['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(10, 6))
        palette = bar_color_list + bar_color_list
        for i in range(len(industry_ts_list) - 1, -1, -1):
            sns.barplot(ax=ax, x='报告日期', y=industry_ts_list[i], data=fund_industry_ts, label=industry_ts_list[i], color=palette[i])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(fund_industry_ts[industry_ts_list[-1]]), linestyles='dashed', color='#959595')
        plt.legend(loc='upper right', bbox_to_anchor=(1.2, 1.03))
        plt.xticks(rotation=90)
        plt.ylabel('行业占比（%）')
        plt.tight_layout()
        plt.savefig('{0}{1}_industry.png'.format(self.file_path, self.target_fund_code))

        fund_industry_text = '该产品在统计区间内（近三年）行业换手率在所有偏股型基金池中排名{0}%，行业集中度在所有偏股型基金池中排名{1}%，属于行业{2}基金，其前五大行业为{3}，相应的暴露占比为{4}。' \
            .format(fund_industry_ratio, fund_industry_cen, industry_type, '、'.join(fund_industry_rank), '、'.join(fund_industry_data_rank))

        compare_industry = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_industry_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_industry_code = FEDB().read_hbs_industry_property_given_code(compare_fund_code)
                compare_industry_code = compare_industry_code[['jjdm', 'yjxymc', 'zsbl_mean']].sort_values('zsbl_mean', ascending=False).groupby(['jjdm']).head(5)
                compare_industry_code['MARK'] = range(1, 6)
                compare_industry_code['MARK'] = compare_industry_code['MARK'].apply(lambda x: 'TOP{0}'.format(x))
                compare_industry_list.append(compare_industry_code)
            compare_industry = pd.concat(compare_industry_list)
            compare_industry['INDUSTRY_DISP'] = compare_industry.apply(lambda x: '{0}（{1}%）'.format(x['yjxymc'], round(x['zsbl_mean'] * 100.0, 2)), axis=1)
            compare_industry['FUND_NAME'] = compare_industry['jjdm'].apply(lambda x: self.compare_fund_name_dic[x])
            compare_industry = compare_industry.pivot(index='FUND_NAME', columns='MARK', values='INDUSTRY_DISP')
            compare_industry = compare_industry.loc[[self.compare_fund_name_dic[code] for code in [self.target_fund_code] + self.compare_fund_code_list]]
            compare_industry = compare_industry.T.reset_index().T.reset_index()
            compare_industry.iloc[0, 0] = ''
        return fund_industry_disp, fund_industry_text, compare_industry

    def get_fund_industry_change(self):
        """
        生成基金行业变动数据
        """
        fund_holding = HBDB().read_fund_holding_given_codes([self.target_fund_code])
        fund_holding = fund_holding.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zqdm': 'TICKER_SYMBOL', 'zqmc': 'SEC_SHORT_NAME', 'ccsz': 'HOLDING_MARKET_VALUE', 'ccsl': 'HOLDING_AMOUNT', 'zjbl': 'MV_IN_NA'})
        fund_holding['REPORT_DATE'] = fund_holding['REPORT_DATE'].astype(str)
        stock_industry = HBDB().read_stock_industry()
        stock_industry = stock_industry.rename(columns={'zqdm': 'TICKER_SYMBOL', 'flmc': 'INDUSTRY_NAME', 'fldm': 'INDUSTRY_ID', 'fljb': 'INDUSTRY_TYPE', 'hyhfbz': 'INDUSTRY_VERSION', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'sfyx': 'IS_NEW'})
        stock_industry = stock_industry[stock_industry['IS_NEW'] == 1]
        stock_industry = stock_industry[stock_industry['INDUSTRY_VERSION'] == 2]
        stock_industry = stock_industry[stock_industry['INDUSTRY_TYPE'] == '1']
        stock_industry = stock_industry.sort_values(['TICKER_SYMBOL', 'credt_etl']).drop_duplicates('TICKER_SYMBOL', keep='last')
        fund_holding = fund_holding.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='left')
        fund_holding = fund_holding[['REPORT_DATE', 'INDUSTRY_NAME', 'MV_IN_NA']].groupby(['REPORT_DATE', 'INDUSTRY_NAME']).sum().reset_index()
        fund_industry = fund_holding.pivot(index='REPORT_DATE', columns='INDUSTRY_NAME', values='MV_IN_NA').fillna(0.0)
        fund_industry = fund_industry.loc[fund_industry.index.str.slice(4, 6).isin(['06', '12'])]
        fund_industry_change = fund_industry.diff().dropna()
        fund_industry_change = fund_industry_change.unstack().reset_index()
        fund_industry_change.columns = ['INDUSTRY_NAME', 'REPORT_DATE', 'MV_IN_NA_DIFF']
        # fund_industry_change = HBDB().read_fund_industry_given_codes([self.target_fund_code])
        # fund_industry_change = fund_industry_change.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zclb': 'IS_ZC', 'hyhfbz': 'INDUSTRY_VERSION', 'fldm': 'INDUSTRY_ID', 'flmc': 'INDUSTRY_NAME', 'zzjbl': 'MV_IN_NA', 'hyzjzbjbd': 'MV_IN_NA_DIFF', 'hyzjzbltlpj': 'HOMO_AVG'})
        # fund_industry_change = fund_industry_change[fund_industry_change['INDUSTRY_VERSION'] == '2']
        # fund_industry_change = fund_industry_change[fund_industry_change['IS_ZC'] == '2']
        # fund_industry_change = fund_industry_change[['REPORT_DATE', 'INDUSTRY_NAME', 'MV_IN_NA_DIFF']]
        fund_industry_change = fund_industry_change.sort_values(['REPORT_DATE', 'MV_IN_NA_DIFF'], ascending=[True, False])
        fund_industry_change = fund_industry_change[fund_industry_change['REPORT_DATE'].isin(sorted(fund_industry_change['REPORT_DATE'].unique().tolist())[-6:])]
        fund_industry_up = fund_industry_change.groupby('REPORT_DATE').head(5).sort_values(['REPORT_DATE', 'MV_IN_NA_DIFF'], ascending=[False, False])
        fund_industry_down = fund_industry_change.groupby('REPORT_DATE').tail(5).sort_values(['REPORT_DATE', 'MV_IN_NA_DIFF'], ascending=[False, True])
        fund_industry_up['MV_IN_NA_DIFF'] = fund_industry_up['MV_IN_NA_DIFF'].apply(lambda x: '{0}%'.format(round(x, 2)))
        fund_industry_down['MV_IN_NA_DIFF'] = fund_industry_down['MV_IN_NA_DIFF'].apply(lambda x: '{0}%'.format(round(x, 2)))
        fund_industry_up_list, fund_industry_down_list = [], []
        date_list = fund_industry_up['REPORT_DATE'].unique().tolist()
        for date in date_list:
            fund_industry_up_date = fund_industry_up[fund_industry_up['REPORT_DATE'] == date]
            fund_industry_up_date.index = range(1, 6)
            fund_industry_up_list.append(fund_industry_up_date[['INDUSTRY_NAME', 'MV_IN_NA_DIFF']])
            fund_industry_down_date = fund_industry_down[fund_industry_down['REPORT_DATE'] == date]
            fund_industry_down_date.index = range(1, 6)
            fund_industry_down_list.append(fund_industry_down_date[['INDUSTRY_NAME', 'MV_IN_NA_DIFF']])
        fund_industry_up = pd.concat(fund_industry_up_list, axis=1)
        fund_industry_down = pd.concat(fund_industry_down_list, axis=1)
        fund_industry_up.columns = [sorted(date_list * 2)[len(sorted(date_list * 2)) - 1 - i] for i, date in enumerate(sorted(date_list * 2))]
        fund_industry_up = fund_industry_up.T.reset_index().T.reset_index().replace({'index': '时间', 1: '第一名', 2: '第二名', 3: '第三名', 4: '第四名', 5: '第五名'})
        fund_industry_down.columns = [sorted(date_list * 2)[len(sorted(date_list * 2)) - 1 - i] for i, date in enumerate(sorted(date_list * 2))]
        fund_industry_down = fund_industry_down.T.reset_index().T.reset_index().replace({'index': '时间', 1: '第一名', 2: '第二名', 3: '第三名', 4: '第四名', 5: '第五名'})

        fund_industry_up_text = '增持前五大行业（净值占比）如下：'
        fund_industry_down_text = '减持前五大行业（净值占比）如下：'
        return fund_industry_up, fund_industry_up_text, fund_industry_down, fund_industry_down_text

    def get_fund_cr(self):
        """
        生成基金持仓集中度数据
        """
        fund_cr = HBDB().read_fund_cr_given_code(self.target_fund_code)
        fund_cr = fund_cr.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'qsdzb': 'CR10', 'qsdzbtlpj': 'CR10_MEAN', 'qwdzb': 'CR5', 'qwdzbtlpj': 'CR5_MEAN'})
        fund_cr = fund_cr.sort_values('REPORT_DATE').iloc[-12:]
        fund_cr = fund_cr[['REPORT_DATE', 'CR5', 'CR5_MEAN', 'CR10', 'CR10_MEAN']]
        fund_cr.columns = ['报告日期', '前五大股票占比（%）', '前五大股票占比同类平均（%）', '前十大股票占比（%）', '前十大股票占比同类平均（%）']

        cr5 = fund_cr[['报告日期', '前五大股票占比（%）']].rename(columns={'前五大股票占比（%）': '占净资产比（%）'})
        cr5['前五大占比'] = '前五大股票占比'
        cr5_mean = fund_cr[['报告日期', '前五大股票占比同类平均（%）']].rename(columns={'前五大股票占比同类平均（%）': '占净资产比（%）'})
        cr5_mean['前五大占比'] = '前五大股票占比同类平均'
        cr5_mean = pd.concat([cr5, cr5_mean]).reset_index().drop('index', axis=1)
        cr10 = fund_cr[['报告日期', '前十大股票占比（%）']].rename(columns={'前十大股票占比（%）': '占净资产比（%）'})
        cr10['前十大占比'] = '前十大股票占比'
        cr10_mean = fund_cr[['报告日期', '前十大股票占比同类平均（%）']].rename(columns={'前十大股票占比同类平均（%）': '占净资产比（%）'})
        cr10_mean['前十大占比'] = '前十大股票占比同类平均'
        cr10_mean = pd.concat([cr10, cr10_mean]).reset_index().drop('index', axis=1)

        cr5_mean['报告日期'] = cr5_mean['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(cr5_mean['报告日期']) and date <= max(cr5_mean['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(cr5_mean['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(ax=ax, x='报告日期', y='占净资产比（%）', data=cr5_mean, hue='前五大占比', palette=['#C94649', '#8588B7'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(cr5_mean['占净资产比（%）']), linestyles='dashed', color='#959595')
        ax.legend(loc=2)
        ax.set_xticklabels(labels=fund_cr['报告日期'], rotation=90)
        plt.tight_layout()
        plt.savefig('{0}{1}_cr5.png'.format(self.file_path, self.target_fund_code))

        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(ax=ax, x='报告日期', y='占净资产比（%）', data=cr10_mean, hue='前十大占比', palette=['#C94649', '#8588B7'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(cr10_mean['占净资产比（%）']), linestyles='dashed', color='#959595')
        ax.legend(loc=2)
        ax.set_xticklabels(labels=fund_cr['报告日期'], rotation=90)
        plt.tight_layout()
        plt.savefig('{0}{1}_cr10.png'.format(self.file_path, self.target_fund_code))

        fund_cr5_mean = round(fund_cr['前五大股票占比（%）'].mean(), 2)
        fund_cr5_mean_mean = round(fund_cr['前五大股票占比同类平均（%）'].mean(), 2)
        fund_cr10_mean = round(fund_cr['前十大股票占比（%）'].mean(), 2)
        fund_cr10_mean_mean = round(fund_cr['前十大股票占比同类平均（%）'].mean(), 2)
        cr5_text = '低' if fund_cr5_mean < fund_cr5_mean_mean else '高'
        cr10_text = '低' if fund_cr10_mean < fund_cr10_mean_mean else '高'
        text = '，持股集中度低'if cr5_text == '低' and cr10_text == '低' else ''

        fund_cr_text = '该产品持股集中度情况如下，近三年前五大权重保持在{0}%左右，整体{1}于同类平均水平，前十大权重保持在{2}%左右，整体{3}于同类平均水平{4}。'\
            .format(fund_cr5_mean, cr5_text, fund_cr10_mean, cr10_text, text)
        return fund_cr_text

    def get_fund_finance(self):
        """
        生成基金持仓财务数据
        """
        fund_finance = HBDB().read_fund_valuation_given_codes([self.target_fund_code])
        fund_finance = fund_finance.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zclb': 'IS_ZC', 'pe': 'PE', 'pb': 'PB', 'roe': 'ROE', 'dividend': 'DIVIDEND'})
        fund_finance['REPORT_DATE'] = fund_finance['REPORT_DATE'].astype(int)
        fund_finance = fund_finance[fund_finance['IS_ZC'] == 1]
        fund_finance = fund_finance.sort_values('REPORT_DATE').iloc[-12:].reset_index().drop('index', axis=1)
        fund_finance = fund_finance[['REPORT_DATE', 'PE', 'PB', 'ROE', 'DIVIDEND']]
        fund_finance = fund_finance[(fund_finance['PE'] != 99999.0) & (fund_finance['PB'] != 99999.0) & (fund_finance['ROE'] != 99999.0) & (fund_finance['DIVIDEND'] != 99999.0)]
        windA_finance = HBDB().read_index_valuation_given_date_and_indexs((datetime.today() - timedelta(365 * 4)).strftime('%Y%m%d'), ['881001'])
        windA_finance = windA_finance.rename(columns={'zqdm': 'INDEX_SYMBOL', 'zqmc': 'INDEX_NAME', 'jyrq': 'REPORT_DATE', 'pe': 'PE_windA', 'pb': 'PB_windA', 'roe': 'ROE_windA', 'gxl': 'DIVIDEND_windA'})
        windA_finance['REPORT_DATE'] = windA_finance['REPORT_DATE'].astype(int)
        windA_finance = windA_finance[['REPORT_DATE', 'PE_windA', 'PB_windA', 'ROE_windA', 'DIVIDEND_windA']]
        fund_finance = pd.merge_asof(fund_finance, windA_finance, on='REPORT_DATE')
        fund_finance['REPORT_DATE'] = fund_finance['REPORT_DATE'].astype(str)
        fund_finance.columns = ['报告日期', 'PE', 'PB', 'ROE（%）', '股息率（%）', '万得全A_PE', '万得全A_PB', '万得全A_ROE（%）', '万得全A_股息率（%）']

        pe_mean = round(fund_finance['PE'].mean(), 2)
        pb_mean = round(fund_finance['PB'].mean(), 2)
        roe_mean = round(fund_finance['ROE（%）'].mean(), 2)
        dividend_mean = round(fund_finance['股息率（%）'].mean(), 2)

        fund_finance['报告日期'] = fund_finance['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_finance['报告日期']) and date <= max(fund_finance['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_finance['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(2, 2, figsize=(6, 6))
        pic_list = [['PE', 'PB'], ['ROE（%）', '股息率（%）']]
        for i in range(2):
            for j in range(2):
                axij = sns.lineplot(ax=ax[i][j], x='报告日期', y=pic_list[i][j], data=fund_finance, color='#F04950', label=self.target_fund_code)
                axij = sns.lineplot(ax=ax[i][j], x='报告日期', y='万得全A_' + pic_list[i][j], data=fund_finance, color='#6268A2', label='万得全A')
                axij.vlines(x=change_dates_loc, ymin=0.0, ymax=max(max(fund_finance[pic_list[i][j]]), max(fund_finance['万得全A_' + pic_list[i][j]])), linestyles='dashed', color='#959595')
                axij.set_title('{0}走势'.format(pic_list[i][j].replace('（%）', '')))
                axij.set_xticklabels(labels=fund_finance['报告日期'], rotation=90)
                axij.legend(loc=2)
        plt.tight_layout()
        plt.savefig('{0}{1}_finance.png'.format(self.file_path, self.target_fund_code))

        fund_finance_text = '该产品PE、PB、ROE、股息率走势如下，在统计区间内（近三年），PE维持在{0}倍左右，PB维持在{1}倍左右，ROE维持在{2}%左右，股息率维持在{3}%左右。'\
            .format(pe_mean, pb_mean, roe_mean, dividend_mean)

        compare_finance = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_finance_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_finance_code = HBDB().read_fund_valuation_given_codes([compare_fund_code])
                compare_finance_code = compare_finance_code.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zclb': 'IS_ZC', 'pe': 'PE', 'pb': 'PB', 'roe': 'ROE', 'dividend': 'DIVIDEND'})
                compare_finance_code['REPORT_DATE'] = compare_finance_code['REPORT_DATE'].astype(int)
                compare_finance_code = compare_finance_code[compare_finance_code['IS_ZC'] == 1]
                compare_finance_code = compare_finance_code.sort_values('REPORT_DATE').iloc[-12:].reset_index().drop('index', axis=1)
                compare_finance_code = compare_finance_code[(compare_finance_code['PE'] != 99999.0) & (compare_finance_code['PB'] != 99999.0) & (compare_finance_code['ROE'] != 99999.0) & (compare_finance_code['DIVIDEND'] != 99999.0)]
                compare_finance_code = pd.DataFrame(compare_finance_code[['PE', 'PB', 'ROE', 'DIVIDEND']].mean()).T
                compare_finance_code['FUND_CODE'] = compare_fund_code
                compare_finance_list.append(compare_finance_code)
            compare_finance = pd.concat(compare_finance_list)
            compare_finance['FUND_NAME'] = compare_finance['FUND_CODE'].apply(lambda x: self.compare_fund_name_dic[x])
            compare_finance = compare_finance[['FUND_NAME', 'PE', 'PB', 'ROE', 'DIVIDEND']].set_index('FUND_NAME')
            compare_finance['PE'] = compare_finance['PE'].apply(lambda x: '{0}'.format(round(x, 2)))
            compare_finance['PB'] = compare_finance['PB'].apply(lambda x: '{0}'.format(round(x, 2)))
            compare_finance['ROE'] = compare_finance['ROE'].apply(lambda x: '{0}%'.format(round(x, 2)))
            compare_finance['DIVIDEND'] = compare_finance['DIVIDEND'].apply(lambda x: '{0}%'.format(round(x, 2)))
            compare_finance.columns = ['PE', 'PB', 'ROE', '股息率']
            compare_finance = compare_finance.loc[[self.compare_fund_name_dic[code] for code in [self.target_fund_code] + self.compare_fund_code_list]]
            compare_finance = compare_finance.T.reset_index().T.reset_index()
            compare_finance.iloc[0, 0] = ''
        return fund_finance_text, compare_finance

    def get_fund_stock_change(self):
        """
        获取基金个股变动数据
        """
        fund_holding = HBDB().read_fund_holding_given_codes([self.target_fund_code])
        fund_holding = fund_holding.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zqdm': 'TICKER_SYMBOL', 'zqmc': 'SEC_SHORT_NAME', 'ccsz': 'HOLDING_MARKET_VALUE', 'ccsl': 'HOLDING_AMOUNT', 'zjbl': 'MV_IN_NA'})
        fund_holding_diff = HBDB().read_fund_holding_diff_given_codes([self.target_fund_code])
        fund_holding_diff = fund_holding_diff.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zqdm': 'TICKER_SYMBOL', 'zqmc': 'SEC_SHORT_NAME', 'zclb': 'IS_ZC', 'zgblbd': 'MV_IN_NA_DIFF', 'sfsqzcg': 'IS_LAST_ZC'})
        stock_industry = HBDB().read_stock_industry()
        stock_industry = stock_industry.rename(columns={'zqdm': 'TICKER_SYMBOL', 'flmc': 'INDUSTRY_NAME', 'fldm': 'INDUSTRY_ID', 'fljb': 'INDUSTRY_TYPE', 'hyhfbz': 'INDUSTRY_VERSION', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'sfyx': 'IS_NEW'})
        stock_industry = stock_industry[stock_industry['IS_NEW'] == 1]
        stock_industry = stock_industry[stock_industry['INDUSTRY_VERSION'] == 2]
        stock_industry = stock_industry[stock_industry['INDUSTRY_TYPE'] == '1']
        stock_industry = stock_industry.sort_values(['TICKER_SYMBOL', 'credt_etl']).drop_duplicates('TICKER_SYMBOL', keep='last')
        fund_holding = fund_holding.sort_values(['REPORT_DATE', 'MV_IN_NA'], ascending=[True, False]).reset_index().drop('index', axis=1)
        fund_holding = fund_holding.groupby('REPORT_DATE').head(10)
        fund_holding_diff = fund_holding_diff[fund_holding_diff['IS_ZC'] == '1']
        fund_holding = fund_holding.merge(fund_holding_diff[['REPORT_DATE', 'TICKER_SYMBOL', 'MV_IN_NA_DIFF']], on=['REPORT_DATE', 'TICKER_SYMBOL'], how='left')
        fund_holding = fund_holding.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='left')

        fund_holding_latest = fund_holding[fund_holding['REPORT_DATE'] == fund_holding['REPORT_DATE'].max()]
        fund_holding_latest = fund_holding_latest[['SEC_SHORT_NAME', 'INDUSTRY_NAME', 'MV_IN_NA_DIFF', 'HOLDING_AMOUNT', 'HOLDING_MARKET_VALUE', 'MV_IN_NA']]
        fund_holding_latest[['MV_IN_NA_DIFF', 'HOLDING_AMOUNT', 'HOLDING_MARKET_VALUE', 'MV_IN_NA']] = round(fund_holding_latest[['MV_IN_NA_DIFF', 'HOLDING_AMOUNT', 'HOLDING_MARKET_VALUE', 'MV_IN_NA']], 2)
        fund_holding_latest['MV_IN_NA_DIFF'] = fund_holding_latest['MV_IN_NA_DIFF'].apply(lambda x: '{0}%'.format(x))
        fund_holding_latest['MV_IN_NA'] = fund_holding_latest['MV_IN_NA'].apply(lambda x: '{0}%'.format(x))
        fund_holding_latest.columns = ['重仓股名称', '所属行业', '占股票市值比较上期变动', '持仓数量（股）', '持仓市值（元）', '占净值比例']
        fund_holding_latest = fund_holding_latest.T.reset_index().T.reset_index().drop('index', axis=1)

        need_dates = sorted(fund_holding['REPORT_DATE'].unique().tolist())[-6:]
        disp_dates = [need_dates[len(need_dates) - 1 - i] for i, date in enumerate(need_dates)]
        fund_holding_disp = fund_holding[fund_holding['REPORT_DATE'].isin(need_dates)]
        fund_holding_disp = fund_holding_disp.pivot(index='REPORT_DATE', columns='SEC_SHORT_NAME', values='MV_IN_NA').T
        fund_holding_disp['平均占净值比例'] = fund_holding_disp.mean(axis=1)
        fund_holding_disp = fund_holding_disp.sort_values(['平均占净值比例'], ascending=False)
        fund_holding_disp_name = fund_holding_disp.index[0]
        fund_holding_disp_data = round(fund_holding_disp['平均占净值比例'].iloc[0], 2)
        fund_holding_disp = fund_holding_disp[disp_dates + ['平均占净值比例']]
        for col in disp_dates + ['平均占净值比例']:
            fund_holding_disp[col] = fund_holding_disp[col].apply(lambda x: '{0}%'.format(round(x, 2)) if not np.isnan(x) else '-')
        fund_holding_disp = fund_holding_disp.T.reset_index().T.reset_index()
        fund_holding_disp.iloc[0, 0] = '股票名称'

        fund_holding_dist = fund_holding.copy(deep=True)
        fund_holding_dist = fund_holding_dist[fund_holding_dist['REPORT_DATE'].isin(sorted(fund_holding['REPORT_DATE'].unique().tolist())[-12:])]
        fund_holding_dist = fund_holding_dist.rename(columns={'MV_IN_NA': '占净值比例'})
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.distplot(fund_holding_dist['占净值比例'], ax=ax, bins=21, hist=True, kde=True, color='#F04950')
        plt.gca().xaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.ylabel('概率分布')
        plt.tight_layout()
        plt.savefig('{0}{1}_dist.png'.format(self.file_path, self.target_fund_code))

        dic = {}
        for i in range(40):
            dic[i] = len(fund_holding_dist[(fund_holding_dist['占净值比例'] >= i) & (fund_holding_dist['占净值比例'] < i + 1)])
        disp = pd.DataFrame([dic]).T.sort_values(0, ascending=False)
        period = '[{0}%, {1}%]'.format(disp.index[0], disp.index[0] + 1)

        fund_holding_latest_text = '该产品最新一期前十大重仓情况如下：'
        fund_holding_disp_text = '该产品近三年前十大重仓分布如下，从概率密度分布图可以看出，该产品持仓权重分布最多的区间是{0}，近六期平均占净值比例最大的股票为{1}，占比为{2}%。'.format(period, fund_holding_disp_name, fund_holding_disp_data)
        return fund_holding_latest, fund_holding_latest_text, fund_holding_disp, fund_holding_disp_text

    def get_fund_stock_compare(self):
        """
        获取基金同细分行业个股比较数据
        """
        fund_holding = HBDB().read_fund_holding_given_codes([self.target_fund_code])
        fund_holding = fund_holding.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'zqdm': 'TICKER_SYMBOL', 'zqmc': 'SEC_SHORT_NAME', 'ccsz': 'HOLDING_MARKET_VALUE', 'ccsl': 'HOLDING_AMOUNT', 'zjbl': 'MV_IN_NA'})
        stock_industry = HBDB().read_stock_industry()
        stock_industry = stock_industry.rename(columns={'zqdm': 'TICKER_SYMBOL', 'flmc': 'INDUSTRY_NAME', 'fldm': 'INDUSTRY_ID', 'fljb': 'INDUSTRY_TYPE', 'hyhfbz': 'INDUSTRY_VERSION', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'sfyx': 'IS_NEW'})
        stock_industry = stock_industry[stock_industry['IS_NEW'] == 1]
        stock_industry = stock_industry[stock_industry['INDUSTRY_VERSION'] == 2]
        stock_industry = stock_industry[stock_industry['INDUSTRY_TYPE'] == '2']
        stock_industry = stock_industry.sort_values(['TICKER_SYMBOL', 'credt_etl']).drop_duplicates('TICKER_SYMBOL', keep='last')
        industry_type = '二级' if stock_industry['INDUSTRY_TYPE'].iloc[0] == '2' else '三级' if stock_industry['INDUSTRY_TYPE'].iloc[0] == '3' else '一级'
        fund_holding = fund_holding.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='left')
        fund_holding = fund_holding[fund_holding['REPORT_DATE'] == fund_holding['REPORT_DATE'].max()].sort_values('MV_IN_NA', ascending=False).reset_index().drop('index', axis=1)
        fund_holding_top5 = fund_holding.iloc[:5]
        fund_holding_top5.index = range(1, 6)
        fund_holding_top5 = fund_holding_top5.reset_index()
        fund_holding_top5['SEC_SHORT_NAME_TOP'] = fund_holding_top5.apply(lambda x: '{0}（TOP{1}）'.format(x['SEC_SHORT_NAME'], x['index']), axis=1)
        fund_holding_indu_list = []
        for indu in fund_holding_top5['INDUSTRY_NAME'].tolist():
            fund_holding_indu = fund_holding[fund_holding['INDUSTRY_NAME'] == indu].sort_values('MV_IN_NA', ascending=False)
            fund_holding_indu = fund_holding_indu[['INDUSTRY_NAME', 'SEC_SHORT_NAME', 'HOLDING_AMOUNT', 'HOLDING_MARKET_VALUE', 'MV_IN_NA']]
            fund_holding_indu_list.append(fund_holding_indu)
        fund_holding_indu = pd.concat(fund_holding_indu_list)
        fund_holding_top5 = fund_holding_top5[['SEC_SHORT_NAME_TOP', 'INDUSTRY_NAME']].merge(fund_holding_indu, on=['INDUSTRY_NAME'], how='left')
        fund_holding_top5[['HOLDING_AMOUNT', 'HOLDING_MARKET_VALUE', 'MV_IN_NA']] = round(fund_holding_top5[['HOLDING_AMOUNT', 'HOLDING_MARKET_VALUE', 'MV_IN_NA']], 2)
        fund_holding_top5['MV_IN_NA'] = fund_holding_top5['MV_IN_NA'].apply(lambda x: '{0}%'.format(x))
        fund_holding_top5.columns = ['前五大持仓', '所属细分行业', '同细分行业个股', '持仓数量（股）', '持仓市值（元）', '占净值比例']
        fund_holding_top5 = fund_holding_top5.T.reset_index().T.reset_index().drop('index', axis=1)
        fund_holding_top5 = fund_holding_top5.dropna(axis=0, how='any').reset_index().drop('index', axis=1)

        fund_holding_top5_text = '该基金中与前五大持仓同细分行业（申万{0}行业）的个股如下：'.format(industry_type)
        return fund_holding_top5, fund_holding_top5_text

    def get_fund_turnover(self):
        """
        生成基金换手数据
        """
        fund_turnover = HBDB().read_fund_turnover_given_code(self.target_fund_code)
        fund_turnover = fund_turnover.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'tjqj': 'TYPE', 'hsl': 'TURNOVER'})
        fund_turnover = fund_turnover[fund_turnover['TYPE'].isin(['1', '3'])]
        fund_turnover['TURNOVER'] = fund_turnover['TURNOVER'] * 2.0
        fund_turnover = fund_turnover.sort_values('REPORT_DATE').iloc[-12:]
        fund_turnover = fund_turnover[['REPORT_DATE', 'TURNOVER']]
        fund_turnover.columns = ['报告日期', '换手率（%）']

        fund_turnover['报告日期'] = fund_turnover['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_turnover['报告日期']) and date <= max(fund_turnover['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_turnover['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(ax=ax, x='报告日期', y='换手率（%）', data=fund_turnover, palette=['#C94649'])
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(fund_turnover['换手率（%）']), linestyles='dashed', color='#959595')
        ax.set_xticklabels(labels=fund_turnover['报告日期'], rotation=90)
        plt.tight_layout()
        plt.savefig('{0}{1}_turnover.png'.format(self.file_path, self.target_fund_code))

        fund_turnover_mean = round(fund_turnover.iloc[-6:]['换手率（%）'].mean() / 100.0, 2)
        latest_turnover = round(fund_turnover.iloc[-1]['换手率（%）'] / 100.0, 2)
        fund_turnover_text = '该产品换手率情况如下，近三年单边年化换手倍数整体稳定在{0}倍左右，最新报告期披露的单边年化换手率为{1}倍。'.format(fund_turnover_mean, latest_turnover)

        compare_turnover = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_turnover_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_turnover_code = HBDB().read_fund_turnover_given_code(compare_fund_code)
                compare_turnover_list.append(compare_turnover_code)
            compare_turnover = pd.concat(compare_turnover_list)
            compare_turnover = compare_turnover.rename(columns={'jjdm': 'FUND_CODE', 'jsrq': 'REPORT_DATE', 'tjqj': 'TYPE', 'hsl': 'TURNOVER'})
            compare_turnover = compare_turnover[compare_turnover['TYPE'].isin(['1', '3'])]
            compare_turnover = compare_turnover[['FUND_CODE', 'REPORT_DATE', 'TURNOVER']]
            compare_turnover['TURNOVER'] = compare_turnover['TURNOVER'].apply(lambda x: '{0}%'.format(round(x, 2)))
            compare_turnover = compare_turnover.pivot(index='REPORT_DATE', columns='FUND_CODE', values='TURNOVER')
            compare_turnover = compare_turnover.dropna(subset=[self.target_fund_code], axis=0).sort_index().iloc[-12:]
            compare_turnover = compare_turnover[[self.target_fund_code] + self.compare_fund_code_list]
            compare_turnover.columns = [self.compare_fund_name_dic[col] for col in compare_turnover.columns]
            compare_turnover = compare_turnover.replace(np.nan, '-')
            compare_turnover = compare_turnover.T.reset_index().T.reset_index()
            compare_turnover.iloc[0, 0] = '换手率'
        return fund_turnover_text, compare_turnover

    def get_fund_performace(self):
        establish_date = str(int(self.target_fund['BEGIN_DATE'].values[0]))

        fund_nav = HBDB().read_fund_cumret_given_code(self.target_fund_code, establish_date, datetime.today().strftime('%Y%m%d'))
        fund_nav = fund_nav.rename(columns={'JJDM': 'FUND_CODE', 'JZRQ': 'TRADE_DATE', 'HBCL': 'CUM_RET'})
        fund_nav['TRADE_DATE'] = fund_nav['TRADE_DATE'].astype(str)
        fund_nav = fund_nav.sort_values('TRADE_DATE')
        fund_nav['FUND_NAME'] = self.target_fund_name
        fund_nav['CUM_RET'] = 0.01 * fund_nav['CUM_RET']
        fund_cumret = fund_nav.pivot(index='TRADE_DATE', columns='FUND_NAME', values='CUM_RET')
        fund_nav['NAV'] = fund_nav['CUM_RET'] + 1
        fund_nav['NAV'] = fund_nav['NAV'] / fund_nav['NAV'].iloc[0]
        fund_nav['IDX'] = range(len(fund_nav))
        fund_nav['DRAWDOWN'] = fund_nav['IDX'].apply(lambda x: fund_nav['NAV'].iloc[x] / max(fund_nav['NAV'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        fund_d = fund_nav.pivot(index='TRADE_DATE', columns='FUND_NAME', values='DRAWDOWN')

        index = '000300'
        index_name = '沪深300'
        index_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(establish_date, [index])
        index_daily_k = index_daily_k[['zqmc', 'jyrq', 'spjg']].rename(columns={'zqmc': 'INDEX_NAME', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        index_daily_k['TRADE_DATE'] = index_daily_k['TRADE_DATE'].astype(str)
        index_daily_k = index_daily_k.sort_values('TRADE_DATE')
        index_daily_k['CUM_RET'] = index_daily_k['CLOSE_INDEX'] / index_daily_k['CLOSE_INDEX'].iloc[0] - 1.0
        index_cumret = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CUM_RET')
        index_daily_k['NAV'] = index_daily_k['CLOSE_INDEX'] / index_daily_k['CLOSE_INDEX'].iloc[0]
        index_daily_k['IDX'] = range(len(index_daily_k))
        index_daily_k['DRAWDOWN'] = index_daily_k['IDX'].apply(lambda x: index_daily_k['CLOSE_INDEX'].iloc[x] / max(index_daily_k['CLOSE_INDEX'].iloc[:x]) - 1.0 if x != 0 else 0.0)
        index_d = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='DRAWDOWN')

        fund_index_cumret = fund_cumret.merge(index_cumret, left_index=True, right_index=True, how='left').reset_index()
        fund_index_cumret['TRADE_DATE'] = fund_index_cumret['TRADE_DATE'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_index_cumret['TRADE_DATE']) and date <= max(fund_index_cumret['TRADE_DATE'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_index_cumret['TRADE_DATE'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.lineplot(x='TRADE_DATE', y=self.target_fund_name, data=fund_index_cumret, color='#F04950', label=self.target_fund_name)
        sns.lineplot(x='TRADE_DATE', y=index_name, data=fund_index_cumret, color='#6268A2', label=index_name)
        ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max(max(fund_index_cumret[self.target_fund_name]), max(fund_index_cumret[index_name])), linestyles='dashed', color='#959595')
        ax.xaxis.set_major_locator(ticker.MultipleLocator(90))
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.xlabel('')
        plt.ylabel('累计收益率')
        plt.xticks(rotation=90)
        plt.title('收益走势图')
        plt.legend(loc=2)
        plt.tight_layout()
        plt.savefig('{0}{1}_return.png'.format(self.file_path, self.target_fund_code))

        fund_index_d = fund_d.merge(index_d, left_index=True, right_index=True, how='left').reset_index()
        fig, ax = plt.subplots(figsize=(6, 3))
        sns.lineplot(x='TRADE_DATE', y=self.target_fund_name, data=fund_index_d, color='#F04950', label=self.target_fund_name)
        sns.lineplot(x='TRADE_DATE', y=index_name, data=fund_index_d, color='#6268A2', label=index_name)
        ax.vlines(x=change_dates_loc, ymin=min(min(fund_index_d[self.target_fund_name]), min(fund_index_d[index_name])), ymax=0.0, linestyles='dashed', color='#959595')
        ax.xaxis.set_major_locator(ticker.MultipleLocator(90))
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.xlabel('')
        plt.ylabel('回撤')
        plt.xticks(rotation=90)
        plt.title('回撤走势图')
        plt.legend(loc=2)
        plt.tight_layout()
        plt.savefig('{0}{1}_drawdown.png'.format(self.file_path, self.target_fund_code))

        fund_return_period = HBDB().read_fund_return_given_code_and_date(self.target_fund_code, (datetime.today() - timedelta(30)).strftime('%Y%m%d'))
        fund_return_period = fund_return_period.rename(columns={'jjdm': 'FUND_CODE', 'zblb': 'RETURN_TYPE', 'rqnp': 'START_DATE', 'jzrq': 'END_DATE', 'zbnp': 'RETURN', 'zbnhnp': 'ANNUAL_RETURN'})
        fund_return_period['END_DATE'] = fund_return_period['END_DATE'].astype(str)
        fund_return_period = fund_return_period[fund_return_period['END_DATE'] == max(fund_return_period['END_DATE'])]
        fund_return_period['RETURN_TYPE'] = fund_return_period['RETURN_TYPE'].astype(int)
        fund_return_period['RETURN_TYPE'] = fund_return_period['RETURN_TYPE'].replace({2998: '今年以来', 2101: '近1月', 2103: '近3月', 2106: '近6月', 2201: '近1年', 2202: '近2年', 2203: '近3年'})
        fund_return_period['RETURN'] = fund_return_period['RETURN'].apply(lambda x: '{0}%'.format(round(x, 2)))
        fund_return_period = fund_return_period[fund_return_period['RETURN_TYPE'].isin(['今年以来', '近1月', '近3月', '近6月', '近1年', '近2年', '近3年'])]
        fund_return_period = fund_return_period.pivot(index='FUND_CODE', columns='RETURN_TYPE', values='RETURN')
        fund_return_period = fund_return_period[['今年以来', '近1月', '近3月', '近6月', '近1年', '近2年', '近3年']]
        fund_samekind_return_period = HBDB().read_fund_samekind_return_given_code_and_date(self.target_fund_code, (datetime.today() - timedelta(30)).strftime('%Y%m%d'))
        fund_samekind_return_period = fund_samekind_return_period.rename(columns={'jjdm': 'FUND_CODE', 'zblb': 'RETURN_TYPE', 'jzrq': 'END_DATE'})
        fund_samekind_return_period['END_DATE'] = fund_samekind_return_period['END_DATE'].astype(str)
        fund_samekind_return_period = fund_samekind_return_period[fund_samekind_return_period['END_DATE'] == max(fund_samekind_return_period['END_DATE'])]
        fund_samekind_return_period['RETURN_TYPE'] = fund_samekind_return_period['RETURN_TYPE'].astype(int)
        fund_samekind_return_period['RETURN_TYPE'] = fund_samekind_return_period['RETURN_TYPE'].replace({2998: '今年以来', 2101: '近1月', 2103: '近3月', 2106: '近6月', 2201: '近1年', 2202: '近2年', 2203: '近3年'})
        fund_samekind_return_period = fund_samekind_return_period[fund_samekind_return_period['RETURN_TYPE'].isin(['今年以来', '近1月', '近3月', '近6月', '近1年', '近2年', '近3年'])]
        fund_samekind_return_period['RANK'] = fund_samekind_return_period.apply(lambda x: '{0}/{1}'.format(x['pmnpej'], x['slnpej']), axis=1)
        fund_samekind_return_period = fund_samekind_return_period.pivot(index='FUND_CODE', columns='RETURN_TYPE', values='RANK')
        fund_samekind_return_period = fund_samekind_return_period[['今年以来', '近1月', '近3月', '近6月', '近1年', '近2年', '近3年']]
        fund_return_period = pd.concat([fund_return_period, fund_samekind_return_period])
        fund_return_period.index = ['产品收益率', '同策略排名']
        fund_return_period = fund_return_period.T.reset_index().T.reset_index()
        fund_return_period.iloc[0, 0] = ''

        fund_year_return = HBDB().read_fund_year_return_given_code(self.target_fund_code)
        fund_year_return = fund_year_return.rename(columns={'jjdm': 'FUND_CODE', 'tjnf': 'YEAR', 'hb1n': 'RETURN'})
        fund_year_return = fund_year_return[fund_year_return['RETURN'] != 99999]
        fund_year_return['RETURN'] = fund_year_return['RETURN'].apply(lambda x: '{0}%'.format(round(x, 2)))
        fund_year_return['RETURN_RANK'] = fund_year_return.apply(lambda x: '{0}/{1}'.format(int(x['hb1npmej']), int(x['hb1npmslej'])), axis=1)
        fund_year_return = fund_year_return[['YEAR', 'RETURN', 'RETURN_RANK']].sort_values('YEAR', ascending=False)

        fund_year_volatility = HBDB().read_fund_year_volatility_given_code(self.target_fund_code)
        fund_year_volatility = fund_year_volatility.rename(columns={'jjdm': 'FUND_CODE', 'tjnf': 'YEAR', 'znhzbnp1n': 'ANNUAL_VOLATILITY'})
        fund_year_volatility = fund_year_volatility[fund_year_volatility['ANNUAL_VOLATILITY'] != 99999]
        fund_year_volatility['ANNUAL_VOLATILITY'] = fund_year_volatility['ANNUAL_VOLATILITY'].apply(lambda x: '{0}%'.format(round(x, 2)))
        fund_year_volatility['ANNUAL_VOLATILITY_RANK'] = fund_year_volatility.apply(lambda x: '{0}/{1}'.format(int(x['zbdpm1nej']), int(x['tlzbdsl1nej'])), axis=1)
        fund_year_volatility = fund_year_volatility[['YEAR', 'ANNUAL_VOLATILITY', 'ANNUAL_VOLATILITY_RANK']].sort_values('YEAR', ascending=False)

        fund_year_sharpratio = HBDB().read_fund_year_sharpratio_given_code(self.target_fund_code)
        fund_year_sharpratio = fund_year_sharpratio.rename(columns={'jjdm': 'FUND_CODE', 'tjnf': 'YEAR', 'znhzbnp1n': 'ANNUAL_SHARPRATIO'})
        fund_year_sharpratio = fund_year_sharpratio[fund_year_sharpratio['ANNUAL_SHARPRATIO'] != 99999]
        fund_year_sharpratio['ANNUAL_SHARPRATIO'] = fund_year_sharpratio['ANNUAL_SHARPRATIO'].apply(lambda x: '{0}'.format(round(x, 2)))
        fund_year_sharpratio['ANNUAL_SHARPRATIO_RANK'] = fund_year_sharpratio.apply(lambda x: '{0}/{1}'.format(int(x['zxppm1nej']), int(x['tlzxpsl1nej'])), axis=1)
        fund_year_sharpratio = fund_year_sharpratio[['YEAR', 'ANNUAL_SHARPRATIO', 'ANNUAL_SHARPRATIO_RANK']].sort_values('YEAR', ascending=False)

        fund_year_maxdrawdown= HBDB().read_fund_year_maxdrawdown_given_code(self.target_fund_code)
        fund_year_maxdrawdown = fund_year_maxdrawdown.rename(columns={'jjdm': 'FUND_CODE', 'tjnf': 'YEAR', 'zbnp1n': 'MAX_DRAWDOWN'})
        fund_year_maxdrawdown = fund_year_maxdrawdown[fund_year_maxdrawdown['MAX_DRAWDOWN'] != 99999]
        fund_year_maxdrawdown['MAX_DRAWDOWN'] = fund_year_maxdrawdown['MAX_DRAWDOWN'].apply(lambda x: '{0}%'.format(round(x, 2)))
        fund_year_maxdrawdown['MAX_DRAWDOWN_RANK'] = fund_year_maxdrawdown.apply(lambda x: '{0}/{1}'.format(int(x['hczdpm1nej']), int(x['tlhczdsl1nej'])), axis=1)
        fund_year_maxdrawdown = fund_year_maxdrawdown[['YEAR', 'MAX_DRAWDOWN', 'MAX_DRAWDOWN_RANK']].sort_values('YEAR', ascending=False)

        fund_year_performance = fund_year_return.merge(fund_year_volatility, on=['YEAR'], how='left').merge(fund_year_sharpratio, on=['YEAR'], how='left').merge(fund_year_maxdrawdown, on=['YEAR'], how='left')
        fund_year_performance.columns = ['年份', '收益率', '同策略排名', '年化波动率', '同策略排名', '年化夏普比率', '同策略排名', '最大回撤', '同策略排名']
        fund_year_performance = fund_year_performance.T.reset_index().T.reset_index().drop('index', axis=1)

        if len(self.compare_fund_code_list) > 0:
            compare_return_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_return_code = HBDB().read_fund_cumret_given_code(compare_fund_code, str(int(max(self.target_fund['BEGIN_DATE'].values[0], self.compare_fund['BEGIN_DATE'].max()))), datetime.today().strftime('%Y%m%d'))
                compare_return_list.append(compare_return_code)
            compare_return = pd.concat(compare_return_list)
            compare_return = compare_return.rename(columns={'JJDM': 'FUND_CODE', 'JZRQ': 'TRADE_DATE', 'HBCL': 'CUM_RET'})
            compare_return['TRADE_DATE'] = compare_return['TRADE_DATE'].astype(str)
            compare_return = compare_return.pivot(index='TRADE_DATE', columns='FUND_CODE', values='CUM_RET')
            compare_return = compare_return.sort_index().interpolate()
            compare_return = compare_return * 0.01 + 1
            compare_return = compare_return / compare_return.iloc[0] - 1.0
            compare_return = compare_return[[self.target_fund_code] + self.compare_fund_code_list]
            compare_return.columns = [self.compare_fund_name_dic[col] for col in compare_return.columns]
            compare_return = compare_return.reset_index()

            change_dates = [date for date in self.change_dates if date >= min(compare_return['TRADE_DATE']) and date <= max(compare_return['TRADE_DATE'])]
            change_dates_loc = get_change_dates_loc(change_dates, sorted(compare_return['TRADE_DATE'].unique().tolist()))
            fig, ax = plt.subplots(figsize=(6, 3))
            max_value = 0.0
            for i, col in enumerate(list(compare_return.columns)[1:]):
                sns.lineplot(x='TRADE_DATE', y=col, data=compare_return, color=line_color_list[i], label=col)
                max_value = max(max(compare_return[col]), max_value)
            ax.vlines(x=change_dates_loc, ymin=0.0, ymax=max_value, linestyles='dashed', color='#959595')
            ax.xaxis.set_major_locator(ticker.MultipleLocator(90))
            ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
            plt.xlabel('')
            plt.ylabel('累计收益率')
            plt.xticks(rotation=90)
            plt.title('收益走势图')
            plt.legend(loc=2)
            plt.tight_layout()
            plt.savefig('{0}{1}_compare_return.png'.format(self.file_path, self.target_fund_code))
        return fund_return_period, fund_year_performance

    def get_fund_achievement(self):
        """
        生成基金业绩数据
        """
        three_years_ago = (datetime.today() - timedelta(365 * 3)).strftime('%Y%m%d')
        three_years_ago = three_years_ago if str(self.target_fund['BEGIN_DATE'].values[0]) < three_years_ago else str(self.target_fund['BEGIN_DATE'].values[0])

        fund_nav = HBDB().read_fund_cumret_given_code(self.target_fund_code, three_years_ago, datetime.today().strftime('%Y%m%d'))
        fund_nav = fund_nav.rename(columns={'JJDM': 'FUND_CODE', 'JZRQ': 'TRADE_DATE', 'HBCL': 'CUM_RET'})
        fund_nav['TRADE_DATE'] = fund_nav['TRADE_DATE'].astype(str)
        fund_nav = fund_nav.sort_values('TRADE_DATE')
        fund_nav['FUND_NAME'] = self.target_fund_name
        fund_nav['ADJ_NAV'] = 0.01 * fund_nav['CUM_RET'] + 1
        fund_nav = fund_nav.pivot(index='TRADE_DATE', columns='FUND_NAME', values='ADJ_NAV')

        fund_benchmark_return = HBDB().read_fund_benchmark_nav_given_code(self.target_fund_code)
        fund_benchmark_return = fund_benchmark_return.rename(columns={'jjdm': 'FUND_CODE', 'jzrq': 'TRADE_DATE', 'jzhbdr': 'FUND_BENCKMARK_RETURN'})
        fund_benchmark_return['TRADE_DATE'] = fund_benchmark_return['TRADE_DATE'].astype(str)
        fund_benchmark_return['FUND_BENCKMARK_RETURN'] = fund_benchmark_return['FUND_BENCKMARK_RETURN'] / 100.0
        fund_benchmark_nav = fund_benchmark_return.copy(deep=True)
        fund_benchmark_nav['FUND_BENCKMARK_NAV'] = (fund_benchmark_nav['FUND_BENCKMARK_RETURN'] + 1.0).cumprod()
        fund_benchmark_nav['FUND_BENCKMARK_NAME'] = self.target_fund_name + '基准'
        fund_benchmark_nav = fund_benchmark_nav[fund_benchmark_nav['TRADE_DATE'] >= three_years_ago]
        fund_benchmark_nav = fund_benchmark_nav.pivot(index='TRADE_DATE', columns='FUND_BENCKMARK_NAME', values='FUND_BENCKMARK_NAV')

        index_type = HBDB().read_fund_wind_size_given_code(self.target_fund_code)
        index_type = index_type.rename(columns={'jjdm': 'FUND_CODE', 'wdszsx': 'INDEX_TYPE'})
        index_type['INDEX_TYPE'] = index_type['INDEX_TYPE'].replace({'1': '小盘', '2': '中盘', '3': '大盘'})
        index_type['INDEX_TYPE'] = index_type['INDEX_TYPE'].replace({'小盘': '000852', '中盘': '000905', '大盘': '000300'})
        index = index_type['INDEX_TYPE'].values[0]
        index = '000300'
        index_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(three_years_ago, [index])
        index_daily_k = index_daily_k[['zqmc', 'jyrq', 'spjg']].rename(columns={'zqmc': 'INDEX_NAME', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        index_daily_k['TRADE_DATE'] = index_daily_k['TRADE_DATE'].astype(str)
        index_daily_k = index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CLOSE_INDEX')

        stock_fund_index_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(three_years_ago, ['885001'])
        stock_fund_index_daily_k = stock_fund_index_daily_k[['zqmc', 'jyrq', 'spjg']].rename(columns={'zqmc': 'INDEX_NAME', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        stock_fund_index_daily_k['TRADE_DATE'] = stock_fund_index_daily_k['TRADE_DATE'].astype(str)
        stock_fund_index_daily_k['INDEX_NAME'] = '偏股混合型基金指数'
        stock_fund_index_daily_k = stock_fund_index_daily_k.pivot(index='TRADE_DATE', columns='INDEX_NAME', values='CLOSE_INDEX')

        # fund_benchmark = HBDB().read_fund_benchmark_given_code(fund_code)
        # fund_benchmark = fund_benchmark.rename(columns={'JJDM': 'FUND_CODE', 'BDRQ': 'CHANGE_DATE', 'BJJZ': 'FUND_BENCHMARK', 'SYZT': 'IS_NEW'})
        # fund_benchmark = fund_benchmark[fund_benchmark['IS_NEW'] == 0]
        # fund_benchmark = fund_benchmark['FUND_BENCHMARK'].values[0]
        # fund_benchmark = fund_benchmark.split('+')
        # fund_benchmark = {b.split('×')[1]: int(b.split('×')[0][:-1]) / 100.0 for b in fund_benchmark}

        nav_df = fund_nav.merge(fund_benchmark_nav, left_index=True, right_index=True, how='left').merge(index_daily_k, left_index=True, right_index=True, how='left').merge(stock_fund_index_daily_k, left_index=True, right_index=True, how='left')
        nav_df = nav_df.replace(0.0, np.nan)
        nav_df = nav_df.fillna(method='ffill')

        achievement = pd.DataFrame(index=['业绩统计（近三年）', '年化收益率', '年化波动率', 'Sharpe比率', '最大回撤', 'Calmar比率', '投资胜率', '平均损益比'], columns=range(len(list(nav_df.columns))))
        for idx, col in enumerate(list(nav_df.columns)):
            col_nav_df = nav_df[col]
            col_ret_df = col_nav_df.pct_change().dropna()
            achievement.loc['业绩统计（近三年）', idx] = col
            achievement.loc['年化收益率', idx] = (col_nav_df.iloc[-1] / col_nav_df.iloc[0]) ** (250.0 / len(col_nav_df)) - 1.0
            achievement.loc['年化波动率', idx] = np.std(col_ret_df, ddof=1) * np.sqrt(250.0)
            achievement.loc['Sharpe比率', idx] = (achievement.loc['年化收益率', idx] - 0.03) / achievement.loc['年化波动率', idx]
            achievement.loc['最大回撤', idx] = -1.0 * max([(min(col_nav_df.iloc[i:]) / col_nav_df.iloc[i] - 1.0) * (-1.0) for i in range(len(col_nav_df))])
            achievement.loc['Calmar比率', idx] = achievement.loc['年化收益率', idx] / achievement.loc['最大回撤', idx]
            achievement.loc['投资胜率', idx] = len(col_ret_df[col_ret_df >= 0]) / float(len(col_ret_df))
            achievement.loc['平均损益比', idx] = col_ret_df[col_ret_df >= 0].mean() / col_ret_df[col_ret_df < 0].mean() * (-1.0)

            achievement.loc['年化收益率', idx] = '{0}%'.format(round(achievement.loc['年化收益率', idx] * 100.0, 2))
            achievement.loc['年化波动率', idx] = '{0}%'.format(round(achievement.loc['年化波动率', idx] * 100.0, 2))
            achievement.loc['Sharpe比率', idx] = '{0}'.format(round(achievement.loc['Sharpe比率', idx], 2))
            achievement.loc['最大回撤', idx] = '{0}%'.format(round(achievement.loc['最大回撤', idx] * 100.0, 2))
            achievement.loc['Calmar比率', idx] = '{0}'.format(round(achievement.loc['Calmar比率', idx], 2))
            achievement.loc['投资胜率', idx] = '{0}%'.format(round(achievement.loc['投资胜率', idx] * 100.0, 2))
            achievement.loc['平均损益比', idx] = '{0}'.format(round(achievement.loc['平均损益比', idx], 2))

        achievement = achievement.reset_index()

        compare_achievement = pd.DataFrame()
        compare_corr = pd.DataFrame()
        if len(self.compare_fund_code_list) > 0:
            compare_nav_list = []
            for compare_fund_code in [self.target_fund_code] + self.compare_fund_code_list:
                compare_nav_code = HBDB().read_fund_cumret_given_code(compare_fund_code, str(int(max(self.target_fund['BEGIN_DATE'].values[0], self.compare_fund['BEGIN_DATE'].max()))), datetime.today().strftime('%Y%m%d'))
                compare_nav_list.append(compare_nav_code)
            compare_nav = pd.concat(compare_nav_list)
            compare_nav = compare_nav.rename(columns={'JJDM': 'FUND_CODE', 'JZRQ': 'TRADE_DATE', 'HBCL': 'CUM_RET'})
            compare_nav['TRADE_DATE'] = compare_nav['TRADE_DATE'].astype(str)
            compare_nav = compare_nav.pivot(index='TRADE_DATE', columns='FUND_CODE', values='CUM_RET')
            compare_nav = compare_nav.sort_index().interpolate()
            compare_nav = compare_nav * 0.01 + 1
            compare_nav = compare_nav / compare_nav.iloc[0]
            compare_nav = compare_nav[[self.target_fund_code] + self.compare_fund_code_list]
            compare_nav.columns = [self.compare_fund_name_dic[col] for col in compare_nav.columns]
            compare_corr = compare_nav.pct_change().dropna().corr()
            for col in list(compare_corr):
                compare_corr[col] = compare_corr[col].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
            for i in range(len(compare_corr)):
                for j in range(len(compare_corr)):
                    if i > j:
                        compare_corr.iloc[i, j] = '-'
            compare_corr = compare_corr.T.reset_index().T.reset_index()
            compare_corr.iloc[0, 0] = ''

            compare_achievement = pd.DataFrame(index=['业绩统计（同时段）', '年化收益率', '年化波动率', 'Sharpe比率', '最大回撤', 'Calmar比率', '投资胜率', '平均损益比'], columns=range(len([self.target_fund_code] + self.compare_fund_code_list)))
            for idx, col in enumerate(list(compare_nav.columns)):
                col_nav_df = compare_nav[col]
                col_ret_df = col_nav_df.pct_change().dropna()
                compare_achievement.loc['业绩统计（同时段）', idx] = col
                compare_achievement.loc['年化收益率', idx] = (col_nav_df.iloc[-1] / col_nav_df.iloc[0]) ** (50.0 / len(col_nav_df)) - 1.0
                compare_achievement.loc['年化波动率', idx] = np.std(col_ret_df, ddof=1) * np.sqrt(250.0)
                compare_achievement.loc['Sharpe比率', idx] = (compare_achievement.loc['年化收益率', idx] - 0.03) / compare_achievement.loc['年化波动率', idx]
                compare_achievement.loc['最大回撤', idx] = -1.0 * max([(min(col_nav_df.iloc[i:]) / col_nav_df.iloc[i] - 1.0) * (-1.0) for i in range(len(col_nav_df))])
                compare_achievement.loc['Calmar比率', idx] = compare_achievement.loc['年化收益率', idx] / compare_achievement.loc['最大回撤', idx]
                compare_achievement.loc['投资胜率', idx] = len(col_ret_df[col_ret_df >= 0]) / float(len(col_ret_df))
                compare_achievement.loc['平均损益比', idx] = col_ret_df[col_ret_df >= 0].mean() / col_ret_df[col_ret_df < 0].mean() * (-1.0)

                compare_achievement.loc['年化收益率', idx] = '{0}%'.format(round(compare_achievement.loc['年化收益率', idx] * 100.0, 2))
                compare_achievement.loc['年化波动率', idx] = '{0}%'.format(round(compare_achievement.loc['年化波动率', idx] * 100.0, 2))
                compare_achievement.loc['Sharpe比率', idx] = '{0}'.format(round(compare_achievement.loc['Sharpe比率', idx], 2))
                compare_achievement.loc['最大回撤', idx] = '{0}%'.format(round(compare_achievement.loc['最大回撤', idx] * 100.0, 2))
                compare_achievement.loc['Calmar比率', idx] = '{0}'.format(round(compare_achievement.loc['Calmar比率', idx], 2))
                compare_achievement.loc['投资胜率', idx] = '{0}%'.format(round(compare_achievement.loc['投资胜率', idx] * 100.0, 2))
                compare_achievement.loc['平均损益比', idx] = '{0}'.format(round(compare_achievement.loc['平均损益比', idx], 2))

            compare_achievement = compare_achievement.reset_index()

        return achievement, compare_achievement, compare_corr

    def get_fund_brison_attribution(self):
        fund_brison_attribution = HBDB().read_fund_brinson_attribution_given_code(self.target_fund_code)
        fund_brison_attribution = fund_brison_attribution.rename(columns={'jjdm': 'FUND_CODE', 'tjrq': 'REPORT_DATE', 'asset_allo': '大类资产配置收益', 'sector_allo': '行业配置收益', 'equity_selection': '行业选股收益', 'trading': '交易收益'})
        fund_brison_attribution['REPORT_DATE'] = fund_brison_attribution['REPORT_DATE'].apply(lambda x: str(x)[:6] + '30' if str(x)[4:6] == '06' else str(x)[:6] + '31')
        fund_brison_attribution = fund_brison_attribution.sort_values('REPORT_DATE')
        fund_brison_attribution['总超额收益'] = fund_brison_attribution['portfolio_return'] - fund_brison_attribution['benchmark_return']
        attr_list = ['大类资产配置收益', '行业配置收益', '行业选股收益', '交易收益']
        fund_brison_attribution = fund_brison_attribution[['REPORT_DATE', '总超额收益'] + attr_list]
        fund_brison_attribution.columns = ['报告日期', '总超额收益'] + attr_list
        fund_brison_attribution_disp = fund_brison_attribution.copy(deep=True)
        for i in range(1, len(attr_list)):
            fund_brison_attribution_disp[attr_list[i]] = fund_brison_attribution_disp[attr_list[i]] + fund_brison_attribution_disp[attr_list[i - 1]]
        fund_brison_attribution_disp = fund_brison_attribution_disp.iloc[-12:]

        fund_brison_attribution_disp['报告日期'] = fund_brison_attribution_disp['报告日期'].astype(str)
        change_dates = [date for date in self.change_dates if date >= min(fund_brison_attribution_disp['报告日期']) and date <= max(fund_brison_attribution_disp['报告日期'])]
        change_dates_loc = get_change_dates_loc(change_dates, sorted(fund_brison_attribution_disp['报告日期'].unique().tolist()))
        fig, ax = plt.subplots(figsize=(8, 3))
        color_list = [bar_color_list[0], bar_color_list[1], bar_color_list[7], bar_color_list[14]]
        for i in range(len(attr_list) - 1, -1, -1):
            sns.barplot(ax=ax, x='报告日期', y=attr_list[i], data=fund_brison_attribution_disp, label=attr_list[i], color=color_list[i])
        sns.lineplot(ax=ax, x='报告日期', y='总超额收益', data=fund_brison_attribution_disp, label='总超额收益', color='#F04950')
        ax.vlines(x=change_dates_loc, ymin=min([min(fund_brison_attribution_disp[col]) for col in attr_list + ['总超额收益']]), ymax=max([max(fund_brison_attribution_disp[col]) for col in attr_list + ['总超额收益']]), linestyles='dashed', color='#959595')
        ax.yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.legend(loc='upper right', bbox_to_anchor=(1.4, 1.02))
        plt.xticks(rotation=90)
        plt.ylabel('Brison归因时序图')
        plt.tight_layout()
        plt.savefig('{0}{1}_brison_attribution.png'.format(self.file_path, self.target_fund_code))

        max_date = max(fund_brison_attribution['报告日期'])
        fund_brison_attribution_latest = fund_brison_attribution[fund_brison_attribution['报告日期'] == max_date].set_index('报告日期')
        excess_return = '{0}%'.format(round(fund_brison_attribution_latest['总超额收益'].values[0] * 100.0, 2))
        fund_brison_attribution_latest = fund_brison_attribution_latest.drop('总超额收益', axis=1).T
        fund_brison_attribution_lateste_sorted = fund_brison_attribution_latest.sort_values(max_date, ascending=False)
        fund_brison_attribution_latest[max_date] = fund_brison_attribution_latest[max_date].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
        fund_brison_attribution_lateste_sorted[max_date] = fund_brison_attribution_lateste_sorted[max_date].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
        brison_attribution_ret = '，'.join(fund_brison_attribution_latest[max_date].unique().tolist())
        max_attribution = fund_brison_attribution_lateste_sorted.index[0]

        fund_brison_attribution_text = '该产品Brison归因结果如下，从最新一期看，总超额收益为{0}，大类资产配置收益、行业配置收益、行业选股收益、交易收益分别为{1}，其中{2}贡献最大。'.format(excess_return, brison_attribution_ret, max_attribution)
        return fund_brison_attribution_text

    def get_fund_barra_attribution(self):
        fund_barra_attribution = HBDB().read_fund_barra_attribution_given_code(self.target_fund_code)
        fund_barra_attribution = fund_barra_attribution.rename(columns={'jjdm': 'FUND_CODE', 'tjrq': 'REPORT_DATE'})
        fund_barra_attribution['REPORT_DATE'] = fund_barra_attribution['REPORT_DATE'].astype(int)
        max_date = max(fund_barra_attribution['REPORT_DATE'])
        fund_barra_attribution = fund_barra_attribution[fund_barra_attribution['REPORT_DATE'] == max_date]
        fund_barra_attribution = fund_barra_attribution[fund_barra_attribution['attr_type'] == 'style']
        fund_barra_exposure = fund_barra_attribution[fund_barra_attribution['data_type'] == 'exposure']
        fund_barra_exposure = fund_barra_exposure[['style_factor', 'data_value', 'data_type']]
        fund_barra_exposure = fund_barra_exposure.replace({'beta': 'Beta', 'btop': '估值', 'earnyield': '盈利', 'growth': '成长性', 'leverage': '杠杆', 'liquidity': '流动性', 'momentum': '动量', 'resvol': '波动性', 'size': '规模', 'sizenl': '非线性市值'})
        fund_barra_return = fund_barra_attribution[fund_barra_attribution['data_type'] == 'return']
        fund_barra_return = fund_barra_return[['style_factor', 'data_value', 'data_type']]
        fund_barra_return = fund_barra_return.replace({'beta': 'Beta', 'btop': '估值', 'earnyield': '盈利', 'growth': '成长性', 'leverage': '杠杆', 'liquidity': '流动性', 'momentum': '动量', 'resvol': '波动性', 'size': '规模', 'sizenl': '非线性市值'})

        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(x='style_factor', y='data_value', data=fund_barra_exposure, color='#C94649')
        plt.xlabel('')
        plt.ylabel('风格暴露')
        plt.tight_layout()
        plt.savefig('{0}{1}_barra_exposure.png'.format(self.file_path, self.target_fund_code))

        fig, ax = plt.subplots(figsize=(6, 3))
        sns.barplot(x='style_factor', y='data_value', data=fund_barra_return, color='#C94649')
        plt.xlabel('')
        plt.ylabel('风格收益')
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_100percent))
        plt.tight_layout()
        plt.savefig('{0}{1}_barra_return.png'.format(self.file_path, self.target_fund_code))

        fund_barra_exposure = fund_barra_exposure.sort_values('data_value', ascending=False)
        fund_barra_return = fund_barra_return.sort_values('data_value', ascending=False)
        fund_barra_text = '该产品Barra风格暴露和风格收益如下，最新报告期在{0}风格上暴露最大，在{1}风格上贡献收益最大。'.format(fund_barra_exposure['style_factor'].iloc[0], fund_barra_return['style_factor'].iloc[0])
        return fund_barra_text

    def get_quantitative_report(self):
        """
        生成公募基金量化分析报告
        """
        document = Document()
        document.styles['Normal'].font.size = Pt(8)
        head = document.add_heading(text='', level=0)
        run_head = head.add_run('{0}（{1}）定量分析报告'.format(self.target_fund_name, self.target_fund_code))
        run_head.font.color.rgb = RGBColor(201, 70, 73)

        risk, statement1, statement2, statement3 = self.report_info()
        document.add_paragraph(text='风险提示').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=risk)
        document.add_paragraph(text='重要声明').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=statement1)
        document.add_paragraph(text=statement2)
        document.add_paragraph(text=statement3)
        document.add_page_break()

        head_1 = document.add_heading(text='', level=1)
        run_head_1 = head_1.add_run('一. 基金概况')
        run_head_1.font.color.rgb = RGBColor(201, 70, 73)
        head_1_1 = document.add_heading(text='', level=2)
        run_head_1_1 = head_1_1.add_run('1. 基金及基金公司简介')
        run_head_1_1.font.color.rgb = RGBColor(201, 70, 73)
        fund_info_text = self.get_fund_info()
        document.add_paragraph(text=fund_info_text)
        fund_company_info_text, company_fund_1st, company_fund_2nd = self.get_fund_company_info()
        document.add_paragraph(text=fund_company_info_text)
        rowc, colc = company_fund_1st.shape[0], company_fund_1st.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(company_fund_1st.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        rowc, colc = company_fund_2nd.shape[0], company_fund_2nd.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(company_fund_2nd.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')

        head_1_2 = document.add_heading(text='', level=2)
        run_head_1_2 = head_1_2.add_run('2. 基金经理简介')
        run_head_1_2.font.color.rgb = RGBColor(201, 70, 73)
        fund_manager_info_text, manager_all_ori, manager_all = self.get_fund_manager_info()
        document.add_paragraph(text=fund_manager_info_text)
        rowc, colc = manager_all.shape[0], manager_all.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(manager_all.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')

        head_1_3 = document.add_heading(text='', level=2)
        run_head_1_3 = head_1_3.add_run('3. 规模变动')
        run_head_1_3.font.color.rgb = RGBColor(201, 70, 73)
        fund_scale_text, compare_scale = self.get_fund_scale()
        document.add_paragraph(text=fund_scale_text)
        document.add_picture('{0}{1}_scale.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        vlines_appendix = '注：本报告中垂直虚线为基金经理变更时间点，若无垂直虚线则代表在定量分析统计期间（近三年）未发生基金经理变更，下同。'
        document.add_paragraph(text=vlines_appendix)
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金规模比较情况如下：')
            rowc, colc = compare_scale.shape[0], compare_scale.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_scale.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        head_1_4 = document.add_heading(text='', level=2)
        run_head_1_4 = head_1_4.add_run('4. 持有人结构')
        run_head_1_4.font.color.rgb = RGBColor(201, 70, 73)
        fund_holder_text, compare_holder = self.get_fund_holder()
        document.add_paragraph(text=fund_holder_text)
        document.add_picture('{0}{1}_holder.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金机构持有比较情况如下：')
            rowc, colc = compare_holder.shape[0], compare_holder.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_holder.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')


        # fund_manager_method_text, \
        # fund_manager_method_shift_text, shift_disp, \
        # fund_manager_method_cen_text, cen_disp, \
        # fund_manager_method_select_text, industry_stock_cen_disp, \
        # fund_manager_method_holding_period_text, holding_period_disp, \
        # fund_manager_method_left_text, left_disp, \
        # fund_manager_method_new_text, new_disp, \
        # compare_shift, compare_cen, compare_industry_stock_cen, \
        # compare_holding_period, compare_left, compare_new = self.get_fund_manager_method()
        # document.add_paragraph(text=fund_manager_method_text)
        #
        # document.add_heading(text='1.1 换手率水平', level=3)
        # document.add_paragraph(text=fund_manager_method_shift_text)
        # rowc, colc = shift_disp.shape[0], shift_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(shift_disp.iloc[i, j])
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_paragraph(text=' ')
        # document.add_picture('{0}{1}_shift.png'.format(self.file_path, self.target_fund_code))
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # shift_appendix = '注：本报告规模、风格、主题、行业换手率定义为统计区间内规模、风格、主题、行业前后两期暴露变化绝对值之和的均值，规模、风格、主题、行业换手率水平定义为规模、风格、主题、行业换手率在偏股型基金整体中的排名，整体换手率水平定义为规模、风格、主题、行业换手率水平的均值。'
        # document.add_paragraph(text=shift_appendix)
        # if len(self.compare_fund_code_list) > 0:
        #     document.add_paragraph(text='与同类型基金换手率比较情况如下：')
        #     rowc, colc = compare_shift.shape[0], compare_shift.shape[1]
        #     table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        #     for i in range(rowc):
        #         for j in range(colc):
        #             table.cell(i, j).text = str(compare_shift.iloc[i, j])
        #     document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     document.add_paragraph(text=' ')
        #
        # document.add_heading(text='1.2 集中度水平', level=3)
        # document.add_paragraph(text=fund_manager_method_cen_text)
        # rowc, colc = cen_disp.shape[0], cen_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(cen_disp.iloc[i, j])
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_paragraph(text=' ')
        # document.add_picture('{0}{1}_cen.png'.format(self.file_path, self.target_fund_code))
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # cen_appendix = '注：本报告规模、风格、主题、行业集中度定义为统计区间内规模、风格、主题、行业前n大暴露占比之和的均值，规模、风格、主题、行业集中度水平定义为规模、风格、主题、行业集中度在偏股型基金整体中的排名，整体集中度水平定义为规模、风格、主题、行业集中度水平的均值。'
        # document.add_paragraph(text=cen_appendix)
        # if len(self.compare_fund_code_list) > 0:
        #     document.add_paragraph(text='与同类型基金集中度比较情况如下：')
        #     rowc, colc = compare_cen.shape[0], compare_cen.shape[1]
        #     table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        #     for i in range(rowc):
        #         for j in range(colc):
        #             table.cell(i, j).text = str(compare_cen.iloc[i, j])
        #     document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     document.add_paragraph(text=' ')
        #
        # document.add_heading(text='1.3 选股方法', level=3)
        # document.add_paragraph(text=fund_manager_method_select_text)
        # rowc, colc = industry_stock_cen_disp.shape[0], industry_stock_cen_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(industry_stock_cen_disp.iloc[i, j])
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_paragraph(text=' ')
        # document.add_picture('{0}{1}_indu_stock_cen.png'.format(self.file_path, self.target_fund_code))
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # industry_stock_cen_appendix = '注：本报告个股集中度定义为统计区间内前n大个股持仓权重之和的均值，个股集中度水平定义为个股集中度在偏股型基金整体中的排名。若行业集中度大于70%，个股集中度小于50%，则定义选股方法为自上而下；若行业集中度小于50%、个股集中度大于70%，则定义选股方法为自下而上；否则定义为无明显的自上而下或者自下而上选股特征。'
        # document.add_paragraph(text=industry_stock_cen_appendix)
        # if len(self.compare_fund_code_list) > 0:
        #     document.add_paragraph(text='与同类型基金选股方法比较情况如下：')
        #     rowc, colc = compare_industry_stock_cen.shape[0], compare_industry_stock_cen.shape[1]
        #     table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        #     for i in range(rowc):
        #         for j in range(colc):
        #             table.cell(i, j).text = str(compare_industry_stock_cen.iloc[i, j])
        #     document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     document.add_paragraph(text=' ')
        #
        # document.add_heading(text='1.4 持仓周期', level=3)
        # document.add_paragraph(text=fund_manager_method_holding_period_text)
        # rowc, colc = holding_period_disp.shape[0], holding_period_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(holding_period_disp.iloc[i, j])
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_paragraph(text=' ')
        # holding_period_appendix = '注：本报告平均持有时间定义为个股进入持仓至离开持仓的时间间隔，持有期内平均收益率定义为个股被持有期间的收益率的均值，平均持有时间水平、持有期内平均收益率水平定义为平均持有时间、持有期内平均收益率在偏股型基金整体中的排名。'
        # document.add_paragraph(text=holding_period_appendix)
        # if len(self.compare_fund_code_list) > 0:
        #     document.add_paragraph(text='与同类型基金持仓周期比较情况如下：')
        #     rowc, colc = compare_holding_period.shape[0], compare_holding_period.shape[1]
        #     table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        #     for i in range(rowc):
        #         for j in range(colc):
        #             table.cell(i, j).text = str(compare_holding_period.iloc[i, j])
        #     document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     document.add_paragraph(text=' ')
        #
        # document.add_heading(text='1.5 左侧特征', level=3)
        # document.add_paragraph(text=fund_manager_method_left_text)
        # rowc, colc = left_disp.shape[0], left_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(left_disp.iloc[i, j])
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_paragraph(text=' ')
        # left_appendix = '注：本报告左侧概率定义为个股进入持仓时股价低于半年线的概率，左侧程度定义为左侧进入时个股半年线/进入价格的均值，左侧概率水平、左侧程度水平定义为左侧概率、左侧程度在偏股型基金整体中的排名。若左侧概率水平大于50%、左侧程度水平大于50%，则定义为深度左侧；若左侧概率水平大于50%，左侧程度水平小于等于50%，则定义为左侧；否则定义为无明显的左侧特征。'
        # document.add_paragraph(text=left_appendix)
        # if len(self.compare_fund_code_list) > 0:
        #     document.add_paragraph(text='与同类型基金左侧特征比较情况如下：')
        #     rowc, colc = compare_left.shape[0], compare_left.shape[1]
        #     table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        #     for i in range(rowc):
        #         for j in range(colc):
        #             table.cell(i, j).text = str(compare_left.iloc[i, j])
        #     document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     document.add_paragraph(text=' ')
        #
        # document.add_heading(text='1.6 新股/次新股偏好', level=3)
        # document.add_paragraph(text=fund_manager_method_new_text)
        # rowc, colc = new_disp.shape[0], new_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(new_disp.iloc[i, j])
        # document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # document.add_paragraph(text=' ')
        # new_appendix = '注：本报告新股概率定义为个股进入持仓时成立是否满半年的概率，次新股概率定义为个股进入持仓时成立是否满一年的概率，新股概率水平、次新股概率水平定义为新股概率、次新股概率在偏股型基金整体中的排名。若新股概率水平大于75%，则定义为偏好新股；若次新股概率水平大于75%，则定义为偏好次新股；否则定义为无明显的新股/次新股偏好。'
        # document.add_paragraph(text=new_appendix)
        # if len(self.compare_fund_code_list) > 0:
        #     document.add_paragraph(text='与同类型基金新股/次新股偏好比较情况如下：')
        #     rowc, colc = compare_new.shape[0], compare_new.shape[1]
        #     table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        #     for i in range(rowc):
        #         for j in range(colc):
        #             table.cell(i, j).text = str(compare_new.iloc[i, j])
        #     document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     document.add_paragraph(text=' ')

        head_2 = document.add_heading(text='', level=1)
        run_head_2 = head_2.add_run('二. 基金定量分析')
        run_head_2.font.color.rgb = RGBColor(201, 70, 73)
        head_2_1 = document.add_heading(text='', level=2)
        run_head_2_1 = head_2_1.add_run('1. 仓位分析')
        run_head_2_1.font.color.rgb = RGBColor(201, 70, 73)
        fund_position_text, compare_position = self.get_fund_position()
        document.add_paragraph(text=fund_position_text)
        document.add_picture('{0}{1}_position.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金股票仓位比较情况如下：')
            rowc, colc = compare_position.shape[0], compare_position.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_position.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        head_2_2 = document.add_heading(text='', level=2)
        run_head_2_2 = head_2_2.add_run('2. 规模分析')
        run_head_2_2.font.color.rgb = RGBColor(201, 70, 73)
        fund_size_disp, fund_size_text, compare_size = self.get_fund_size()
        document.add_paragraph(text=fund_size_text)
        rowc, colc = fund_size_disp.shape[0], fund_size_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_size_disp.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        document.add_picture('{0}{1}_size.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fund_size_appendix = '注：本报告大中小盘根据持仓情况进行划分，对标巨潮大中小盘指数成分股确定大中小盘划分及暴露。'
        document.add_paragraph(text=fund_size_appendix)
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金大中小盘暴露比较情况如下：')
            rowc, colc = compare_size.shape[0], compare_size.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_size.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        head_2_3 = document.add_heading(text='', level=2)
        run_head_2_3 = head_2_3.add_run('3. 风格分析')
        run_head_2_3.font.color.rgb = RGBColor(201, 70, 73)
        fund_style_disp, fund_style_text, compare_style = self.get_fund_style()
        document.add_paragraph(text=fund_style_text)
        rowc, colc = fund_style_disp.shape[0], fund_style_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_style_disp.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        document.add_picture('{0}{1}_style.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fund_style_appendix = '注：本报告成长价值风格根据持仓情况进行划分，对标巨潮成长价值指数成分股确定成长价值划分及暴露。'
        document.add_paragraph(text=fund_style_appendix)
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金成长价值暴露比较情况如下：')
            rowc, colc = compare_style.shape[0], compare_style.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_style.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        head_2_4 = document.add_heading(text='', level=2)
        run_head_2_4 = head_2_4.add_run('4. 主题分析')
        run_head_2_4.font.color.rgb = RGBColor(201, 70, 73)
        fund_theme_disp, fund_theme_text, compare_theme = self.get_fund_theme()
        document.add_paragraph(text=fund_theme_text)
        rowc, colc = fund_theme_disp.shape[0], fund_theme_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_theme_disp.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        document.add_picture('{0}{1}_theme.png'.format(self.file_path, self.target_fund_code), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fund_theme_appendix = '注：本报告主题分类如下：大金融：银行、非银金融、房地产；消费：食品饮料、家用电器、医药生物、社会服务、农林牧渔、商贸零售，美容护理；TMT：通信、计算机、电子、传媒、国防军工；制造：交通运输、机械设备、汽车、纺织服饰、轻工制造，电力设备；周期：钢铁、有色金属、建筑装饰、建筑材料、基础化工，石油石化，煤炭，公用事业，环保。'
        document.add_paragraph(text=fund_theme_appendix)
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金主题暴露比较情况如下：')
            rowc, colc = compare_theme.shape[0], compare_theme.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_theme.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        head_2_5 = document.add_heading(text='', level=2)
        run_head_2_5 = head_2_5.add_run('5. 行业分析')
        run_head_2_5.font.color.rgb = RGBColor(201, 70, 73)
        fund_industry_disp, fund_industry_text, compare_industry = self.get_fund_industry()
        document.add_paragraph(text=fund_industry_text)
        rowc, colc = fund_industry_disp.shape[0], fund_industry_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_industry_disp.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        document.add_picture('{0}{1}_industry.png'.format(self.file_path, self.target_fund_code), width=Cm(15), height=Cm(8))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金前五大行业暴露比较情况如下：')
            rowc, colc = compare_industry.shape[0], compare_industry.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_industry.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')
        fund_industry_up, fund_industry_up_text, fund_industry_down, fund_industry_down_text = self.get_fund_industry_change()
        document.add_paragraph(text=fund_industry_up_text)
        rowc, colc = fund_industry_up.shape[0], fund_industry_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_industry_up.iloc[i, j])
                if j % 2 == 1:
                    table.cell(0, j).merge(table.cell(0, j + 1))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        document.add_paragraph(text=fund_industry_down_text)
        rowc, colc = fund_industry_down.shape[0], fund_industry_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_industry_down.iloc[i, j])
                if j % 2 == 1:
                    table.cell(0, j).merge(table.cell(0, j + 1))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        fund_industry_appendix = '注：本报告行业采用2021版申万一级行业。'
        document.add_paragraph(text=fund_industry_appendix)

        head_2_6 = document.add_heading(text='', level=2)
        run_head_2_6 = head_2_6.add_run('6. 持股分析')
        run_head_2_6.font.color.rgb = RGBColor(201, 70, 73)
        fund_cr_text = self.get_fund_cr()
        document.add_paragraph(text=fund_cr_text)
        document.add_picture('{0}{1}_cr5.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_cr10.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fund_finance_text, compare_finance = self.get_fund_finance()
        document.add_paragraph(text=fund_finance_text)
        document.add_picture('{0}{1}_finance.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：wind，好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金估值盈利比较情况如下：')
            rowc, colc = compare_finance.shape[0], compare_finance.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_finance.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        fund_holding_latest, fund_holding_latest_text, fund_holding_disp, fund_holding_disp_text = self.get_fund_stock_change()
        document.add_paragraph(text=fund_holding_latest_text)
        rowc, colc = fund_holding_latest.shape[0], fund_holding_latest.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_holding_latest.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        document.add_paragraph(text=fund_holding_disp_text)
        document.add_picture('{0}{1}_dist.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        rowc, colc = fund_holding_disp.shape[0], fund_holding_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_holding_disp.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        fund_stock_compare, fund_stock_compare_text = self.get_fund_stock_compare()
        document.add_paragraph(text=fund_stock_compare_text)
        rowc, colc = fund_stock_compare.shape[0], fund_stock_compare.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_stock_compare.iloc[i, j])
        for i, stock in enumerate(fund_stock_compare[0].unique().tolist()[1:]):
            stock_index = list(fund_stock_compare[fund_stock_compare[0] == stock].index)
            table.cell(stock_index[0], 0).merge(table.cell(stock_index[-1], 0))
            table.cell(stock_index[0], 1).merge(table.cell(stock_index[-1], 1))
            table.rows[stock_index[0]].cells[0].text = stock
            table.rows[stock_index[0]].cells[1].text = fund_stock_compare[fund_stock_compare[0] == stock][1].values[0]
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        # HoldingCompare(self.target_fund_code, self.compare_fund_code_list).get_holding_compare()
        # p = document.add_paragraph(text='该产品持仓情况以及与其他同类型产品持仓比较情况（若有）见此excel：', style=None)
        # add_hyperlink(p, 'D:/Git/hbshare/hbshare/fe/xwq/data/quantitative_report/{0}/持仓情况.xlsx'.format(self.target_fund_code), '持仓情况。', '#3FAEEE', True)

        head_2_7 = document.add_heading(text='', level=2)
        run_head_2_7 = head_2_7.add_run('7. 换手率分析')
        run_head_2_7.font.color.rgb = RGBColor(201, 70, 73)
        fund_turnover_text, compare_turnover = self.get_fund_turnover()
        document.add_paragraph(text=fund_turnover_text)
        document.add_picture('{0}{1}_turnover.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金换手率比较情况如下：')
            rowc, colc = compare_turnover.shape[0], compare_turnover.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_turnover.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        head_3 = document.add_heading(text='', level=1)
        run_head_3 = head_3.add_run('三. 业绩概况')
        run_head_3.font.color.rgb = RGBColor(201, 70, 73)
        fund_return_period, fund_year_performance = self.get_fund_performace()
        document.add_paragraph(text='该产品业绩收益和回撤情况统计如下：', style=None)
        document.add_picture('{0}{1}_return.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_drawdown.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        achievement, compare_achievement, compare_corr =self.get_fund_achievement()
        rowc, colc = achievement.shape[0], achievement.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(achievement.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        document.add_paragraph(text='从不同时间区间和不同年度对该产品业绩进行统计，情况如下：', style=None)
        rowc, colc = fund_return_period.shape[0], fund_return_period.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_return_period.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        rowc, colc = fund_year_performance.shape[0], fund_year_performance.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(fund_year_performance.iloc[i, j])
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(text=' ')
        if len(self.compare_fund_code_list) > 0:
            document.add_paragraph(text='与同类型基金业绩比较情况如下：')
            document.add_picture('{0}{1}_compare_return.png'.format(self.file_path, self.target_fund_code))
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            rowc, colc = compare_achievement.shape[0], compare_achievement.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_achievement.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')
            document.add_paragraph(text='与同类型基金业绩相关性如下：')
            rowc, colc = compare_corr.shape[0], compare_corr.shape[1]
            table = document.add_table(rows=rowc, cols=colc, style='Medium Shading 1 Accent 2')
            for i in range(rowc):
                for j in range(colc):
                    table.cell(i, j).text = str(compare_corr.iloc[i, j])
            document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph(text=' ')

        head_4 = document.add_heading(text='', level=1)
        run_head_4 = head_4.add_run('四. 绩效归因')
        run_head_4.font.color.rgb = RGBColor(201, 70, 73)
        fund_brison_attribution_text = self.get_fund_brison_attribution()
        document.add_paragraph(text=fund_brison_attribution_text, style=None)
        document.add_picture('{0}{1}_brison_attribution.png'.format(self.file_path, self.target_fund_code), width=Cm(18), height=Cm(10))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fund_barra_text = self.get_fund_barra_attribution()
        document.add_paragraph(text=fund_barra_text, style=None)
        document.add_picture('{0}{1}_barra_exposure.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture('{0}{1}_barra_return.png'.format(self.file_path, self.target_fund_code))
        document.add_paragraph(text='数据来源：好买基金研究中心').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        head_5 = document.add_heading(text='', level=1)
        run_head_5 = head_5.add_run('五. 总结')
        run_head_5.font.color.rgb = RGBColor(201, 70, 73)
        fund_summary = '1. 规模：该产品最新规模为{0}'.format(fund_scale_text.split('最新规模为')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '2. 持有人结构：该产品目前{0}。'.format(fund_holder_text.split('从最新持有人结构看，')[1].split('，占比分别为')[0])
        document.add_paragraph(text=fund_summary)
        # fund_summary = '※ 投资理念：该产品{0}，{1}，{2}，{3}（{4}），{5}，{6}。'.format(fund_manager_method_shift_text.split('该产品')[1].split('，规模')[0],
        #                                                 fund_manager_method_cen_text.split('该产品')[1].split('，规模')[0],
        #                                                 fund_manager_method_select_text.split('，')[2].split('。')[0],
        #                                                 fund_manager_method_holding_period_text.split('该产品')[1].split('（')[0],
        #                                                 fund_manager_method_holding_period_text.split('（')[1].split('）')[0],
        #                                                 fund_manager_method_left_text.split('，')[2].split('。')[0],
        #                                                 fund_manager_method_new_text.split('，')[2].split('。')[0])
        # document.add_paragraph(text=fund_summary)
        fund_summary = '3. 仓位：近三年股票平均仓位为{0}。'.format(fund_position_text.split('其中股票资产近三年平均配置比例为')[1].split('，最新一期')[0])
        document.add_paragraph(text=fund_summary)
        fund_summary = '4. 投资风格：从产品规模看，该产品属于{0}从产品风格看，该产品属于{1}'.format(fund_size_text.split('属于')[1], fund_style_text.split('属于')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '5. 主题配置：该产品属于{0}'.format(fund_theme_text.split('属于')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '6. 行业配置：该产品属于{0}'.format(fund_industry_text.split('属于')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '7. 持股：该产品{0}，{1}。'.format(fund_cr_text.split('该产品持股集中度情况如下，')[1].split('。')[0],
                                                    fund_holding_disp_text.split('从概率密度分布图可以看出，该产品')[1].split('，近六期')[0])
        document.add_paragraph(text=fund_summary)
        fund_summary = '8. 换手率：该产品{0}'.format(fund_turnover_text.split('该产品换手率情况如下，')[1])
        document.add_paragraph(text=fund_summary)
        fund_summary = '9. 绩效归因：从Brison绩效归因看，该产品最新一期{0}；从Barra风格归因看，该产品最新一期{1}'.format(fund_brison_attribution_text.split('从最新一期看，')[1].split('。')[0], fund_barra_text.split('如下，')[1])
        document.add_paragraph(text=fund_summary)

        for paragraph in document.paragraphs:
            if paragraph.style.name == 'Normal' and paragraph.text != '':
                paragraph.paragraph_format.first_line_indent = paragraph.style.font.size * 2

        document.save('{0}{1}（{2}）定量分析报告.docx'.format(self.file_path, self.target_fund_name, self.target_fund_code))

if __name__ == "__main__":
    file_path = 'D:/Git/hbshare/hbshare/fe/xwq/data/quantitative_report/'
    target_fund_code = '001476'  # 001476 270028 377240 008657
    compare_fund_code_list = ['001856', '005968', '001975', '001410', '688888']
    QuantitativeReport(file_path, target_fund_code, compare_fund_code_list).get_quantitative_report()
