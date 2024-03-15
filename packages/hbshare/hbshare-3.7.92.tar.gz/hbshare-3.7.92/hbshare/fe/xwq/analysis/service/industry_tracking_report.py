# -*- coding: utf-8 -*-

from hbshare.fe.xwq.analysis.orm.fedb import FEDB
from hbshare.fe.xwq.analysis.orm.hbdb import HBDB
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.ticker import FuncFormatter
import docx
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False
sns.set_style('white', {'font.sans-serif': ['simhei', 'Arial']})
line_color_list = ['#F04950', '#6268A2', '#959595', '#333335', '#EE703F', '#7E4A9B', '#8A662C',
                   '#44488E', '#BA67E9', '#3FAEEE']
bar_color_list = ['#C94649', '#EEB2B4', '#E1777A', '#D57C56', '#E39A79', '#DB8A66', '#E5B88C',
                  '#8588B7', '#B4B6D1', '#55598D', '#628497', '#A9C6CB', '#866EA9', '#B79BC7',
                  '#7D7D7E', '#CACACA', '#A7A7A8', '#606063', '#C4C4C4', '#99999B', '#B7B7B7']


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def to_percent(temp, position):
    return '%1.0f'%(temp) + '%'

def to_100percent(temp, position):
    return '%1.0f'%(temp * 100) + '%'

def filter_extreme_mad(ser, n=3):
    median = ser.quantile(0.5)
    new_median = ((ser - median).abs()).quantile(0.5)
    max_range = median + n * new_median
    min_range = median - n * new_median
    ser = np.clip(ser, min_range, max_range)
    return ser

def get_date(start_date, end_date):
    calendar_df = HBDB().read_cal(start_date, end_date)
    calendar_df = calendar_df.rename(columns={'JYRQ': 'CALENDAR_DATE', 'SFJJ': 'IS_OPEN', 'SFZM': 'IS_WEEK_END', 'SFYM': 'IS_MONTH_END'})
    calendar_df['CALENDAR_DATE'] = calendar_df['CALENDAR_DATE'].astype(str)
    calendar_df = calendar_df.sort_values('CALENDAR_DATE')
    calendar_df['IS_OPEN'] = calendar_df['IS_OPEN'].astype(int).replace({0: 1, 1: 0})
    calendar_df['YEAR_MONTH'] = calendar_df['CALENDAR_DATE'].apply(lambda x: x[:6])
    calendar_df['MONTH'] = calendar_df['CALENDAR_DATE'].apply(lambda x: x[4:6])
    calendar_df['MONTH_DAY'] = calendar_df['CALENDAR_DATE'].apply(lambda x: x[4:])
    calendar_df = calendar_df[(calendar_df['CALENDAR_DATE'] >= start_date) & (calendar_df['CALENDAR_DATE'] <= end_date)]
    trade_df = calendar_df[calendar_df['IS_OPEN'] == 1].rename(columns={'CALENDAR_DATE': 'TRADE_DATE'})
    trade_df = trade_df[(trade_df['TRADE_DATE'] >= start_date) & (trade_df['TRADE_DATE'] <= end_date)]
    report_df = calendar_df.drop_duplicates('YEAR_MONTH', keep='last').rename(columns={'CALENDAR_DATE': 'REPORT_DATE'})
    report_df = report_df[report_df['MONTH_DAY'].isin(['0331', '0630', '0930', '1231'])]
    report_df = report_df[(report_df['REPORT_DATE'] >= start_date) & (report_df['REPORT_DATE'] <= end_date)]
    report_trade_df = calendar_df[calendar_df['IS_OPEN'] == 1].rename(columns={'CALENDAR_DATE': 'TRADE_DATE'})
    report_trade_df = report_trade_df.sort_values('TRADE_DATE').drop_duplicates('YEAR_MONTH', keep='last')
    report_trade_df = report_trade_df[report_trade_df['MONTH'].isin(['03', '06', '09', '12'])]
    report_trade_df = report_trade_df[(report_trade_df['TRADE_DATE'] >= start_date) & (report_trade_df['TRADE_DATE'] <= end_date)]
    calendar_trade_df = calendar_df[['CALENDAR_DATE']].merge(trade_df[['TRADE_DATE']], left_on=['CALENDAR_DATE'], right_on=['TRADE_DATE'], how='left')
    calendar_trade_df['TRADE_DATE'] = calendar_trade_df['TRADE_DATE'].fillna(method='ffill')
    calendar_trade_df = calendar_trade_df[(calendar_trade_df['TRADE_DATE'] >= start_date) & (calendar_trade_df['TRADE_DATE'] <= end_date)]
    return calendar_df, report_df, trade_df, report_trade_df, calendar_trade_df

def get_stock_industry():
    # stock_industry = HBDB().read_stock_industry()
    # stock_industry.to_hdf('D:/Git/hbshare/hbshare/fe/xwq/data/industry_tracking/stock_industry.hdf', key='table', mode='w')
    stock_industry = pd.read_hdf('D:/Git/hbshare/hbshare/fe/xwq/data/industry_tracking/stock_industry.hdf', key='table')
    stock_industry = stock_industry.rename(columns={'zqdm': 'TICKER_SYMBOL', 'flmc': 'INDUSTRY_NAME', 'fldm': 'INDUSTRY_ID', 'hyhfbz': 'INDUSTRY_VERSION', 'fljb': 'INDUSTRY_TYPE', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'sfyx': 'IS_NEW'})
    stock_industry = stock_industry.dropna(subset=['BEGIN_DATE'])
    stock_industry['END_DATE'] = stock_industry['END_DATE'].fillna('20990101')
    stock_industry['BEGIN_DATE'] = stock_industry['BEGIN_DATE'].astype(int).astype(str)
    stock_industry['END_DATE'] = stock_industry['END_DATE'].astype(int).astype(str)
    stock_industry['INDUSTRY_VERSION'] = stock_industry['INDUSTRY_VERSION'].astype(int)
    stock_industry['INDUSTRY_TYPE'] = stock_industry['INDUSTRY_TYPE'].astype(int)
    stock_industry['IS_NEW'] = stock_industry['IS_NEW'].astype(int)
    stock_industry = stock_industry[stock_industry['INDUSTRY_VERSION'] == 2]
    stock_industry = stock_industry.loc[stock_industry['TICKER_SYMBOL'].str.len() == 6]
    stock_industry = stock_industry.loc[stock_industry['TICKER_SYMBOL'].astype(str).str.slice(0, 1).isin(['0', '3', '6'])]
    return stock_industry

def get_industry_info():
    industry_info = HBDB().read_industry_info()
    industry_info = industry_info.rename(columns={'flmc': 'INDUSTRY_NAME', 'zsdm': 'INDUSTRY_ID', 'hyhfbz': 'INDUSTRY_VERSION', 'fljb': 'INDUSTRY_TYPE', 'qsrq': 'BEGIN_DATE', 'jsrq': 'END_DATE', 'sfyx': 'IS_NEW'})
    industry_info = industry_info.dropna(subset=['BEGIN_DATE'])
    industry_info['END_DATE'] = industry_info['END_DATE'].replace('', np.nan).fillna('20990101')
    industry_info['BEGIN_DATE'] = industry_info['BEGIN_DATE'].apply(lambda x: datetime.strptime(x, '%Y/%m/%d').strftime('%Y%m%d') if '/' in x else x)
    industry_info['END_DATE'] = industry_info['END_DATE'].apply(lambda x: datetime.strptime(x, '%Y/%m/%d').strftime('%Y%m%d') if '/' in x else x)
    industry_info['BEGIN_DATE'] = industry_info['BEGIN_DATE'].astype(int).astype(str)
    industry_info['END_DATE'] = industry_info['END_DATE'].astype(int).astype(str)
    industry_info['INDUSTRY_VERSION'] = industry_info['INDUSTRY_VERSION'].astype(int)
    industry_info['INDUSTRY_TYPE'] = industry_info['INDUSTRY_TYPE'].astype(int)
    industry_info['IS_NEW'] = industry_info['IS_NEW'].astype(int)
    industry_info = industry_info[industry_info['INDUSTRY_VERSION'] == 3]
    return industry_info

def get_stock_info():
    stock_info = HBDB().read_stock_info()
    stock_info = stock_info.rename(columns={'zqdm': 'TICKER_SYMBOL', 'zqjc': 'SEC_SHORT_NAME', 'ssrq': 'ESTABLISH_DATE'})
    stock_info['ESTABLISH_DATE'] = stock_info['ESTABLISH_DATE'].dropna()
    stock_info['ESTABLISH_DATE'] = stock_info['ESTABLISH_DATE'].astype(int).astype(str)
    stock_info = stock_info.loc[stock_info['TICKER_SYMBOL'].str.len() == 6]
    stock_info = stock_info.loc[stock_info['TICKER_SYMBOL'].astype(str).str.slice(0, 1).isin(['0', '3', '6'])]
    stock_info['SAMPLE_DATE'] = stock_info['ESTABLISH_DATE'].apply(lambda x: (datetime.strptime(x, '%Y%m%d') + timedelta(365)).strftime('%Y%m%d'))
    return stock_info

class IndustryTracking:
    def __init__(self, file_path, date, last_date, report_date, report_date_list, head):
        self.file_path = file_path
        self.date = date
        self.last_date = last_date
        self.report_date = report_date
        self.report_date_list = report_date_list
        self.head = head
        self.load_data()

    def load_data(self):
        self.calendar_df, self.report_df, self.trade_df, self.report_trade_df, self.calendar_trade_df = get_date('20190101', self.date)

        # 申万行业信息
        self.industry_info = get_industry_info()
        self.industry_info = self.industry_info[self.industry_info['IS_NEW'] == 1]
        self.industry_info = self.industry_info[['INDUSTRY_NAME', 'INDUSTRY_ID', 'INDUSTRY_TYPE', 'BEGIN_DATE', 'END_DATE']]
        self.industry_id_name_dic = self.industry_info[['INDUSTRY_ID', 'INDUSTRY_NAME']].set_index('INDUSTRY_ID')['INDUSTRY_NAME'].to_dict()
        self.industry_name_id_dic = self.industry_info[['INDUSTRY_ID', 'INDUSTRY_NAME']].set_index('INDUSTRY_NAME')['INDUSTRY_ID'].to_dict()
        self.sw1_industry_list = self.industry_info[self.industry_info['INDUSTRY_TYPE'] == 1]['INDUSTRY_NAME'].unique().tolist()
        self.sw2_industry_list = self.industry_info[self.industry_info['INDUSTRY_TYPE'] == 2]['INDUSTRY_NAME'].unique().tolist()
        self.sw3_industry_list = self.industry_info[self.industry_info['INDUSTRY_TYPE'] == 3]['INDUSTRY_NAME'].unique().tolist()
        self.sw3_industry_list = ['钴', '锂', '铝', '锂电池', '电池化学品', '硅料硅片', '光伏电池组件', '逆变器', '航空装备', '军工电子', '电动乘用车', '车身附件及饰件', '底盘与发动机系统', '空调', '冰洗', '酒店', '白酒', '数字芯片设计', '模拟芯片设计', '半导体设备', '医疗研发外包']

        # 个股与申万行业对应关系
        self.stock_industry = get_stock_industry()
        self.stock_industry = self.stock_industry[self.stock_industry['IS_NEW'] == 1]
        self.stock_industry = self.stock_industry[['INDUSTRY_NAME', 'TICKER_SYMBOL', 'INDUSTRY_TYPE', 'BEGIN_DATE', 'END_DATE']]
        self.sw1_stock_industry = self.stock_industry[self.stock_industry['INDUSTRY_TYPE'] == 1]
        self.sw2_stock_industry = self.stock_industry[self.stock_industry['INDUSTRY_TYPE'] == 2]
        self.sw3_stock_industry = self.stock_industry[self.stock_industry['INDUSTRY_TYPE'] == 3]
        self.sw3_stock_industry = self.stock_industry[self.stock_industry['INDUSTRY_NAME'].isin(self.sw3_industry_list)]

        # 个股信息
        self.stock_info = get_stock_info()

        # 个股市值数据
        # stock_market_value = HBDB().read_stock_market_value_given_date(self.date)
        # star_stock_market_value = HBDB().read_star_stock_market_value_given_date(self.date)
        # self.stock_market_value = pd.concat([stock_market_value, star_stock_market_value])
        # self.stock_market_value.to_hdf('{0}stock_market_value.hdf'.format(self.file_path), key='table', mode='w')
        self.stock_market_value = pd.read_hdf('{0}stock_market_value.hdf'.format(self.file_path),  key='table')

        # 个股估值数据
        stock_valuation_path = '{0}stock_valuation.hdf'.format(self.file_path)
        # if os.path.isfile(stock_valuation_path):
        #     existed_stock_valuation = pd.read_hdf(stock_valuation_path, key='table')
        #     max_date = max(existed_stock_valuation['TRADE_DATE'])
        #     start_date = max(str(max_date), '20190101')
        # else:
        #     existed_stock_valuation = pd.DataFrame()
        #     start_date = '20190101'
        # trade_df = self.trade_df[(self.trade_df['TRADE_DATE'] > start_date) & (self.trade_df['TRADE_DATE'] < datetime.today().strftime('%Y%m%d'))]
        # stock_valuation_list = []
        # for date in trade_df['TRADE_DATE'].unique().tolist():
        #     stock_valuation_date = HBDB().read_stock_valuation_given_date(date)
        #     stock_valuation_date = stock_valuation_date[['TRADE_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'MARKET_VALUE', 'PE_TTM', 'PB_LF', 'PEG', 'DIVIDEND_RATIO_TTM']] if len(stock_valuation_date) > 0 else pd.DataFrame(columns=['TRADE_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'MARKET_VALUE', 'PE_TTM', 'PB_LF', 'PEG', 'DIVIDEND_RATIO_TTM'])
        #     star_stock_valuation_date = HBDB().read_star_stock_valuation_given_date(date)
        #     star_stock_valuation_date = star_stock_valuation_date[['TRADE_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'MARKET_VALUE', 'PE_TTM', 'PB_LF', 'PEG', 'DIVIDEND_RATIO_TTM']] if len(star_stock_valuation_date) > 0 else pd.DataFrame(columns=['TRADE_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'MARKET_VALUE', 'PE_TTM', 'PB_LF', 'PEG', 'DIVIDEND_RATIO_TTM'])
        #     stock_valuation_date = pd.concat([stock_valuation_date, star_stock_valuation_date])
        #     stock_valuation_list.append(stock_valuation_date)
        #     print(date)
        # self.stock_valuation = pd.concat([existed_stock_valuation] + stock_valuation_list, ignore_index=True)
        # self.stock_valuation = self.stock_valuation.sort_values(['TRADE_DATE', 'TICKER_SYMBOL'])
        # self.stock_valuation = self.stock_valuation.reset_index().drop('index', axis=1)
        # self.stock_valuation.to_hdf(stock_valuation_path, key='table', mode='w')
        self.stock_valuation = pd.read_hdf(stock_valuation_path, key='table')

        # 个股财务数据
        stock_finance_path = '{0}stock_finance.hdf'.format(self.file_path)
        # if os.path.isfile(stock_finance_path):
        #     existed_stock_finance = pd.read_hdf(stock_finance_path, key='table')
        #     max_date = max(existed_stock_finance['END_DATE'])
        #     start_date = max(str(max_date), '20190101')
        # else:
        #     existed_stock_finance = pd.DataFrame()
        #     start_date = '20190101'
        # report_df = self.report_df[(self.report_df['REPORT_DATE'] > start_date) & (self.report_df['REPORT_DATE'] < self.report_date)]
        # stock_finance_list = []
        # for date in report_df['REPORT_DATE'].unique().tolist():
        #     stock_finance_date = HBDB().read_stock_finance_given_date(date)
        #     stock_finance_date = stock_finance_date[['END_DATE', 'PUBLISH_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'NET_PROFIT', 'MAIN_INCOME_PS', 'ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM', 'EPS_TTM', 'OPER_CASH_FLOW_PS_TTM', 'NET_ASSET_PS']] if len(stock_finance_date) > 0 else pd.DataFrame(columns=['END_DATE', 'PUBLISH_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'NET_PROFIT', 'MAIN_INCOME_PS', 'ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM', 'EPS_TTM', 'OPER_CASH_FLOW_PS_TTM', 'NET_ASSET_PS'])
        #     star_stock_finance_date = HBDB().read_star_stock_finance_given_date(date)
        #     star_stock_finance_date = star_stock_finance_date[['END_DATE', 'PUBLISH_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'NET_PROFIT', 'MAIN_INCOME_PS', 'ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM', 'EPS_TTM', 'OPER_CASH_FLOW_PS_TTM', 'NET_ASSET_PS']] if len(star_stock_finance_date) > 0 else pd.DataFrame(columns=['END_DATE', 'PUBLISH_DATE', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'NET_PROFIT', 'MAIN_INCOME_PS', 'ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM', 'EPS_TTM', 'OPER_CASH_FLOW_PS_TTM', 'NET_ASSET_PS'])
        #     stock_finance_date = pd.concat([stock_finance_date, star_stock_finance_date])
        #     stock_finance_list.append(stock_finance_date)
        #     print(date)
        # self.stock_finance = pd.concat([existed_stock_finance] + stock_finance_list, ignore_index=True)
        # self.stock_finance = self.stock_finance.sort_values(['END_DATE', 'TICKER_SYMBOL', 'PUBLISH_DATE'])
        # self.stock_finance = self.stock_finance.reset_index().drop('index', axis=1)
        # self.stock_finance.to_hdf(stock_finance_path, key='table', mode='w')
        self.stock_finance = pd.read_hdf(stock_finance_path, key='table')

        # 个股一致预期数据
        # stock_consensus_list = []
        # for type in ['FY0', 'FY1', 'FY2', 'FY3']:
        #     stock_consensus_type = HBDB().read_consensus_given_date(self.date, type)
        #     stock_consensus_list.append(stock_consensus_type)
        # self.stock_consensus = pd.concat(stock_consensus_list)
        # self.stock_consensus['TICKER_SYMBOL'] = self.stock_consensus['TICKER_SYMBOL'].apply(lambda x: x.split('.')[0])
        # self.stock_consensus.to_hdf('{0}stock_consensus.hdf'.format(self.file_path), key='table', mode='w')
        self.stock_consensus = pd.read_hdf('{0}stock_consensus.hdf'.format(self.file_path), key='table')

    def report_info(self):
        """
        风险提示及重要声明
        """
        risk = '投资有风险。基金的过往业绩并不预示其未来表现。基金管理人管理的其他基金的业绩并不构成基金业绩表现的保证。相关数据仅供参考，不构成投资建议。投资人请详阅基金合同等法律文件，了解产品风险收益特征，根据自身资产状况、风险承受能力审慎决策，独立承担投资风险。本资料仅为宣传用品，本机构及工作人员不存在直接或间接主动推介相关产品的行为。'
        statement1 = '本文件中的信息基于已公开的信息、数据及尽调访谈等，好买基金或好买基金研究中心（以下简称“本公司”）对这些信息的及时性、准确性及完整性不做任何保证，也不保证所包含的信息不会发生变更。文件中的内容仅供参考，不代表任何确定性的判断。本文件及其内容均不构成投资建议，也没有考虑个别客户特殊的投资目标、财务状况或需要。获得本文件的机构或个人据此做出投资决策，应自行承担投资风险。'
        statement2 = '本文件版权为本公司所有。未经本公司书面许可，任何机构或个人不得以翻版、复制、 发表、引用或再次分发他人等任何形式侵犯本公司版权。本文件中的信息均为保密信息，未经本公司事先同意，不得以任何目的，复制或传播本文本中所含信息，亦不可向任何第三方披露。'
        statement3 = '本文件系为好买基金备制并仅以非公开方式提交给符合监管规定之合格投资者。'
        return risk, statement1, statement2, statement3

    def get_industry_wret(self, sw_type):
        industry = '申万一级行业' if sw_type == 1 else '申万二级行业' if sw_type == 2 else '申万三级关注行业'
        industry_list = self.sw1_industry_list if sw_type == 1 else self.sw2_industry_list if sw_type == 2 else self.sw3_industry_list
        industry_id_list = [self.industry_name_id_dic[industry] for industry in industry_list]
        self.industry_daily_k = HBDB().read_index_daily_k_given_date_and_indexs(self.last_date, industry_id_list)
        self.industry_daily_k = self.industry_daily_k.rename(columns={'zqdm': 'INDUSTRY_ID', 'jyrq': 'TRADE_DATE', 'spjg': 'CLOSE_INDEX'})
        self.industry_daily_k['TRADE_DATE'] = self.industry_daily_k['TRADE_DATE'].astype(str)
        self.industry_daily_k = self.industry_daily_k[(self.industry_daily_k['TRADE_DATE'] >= self.last_date) & (self.industry_daily_k['TRADE_DATE'] <= self.date)]
        self.industry_daily_k = self.industry_daily_k.pivot(index='TRADE_DATE', columns='INDUSTRY_ID', values='CLOSE_INDEX').sort_index()
        self.industry_daily_k.columns = [self.industry_id_name_dic[col] for col in self.industry_daily_k]
        industry_wret = self.industry_daily_k.iloc[-1] / self.industry_daily_k.iloc[0] - 1.0
        industry_wret = industry_wret.reset_index()
        industry_wret.columns = ['INDUSTRY_NAME', 'WRET']
        industry_wret['WRET'] = industry_wret['WRET'].apply(lambda x: round(x * 100.0, 2))
        industry_wret = industry_wret.sort_values('WRET', ascending=False)
        industry_wret = industry_wret[['INDUSTRY_NAME', 'WRET']]
        industry_wret.columns = [industry, '周度收益率（%）']

        # 收益柱状图
        plt.figure(figsize=(6, 3))
        sns.barplot(x=industry, y='周度收益率（%）', data=industry_wret, palette=[bar_color_list[0]])
        plt.xlabel(industry)
        plt.ylabel('')
        plt.xticks(rotation=90)
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.tight_layout()
        plt.savefig('{0}sw{1}_wret.png'.format(self.file_path, sw_type))

        industry_wret_disp = industry_wret.copy(deep=True)
        industry_wret_disp['周度收益率（%）_展示'] = industry_wret_disp.apply(lambda x: '{0}（{1}%）'.format(x[industry], x['周度收益率（%）']), axis=1)
        industry_wret_text = '下图展示了{0}指数本周收益率排名情况，其中{1}行业收益排名较为靠前，而{2}行业收益排名较为靠后。'\
            .format(industry, '、'.join(industry_wret_disp['周度收益率（%）_展示'].unique().tolist()[:3]), '、'.join(industry_wret_disp['周度收益率（%）_展示'].unique().tolist()[-3:][::-1]))
        return industry_wret_text

    def get_industry_pe(self, sw_type):
        industry = '申万一级行业' if sw_type == 1 else '申万二级行业' if sw_type == 2 else '申万三级关注行业'
        stock_industry = self.sw1_stock_industry if sw_type == 1 else self.sw2_stock_industry if sw_type == 2 else self.sw3_stock_industry

        # 历史估值水平
        three_years_ago = (datetime.strptime(self.date, '%Y%m%d') - timedelta(365 * 3)).strftime('%Y%m%d')
        stock_valuation = self.stock_valuation[(self.stock_valuation['TRADE_DATE'] >= three_years_ago) & (self.stock_valuation['TRADE_DATE'] <= self.date)]
        stock_valuation = stock_valuation.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='inner')
        stock_valuation = stock_valuation.merge(self.stock_info[['TICKER_SYMBOL', 'SAMPLE_DATE']], on=['TICKER_SYMBOL'], how='inner')
        stock_valuation = stock_valuation[stock_valuation['TRADE_DATE'] >= stock_valuation['SAMPLE_DATE']]
        stock_valuation = stock_valuation.dropna(subset=['TRADE_DATE', 'TICKER_SYMBOL', 'INDUSTRY_NAME', 'MARKET_VALUE', 'PE_TTM'])
        stock_valuation['PE_TTM'] = filter_extreme_mad(stock_valuation['PE_TTM'])
        industry_market_value = stock_valuation[['TRADE_DATE', 'INDUSTRY_NAME', 'MARKET_VALUE']].groupby(['TRADE_DATE', 'INDUSTRY_NAME']).sum().reset_index().rename(columns={'MARKET_VALUE': 'TOTAL_MARKET_VALUE'})
        stock_valuation = stock_valuation.merge(industry_market_value, on=['TRADE_DATE', 'INDUSTRY_NAME'], how='inner')
        stock_valuation['WEIGHT_' + 'PE_TTM'] = stock_valuation['PE_TTM'] * stock_valuation['MARKET_VALUE'] / stock_valuation['TOTAL_MARKET_VALUE']
        industry_pe = stock_valuation[['TRADE_DATE', 'INDUSTRY_NAME', 'WEIGHT_' + 'PE_TTM']].groupby(['TRADE_DATE', 'INDUSTRY_NAME']).sum().reset_index().rename(columns={'WEIGHT_' + 'PE_TTM': 'PE_TTM'})
        industry_pe_latest = industry_pe[industry_pe['TRADE_DATE'] == self.date]
        industry_pe_quantile = industry_pe[['INDUSTRY_NAME', 'TRADE_DATE', 'PE_TTM']].groupby('INDUSTRY_NAME').apply(lambda x: (1.0 - np.count_nonzero(x.sort_values('TRADE_DATE')['PE_TTM'].iloc[-1] <= x['PE_TTM']) / x['PE_TTM'].size) * 100.0)
        industry_pe_quantile = pd.DataFrame(industry_pe_quantile).reset_index().rename(columns={0: 'PE_TTM_QUANTILE'}).sort_values('PE_TTM_QUANTILE')
        industry_list = industry_pe_quantile['INDUSTRY_NAME'].unique().tolist()[::-1]
        industry_pe = industry_pe.pivot(index='TRADE_DATE', columns='INDUSTRY_NAME', values='PE_TTM').sort_index()
        industry_pe_list = [industry_pe[ind] for ind in industry_list]
        industry_pe_latest['INDUSTRY_NAME'] = industry_pe_latest['INDUSTRY_NAME'].astype('category')
        industry_pe_latest['INDUSTRY_NAME'].cat.reorder_categories(industry_list, inplace=True)
        industry_pe_latest = industry_pe_latest.sort_values('INDUSTRY_NAME')
        plt.figure(figsize=(6, 6))
        plt.boxplot(industry_pe_list, labels=industry_list, vert=False, flierprops={'marker': 'o', 'markersize': 1}, meanline=True, showmeans=True)
        plt.scatter(industry_pe_latest['PE_TTM'], range(1, len(industry_pe_latest) + 1), marker='o')
        plt.xlabel(industry)
        plt.tight_layout()
        plt.savefig('{0}sw{1}_pe.png'.format(self.file_path, sw_type))

        industry_pe_quantile_disp = industry_pe_quantile.copy(deep=True)
        industry_pe_quantile_disp['PE_TTM_QUANTILE'] = industry_pe_quantile_disp['PE_TTM_QUANTILE'].apply(lambda x: round(x, 2))
        industry_pe_quantile_disp.columns = [industry, '历史分位水平（%）']
        industry_pe_quantile_disp['历史分位水平（%）_展示'] = industry_pe_quantile_disp.apply(lambda x: '{0}（{1}%）'.format(x[industry], x['历史分位水平（%）']), axis=1)
        industry_pe_text = '下图展示了{0}最新估值在近三年中的分位水平，其中{1}行业PE_TTM处于近三年以来较低水平，而{2}行业PE_TTM处于近三年以来较高水平。' \
            .format(industry, '、'.join(industry_pe_quantile_disp['历史分位水平（%）_展示'].unique().tolist()[:3]), '、'.join(industry_pe_quantile_disp['历史分位水平（%）_展示'].unique().tolist()[-3:][::-1]))

        # 预期估值水平
        stock_consensus = self.stock_consensus[self.stock_consensus['EST_DT'] == self.date].rename(columns={'EST_DT': 'TRADE_DATE'})
        stock_market_value = self.stock_market_value[self.stock_market_value['TRADE_DATE'] == self.date][['TICKER_SYMBOL', 'MARKET_VALUE']]
        stock_consensus = stock_consensus.merge(stock_market_value, on=['TICKER_SYMBOL'], how='inner')
        stock_consensus = stock_consensus.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='inner')
        stock_consensus = stock_consensus.merge(self.stock_info[['TICKER_SYMBOL', 'SAMPLE_DATE']], on=['TICKER_SYMBOL'], how='inner')
        stock_consensus = stock_consensus[stock_consensus['TRADE_DATE'] >= stock_consensus['SAMPLE_DATE']]
        stock_consensus = stock_consensus.dropna(subset=['TRADE_DATE', 'TICKER_SYMBOL', 'INDUSTRY_NAME', 'ROLLING_TYPE', 'MARKET_VALUE', 'EST_PE'])
        industry_cpe_list = []
        for type in ['FY0', 'FY1', 'FY2', 'FY3']:
            stock_consensus_type = stock_consensus[stock_consensus['ROLLING_TYPE'] == type]
            stock_consensus_type['EST_PE'] = filter_extreme_mad(stock_consensus['EST_PE'])
            industry_market_value = stock_consensus_type[['TRADE_DATE', 'INDUSTRY_NAME', 'MARKET_VALUE']].groupby(['TRADE_DATE', 'INDUSTRY_NAME']).sum().reset_index().rename(columns={'MARKET_VALUE': 'TOTAL_MARKET_VALUE'})
            stock_consensus_type = stock_consensus_type.merge(industry_market_value, on=['TRADE_DATE', 'INDUSTRY_NAME'], how='inner')
            stock_consensus_type['WEIGHT_' + 'EST_PE'] = stock_consensus_type['EST_PE'] * stock_consensus_type['MARKET_VALUE'] / stock_consensus_type['TOTAL_MARKET_VALUE']
            industry_cpe_type = stock_consensus_type[['TRADE_DATE', 'INDUSTRY_NAME', 'WEIGHT_' + 'EST_PE']].groupby(['TRADE_DATE', 'INDUSTRY_NAME']).sum().reset_index().rename(columns={'WEIGHT_' + 'EST_PE': 'EST_PE'})
            industry_cpe_type['EST_PE'] = industry_cpe_type['EST_PE'].apply(lambda x: round(x, 2))
            industry_cpe_type['ROLLING_TYPE'] = type
            industry_cpe_list.append(industry_cpe_type)
        industry_cpe = pd.concat(industry_cpe_list)
        industry_cpe = industry_cpe.pivot(index='INDUSTRY_NAME', columns='ROLLING_TYPE', values='EST_PE')
        industry_cpe = industry_cpe.reset_index()
        industry_cpe['INDUSTRY_NAME'] = industry_cpe['INDUSTRY_NAME'].astype('category')
        industry_cpe['INDUSTRY_NAME'].cat.reorder_categories(industry_list[::-1], inplace=True)
        industry_cpe = industry_cpe.sort_values('INDUSTRY_NAME')
        industry_cpe = industry_cpe.set_index('INDUSTRY_NAME')

        industry_cpe_disp = industry_cpe.T.reset_index().T.reset_index()
        industry_cpe_disp.iloc[0, 0] = ''
        industry_cpe_text = '下表列示了分析师对{0}PE水平的未来三年一致预期情况，其中FY0表示最新公告年报对应年份的估值情况，FY1、FY2、FY3分布代表未来一年、未来两年、未来三年的估值情况'\
            .format(industry)

        # 龙头估值情况
        stock_market_value = self.stock_market_value.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='inner')
        stock_market_value = stock_market_value.dropna(subset=['TICKER_SYMBOL', 'INDUSTRY_NAME', 'MARKET_VALUE'])
        industry_head_stocks = stock_market_value.sort_values(['INDUSTRY_NAME', 'MARKET_VALUE'], ascending=[True, False]).groupby(['INDUSTRY_NAME']).head(self.head)
        industry_head_stocks = industry_head_stocks[['INDUSTRY_NAME', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'MARKET_VALUE']]
        industry_head_stocks['MARKET_VALUE'] = industry_head_stocks['MARKET_VALUE'].apply(lambda x: round(x / 100000000.0, 2))
        industry_head_stocks['INDUSTRY_NAME'] = industry_head_stocks['INDUSTRY_NAME'].astype('category')
        industry_head_stocks['INDUSTRY_NAME'].cat.reorder_categories(industry_list[::-1], inplace=True)
        industry_head_stocks = industry_head_stocks.sort_values(['INDUSTRY_NAME', 'MARKET_VALUE'],  ascending=[True, False])
        stock_pe = self.stock_valuation[(self.stock_valuation['TRADE_DATE'] >= three_years_ago) & (self.stock_valuation['TRADE_DATE'] <= self.date)]
        stock_pe = stock_pe[stock_pe['TICKER_SYMBOL'].isin(industry_head_stocks['TICKER_SYMBOL'].unique().tolist())]
        stock_pe_latest = stock_pe[stock_pe['TRADE_DATE'] == self.date]
        stock_pe_latest['PE_TTM'] = stock_pe_latest['PE_TTM'].apply(lambda x: round(x, 2))
        stock_pe_quantile = stock_pe[['TICKER_SYMBOL', 'TRADE_DATE', 'PE_TTM']].groupby('TICKER_SYMBOL').apply(lambda x: (1.0 - np.count_nonzero(x.sort_values('TRADE_DATE')['PE_TTM'].iloc[-1] <= x['PE_TTM']) / x['PE_TTM'].size) * 100.0)
        stock_pe_quantile = pd.DataFrame(stock_pe_quantile).reset_index().rename(columns={0: 'PE_TTM_QUANTILE'}).sort_values('PE_TTM_QUANTILE')
        stock_pe_quantile['PE_TTM_QUANTILE'] = stock_pe_quantile['PE_TTM_QUANTILE'].apply(lambda x: round(x, 2))
        stock_cpe = self.stock_consensus[self.stock_consensus['EST_DT'] == self.date].rename(columns={'EST_DT': 'TRADE_DATE'})
        stock_cpe = stock_cpe[stock_cpe['TICKER_SYMBOL'].isin(industry_head_stocks['TICKER_SYMBOL'].unique().tolist())]
        stock_cpe = stock_cpe.pivot(index='TICKER_SYMBOL', columns='ROLLING_TYPE', values='EST_PE')
        for type in ['FY0', 'FY1', 'FY2', 'FY3']:
            stock_cpe[type] = stock_cpe[type].apply(lambda x: round(x, 2))
        stock_cpe = stock_cpe.reset_index()

        industry_head_pe_disp = industry_head_stocks.merge(stock_pe_latest[['TICKER_SYMBOL', 'PE_TTM']], on=['TICKER_SYMBOL'], how='left') \
                                                    .merge(stock_pe_quantile[['TICKER_SYMBOL', 'PE_TTM_QUANTILE']], on=['TICKER_SYMBOL'], how='left') \
                                                    .merge(stock_cpe, on=['TICKER_SYMBOL'], how='left')
        industry_head_pe_disp['PE_TTM_QUANTILE'] = industry_head_pe_disp['PE_TTM_QUANTILE'].apply(lambda x: '{0}%'.format(x))
        industry_head_pe_disp.columns = [industry, '股票代码', '股票名称', '市值（亿）', 'PE（TTM）', '历史分位水平', 'PE（TTM）FY0', 'PE（TTM）FY1', 'PE（TTM）FY2', 'PE（TTM）FY3']
        industry_head_pe_text = '以行业内市值排名前{0}的个股作为行业龙头个股，{1}内龙头个股的估值水平详见此excel：'.format(self.head, industry)
        return industry_pe_latest, industry_pe_quantile_disp, industry_pe_text, industry_cpe_text, industry_cpe_disp, industry_head_pe_text, industry_head_pe_disp

    def get_industry_fmt(self, sw_type):
        industry = '申万一级行业' if sw_type == 1 else '申万二级行业' if sw_type == 2 else '申万三级关注行业'
        stock_industry = self.sw1_stock_industry if sw_type == 1 else self.sw2_stock_industry if sw_type == 2 else self.sw3_stock_industry

        # 历史盈利水平
        industry_np_yoy = FEDB().read_industry_data(['REPORT_DATE', 'INDUSTRY_ID', 'INDUSTRY_NAME', 'INDUSTRY_TYPE', 'NET_PROFIT_ACCUM_YOY'], 'industry_fundamental_derive', sw_type)
        industry_np_yoy = industry_np_yoy[industry_np_yoy['REPORT_DATE'].isin(self.report_date_list)]
        industry_np_yoy = industry_np_yoy.pivot(index='INDUSTRY_NAME', columns='REPORT_DATE', values='NET_PROFIT_ACCUM_YOY').reset_index()
        industry_np_yoy = industry_np_yoy.sort_values(self.report_date, ascending=False)
        for date in self.report_date_list:
            industry_np_yoy[date] = industry_np_yoy[date].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
        industry_np_yoy = industry_np_yoy[['INDUSTRY_NAME'] + self.report_date_list]
        industry_list = industry_np_yoy['INDUSTRY_NAME'].unique().tolist()
        industry_np_yoy_head = industry_np_yoy['INDUSTRY_NAME'].unique().tolist()[:5]
        industry_np_yoy_tail = industry_np_yoy['INDUSTRY_NAME'].unique().tolist()[-5:]
        industry_np_yoy.columns = [industry] + [date.replace('0331', 'Q1').replace('0630', 'Q2').replace('0930', 'Q3').replace('1231', 'A') for date in self.report_date_list]
        industry_np_yoy = industry_np_yoy.T.reset_index().T
        industry_np_yoy.columns = ['行业名称'] + ['归母净利润同比增速'] * (industry_np_yoy.shape[1] - 1)

        industry_cnp_yoy = FEDB().read_industry_data(['REPORT_DATE', 'INDUSTRY_ID', 'INDUSTRY_NAME', 'INDUSTRY_TYPE', 'EST_NET_PROFIT_YOY'], 'industry_consensus', sw_type)
        industry_cnp_yoy = industry_cnp_yoy[industry_cnp_yoy['REPORT_DATE'] == self.report_date]
        industry_cnp_yoy = industry_cnp_yoy.pivot(index='INDUSTRY_NAME', columns='REPORT_DATE', values='EST_NET_PROFIT_YOY').reset_index()
        industry_cnp_yoy[self.report_date] = industry_cnp_yoy[self.report_date].apply(lambda x: '{0}%'.format(round(x, 2)))
        industry_cnp_yoy = industry_cnp_yoy[['INDUSTRY_NAME', self.report_date]]
        industry_cnp_yoy.columns = [industry, 'FYOY']
        industry_cnp_yoy = industry_cnp_yoy.T.reset_index().T
        industry_cnp_yoy.columns = ['行业名称'] + ['一致预期归母净利润同比增速'] * (industry_cnp_yoy.shape[1] - 1)

        index_list = ['ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM']
        industry_fmt = FEDB().read_industry_data(['REPORT_DATE', 'INDUSTRY_ID', 'INDUSTRY_NAME', 'INDUSTRY_TYPE'] + index_list, 'industry_fundamental', sw_type)
        industry_fmt = industry_fmt[industry_fmt['REPORT_DATE'] == self.report_date]
        industry_fmt = industry_fmt[['INDUSTRY_NAME'] + index_list]
        for index in index_list:
            industry_fmt[index] = industry_fmt[index].apply(lambda x: '{0}%'.format(round(x, 2)))
        industry_fmt = industry_fmt[['INDUSTRY_NAME'] + index_list]
        industry_fmt.columns = [industry] + [self.report_date.replace('0331', 'Q1').replace('0630', 'Q2').replace('0930', 'Q3').replace('1231', 'A')] * (industry_fmt.shape[1] - 1)
        industry_fmt = industry_fmt.T.reset_index().T
        industry_fmt.columns = ['行业名称'] + ['ROE（TTM）', '毛利率（TTM）', '净利率（TTM）']

        industry_fmt_disp = industry_np_yoy.merge(industry_cnp_yoy, on=['行业名称'], how='left') \
                                           .merge(industry_fmt, on=['行业名称'], how='left')
        industry_fmt_disp = industry_fmt_disp.T.reset_index().T.reset_index().drop('index', axis=1)
        industry_fmt_disp.iloc[0, 0] = ''
        industry_fmt_disp_text = '下表展示了{0}的财务相关数据，最新报告期{1}行业的净利润同比增速居前，而{2}行业的净利润同比增速居后。'\
            .format(industry, '、'.join(industry_np_yoy_head), '、'.join(industry_np_yoy_tail))

        # 龙头盈利情况
        stock_market_value = self.stock_market_value.merge(stock_industry[['TICKER_SYMBOL', 'INDUSTRY_NAME']], on=['TICKER_SYMBOL'], how='inner')
        stock_market_value = stock_market_value.dropna(subset=['TICKER_SYMBOL', 'INDUSTRY_NAME', 'MARKET_VALUE'])
        industry_head_stocks = stock_market_value.sort_values(['INDUSTRY_NAME', 'MARKET_VALUE'], ascending=[True, False]).groupby(['INDUSTRY_NAME']).head(self.head)
        industry_head_stocks = industry_head_stocks[['INDUSTRY_NAME', 'TICKER_SYMBOL', 'SEC_SHORT_NAME', 'MARKET_VALUE']]
        industry_head_stocks['MARKET_VALUE'] = industry_head_stocks['MARKET_VALUE'].apply(lambda x: round(x / 100000000.0, 2))
        industry_head_stocks['INDUSTRY_NAME'] = industry_head_stocks['INDUSTRY_NAME'].astype('category')
        industry_head_stocks['INDUSTRY_NAME'].cat.reorder_categories(industry_list, inplace=True)
        industry_head_stocks = industry_head_stocks.sort_values(['INDUSTRY_NAME', 'MARKET_VALUE'], ascending=[True, False])

        industry_head_fmt = self.stock_finance[self.stock_finance['TICKER_SYMBOL'].isin(industry_head_stocks['TICKER_SYMBOL'].unique().tolist())]
        industry_head_fmt = industry_head_fmt.sort_values(['TICKER_SYMBOL', 'END_DATE', 'PUBLISH_DATE']).drop_duplicates(['TICKER_SYMBOL', 'END_DATE'], keep='last')
        industry_head_np = industry_head_fmt.pivot(index='END_DATE', columns='TICKER_SYMBOL', values='NET_PROFIT')
        industry_head_np_yoy = (industry_head_np - industry_head_np.shift(4)) / abs(industry_head_np.shift(4))
        industry_head_np_yoy = industry_head_np_yoy.unstack().reset_index()
        industry_head_np_yoy.columns = ['TICKER_SYMBOL', 'END_DATE', 'NET_PROFIT_YOY']
        industry_head_np_yoy = industry_head_np_yoy[industry_head_np_yoy['END_DATE'] == self.report_date]
        industry_head_fmt = industry_head_fmt[industry_head_fmt['END_DATE'] == self.report_date]
        industry_head_fmt = industry_head_fmt[['TICKER_SYMBOL', 'NET_PROFIT', 'ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM']]

        stock_cnp = self.stock_consensus[self.stock_consensus['EST_DT'] == self.date].rename(columns={'EST_DT': 'TRADE_DATE'})
        stock_cnp = stock_cnp[stock_cnp['TICKER_SYMBOL'].isin(industry_head_stocks['TICKER_SYMBOL'].unique().tolist())]
        stock_cnp = stock_cnp.pivot(index='TICKER_SYMBOL', columns='ROLLING_TYPE', values='EST_NET_PROFIT')
        stock_cnp = stock_cnp.reset_index()

        industry_head_fmt_disp = industry_head_stocks.merge(industry_head_np_yoy[['TICKER_SYMBOL', 'NET_PROFIT_YOY']], on=['TICKER_SYMBOL'], how='left') \
                                                     .merge(industry_head_fmt, on=['TICKER_SYMBOL'], how='left') \
                                                     .merge(stock_cnp, on=['TICKER_SYMBOL'], how='left')
        industry_head_fmt_disp['NET_PROFIT_YOY'] = industry_head_fmt_disp['NET_PROFIT_YOY'].apply(lambda x: '{0}%'.format(round(x * 100.0, 2)))
        for col in ['NET_PROFIT', 'FY0', 'FY1', 'FY2', 'FY3']:
            industry_head_fmt_disp[col] = industry_head_fmt_disp[col].apply(lambda x: '{0}'.format(round(x / 10000000000.0, 2)))
        for col in ['ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM']:
            industry_head_fmt_disp[col] = industry_head_fmt_disp[col].apply(lambda x: '{0}%'.format(round(x, 2)))
        industry_head_fmt_disp = industry_head_fmt_disp[['TICKER_SYMBOL', 'NET_PROFIT', 'NET_PROFIT_YOY', 'FY0', 'FY1', 'FY2', 'FY3', 'ROE_TTM', 'GROSS_INCOME_RATIO_TTM', 'NET_PROFIT_RATIO_TTM']]
        industry_head_fmt_disp.columns = ['股票代码', '归母净利润（百亿）', '归母净利润（百亿）FY0', '归母净利润（百亿）FY1', '归母净利润（百亿）FY2', '归母净利润（百亿）FY3', '归母净利润同比增长率', 'ROE（TTM）', '毛利率（TTM）', '净利率（TTM）']
        industry_head_fmt_text = '以行业内市值排名前{0}的个股作为行业龙头个股，{1}内龙头个股的盈利水平详见此excel：'.format(self.head, industry)
        return industry_fmt_disp_text, industry_fmt_disp, industry_head_fmt_text, industry_head_fmt_disp

    def get_report(self):
        document = Document()
        document.styles['Normal'].font.size = Pt(10)
        document.add_heading(text='【{0}】行业数据跟踪周报'.format(self.date), level=0)

        risk, statement1, statement2, statement3 = self.report_info()
        document.add_paragraph(text='风险提示')
        document.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(text=risk)
        document.add_paragraph(text='重要声明')
        document.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(text=statement1)
        document.add_paragraph(text=statement2)
        document.add_paragraph(text=statement3)
        document.add_page_break()

        document.add_heading(text='一. 本周行业表现', level=1)
        sw1_wret_text = self.get_industry_wret(1)
        document.add_paragraph(text=sw1_wret_text)
        document.add_picture('{0}sw{1}_wret.png'.format(self.file_path, 1))
        # sw3_wret_text = self.get_industry_wret(3)
        # document.add_paragraph(text=sw3_wret_text)
        # document.add_picture('{0}sw{1}_wret.png'.format(self.file_path, 3))

        document.add_heading(text='二. 本周行业估值数据跟踪', level=1)
        sw1_pe_latest, sw1_pe_quantile_disp, sw1_pe_text, sw1_cpe_text, sw1_cpe_disp, sw1_head_pe_text, sw1_head_pe_disp = self.get_industry_pe(1)
        document.add_paragraph(text=sw1_pe_text)
        document.add_picture('{0}sw{1}_pe.png'.format(self.file_path, 1))
        document.add_paragraph(text='注：蓝色点代表行业当前估值水平，黄色实线代表行业中位水平，绿色虚线代表行业平均水平。')
        document.add_paragraph(text=sw1_cpe_text)
        rowc, colc = sw1_cpe_disp.shape[0], sw1_cpe_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Grid 3')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(sw1_cpe_disp.iloc[i, j])
        document.add_paragraph(text=' ')
        p = document.add_paragraph(text=sw1_head_pe_text)
        add_hyperlink(p, '{0}行业数据_{1}.xlsx'.format(self.file_path, self.date), '行业数据。', '#3FAEEE', True)
        # sw3_pe_latest, sw3_pe_quantile_disp, sw3_pe_text, sw3_cpe_text, sw3_cpe_disp, sw3_head_pe_text, sw3_head_pe_disp = self.get_industry_pe(3)
        # document.add_paragraph(text=sw3_pe_text)
        # document.add_picture('{0}sw{1}_pe.png'.format(self.file_path, 3))
        # document.add_paragraph(text='注：蓝色点代表行业当前估值水平，黄色实线代表行业中位水平，绿色虚线代表行业平均水平。')
        # document.add_paragraph(text=sw3_cpe_text)
        # rowc, colc = sw3_cpe_disp.shape[0], sw3_cpe_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Grid 3')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(sw3_cpe_disp.iloc[i, j])
        # document.add_paragraph(text=' ')
        # document.add_paragraph(text=sw3_head_pe_text)
        # p = document.add_paragraph(text=sw1_head_pe_text)
        # add_hyperlink(p, '{0}行业数据_{1}.xlsx'.format(self.file_path, self.date), '行业数据。', '#3FAEEE', True)

        document.add_heading(text='三. 本周行业基本面数据跟踪', level=1)
        sw1_fmt_text, sw1_fmt_disp, sw1_head_fmt_text, sw1_head_fmt_disp = self.get_industry_fmt(1)
        document.add_paragraph(text=sw1_fmt_text)
        rowc, colc = sw1_fmt_disp.shape[0], sw1_fmt_disp.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Grid 3')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(sw1_fmt_disp.iloc[i, j])
        for i, col in enumerate(sw1_fmt_disp.iloc[0].unique().tolist()[1:]):
            col_index = list(sw1_fmt_disp.iloc[0][sw1_fmt_disp.iloc[0] == col].index)
            table.cell(0, col_index[0]).merge(table.cell(0, col_index[-1]))
            table.cell(0, col_index[0]).text = col
        document.add_paragraph(text='')
        p = document.add_paragraph(text=sw1_head_fmt_text)
        add_hyperlink(p, '{0}行业数据_{1}.xlsx'.format(self.file_path, self.date), '行业数据。', '#3FAEEE', True)
        # sw3_fmt_text, sw3_fmt_disp, sw3_head_fmt_text, sw3_head_fmt_disp = self.get_industry_fmt(3)
        # document.add_paragraph(text=sw3_fmt_text)
        # rowc, colc = sw3_fmt_disp.shape[0], sw3_fmt_disp.shape[1]
        # table = document.add_table(rows=rowc, cols=colc, style='Medium Grid 3')
        # for i in range(rowc):
        #     for j in range(colc):
        #         table.cell(i, j).text = str(sw3_fmt_disp.iloc[i, j])
        # for i, col in enumerate(sw3_fmt_disp.iloc[0].unique().tolist()[1:]):
        #     col_index = list(sw3_fmt_disp.iloc[0][sw3_fmt_disp.iloc[0] == col].index)
        #     table.cell(0, col_index[0]).merge(table.cell(0, col_index[-1]))
        #     table.cell(0, col_index[0]).text = col
        # document.add_paragraph(text='')
        # p = document.add_paragraph(text=sw3_head_fmt_text)
        # add_hyperlink(p, '{0}行业数据_{1}.xlsx'.format(self.file_path, self.date), '行业数据。', '#3FAEEE', True)

        document.add_heading(text='四. 行业观点', level=1)
        sw1_pe_q = sw1_pe_quantile_disp[['申万一级行业', '历史分位水平（%）']].rename(columns={'历史分位水平（%）': 'PE（TTM）历史分位水平（%）'})
        sw1_roe = sw1_fmt_disp.iloc[2:, [0, 9]]
        sw1_roe[9] = sw1_roe[9].apply(lambda x: float(x.split('%')[0]))
        sw1_roe.columns = ['申万一级行业', 'ROE（TTM）（%）']
        sw1_pe_roe = sw1_pe_q.merge(sw1_roe, on=['申万一级行业'], how='left')
        sw1_pe_roe_select = sw1_pe_roe[(sw1_pe_roe['PE（TTM）历史分位水平（%）'] < 50) & (sw1_pe_roe['ROE（TTM）（%）'] > sw1_pe_roe['ROE（TTM）（%）'].quantile(0.75))]
        sw1_pe_roe_select = sw1_pe_roe_select['申万一级行业'].unique().tolist()
        plt.figure(figsize=(6, 6))
        plt.scatter(sw1_pe_roe['PE（TTM）历史分位水平（%）'], sw1_pe_roe['ROE（TTM）（%）'])
        for i in range(len(sw1_pe_roe)):
            plt.annotate(sw1_pe_roe['申万一级行业'][i], xy=(sw1_pe_roe['PE（TTM）历史分位水平（%）'][i], sw1_pe_roe['ROE（TTM）（%）'][i]), xytext=(sw1_pe_roe['PE（TTM）历史分位水平（%）'][i] + 0.1, sw1_pe_roe['ROE（TTM）（%）'][i] + 0.1))  # 这里xy是需要标记的坐标，xytext是对应的标签坐标
        plt.xlabel('PE（TTM）历史分位水平')
        plt.ylabel('ROE（TTM）')
        plt.hlines(y=sw1_roe['ROE（TTM）（%）'].mean(), xmin=0.0, xmax=100.0, linestyles='dashed', color='#959595')
        plt.vlines(x=50, ymin=sw1_roe['ROE（TTM）（%）'].min(), ymax=sw1_roe['ROE（TTM）（%）'].max(), linestyles='dashed', color='#959595')
        plt.gca().xaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.gca().yaxis.set_major_formatter(FuncFormatter(to_percent))
        plt.tight_layout()
        plt.savefig('{0}sw{1}_pe_roe.png'.format(self.file_path, 1))
        sw1_pe_roe_text = '从下图可以看出，估值低位且ROE水平相对较高的行业是我们重点关注的行业，包括{0}，这些行业的PE（TTM）历史分位水平小于50%，且ROE（TTM）在所有行业中处于前1/4的水平。'.format('、'.join(sw1_pe_roe_select))
        document.add_paragraph(text=sw1_pe_roe_text)
        document.add_picture('{0}sw{1}_pe_roe.png'.format(self.file_path, 1))

        pe = sw1_pe_latest[['INDUSTRY_NAME', 'PE_TTM']].set_index('INDUSTRY_NAME')
        pe['PE'] = pe['PE_TTM'].rank(ascending=False)
        pe_q = sw1_pe_quantile_disp[['申万一级行业', '历史分位水平（%）']].set_index('申万一级行业')
        pe_q['PE_Q'] = pe_q['历史分位水平（%）'].rank(ascending=False)
        pe_score = pe[['PE']].merge(pe_q[['PE_Q']], left_index=True, right_index=True, how='left')
        pe_score['PE_SCORE'] = (pe_score['PE'] + pe_score['PE_Q']) / 2.0
        pe_score['PE_SCORE'] = pe_score['PE_SCORE'].rank()
        fmt_score = sw1_fmt_disp.iloc[2:, [0, 7, 8, 9, 10, 11]]
        fmt_score = fmt_score.set_index(0)
        fmt_score.columns = ['NP_YOY', 'CNP_YOY', 'ROE', 'GIR', 'NPR']
        for col in ['NP_YOY', 'CNP_YOY', 'ROE', 'GIR', 'NPR']:
            fmt_score[col] = fmt_score[col].apply(lambda x: float(x.split('%')[0]))
            fmt_score[col] = fmt_score[col].rank()
        fmt_score['FMT_SCORE'] = fmt_score.mean(axis=1)
        score = pe_score[['PE_SCORE']].merge(fmt_score[['FMT_SCORE']], left_index=True, right_index=True, how='left')
        score['SCORE'] = (score['PE_SCORE'] + score['FMT_SCORE']) / 2.0
        score['SCORE'] = score['SCORE'].rank()
        score = score.sort_values('SCORE', ascending=False)
        pe_induatry = list(score.sort_values('PE_SCORE', ascending=False).index)[:5]
        fmt_induatry = list(score.sort_values('FMT_SCORE', ascending=False).index)[:5]
        induatry = list(score.sort_values('SCORE', ascending=False).index)[:5]
        score.columns = ['估值面打分', '基本面打分', '总分']
        score = score.T.reset_index().T.reset_index()
        score.iloc[0, 0] = ''
        score_text = '从打分情况可以看出，估值维度打分前5的行业为{0}，基本面维度打分前5的行业为{1}，从总打分排名看，我们重点关注的行业为{2}'.format('、'.join(pe_induatry), '、'.join(fmt_induatry),'、'.join(induatry))
        document.add_paragraph(text=score_text)
        rowc, colc = score.shape[0], score.shape[1]
        table = document.add_table(rows=rowc, cols=colc, style='Medium Grid 3')
        for i in range(rowc):
            for j in range(colc):
                table.cell(i, j).text = str(score.iloc[i, j])
        document.add_paragraph(text='')

        for paragraph in document.paragraphs:
            if paragraph.style.name == 'Normal' and paragraph.text != '':
                paragraph.paragraph_format.first_line_indent = paragraph.style.font.size * 2

        document.save('{0}行业数据跟踪周报_{1}.docx'.format(self.file_path, self.date))

        sw1_head_disp = sw1_head_pe_disp.merge(sw1_head_fmt_disp, on=['股票代码'], how='left')
        # sw3_head_disp = sw3_head_pe_disp.merge(sw3_head_fmt_disp, on=['TICKER_SYMBOL'], how='left')
        excel_writer = pd.ExcelWriter('{0}行业数据_{1}.xlsx'.format(self.file_path, self.date), engine='xlsxwriter')
        sw1_head_disp.to_excel(excel_writer, sheet_name='申万一级行业', index=False)
        # sw3_head_disp.to_excel(excel_writer, sheet_name='申万三级关注行业', index=False)
        excel_writer.save()
        return

if __name__ == '__main__':
    today = '20220715'#datetime.today().strftime('%Y%m%d')
    calendar_df = HBDB().read_cal((datetime.strptime(today, '%Y%m%d') - timedelta(20)).strftime('%Y%m%d'), today)
    calendar_df = calendar_df.rename(columns={'JYRQ': 'CALENDAR_DATE', 'SFJJ': 'IS_OPEN', 'SFZM': 'IS_WEEK_END', 'SFYM': 'IS_MONTH_END'})
    calendar_df['CALENDAR_DATE'] = calendar_df['CALENDAR_DATE'].astype(str)
    calendar_df = calendar_df.sort_values('CALENDAR_DATE')
    calendar_df['IS_WEEK_END'] = calendar_df['IS_WEEK_END'].fillna(0).astype(int)
    week_trade_df = calendar_df[calendar_df['IS_WEEK_END'] == 1].rename(columns={'CALENDAR_DATE': 'TRADE_DATE'})
    run_date = week_trade_df[week_trade_df['TRADE_DATE'] <= today]['TRADE_DATE'].iloc[-1]
    last_run_date = week_trade_df[week_trade_df['TRADE_DATE'] <= today]['TRADE_DATE'].iloc[-2]
    if today == run_date:
        file_path = 'D:/Git/hbshare/hbshare/fe/xwq/data/industry_tracking/{}'.format(run_date)
        if not os.path.exists(file_path):
            os.makedirs(file_path)
        head = 5
        report_date = '20220331'
        report_date_list = ['20191231', '20201231', '20210331', '20210630', '20210930', '20211231', '20220331']
        IndustryTracking(file_path, run_date, last_run_date, report_date, report_date_list, head).get_report()